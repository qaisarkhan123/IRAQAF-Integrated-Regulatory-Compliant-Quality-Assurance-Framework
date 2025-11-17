# --- bootstrap: make project root importable ---
from __future__ import annotations
from contextlib import contextmanager
import os
import sys
from helpers import (
    _delta_badge,
    _human_size,
    _label_from_path,
    _hash_file,
    _normalize_clause_evidence,
    _sanitize_name,
    _evidence_count,
    preview_widget,
    ribbon_chart,
    _nearest_by_time,
    _load_json,
)
from audit_utils import audit_locked, hash_file, record_audit_event
from validation import SCHEMA_VERSION, validate_incidents, validate_report
from sklearn.inspection import permutation_importance
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score, roc_auc_score
from sklearn.model_selection import train_test_split
from streamlit_autorefresh import st_autorefresh
import yaml
import streamlit as st
import requests
import pandas as pd
import numpy as np
import altair as alt
from dotenv import load_dotenv
import io
import re
import time
import math
import json
import json as _json
import shutil
import glob
import hashlib
import random
import logging
import traceback
import importlib.util
from typing import Dict, Any, Optional, List, Iterable, Tuple, Union
from logging.handlers import RotatingFileHandler
from collections import defaultdict
from datetime import datetime
from io import BytesIO
from contextlib import contextmanager  # ⬅️ ADD THIS LINE
from pathlib import Path
_CURR = os.path.dirname(os.path.abspath(__file__))
_ROOT = os.path.dirname(_CURR)
if _ROOT not in sys.path:
    sys.path.insert(0, _ROOT)

# Initialize logger early so it exists before performance import
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

# Try to import UI utilities (optional, will create later)
try:
    from ui_utils import show_success, show_error, show_warning, show_info
except ImportError:
    # Fallback to standard streamlit functions
    show_success = st.success
    def show_error(msg, details=None): return st.error(msg)
    show_warning = st.warning
    show_info = st.info
# Add to top of dashboard.py (after imports)


class DashboardError(Exception):
    """Base exception for dashboard errors"""
    pass


class DataLoadError(DashboardError):
    """Error loading data files"""
    pass


class ExportError(DashboardError):
    """Error during export operations"""
    pass

# Global error handler


def handle_dashboard_error(error: Exception, context: str = ""):
    """Centralized error handler with user-friendly messages"""
    error_id = f"ERR-{datetime.now().strftime('%Y%m%d%H%M%S')}"

    # Log error
    logger.error(f"{error_id} | {context}: {error}", exc_info=True)

    # User-friendly message
    st.error(f"❌ **Error {error_id}**")

    with st.expander("🔍 Error Details & Recovery", expanded=False):
        st.code(str(error), language="text")

        # Context-specific recovery tips
        if isinstance(error, FileNotFoundError):
            st.info("""
            **Recovery Tips:**
            - Verify the file exists in the expected location
            - Check file permissions (read access required)
            - Re-run evaluation to regenerate missing files
            """)
        elif isinstance(error, json.JSONDecodeError):
            st.info("""
            **Recovery Tips:**
            - File may be corrupted - try regenerating
            - Check for manual edits that broke JSON syntax
            - Use a JSON validator to identify syntax errors
            """)
        elif isinstance(error, PermissionError):
            st.info("""
            **Recovery Tips:**
            - Close programs using the file
            - Run with administrator privileges (Windows) or sudo (Linux/Mac)
            - Check folder write permissions
            """)
        else:
            st.info("""
            **Recovery Tips:**
            - Check logs/dashboard.log for details
            - Try refreshing the page
            - Contact support if issue persists
            """)

        st.caption(
            f"Error ID: `{error_id}` • Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # Copy error ID button
        if st.button("📋 Copy Error ID", key=f"copy_err_{error_id}"):
            st.code(error_id)
            st.caption("Share this ID when reporting the issue")


# -------------------------------------------------
# Guided Tour Step Definitions (must be defined once)
# -------------------------------------------------
TOUR_STEPS = [
    {"id": "module-summary", "title": "Module Summary"},
    {"id": "evidence-tray", "title": "📎 Evidence Tray"},
    {"id": "l1-governance", "title": "🏛️ L1 Governance – Compliance Coverage"},
    {"id": "l2-privacy", "title": "🔒 L2 Privacy/Security – Trends"},
    {"id": "l3-fairness", "title": "🧮 L3 Fairness – Evolution Across Runs"},
    {"id": "l4-explainability", "title": "🔍 L4 Explainability – Interactive"},
    {"id": "l5-operations", "title": "⚙️ L5 — Operations & Live Monitoring"},
]
# ============================================================================
# CORE UTILITY FUNCTIONS (Consolidated - Single Source of Truth)
# ============================================================================


def load_json(path: str) -> dict:
    """Load JSON file safely with error handling"""
    try:
        with open(path, "r", encoding="utf-8") as fh:
            return json.load(fh)
    except FileNotFoundError:
        logger.warning(f"File not found: {path}")
        return {}
    except json.JSONDecodeError as e:
        logger.error(f"Invalid JSON in {path}: {e}")
        return {}
    except Exception as e:
        logger.error(f"Failed to load {path}: {e}", exc_info=True)
        return {}


def show_error_inline(e: Exception, context: str):
    """
    Show a concise inline error message with helpful recovery tips.

    Args:
        e: The exception that occurred
        context: User-friendly context describing what failed
    """
    tb = "".join(traceback.format_exception_only(type(e), e)).strip()
    error_msg = f"{context}: {tb}"

    # Log full traceback
    logger.error(error_msg, exc_info=True)

    # Show user-friendly message with recovery tips
    st.error(f"❌ **{context}**")

    with st.expander("🔍 Error Details & Recovery Tips"):
        st.code(tb, language="text")

        # Provide context-specific recovery tips
        if "FileNotFoundError" in str(type(e)):
            st.info("""
            **Recovery Tips:**
            - Check if the file exists in the expected location
            - Verify file permissions
            - Re-run the evaluation to generate missing reports
            """)
        elif "JSONDecodeError" in str(type(e)):
            st.info("""
            **Recovery Tips:**
            - The file may be corrupted
            - Try regenerating the report
            - Check if the file was manually edited
            """)
        elif "PermissionError" in str(type(e)):
            st.info("""
            **Recovery Tips:**
            - Check file permissions
            - Close any programs using the file
            - Run as administrator (Windows) or with sudo (Linux/Mac)
            """)
        else:
            st.info("""
            **Recovery Tips:**
            - Check the logs for more details
            - Try refreshing the dashboard
            - Contact support if the issue persists
            """)

        # Show timestamp
        st.caption(
            f"Error occurred at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")


def load_evidence_index(path: str = "configs/evidence_index.json") -> dict:
    """Load evidence index with error handling"""
    try:
        p = Path(path)
        if not p.exists():
            logger.info(
                f"Evidence index not found at {path}, creating empty index")
            return {}
        with p.open("r", encoding="utf-8") as fh:
            data = json.load(fh)
        return data if isinstance(data, dict) else {}
    except Exception as e:
        logger.warning(f"Failed to load evidence index {path}: {e}")
        return {}


# Performance monitoring
try:
    from performance import (
        get_monitor,
        timed,
        measure_section,
        show_performance_panel,
        get_memory_usage
    )
    PERFORMANCE_MONITORING = True
    logger.info("Performance monitoring enabled")
except ImportError:
    # Fallback: no-op decorators if performance.py not available
    def timed(name=None):
        def decorator(func):
            return func
        return decorator

    class measure_section:
        def __init__(self, name):
            pass

        def __enter__(self):
            return self

        def __exit__(self, *args):
            pass

    def show_performance_panel():
        st.info("Performance monitoring module not available")

    def get_memory_usage():
        return {}

    PERFORMANCE_MONITORING = False
    logger.warning(
        "Performance monitoring disabled (performance.py not found)")

# create or configure logger
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

logger.info("Logger initialized")

# stdlib

# third-party

# local

# legacy alias kept for compatibility

# yaml alias used by legacy code
try:
    import yaml as _yaml
except Exception:
    _yaml = None

load_dotenv()

# logging
os.makedirs("logs", exist_ok=True)
logger = logging.getLogger("iraqaf_dashboard")
logger.setLevel(logging.INFO)

file_handler = RotatingFileHandler(
    "logs/dashboard.log",
    maxBytes=10 * 1024 * 1024,
    backupCount=5,
    encoding="utf-8",
)
file_handler.setLevel(logging.INFO)

console_handler = logging.StreamHandler()
console_handler.setLevel(logging.WARNING)

formatter = logging.Formatter(
    "%(asctime)s - %(name)s - %(levelname)s - %(funcName)s:%(lineno)d - %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
file_handler.setFormatter(formatter)
console_handler.setFormatter(formatter)

logger.addHandler(file_handler)
logger.addHandler(console_handler)

logger.info("=" * 50)
logger.info("IRAQAF Dashboard started")
logger.info("=" * 50)
# ============================================================================
# INPUT VALIDATION & SANITIZATION
# ============================================================================


def validate_and_sanitize_input(
    value: str,
    input_type: str = "general",
    max_length: int = 1000,
    allow_special_chars: bool = False
) -> tuple[bool, str, str]:
    """
    Validate and sanitize user input with security checks.

    Args:
        value: Input string to validate
        input_type: Type of input (email, url, filename, general)
        max_length: Maximum allowed length
        allow_special_chars: Whether to allow special characters

    Returns:
        Tuple of (is_valid, sanitized_value, error_message)
    """
    import re
    from urllib.parse import urlparse

    if not value:
        return True, "", ""

    # Check length
    if len(value) > max_length:
        return False, value[:max_length], f"Input too long (max {max_length} characters)"

    # Type-specific validation
    if input_type == "email":
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if not re.match(email_pattern, value):
            return False, value, "Invalid email format"
        return True, value.strip().lower(), ""

    elif input_type == "url":
        try:
            result = urlparse(value)
            if not all([result.scheme, result.netloc]):
                return False, value, "Invalid URL format"
            if result.scheme not in ['http', 'https']:
                return False, value, "Only HTTP/HTTPS URLs allowed"
            return True, value.strip(), ""
        except Exception:
            return False, value, "Invalid URL"

    elif input_type == "filename":
        # Remove path separators and dangerous characters
        dangerous_chars = ['/', '\\', '..', '<',
                           '>', '|', ':', '*', '?', '"', '\x00']
        sanitized = value
        for char in dangerous_chars:
            sanitized = sanitized.replace(char, '_')

        # Check for reserved Windows names
        reserved_names = ['CON', 'PRN', 'AUX',
                          'NUL', 'COM1', 'COM2', 'LPT1', 'LPT2']
        if sanitized.upper() in reserved_names:
            sanitized = f"_{sanitized}"

        return True, sanitized[:255], ""  # Max filename length

    elif input_type == "general":
        # Remove potential XSS/injection characters
        dangerous_patterns = [
            r'<script[^>]*>.*?</script>',  # Script tags
            r'javascript:',                 # JavaScript protocol
            r'on\w+\s*=',                  # Event handlers
            r'<iframe[^>]*>',              # Iframes
        ]

        sanitized = value
        for pattern in dangerous_patterns:
            sanitized = re.sub(pattern, '', sanitized, flags=re.IGNORECASE)

        if not allow_special_chars:
            # Keep only alphanumeric, spaces, and basic punctuation
            sanitized = re.sub(r'[^\w\s.,;:!?()\-]', '', sanitized)

        return True, sanitized.strip(), ""

    return True, value, ""

# ============================================================================
# RATE LIMITING
# ============================================================================


class RateLimiter:
    """Simple rate limiter for API calls."""

    def __init__(self):
        self.calls = {}
        self.limits = {
            'prometheus': {'max_calls': 60, 'window_seconds': 60},
            'file_upload': {'max_calls': 10, 'window_seconds': 60},
            'export': {'max_calls': 20, 'window_seconds': 60},
        }

    def is_allowed(self, action: str, identifier: str = "default") -> tuple[bool, str]:
        """
        Check if action is allowed under rate limits.

        Args:
            action: Type of action (prometheus, file_upload, export)
            identifier: Unique identifier for this user/session

        Returns:
            Tuple of (allowed, error_message)
        """
        if action not in self.limits:
            return True, ""

        limit_config = self.limits[action]
        max_calls = limit_config['max_calls']
        window = limit_config['window_seconds']

        key = f"{action}:{identifier}"
        current_time = time.time()

        # Initialize or clean old entries
        if key not in self.calls:
            self.calls[key] = []

        # Remove calls outside the window
        self.calls[key] = [t for t in self.calls[key]
                           if current_time - t < window]

        # Check limit
        if len(self.calls[key]) >= max_calls:
            retry_after = int(window - (current_time - self.calls[key][0]))
            return False, f"Rate limit exceeded. Retry in {retry_after}s"

        # Record this call
        self.calls[key].append(current_time)
        return True, ""


# Initialize global rate limiter
if 'rate_limiter' not in st.session_state:
    st.session_state.rate_limiter = RateLimiter()

rate_limiter = st.session_state.rate_limiter


def safe_json_loads(json_string: str, max_size_mb: int = 10) -> tuple[bool, Any, str]:
    """
    Safely load JSON with size limits and error handling.

    Args:
        json_string: JSON string to parse
        max_size_mb: Maximum size in megabytes

    Returns:
        Tuple of (success, data, error_message)
    """
    try:
        # Check size
        size_mb = len(json_string.encode('utf-8')) / (1024 * 1024)
        if size_mb > max_size_mb:
            return False, None, f"JSON too large ({size_mb:.1f}MB > {max_size_mb}MB limit)"

        # Parse JSON
        data = json.loads(json_string)
        return True, data, ""

    except json.JSONDecodeError as e:
        return False, None, f"Invalid JSON: {str(e)}"
    except Exception as e:
        return False, None, f"Error loading JSON: {str(e)}"


st.set_page_config(page_title="IRAQAF Dashboard", layout="wide")
# ============================================================================
# DATA PROCESSING HELPERS
# ============================================================================


def collect_module_history(files: list, module_prefix: str, extract_metrics: callable) -> pd.DataFrame:
    """
    Efficiently collect historical data for a module.

    Args:
        files: List of report file paths
        module_prefix: Module identifier (e.g., "L3", "L1")
        extract_metrics: Function to extract metrics from report dict

    Returns:
        DataFrame with historical data, sorted by time
    """
    rows = []

    for f in files:
        if not os.path.basename(f).startswith(f"{module_prefix}-"):
            continue

        try:
            with open(f, "r", encoding="utf-8") as fh:
                report = json.load(fh)

            if report.get("module") != module_prefix:
                continue

            # Extract common metadata
            base = os.path.basename(f).replace(".json", "")
            ts_match = re.search(r"(\d{8}-\d{6})", base)
            ts_label = ts_match.group(1) if ts_match else base
            mtime = os.path.getmtime(f)

            # Use custom extractor for module-specific metrics
            row_data = extract_metrics(report)
            row_data.update({
                "Run": base,
                "Label": ts_label,
                "time": mtime,
            })

            rows.append(row_data)

        except Exception as e:
            logger.debug(f"Skipped {f}: {e}")
            continue

    if not rows:
        return pd.DataFrame()

    df = pd.DataFrame(rows).sort_values("time").reset_index(drop=True)
    df["Run #"] = df.index + 1

    return df


@contextmanager
def progress_tracker(message: str, total: int = None):
    """
    Context manager for showing progress during long operations.

    Usage:
        with progress_tracker("Processing files", total=100) as tracker:
            for i in range(100):
                tracker.update(i + 1)
                # do work
    """
    if total:
        progress_bar = st.progress(0, text=message)
        status_text = st.empty()

        class Tracker:
            def update(self, current: int, custom_message: str = None):
                progress = current / total
                progress_bar.progress(
                    progress, text=custom_message or f"{message} ({current}/{total})")

        try:
            yield Tracker()
        finally:
            progress_bar.empty()
            status_text.empty()
    else:
        # Simple spinner for indeterminate progress
        with st.spinner(message):
            yield type('Tracker', (), {'update': lambda *args: None})()


@contextmanager
def progress_tracker(message: str, total: int = None):
    """
    Context manager for showing progress during long operations.

    Usage:
        with progress_tracker("Processing files", total=100) as tracker:
            for i in range(100):
                tracker.update(i + 1)
                # do work
    """
    if total:
        progress_bar = st.progress(0, text=message)
        status_text = st.empty()

        class Tracker:
            def update(self, current: int, custom_message: str = None):
                progress = current / total
                progress_bar.progress(
                    progress, text=custom_message or f"{message} ({current}/{total})")

        try:
            yield Tracker()
        finally:
            progress_bar.empty()
            status_text.empty()
    else:
        # Simple spinner for indeterminate progress
        with st.spinner(message):
            yield type('Tracker', (), {'update': lambda *args: None})()
# ============================================================================
# SESSION STATE INITIALIZATION
# ============================================================================


def init_session_state():
    """Initialize all session state variables with defaults"""
    defaults = {
        "__audit_mode__": False,
        "theme_mode": "Auto",
        "high_contrast": False,
        "search_history": [],
        "last_refresh": None,
        "user_preferences": {},
        "notification_settings": {
            "slack_url": "",
            "email_to": "",
            "smtp_host": "",
            "smtp_user": "",
        },
        "_incident_lock": None,  # Will be set to threading.Lock() when needed
    }

    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

    # Initialize threading lock for incident logging
    if st.session_state["_incident_lock"] is None:
        import threading
        st.session_state["_incident_lock"] = threading.Lock()


# Initialize session state BEFORE anything else
init_session_state()

# NOW the audit mode check works reliably


def get_audit_mode():
    """
    Placeholder for audit mode check.
    Returns False by default (i.e., unlocked / editable).
    """
    return False


_LOCK = get_audit_mode()

# Add after session state initialization


@contextmanager
def show_loading(message: str, context: str = ""):
    """Show loading spinner with optional context inside a context manager."""
    placeholder = st.empty()

    if context:
        placeholder.caption(context)

    try:
        with st.spinner(f"{message}..."):
            yield
    finally:
        placeholder.empty()

# ===== Status Banner =====


def show_status_banner():
    """Show system status banner at the top."""

    # Check if reports are recent (< 24 hours old)
    import glob
    from datetime import datetime, timedelta

    try:
        report_files = glob.glob("reports/*.json")
        if report_files:
            newest_report = max(report_files, key=os.path.getmtime)
            mtime = datetime.fromtimestamp(os.path.getmtime(newest_report))
            age = datetime.now() - mtime

            if age < timedelta(hours=24):
                status_color = "#d4edda"
                status_icon = "✅"
                status_text = f"Reports are up to date (last updated {age.seconds // 3600}h ago)"
            elif age < timedelta(days=7):
                status_color = "#fff3cd"
                status_icon = "⚠️"
                status_text = f"Reports are {age.days} days old - consider re-running evaluation"
            else:
                status_color = "#f8d7da"
                status_icon = "❌"
                status_text = f"Reports are {age.days} days old - please re-run evaluation"

            st.markdown(
                f"""
                <div style='
                    background: {status_color};
                    padding: 10px 20px;
                    border-radius: 8px;
                    margin-bottom: 20px;
                    text-align: center;
                    border: 1px solid rgba(0,0,0,0.1);
                '>
                    <span style='font-size: 1.1rem;'>{status_icon} {status_text}</span>
                </div>
                """,
                unsafe_allow_html=True
            )
    except Exception as e:
        logger.error(f"Failed to show status banner: {e}")


show_status_banner()

trace_map: Dict[str, Any] = {}


def _load_evidence_index(path: str = "configs/evidence_index.json") -> dict:
    try:
        p = Path(path)
        if not p.exists():
            return {}
        with p.open("r", encoding="utf-8") as fh:
            data = _json.load(fh)
        return data if isinstance(data, dict) else {}
    except Exception as e:
        try:
            logger.warning(f"Failed to load evidence index {path}: {e}")
        except Exception:
            pass
        return {}


def _show_error_inline(e: Exception, context: str):
    tb = "".join(traceback.format_exception_only(type(e), e)).strip()
    st.error(f"⚠️ {context}: {tb}")


# --------------------------------------------------------------------
# ⚙️ LOAD POLICIES (thresholds, refresh interval)
# --------------------------------------------------------------------


def load_policies(path: str = "configs/policies.yaml") -> Dict[str, Any]:
    """
    Load and validate policy configuration.
    """
    defaults = {
        "latency_slo": 1.0,
        "error_rate_threshold": 5.0,
        "refresh_interval_s": 60
    }

    if not os.path.exists(path):
        logger.warning(f"Policy file not found: {path}. Using defaults.")
        return defaults

    try:
        with open(path, "r", encoding="utf-8") as fh:
            raw = yaml.safe_load(fh) or {}

        policy: Dict[str, Any] = {}

        # latency_slo: must be positive float
        try:
            policy["latency_slo"] = float(
                raw.get("latency_slo", defaults["latency_slo"]))
            if policy["latency_slo"] <= 0:
                logger.warning(
                    f"Invalid latency_slo: {policy['latency_slo']}. Using default.")
                policy["latency_slo"] = defaults["latency_slo"]
        except (ValueError, TypeError):
            logger.warning("Invalid latency_slo type. Using default.")
            policy["latency_slo"] = defaults["latency_slo"]

        # error_rate_threshold: must be 0-100
        try:
            policy["error_rate_threshold"] = float(
                raw.get("error_rate_threshold",
                        defaults["error_rate_threshold"])
            )
            if not (0 <= policy["error_rate_threshold"] <= 100):
                logger.warning(
                    f"Invalid error_rate_threshold: {policy['error_rate_threshold']}. Using default."
                )
                policy["error_rate_threshold"] = defaults["error_rate_threshold"]
        except (ValueError, TypeError):
            logger.warning("Invalid error_rate_threshold type. Using default.")
            policy["error_rate_threshold"] = defaults["error_rate_threshold"]

        # refresh_interval_s: must be 5-3600
        try:
            policy["refresh_interval_s"] = int(
                raw.get("refresh_interval_s", defaults["refresh_interval_s"])
            )
            if not (5 <= policy["refresh_interval_s"] <= 3600):
                logger.warning(
                    f"Invalid refresh_interval_s: {policy['refresh_interval_s']}. Using default."
                )
                policy["refresh_interval_s"] = defaults["refresh_interval_s"]
        except (ValueError, TypeError):
            logger.warning("Invalid refresh_interval_s type. Using default.")
            policy["refresh_interval_s"] = defaults["refresh_interval_s"]

        logger.info(f"Loaded policies: {policy}")
        return policy

    except yaml.YAMLError as e:
        logger.error(f"YAML parse error in {path}: {e}")
        st.error(
            f"⚠️ Failed to parse policies.yaml. Using defaults. Error: {e}")
        return defaults
    except Exception as e:
        logger.error(f"Unexpected error loading policies: {e}", exc_info=True)
        return defaults


POLICY = load_policies()
logger.info(f"Active policy configuration: {POLICY}")

# --------------------------------------------------------------------
# 🧩 OPTIONAL DEPENDENCY CHECK
# --------------------------------------------------------------------
OPTIONAL_DEPS = {
    "matplotlib": "pip install matplotlib",
    "pdfkit": "pip install pdfkit",
    "yaml": "pip install pyyaml",
    "docx": "pip install python-docx"
}

with st.sidebar.expander("🧩 Dependencies"):
    for lib, cmd in OPTIONAL_DEPS.items():
        spec = importlib.util.find_spec(lib)
        if spec is None:
            st.error(f"❌ {lib} not installed — `{cmd}`")
        else:
            st.success(f"✅ {lib}")


# ============================================================================
# AUDIT MODE MANAGEMENT
# ============================================================================

def get_audit_mode() -> bool:
    """Centralized audit mode state getter"""
    return st.session_state.get("__audit_mode__", False)


def set_audit_mode(enabled: bool):
    """Centralized audit mode state setter"""
    st.session_state["__audit_mode__"] = enabled
    record_audit_event("ENABLED" if enabled else "DISABLED",
                       get_current_run_hash())


def get_current_run_hash() -> str:
    """Get hash of current aggregate report for audit trail"""
    try:
        agg_paths = sorted(glob.glob("reports/AGG-*.json"),
                           key=os.path.getmtime)
        if agg_paths:
            latest_agg = agg_paths[-1]
            return _hash_file(latest_agg, _mtime=os.path.getmtime(latest_agg))
    except Exception as e:
        logger.error(f"Failed to get run hash: {e}")
    return "unknown"


# Alias for backward compatibility
def audit_locked() -> bool:
    """Check if audit mode is active (locks all inputs)"""
    return get_audit_mode()


# --------------------------------------------------------------------
# 🔒 AUDIT LOCK BOOTSTRAP
# --------------------------------------------------------------------
# Ensure the audit flag exists before any widgets reference _LOCK.
if "__audit_mode__" not in st.session_state:
    st.session_state["__audit_mode__"] = False

# Convenience variable used throughout the app to disable widgets.
_LOCK = get_audit_mode()

# --------------------------------------------------------------------
# 🧩 DATA VALIDATION + THEME + AUDIT LOG (SIDEBAR CONFIG)
# --------------------------------------------------------------------
# This unified sidebar section handles:
#   1. Schema version display
#   2. Theme switcher (Auto / Light / Dark + High Contrast)
#   3. Audit log viewer with download option
# --------------------------------------------------------------------

# ===== 1️⃣ DATA VALIDATION INFO =====
with st.sidebar.expander("🧩 Data Validation", expanded=False):
    st.markdown(f"Schema version: **{SCHEMA_VERSION}**")
    st.caption("Reports automatically validated on load.")

# ===== 2️⃣ THEME SWITCHER (Auto / Light / Dark + High Contrast) =====


def _inject_theme_css(mode: str):
    """Inject CSS variables for light/dark themes with system auto support."""
    base_css = """
    <style>
      :root {
        --bg: #ffffff;
        --text: #0f172a;
        --muted: #6b7280;
        --card: #ffffff;
        --border: #e5e7eb;
        --shadow: 0 1px 2px rgba(0,0,0,0.04);
        --btn-bg: #f3f4f6;
        --btn-text: #111827;
        --btn-border: #d1d5db;
        --link: #2563eb;
        --metric-bg: transparent;
        --df-header-bg: #f8fafc;
        --df-border: #e5e7eb;
        --expander-bg: #ffffff;
      }
      .stApp, body { background: var(--bg); color: var(--text); }
      a, .markdown-text-container a { color: var(--link); }
      .iraqaf-card {
        background: var(--card);
        border: 1px solid var(--border);
        border-radius: 12px;
        padding: 14px;
        margin-bottom: 12px;
        box-shadow: var(--shadow);
      }
      .stButton>button, .stDownloadButton>button {
        background: var(--btn-bg) !important;
        color: var(--btn-text) !important;
        border: 1px solid var(--btn-border) !important;
        border-radius: 10px !important;
      }
      [data-testid="stTextInput"] input,
      [data-testid="stNumberInput"] input,
      [data-testid="stSelectbox"] div[data-baseweb="select"] > div,
      [data-testid="stFileUploader"] {
        color: var(--text);
        background: var(--card);
        border: 1px solid var(--border);
      }
      [data-testid="stMetricValue"], [data-testid="stMetricDelta"] { color: var(--text); }
      .stDataFrame thead tr th { background: var(--df-header-bg); color: var(--text); }
      .stDataFrame table, .stDataFrame td, .stDataFrame th { border-color: var(--df-border) !important; }
      [data-testid="stExpander"] { background: var(--expander-bg); border-radius:12px; }
      .vega-embed, canvas { background: transparent !important; }
    </style>
    """
    dark_overrides = """
    <style>
      :root {
        --bg: #0f1116;
        --text: #e6e6e6;
        --muted: #9aa3b2;
        --card: #151922;
        --border: #2a2f3a;
        --shadow: 0 1px 2px rgba(0,0,0,0.3);
        --btn-bg: #1f2937;
        --btn-text: #e6e6e6;
        --btn-border: #374151;
        --link: #8ab4ff;
        --metric-bg: rgba(255,255,255,0.02);
        --df-header-bg: #1e2430;
        --df-border: #2a2f3a;
        --expander-bg: #151922;
      }
    </style>
    """
    auto_block = f"""
    <style>
      @media (prefers-color-scheme: dark) {{
        {dark_overrides.replace('<style>', '').replace('</style>', '')}
      }}
    </style>
    """
    if mode == "Light":
        st.markdown(base_css, unsafe_allow_html=True)
    elif mode == "Dark":
        st.markdown(base_css + dark_overrides, unsafe_allow_html=True)
    else:
        st.markdown(base_css + auto_block, unsafe_allow_html=True)


# --- Theme Expander ---
with st.sidebar.expander("🌓 Theme", expanded=False):
    mode = st.radio(
        "Mode", ["Auto", "Light", "Dark"],
        horizontal=True,
        index=0,
        disabled=_LOCK,
        key="theme_mode_selector"
    )

    # Apply base theme
    _inject_theme_css(mode)

    # Accessibility: High contrast mode toggle
    high_contrast = st.checkbox(
        "🔳 High contrast mode",
        key="high_contrast_toggle",
        disabled=_LOCK,
        help="Boosts text contrast and slightly increases base font size for readability."
    )
    if high_contrast:
        st.markdown("""
        <style>
          :root {
            --text: #ffffff !important;
            --muted: #d9d9d9 !important;
          }
          html, body, .stApp { font-size: 1.05rem !important; }
          [data-testid="stTextInput"] input,
          [data-testid="stNumberInput"] input,
          [data-testid="stSelectbox"] div[data-baseweb="select"] > div {
            border-color: var(--link) !important;
          }
        </style>
        """, unsafe_allow_html=True)
# ===== Keyboard Shortcuts =====
st.markdown("""
<script>
// Keyboard shortcuts for IRAQAF dashboard
document.addEventListener('keydown', function(e) {
    // Ctrl/Cmd + K: Focus search
    if ((e.ctrlKey || e.metaKey) && e.key === 'k') {
        e.preventDefault();
        const searchInputs = document.querySelectorAll('input[type="text"]');
        if (searchInputs.length > 0) {
            searchInputs[0].focus();
        }
    }

    // Ctrl/Cmd + R: Refresh (with confirmation)
    if ((e.ctrlKey || e.metaKey) && e.key === 'r') {
        e.preventDefault();
        if (confirm('Refresh dashboard?')) {
            window.location.reload();
        }
    }

    // Escape: Close modals/expandable sections
    if (e.key === 'Escape') {
        const expandedSections = document.querySelectorAll('[data-testid="stExpander"][aria-expanded="true"]');
        expandedSections.forEach(section => {
            const button = section.querySelector('button');
            if (button) button.click();
        });
    }

    // ? key: Show keyboard shortcuts help
    if (e.key === '?') {
        e.preventDefault();
        alert(
            'IRAQAF Keyboard Shortcuts:\\n\\n' +
            'Ctrl/Cmd + K  →  Focus search\\n' +
            'Ctrl/Cmd + R  →  Refresh dashboard\\n' +
            'Escape        →  Close expanded sections\\n' +
            '?             →  Show this help'
        );
    }
});

// Show keyboard shortcut hint on first load
if (!sessionStorage.getItem('shortcuts_shown')) {
    sessionStorage.setItem('shortcuts_shown', 'true');

    // Create tooltip
    const tooltip = document.createElement('div');
    tooltip.innerHTML = '💡 Press <kbd>?</kbd> for keyboard shortcuts';
    tooltip.style.cssText = `
        position: fixed;
        bottom: 20px;
        left: 20px;
        background: #1f2937;
        color: white;
        padding: 12px 16px;
        border-radius: 8px;
        font-size: 14px;
        z-index: 9999;
        animation: fadeInOut 4s ease-in-out;
    `;

    document.body.appendChild(tooltip);

    setTimeout(() => tooltip.remove(), 4000);
}
</script>

<style>
@keyframes fadeInOut {
    0%, 100% { opacity: 0; }
    10%, 90% { opacity: 1; }
}

kbd {
    background: #374151;
    padding: 2px 6px;
    border-radius: 4px;
    font-family: monospace;
    font-size: 0.9em;
}
</style>
""", unsafe_allow_html=True)


# ===== Responsive Design Enhancements =====
st.markdown("""
<style>
/* Mobile-friendly adjustments */
@media (max-width: 768px) {
    /* Stack columns on mobile */
    [data-testid="column"] {
        width: 100% !important;
        flex: none !important;
    }

    /* Larger tap targets for mobile */
    button {
        min-height: 44px !important;
        font-size: 16px !important;
    }

    /* Better text readability */
    .stMarkdown {
        font-size: 16px !important;
    }

    /* Hide less important elements on small screens */
    .mobile-hide {
        display: none !important;
    }
}

/* Tablet adjustments */
@media (min-width: 769px) and (max-width: 1024px) {
    [data-testid="stSidebar"] {
        width: 250px !important;
    }
}

/* Better scrolling for tables */
.stDataFrame {
    max-height: 600px;
    overflow-y: auto;
}

/* Smooth scrolling */
html {
    scroll-behavior: smooth;
}

/* Better focus indicators for accessibility */
*:focus-visible {
    outline: 2px solid var(--link) !important;
    outline-offset: 2px !important;
    border-radius: 4px;
}

/* Loading animation */
@keyframes pulse {
    0%, 100% { opacity: 1; }
    50% { opacity: 0.5; }
}

.loading {
    animation: pulse 2s cubic-bezier(0.4, 0, 0.6, 1) infinite;
}

/* Improved card shadows on hover */
.iraqaf-card:hover {
    box-shadow: 0 4px 12px rgba(0,0,0,0.1);
    transition: box-shadow 0.3s ease;
}

/* Better spacing for metrics */
[data-testid="stMetricValue"] {
    font-size: 2rem !important;
    font-weight: 600 !important;
}

/* Prettier scrollbars (Webkit browsers) */
::-webkit-scrollbar {
    width: 10px;
    height: 10px;
}

::-webkit-scrollbar-track {
    background: var(--bg);
}

::-webkit-scrollbar-thumb {
    background: var(--border);
    border-radius: 5px;
}

::-webkit-scrollbar-thumb:hover {
    background: var(--muted);
}
</style>
""", unsafe_allow_html=True)
# --- Focus outlines for accessibility (global, after theme) ---
st.markdown("""
<style>
button:focus, input:focus, textarea:focus, select:focus,
[data-testid="stTextInput"] input:focus,
[data-testid="stNumberInput"] input:focus,
[data-testid="stSelectbox"] [data-baseweb="select"]:focus,
[data-testid="stFileUploader"] [data-testid="stDropzone"] {
  outline: 2px solid var(--link) !important;
  outline-offset: 2px !important;
}
</style>
""", unsafe_allow_html=True)

# ===== 3️⃣ AUDIT LOG VIEWER =====
with st.sidebar.expander("🪵 Audit Log", expanded=False):
    log_path = "logs/audit.log"
    if os.path.exists(log_path):
        with open(log_path, "r", encoding="utf-8") as fh:
            content = fh.read().strip()
        if content:
            st.text(content)
            st.download_button(
                "⬇️ Download audit log",
                data=content.encode("utf-8"),
                file_name="audit.log",
                mime="text/plain",
                disabled=_LOCK
            )
        else:
            st.caption("Audit log is empty.")
    else:
        st.caption("No audit events recorded yet.")


# ===== 4️⃣ APPLICATION LOG VIEWER =====
with st.sidebar.expander("📜 Application Logs", expanded=False):
    log_file = "logs/app.log"
    if os.path.exists(log_file):
        with open(log_file, "r", encoding="utf-8") as fh:
            content = fh.read().strip()
        if content:
            st.text(content[-4000:])  # show last few KB for readability
            st.download_button(
                "⬇️ Download app.log",
                data=open(log_file, "rb").read(),
                file_name="app.log",
                mime="text/plain",
                disabled=_LOCK,
            )
        else:
            st.caption("Log file is empty.")
    else:
        st.caption("No logs yet.")
# =============================================================================
#  ⚙️ SETTINGS PANEL
# =============================================================================

with st.sidebar.expander("⚙️ Settings", expanded=False):
    st.markdown("**Display Preferences**")

    # Auto-refresh setting
    auto_refresh = st.checkbox(
        "Auto-refresh data",
        value=st.session_state.get('auto_refresh_enabled', False),
        key="auto_refresh_toggle",
        disabled=_LOCK,
        help="Automatically refresh data every 5 minutes"
    )
    st.session_state['auto_refresh_enabled'] = auto_refresh

    if auto_refresh:
        st.caption("⏱️ Data refreshes every 5 minutes")
        # Trigger auto-refresh
        st_autorefresh(interval=300000, key="settings_autorefresh")

    # Compact mode
    compact_mode = st.checkbox(
        "Compact mode",
        value=st.session_state.get('compact_mode', False),
        key="compact_mode_toggle",
        disabled=_LOCK,
        help="Reduce spacing for more content on screen"
    )
    st.session_state['compact_mode'] = compact_mode

    # Apply compact mode CSS
    if compact_mode:
        st.markdown("""
        <style>
        /* Compact mode overrides */
        .stMarkdown, .stText { margin-bottom: 0.5rem !important; }
        .stExpander { margin-bottom: 0.5rem !important; }
        [data-testid="stVerticalBlock"] > div { gap: 0.5rem !important; }
        .iraqaf-card { padding: 10px !important; margin-bottom: 8px !important; }
        h1, h2, h3, h4 { margin-top: 0.5rem !important; margin-bottom: 0.5rem !important; }
        .stDataFrame { max-height: 400px !important; }
        </style>
        """, unsafe_allow_html=True)

    # Show timestamps
    show_timestamps = st.checkbox(
        "Show timestamps",
        value=st.session_state.get('show_timestamps', True),
        key="show_timestamps_toggle",
        disabled=_LOCK,
        help="Display when each section was last updated"
    )
    st.session_state['show_timestamps'] = show_timestamps

    # Default chart style
    st.markdown("**Chart Settings**")
    chart_theme = st.selectbox(
        "Chart theme",
        options=["Default", "Dark", "Light"],
        index=0,
        key="chart_theme_select",
        disabled=_LOCK
    )
    st.session_state['chart_theme'] = chart_theme

    # Animation toggle
    enable_animations = st.checkbox(
        "Enable animations",
        value=st.session_state.get('enable_animations', True),
        key="animations_toggle",
        disabled=_LOCK,
        help="Smooth transitions and hover effects"
    )
    st.session_state['enable_animations'] = enable_animations

    if not enable_animations:
        st.markdown("""
        <style>
        * { transition: none !important; animation: none !important; }
        </style>
        """, unsafe_allow_html=True)

    # Notification preferences
    st.markdown("**Notifications**")
    notify_on_breach = st.checkbox(
        "Alert on SLO breaches",
        value=st.session_state.get('notify_on_breach', True),
        key="notify_breach_toggle",
        disabled=_LOCK
    )
    st.session_state['notify_on_breach'] = notify_on_breach

    # Data display preferences
    st.markdown("**Data Display**")
    decimal_places = st.slider(
        "Decimal places",
        min_value=0, max_value=4, value=2,
        key="decimal_places_slider",
        disabled=_LOCK,
        help="Number of decimals to show in metrics"
    )
    st.session_state['decimal_places'] = decimal_places

    # Export settings to JSON
    st.markdown("---")
    if st.button("💾 Save Settings", key="save_settings_btn", disabled=_LOCK, use_container_width=True):
        settings = {
            'auto_refresh_enabled': auto_refresh,
            'compact_mode': compact_mode,
            'show_timestamps': show_timestamps,
            'chart_theme': chart_theme,
            'enable_animations': enable_animations,
            'notify_on_breach': notify_on_breach,
            'decimal_places': decimal_places,
            'theme_mode': st.session_state.get('theme_mode', 'Auto'),
            'high_contrast': st.session_state.get('high_contrast', False),
            'saved_at': datetime.now().isoformat()
        }

        settings_file = Path("configs/user_settings.json")
        settings_file.parent.mkdir(exist_ok=True)

        try:
            with open(settings_file, "w", encoding="utf-8") as f:
                json.dump(settings, f, indent=2)
            st.success("✅ Settings saved!")
            logger.info(f"User settings saved: {settings}")
        except Exception as e:
            st.error(f"Failed to save settings: {e}")

    # Load settings
    if st.button("📂 Load Settings", key="load_settings_btn", disabled=_LOCK, use_container_width=True):
        settings_file = Path("configs/user_settings.json")
        if settings_file.exists():
            try:
                with open(settings_file, "r", encoding="utf-8") as f:
                    settings = json.load(f)

                # Apply settings
                for key, value in settings.items():
                    if key != 'saved_at':
                        st.session_state[key] = value

                st.success(
                    f"✅ Settings loaded from {settings.get('saved_at', 'unknown time')}")
                st.rerun()
            except Exception as e:
                st.error(f"Failed to load settings: {e}")
        else:
            st.info("No saved settings found")

    # Reset to defaults
    if st.button("🔄 Reset to Defaults", key="reset_settings_btn", disabled=_LOCK, use_container_width=True):
        defaults = {
            'auto_refresh_enabled': False,
            'compact_mode': False,
            'show_timestamps': True,
            'chart_theme': 'Default',
            'enable_animations': True,
            'notify_on_breach': True,
            'decimal_places': 2,
            'theme_mode': 'Auto',
            'high_contrast': False
        }
        for key, value in defaults.items():
            st.session_state[key] = value
        st.success("✅ Settings reset to defaults")
        st.rerun()
# Add to sidebar (after Settings)

with st.sidebar.expander("❓ Help & Documentation", expanded=False):
    st.markdown("""
    ### Quick Reference
    
    **Navigation:**
    - Use the search bar (Ctrl+K) to find specific metrics
    - Click module names to jump to detailed sections
    - Enable tour mode for guided walkthrough
    
    **Keyboard Shortcuts:**
    - `Ctrl+K` - Focus search
    - `Ctrl+R` - Refresh dashboard
    - `Esc` - Close expanded sections
    - `?` - Show all shortcuts
    
    **Troubleshooting:**
    - If data is missing, check that reports were generated
    - Red modules indicate issues - see details in module section
    - Use Debug Mode (sidebar) for performance insights
    
    **Getting Help:**
    """)

    col1, col2 = st.columns(2)
    with col1:
        if st.button("📖 Full Docs", use_container_width=True):
            st.info("Opening documentation...")
            # Link to your docs

    with col2:
        if st.button("🐛 Report Issue", use_container_width=True):
            st.info("Opening issue tracker...")
            # Link to GitHub issues
# =============================================================================
#  🎯 TOUR MODE
# =============================================================================
with st.sidebar.expander("🎯 Interactive Tour", expanded=False):
    st.markdown("Take a guided tour of the dashboard features")

    if st.button("🚀 Start Tour", key="sidebar_start_tour", use_container_width=True, type="primary"):
        st.session_state['tour_active'] = True
        st.session_state['tour_step'] = 0
        st.rerun()

    if st.session_state.get('tour_active', False):
        st.info(
            f"📍 Tour in progress: Step {st.session_state.get('tour_step', 0) + 1}/{len(TOUR_STEPS)}")
# --------------------------------------------------------------------
# ⚠️ ERROR SURFACING HELPER
# --------------------------------------------------------------------


def show_error_inline(e: Exception, context: str):
    """
    Show a concise inline error message with helpful recovery tips.

    Args:
        e: The exception that occurred
        context: User-friendly context describing what failed
    """
    tb = "".join(traceback.format_exception_only(type(e), e)).strip()
    error_msg = f"{context}: {tb}"

    # Log full traceback
    logger.error(error_msg, exc_info=True)

    # Show user-friendly message with recovery tips
    st.error(f"❌ **{context}**")

    with st.expander("🔍 Error Details & Recovery Tips"):
        st.code(tb, language="text")

        # Provide context-specific recovery tips
        if "FileNotFoundError" in str(type(e)):
            st.info("""
            **Recovery Tips:**
            - Check if the file exists in the expected location
            - Verify file permissions
            - Re-run the evaluation to generate missing reports
            """)
        elif "JSONDecodeError" in str(type(e)):
            st.info("""
            **Recovery Tips:**
            - The file may be corrupted
            - Try regenerating the report
            - Check if the file was manually edited
            """)
        elif "PermissionError" in str(type(e)):
            st.info("""
            **Recovery Tips:**
            - Check file permissions
            - Close any programs using the file
            - Run as administrator (Windows) or with sudo (Linux/Mac)
            """)
        else:
            st.info("""
            **Recovery Tips:**
            - Check the logs for more details
            - Try refreshing the dashboard
            - Contact support if the issue persists
            """)

        # Show timestamp
        st.caption(
            f"Error occurred at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")


# --- Simple clause tagging keywords (extend as needed) ---
_CLAUSE_KEYWORDS = {
    "healthcare": ["SaMD", "ISO_IEC_42001"],
    "medical": ["SaMD"],
    "privacy": ["GDPR", "L2", "Art5_Data_Principles", "Art22_Automated_Decisions"],
    "fairness": ["EU_AI_ACT", "Art9_Data_Governance", "L3"],
    "bias": ["EU_AI_ACT", "Art9_Data_Governance", "L3"],
    "interpretability": ["EU_AI_ACT", "Art13_Transparency", "L4"],
    "explainability": ["EU_AI_ACT", "Art13_Transparency", "L4"],
    "oversight": ["EU_AI_ACT", "Art14_Human_Oversight"],
    "monitoring": ["EU_AI_ACT", "Art15_Robustness", "L5"],
    "logging": ["EU_AI_ACT", "Art15_Robustness", "HIPAA", "L5"],
    "security": ["HIPAA", "L2", "L5"],
    "governance": ["L1", "EU_AI_ACT", "ISO_IEC_42001"],
}

# ===== Automated Clause Tagging (suggestions) =====
st.markdown("### 🧩 Clause Suggestions (project description → suggested clauses)")

desc = st.text_area(
    "Paste a short project description (offline keyword-based matcher):",
    key="proj_desc",
    height=120,
    placeholder="e.g., Healthcare device predicting risks, must ensure privacy & fairness, provide explanations, monitor incidents..."
)

if st.button("Suggest clauses", key="btn_suggest_clauses", disabled=_LOCK):
    text = (desc or "").lower()
    hits = []
    for kw, tags in _CLAUSE_KEYWORDS.items():
        if kw in text:
            hits.extend(tags)
    # also map trace_map keys if we have it loaded
    _extra = []
    try:
        for fw in (trace_map or {}):
            if fw.lower() in text:
                _extra.append(fw)
        hits.extend(_extra)
    except Exception:
        pass
    if not hits:
        st.info("No suggestions found. Try adding more context words (privacy, fairness, explainability, monitoring, healthcare, etc.).")
    else:
        unique = sorted(set(hits))
        st.success(f"Suggested tags/clauses: {', '.join(unique)}")

        # Allow storing to configs/suggested_clauses.json
        save_it = st.checkbox(
            "Save suggestions to configs/suggested_clauses.json", value=False)
        if save_it:
            outp = os.path.join("configs", "suggested_clauses.json")
            payload = {"description": desc, "suggestions": unique}
            try:
                with open(outp, "w", encoding="utf-8") as f:
                    json.dump(payload, f, indent=2)
                st.success(f"Saved → {outp}")
            except Exception as e:
                st.error(f"Failed to write suggestions: {e}")


# ===== Audit Mode (hardened) =====
with st.container():
    audit_mode = st.toggle(
        "🔒 Audit Mode",
        value=st.session_state.get("__audit_mode__", False),
        help="Locks inputs and records a content hash for this snapshot."
    )
    files = sorted(glob.glob("reports/*.json"))
    agg_paths = [p for p in files if os.path.basename(p).startswith("AGG-")]
    agg_path = max(agg_paths, key=os.path.getmtime) if agg_paths else None
    # NEW:
if agg_path:
    try:
        mtime = os.path.getmtime(agg_path)
        run_hash = _hash_file(agg_path, _mtime=mtime)
    except Exception as e:
        logger.error(f"Failed to hash {agg_path}: {e}")
        run_hash = None
else:
    run_hash = None

    # display hash + time
    st.caption(
        f"Run hash: **{(run_hash or 'n/a')[:16]}**  |  Time UTC: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}"
    )

st.divider()

# --- Persist + log changes ---
if audit_mode != st.session_state.get("__audit_mode__", False):
    st.session_state["__audit_mode__"] = audit_mode
    record_audit_event("ENABLED" if audit_mode else "DISABLED", run_hash)

_LOCK = audit_locked()


# --- Audit lock overlay ---
if _LOCK:
    st.markdown("""
    <style>
    .stApp::after {
        content: "🔒 Audit Mode Active – Read-only";
        position: fixed;
        top: 50%%;
        left: 50%%;
        transform: translate(-50%%,-50%%);
        background: rgba(15,17,22,0.8);
        color: #e5e7eb;
        padding: 20px 40px;
        border-radius: 12px;
        font-size: 1.2rem;
        z-index: 9999;
    }
    .stButton>button, .stFileUploader, .stSelectbox, input, textarea {
        filter: grayscale(100%%) brightness(0.7);
        pointer-events: none;
    }
    </style>
    """, unsafe_allow_html=True)


# ====== STREAMLIT DASHBOARD FOR IRAQAF REPORTS ======
st.title("🧠 IRAQAF – Quality Assurance Scores")

# Check if this is first visit
if 'first_visit' not in st.session_state:
    st.session_state['first_visit'] = True
    st.session_state['hide_quickstart'] = False

# Auto-close guide after first visit
if not st.session_state['hide_quickstart']:
    # Expanded only on first visit
    is_expanded = st.session_state.get('first_visit', False)

    with st.expander("👋 Quick Start Guide", expanded=is_expanded):
        # Mark as not first visit anymore
        if st.session_state.get('first_visit'):
            st.session_state['first_visit'] = False

        # Styled banner
        st.markdown("""
        <div style='
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            padding: 20px;
            border-radius: 12px;
            margin-bottom: 20px;
            color: white;
            text-align: center;
        '>
            <h2 style='margin: 0; color: white;'>Welcome to IRAQAF Dashboard! 🎯</h2>
            <p style='margin: 8px 0 0 0; opacity: 0.95; font-size: 1.1rem;'>
                Your AI Quality Assurance companion
            </p>
        </div>
        """, unsafe_allow_html=True)

        col1, col2 = st.columns([4, 1])

        with col1:
            # Step 1
            st.markdown("""
            <div style='
                background: #f0f9ff;
                border-left: 4px solid #0284c7;
                padding: 16px;
                border-radius: 8px;
                margin-bottom: 16px;
            '>
                <h4 style='margin: 0 0 8px 0; color: #0c4a6e;'>
                    1️⃣ 📊 View Your Scores
                </h4>
                <ul style='margin: 0; color: #0c4a6e;'>
                    <li>Check the <b>Module Summary</b> table below</li>
                    <li>🟢 Green = Excellent (≥90) | 🟡 Yellow = Good (75-89) | 🔴 Red = Needs Work (<75)</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)

            # Step 2
            st.markdown("""
            <div style='
                background: #f0fdf4;
                border-left: 4px solid #16a34a;
                padding: 16px;
                border-radius: 8px;
                margin-bottom: 16px;
            '>
                <h4 style='margin: 0 0 8px 0; color: #14532d;'>
                    2️⃣ 🔍 Explore Modules
                </h4>
                <ul style='margin: 0; color: #14532d;'>
                    <li>Click on <b>L1-L5 sections</b> to see detailed metrics</li>
                    <li>Use the <b>search bar</b> (Ctrl+K) to find specific items</li>
                    <li>Check <b>Evidence Tray 📎</b> for supporting documents</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)

            # Step 3
            st.markdown("""
            <div style='
                background: #fef3c7;
                border-left: 4px solid #f59e0b;
                padding: 16px;
                border-radius: 8px;
                margin-bottom: 16px;
            '>
                <h4 style='margin: 0 0 8px 0; color: #78350f;'>
                    3️⃣ 📈 Track Progress
                </h4>
                <ul style='margin: 0; color: #78350f;'>
                    <li>Compare runs over time</li>
                    <li>View <b>GQAS</b> (Global Quality Score) at the bottom</li>
                    <li>Export reports for stakeholders</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)

            # Keyboard shortcuts
            st.markdown("""
            <div style='
                background: #f5f3ff;
                border: 1px solid #c4b5fd;
                padding: 12px;
                border-radius: 8px;
                margin-bottom: 12px;
            '>
                <h4 style='margin: 0 0 8px 0; color: #5b21b6;'>⌨️ Keyboard Shortcuts</h4>
                <div style='display: grid; grid-template-columns: 1fr 1fr; gap: 8px; color: #5b21b6;'>
                    <div><kbd style='background: #ddd6fe; padding: 2px 8px; border-radius: 4px;'>Ctrl+K</kbd> → Focus search</div>
                    <div><kbd style='background: #ddd6fe; padding: 2px 8px; border-radius: 4px;'>Ctrl+R</kbd> → Refresh</div>
                    <div><kbd style='background: #ddd6fe; padding: 2px 8px; border-radius: 4px;'>Esc</kbd> → Close sections</div>
                    <div><kbd style='background: #ddd6fe; padding: 2px 8px; border-radius: 4px;'>?</kbd> → Show all shortcuts</div>
                </div>
            </div>
            """, unsafe_allow_html=True)

            # Help section
            st.markdown("""
            <div style='
                text-align: center;
                padding: 12px;
                color: #6b7280;
                font-size: 0.9rem;
            '>
                💡 <b>Need help?</b> Enable <b>Debug Mode</b> in the sidebar or press <kbd style='background: #e5e7eb; padding: 2px 6px; border-radius: 4px;'>?</kbd> for shortcuts
            </div>
            """, unsafe_allow_html=True)

        with col2:
            st.markdown("<br>" * 2, unsafe_allow_html=True)
            if st.button("🎯 Start Interactive Tour", key="start_tour_btn", help="Take a guided tour of the dashboard", use_container_width=True, type="primary"):
                st.session_state['tour_active'] = True
                st.session_state['tour_step'] = 0
                # Hide guide during tour
                st.session_state['hide_quickstart'] = True
                st.rerun()
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("✕ Don't show again", key="hide_quickstart_btn", help="Hide this guide permanently", use_container_width=True):
                st.rerun()

        st.markdown("---")

# Show a subtle button to bring back the guide if hidden
if st.session_state['hide_quickstart']:
    col_empty, col_btn = st.columns([5, 1])
    with col_btn:
        if st.button("👋 Quick Start", key="show_quickstart_btn", help="Show the Quick Start Guide"):
            st.session_state['hide_quickstart'] = False
            # Expand it when shown again
            st.session_state['first_visit'] = True
            st.rerun()
# ============================================================================
# INTERACTIVE TOUR MODE
# ============================================================================

# Initialize tour state
if 'tour_active' not in st.session_state:
    st.session_state['tour_active'] = False
    st.session_state['tour_step'] = 0

# Tour steps configuration
TOUR_STEPS = [
    {
        'title': '📊 Module Summary',
        'description': 'This table shows the health of all 5 quality modules. Green = Excellent, Yellow = Good, Red = Needs Work.',
        'target': 'module-summary',
        'position': 'bottom'
    },
    {
        'title': '📎 Evidence Tray',
        'description': 'Upload and manage supporting documents for each module. Files are organized by module and can be previewed inline.',
        'target': 'evidence-tray',
        'position': 'bottom'
    },
    {
        'title': '🏛️ L1 Governance',
        'description': 'Track compliance with regulations like GDPR, EU AI Act, HIPAA. See which clauses pass or fail.',
        'target': 'l1-governance',
        'position': 'bottom'
    },
    {
        'title': '🔒 L2 Privacy/Security',
        'description': 'Monitor encryption, data protection, and security metrics. Track incidents and access reviews.',
        'target': 'l2-privacy',
        'position': 'bottom'
    },
    {
        'title': '🧮 L3 Fairness',
        'description': 'Measure bias metrics like DPG (Demographic Parity Gap) and EOD (Equal Opportunity Difference) across groups.',
        'target': 'l3-fairness',
        'position': 'bottom'
    },
    {
        'title': '🔍 L4 Explainability',
        'description': 'View feature importance, SHAP values, and model interpretability metrics. Compare models side-by-side.',
        'target': 'l4-explainability',
        'position': 'bottom'
    },
    {
        'title': '🛠️ L5 Operations',
        'description': 'Monitor live metrics, SLO breaches, and incident timelines. Set up alerting for quality drift.',
        'target': 'l5-operations',
        'position': 'bottom'
    },
    {
        'title': '⚙️ GQAS Score',
        'description': 'The Global Quality Assurance Score aggregates all modules. Check if you meet shipping floors for your risk profile.',
        'target': 'gqas-aggregate',
        'position': 'top'
    }
]

# Tour Mode CSS and JavaScript
if st.session_state['tour_active']:
    current_step = st.session_state['tour_step']

    if current_step < len(TOUR_STEPS):
        step = TOUR_STEPS[current_step]

        st.markdown(f"""
        <style>
        /* Overlay backdrop */
        .tour-overlay {{
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background: rgba(0, 0, 0, 0.7);
            z-index: 9998;
            pointer-events: none;
        }}

        /* Highlight specific section */
        [data-tour-target="{step['target']}"] {{
            position: relative;
            z-index: 9999 !important;
            #667eea, 0 0 0 8px rgba(102, 126, 234, 0.3) !important;
            box-shadow: 0 0 0 4px
            border-radius: 12px !important;
            background: white !important;
        }}

        /* Tour tooltip */
        .tour-tooltip {{
            position: fixed;
            {step['position']}: 20px;
            left: 50%;
            transform: translateX(-50%);
            max-width: 500px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 24px;
            border-radius: 16px;
            box-shadow: 0 20px 60px rgba(0, 0, 0, 0.4);
            z-index: 10000;
            animation: slideIn 0.3s ease-out;
        }}

        @keyframes slideIn {{
            from {{
                opacity: 0;
                transform: translateX(-50%) translateY(-20px);
            }}
            to {{
                opacity: 1;
                transform: translateX(-50%) translateY(0);
            }}
        }}

        .tour-tooltip h3 {{
            margin: 0 0 12px 0;
            font-size: 1.4rem;
            color: white;
        }}

        .tour-tooltip p {{
            margin: 0 0 20px 0;
            font-size: 1rem;
            line-height: 1.6;
            opacity: 0.95;
        }}

        .tour-progress {{
            margin-bottom: 16px;
            font-size: 0.875rem;
            opacity: 0.8;
        }}

        /* Arrow pointing to element */
        .tour-arrow {{
            position: fixed;
            left: 50%;
            transform: translateX(-50%);
            {('bottom' if step['position'] == 'top' else 'top')}: 180px;
            font-size: 3rem;
            color: #667eea;
            z-index: 9999;
            animation: bounce 1s infinite;
            text-shadow: 0 4px 8px rgba(0, 0, 0, 0.3);
        }}

        @keyframes bounce {{
            0%, 100% {{ transform: translateX(-50%) translateY(0); }}
            50% {{ transform: translateX(-50%) translateY({'-10px' if step['position'] == 'top' else '10px'}); }}
        }}
        </style>

        <div class="tour-overlay"></div>

        <div class="tour-arrow">
            {'⬆️' if step['position'] == 'top' else '⬇️'}
        </div>

        <div class="tour-tooltip">
            <div class="tour-progress">
                Step {current_step + 1} of {len(TOUR_STEPS)}
            </div>
            <h3>{step['title']}</h3>
            <p>{step['description']}</p>
        </div>

        <script>
        // Auto-scroll to highlighted element
        setTimeout(function() {{
            var target = document.querySelector('[data-tour-target="{step['target']}"]');
            if (target) {{
                target.scrollIntoView({{
                    behavior: 'smooth',
                    block: 'center'
                }});
            }}
        }}, 100);
        </script>
        """, unsafe_allow_html=True)

        # Tour controls (floating buttons)
        st.markdown("""
        <style>
        .tour-controls {
            position: fixed;
            bottom: 120px;
            left: 50%;
            transform: translateX(-50%);
            display: flex;
            gap: 12px;
            z-index: 10001;
        }
        </style>
        """, unsafe_allow_html=True)

# Tour control buttons
if st.session_state['tour_active']:
    col1, col2, col3, col4, col5 = st.columns([1, 1, 1, 1, 1])

    with col1:
        st.write("")  # Spacer

    with col2:
        if st.button("⏮️ Previous", key="tour_prev", disabled=st.session_state['tour_step'] == 0):
            st.session_state['tour_step'] -= 1
            st.rerun()

    with col3:
        if st.session_state['tour_step'] < len(TOUR_STEPS) - 1:
            if st.button("Next ⏭️", key="tour_next", type="primary"):
                st.session_state['tour_step'] += 1
                st.rerun()
        else:
            if st.button("🎉 Finish Tour", key="tour_finish", type="primary"):
                st.session_state['tour_active'] = False
                st.session_state['tour_step'] = 0
                st.success("✅ Tour completed! You're ready to use IRAQAF.")
                st.rerun()

    with col4:
        if st.button("✕ Exit Tour", key="tour_exit"):
            st.session_state['tour_active'] = False
            st.session_state['tour_step'] = 0
            st.rerun()

    with col5:
        st.write("")  # Spacer
# ===== Quick Actions Panel =====
with st.expander("⚡ Quick Actions", expanded=False):
    col1, col2, col3, col4 = st.columns(4)

    with col1:
        if st.button("🔄 Refresh Data", use_container_width=True, disabled=_LOCK, help="Clear cache and reload all reports"):
            st.session_state["force_refresh"] = True
            load_all_reports_cached.clear()
            st.rerun()

    with col2:
        if st.button("📊 Jump to GQAS", use_container_width=True):
            st.markdown(
                '<a href="#aggregate-global-qa-score-gqas"></a>', unsafe_allow_html=True)
            st.info("Scroll down to see GQAS section")

    with col3:
        if st.button("📎 Evidence Tray", use_container_width=True):
            st.markdown('<a href="#evidence-tray"></a>',
                        unsafe_allow_html=True)
            st.info("Scroll down to Evidence Tray")

    with col4:
        if st.button("🧭 Navigation", use_container_width=True):
            st.info("""
            **Quick Navigation:**
            - L1: Governance
            - L2: Privacy/Security
            - L3: Fairness
            - L4: Explainability
            - L5: Operations
            """)
# ===== Global Search =====
st.markdown("---")
search_col1, search_col2 = st.columns([3, 1])

with search_col1:
    search_query = st.text_input(
        "🔍 Search across all modules",
        placeholder="Search for metrics, clauses, or evidence...",
        key="global_search",
        disabled=_LOCK,
        help="Press Ctrl+K to focus"
    )

with search_col2:
    search_in = st.multiselect(
        "Search in",
        ["Modules", "Evidence", "Clauses", "Metrics"],
        default=["Modules"],
        key="search_scope"
    )

if search_query:
    st.markdown(f"### 🔍 Search Results for '{search_query}'")

    results = []

    # Search in modules
    if "Modules" in search_in:
        for mid, rep in latest.items():
            if rep and search_query.lower() in str(rep).lower():
                results.append({
                    "Type": "Module",
                    "Location": NAMES.get(mid, mid),
                    "Match": f"Found in {mid} report"
                })

    # Search in evidence
    if "Evidence" in search_in:
        eidx = _load_evidence_index()
        for mid, files in eidx.items():
            for f in files:
                if search_query.lower() in f.lower():
                    results.append({
                        "Type": "Evidence",
                        "Location": NAMES.get(mid, mid),
                        "Match": f
                    })

    if results:
        st.dataframe(pd.DataFrame(results), use_container_width=True)
    else:
        st.info("No results found. Try different keywords.")

st.markdown("---")

# =============================================================================
# TIMESTAMP HELPER (Use throughout dashboard)
# =============================================================================


def render_timestamp(label: str, show: bool = True):
    """Render a timestamp badge if enabled in settings"""
    if not show or not st.session_state.get('show_timestamps', True):
        return

    st.markdown(f"""
    <div style='
        display: inline-block;
        background: #f3f4f6;
        padding: 4px 12px;
        border-radius: 999px;
        font-size: 0.75rem;
        color: #6b7280;
        margin-left: 8px;
    '>
        🕐 {label}: {datetime.now().strftime('%H:%M:%S')}
    </div>
    """, unsafe_allow_html=True)
# Helper: normalize JSON payload into list[dict]


def _as_dict_list(obj):
    if isinstance(obj, dict):
        return [obj]
    if isinstance(obj, list):
        return [x for x in obj if isinstance(x, dict)]
    return []


# Find latest reports (define early for audit hash)
files = sorted(glob.glob("reports/*.json"))
if not files:
    st.info("No reports found. Run the CLI to generate JSON reports in ./reports")
    st.stop()

# Module names and emoji map
NAMES = {
    "L1": "L1 Governance & Regulatory",
    "L2": "L2 Privacy & Security",
    "L3": "L3 Fairness & Ethics",
    "L4": "L4 Explainability & Transparency",
    "L5": "L5 Operations & Monitoring",
}


def name_for(module_id: str) -> str:
    return NAMES.get(module_id, module_id)


COLOR_EMOJI = {"green": "🟢", "yellow": "🟡", "red": "🔴"}

# ============================================================================
# CACHED REPORT LOADING (Performance Optimized)
# ============================================================================


@st.cache_data(ttl=300, show_spinner=False)  # Cache for 5 minutes
def load_all_reports_cached(file_list: tuple, force_reload: bool = False) -> dict:
    """
    Load all reports with caching for performance.

    Args:
        file_list: Tuple of file paths (must be tuple for hashing)
        force_reload: Force cache bypass

    Returns:
        Dict mapping module names to their latest reports
    """
    latest_by_module = {k: None for k in ["L1", "L2", "L3", "L4", "L5", "AGG"]}
    seen_mtime = {k: -1 for k in latest_by_module.keys()}

    errors = []

    for f in file_list:
        try:
            with open(f, "r", encoding="utf-8") as fh:
                raw = json.load(fh)

            mtime = os.path.getmtime(f)

            # Handle both single report and list of reports
            for rec in (raw if isinstance(raw, list) else [raw]):
                if not isinstance(rec, dict):
                    continue

                mod = str(rec.get("module", "")).upper()
                if mod in latest_by_module and mtime >= seen_mtime[mod]:
                    latest_by_module[mod] = rec
                    seen_mtime[mod] = mtime

        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON in {f}: {e}")
            errors.append({
                "file": os.path.basename(f),
                "error": "Invalid JSON format",
                "type": "JSONDecodeError"
            })
        except FileNotFoundError:
            logger.error(f"File not found: {f}")
            errors.append({
                "file": os.path.basename(f),
                "error": "File not found",
                "type": "FileNotFoundError"
            })
        except Exception as e:
            logger.error(f"Failed to load {f}: {e}", exc_info=True)
            errors.append({
                "file": os.path.basename(f),
                "error": str(e),
                "type": type(e).__name__
            })

    # Attach error log to return value
    latest_by_module["_errors"] = errors
    latest_by_module["_load_time"] = datetime.now().isoformat()

    return latest_by_module


# ========================================
# Load Reports with Progress UI
# ========================================

# Get audit lock state for UI controls
_LOCK = get_audit_mode()

with measure_section("load_all_reports"):
    files = sorted(glob.glob("reports/*.json"))

    if not files:
        st.info("No reports found. Run the CLI to generate JSON reports in ./reports")
        st.stop()

    # Check if user wants to force refresh
    force_refresh = st.session_state.get("force_refresh", False)
    if force_refresh:
        st.session_state["force_refresh"] = False
        load_all_reports_cached.clear()  # Clear cache

    # Show loading indicator for first load or many files
    if len(files) > 5 and st.session_state.get("last_refresh") is None:
        with show_loading("Loading reports", f"Loading {len(files)} reports..."):
            result = load_all_reports_cached(tuple(files), force_refresh)
    else:
        result = load_all_reports_cached(tuple(files), force_refresh)

    # Extract errors and actual data
    load_errors = result.pop("_errors", [])
    load_time = result.pop("_load_time", None)
    latest = result  # ✅ This is the correct assignment

    # Update last refresh time
    st.session_state["last_refresh"] = datetime.now().isoformat()

    # Show loading stats
    loaded_count = sum(
        1 for k, v in latest.items()
        if v is not None and k not in ["_errors", "_load_time"]
    )

    if PERFORMANCE_MONITORING:
        monitor = get_monitor()
        stats = monitor.get_stats("load_all_reports")
        if stats:
            st.caption(
                f"✅ Loaded {loaded_count}/6 modules in {stats.get('last', 0):.2f}s "
                f"(cached: {load_time})"
            )

    # Show errors if any
    if load_errors:
        with st.expander(f"⚠️ {len(load_errors)} file(s) failed to load", expanded=False):
            error_df = pd.DataFrame(load_errors)
            st.dataframe(error_df, use_container_width=True)

            st.markdown("**Recovery Tips:**")
            if any(e["type"] == "JSONDecodeError" for e in load_errors):
                st.info(
                    "• **Invalid JSON**: Re-run evaluation to regenerate corrupt files")
            if any(e["type"] == "PermissionError" for e in load_errors):
                st.info(
                    "• **Permission Denied**: Check file permissions or close programs using the files")
            if any(e["type"] == "FileNotFoundError" for e in load_errors):
                st.info("• **File Not Found**: Verify reports directory path")

# Show summary of loaded modules
loaded_modules = [k for k, v in latest.items() if v is not None]
missing_modules = [k for k in ["L1", "L2", "L3",
                               "L4", "L5", "AGG"] if k not in loaded_modules]

if missing_modules:
    st.warning(
        f"⚠️ Missing reports for: {', '.join(missing_modules)}. "
        f"Run evaluation to generate these modules."
    )

    # Provide quick action button
    col1, col2, col3 = st.columns([1, 1, 2])
    with col1:
        if st.button("📖 View Documentation", use_container_width=True):
            st.info("""
            **To generate missing reports:**
```bash
            # Generate all modules
            python -m cli.iraqaf_cli run --module ALL --config configs/project.example.yaml --out reports

            # Generate specific module
            python -m cli.iraqaf_cli run --module L1 --config configs/project.example.yaml --out reports
        """)
with col2:
    if st.button("🔄 Refresh", use_container_width=True, disabled=_LOCK):
        st.rerun()
st.divider()

# ========== MODULE SUMMARY ==========
with measure_section("build_module_summary"):
    rows = []
    for mid in ["L1", "L2", "L3", "L4", "L5"]:
        rep = latest.get(mid)
        if rep:
            emoji = COLOR_EMOJI.get(str(rep.get("band", "")).lower(), "⚪")
            e_cnt = _evidence_count(mid)  # uses your existing helper/alias
            rows.append([
                NAMES[mid],
                f"{emoji} {str(rep.get('band', '')).capitalize()}",
                rep.get("score", "-"),
                json.dumps(rep.get("metrics", {}), indent=2),
                e_cnt,
            ])
        else:
            rows.append([NAMES[mid], "❔ Missing", "-", "-", 0])

    df = pd.DataFrame(
        rows,
        columns=["Module", "Status", "Score", "Metrics", "Evidence 📎"]
    )

# ===== Module Summary =====
st.markdown('<div data-tour-target="module-summary">', unsafe_allow_html=True)

st.markdown("### 🧾 Module Summary")

# Simple help text right below the title
st.markdown("""
<div style='
    background: #f0f9ff;
    border-left: 4px solid #0284c7;
    padding: 12px 16px;
    margin-bottom: 16px;
    border-radius: 4px;
'>
    <p style='margin: 0; font-size: 0.9rem; color: #0c4a6e;'>
        <b>Quick Guide:</b>
        🟢 Excellent (≥90) •
        🟡 Good (75-89) •
        🔴 Needs Work (&lt;75) •
        📎 Evidence Count
        <span style='float: right; cursor: pointer;' onclick='this.parentElement.parentElement.style.display="none"'>✕</span>
    </p>
</div>
""", unsafe_allow_html=True)

st.dataframe(df, width="stretch")

st.markdown('</div>', unsafe_allow_html=True)  # Close tour target

# ===== 📎 Evidence Tray (per module) with Search + Global Preview =====
st.markdown("""
<div data-tour-target="evidence-tray">
<h3>📎 Evidence Tray</h3>
""", unsafe_allow_html=True)

with measure_section("evidence_tray_render"):

    eidx = _load_evidence_index()
    if not eidx:
        st.info(
            "No evidence index found yet. Create **configs/evidence_index.json** to pin files per module."
        )
    else:
        # Global controls
        ctrl1, ctrl2, ctrl3 = st.columns([0.5, 0.25, 0.25])
        with ctrl1:
            q = st.text_input(
                "Filter files (name or path)",
                value="",
                placeholder="e.g., fairness, .csv, security.yaml",
                disabled=_LOCK
            )
        with ctrl2:
            expand_all = st.toggle(
                "Expand all",
                value=False,
                disabled=_LOCK
            )
        with ctrl3:
            preview_all = st.toggle(
                "Preview all",
                value=False,
                disabled=_LOCK
            )

        # Normalize query
        q_norm = (q or "").strip().lower()

        for mid in ["L1", "L2", "L3", "L4", "L5"]:
            files_for_mid = eidx.get(mid, [])
            # Apply filter
            if q_norm:
                files_for_mid = [
                    p for p in files_for_mid
                    if q_norm in p.lower() or q_norm in os.path.basename(p).lower()
                ]

            count = len(files_for_mid)
            title = f"{NAMES[mid]} — {count} file{'s' if count != 1 else ''}"
            with st.expander(title, expanded=expand_all):
                if count == 0:
                    st.caption(
                        "No files match the current filter." if q_norm else "No files pinned for this module."
                    )
                    continue

                # Header row
                h1, h2, h3, h4, h5, h6 = st.columns(
                    [0.36, 0.12, 0.16, 0.16, 0.10, 0.10])
                h1.markdown("**File**")
                h2.markdown("**Status**")
                h3.markdown("**Size**")
                h4.markdown("**Modified**")
                h5.markdown("**Preview**")
                h6.markdown("**Download**")

                for i, p in enumerate(files_for_mid, start=1):
                    fp = Path(p)
                    exists = fp.exists()
                    size = _human_size(fp.stat().st_size) if exists else "—"
                    mtime = time.strftime(
                        "%Y-%m-%d %H:%M:%S", time.localtime(fp.stat().st_mtime)
                    ) if exists else "—"

                    c1, c2, c3, c4, c5, c6 = st.columns(
                        [0.36, 0.12, 0.16, 0.16, 0.10, 0.10])
                    c1.markdown(
                        f"**{fp.name}**  \n<small style='color:#888'>{fp.as_posix()}</small>",
                        unsafe_allow_html=True
                    )
                    c2.markdown("✅ Exists" if exists else "❌ Missing")
                    c3.markdown(f"{size}")
                    c4.markdown(mtime)

                    # Row toggle honors the global "Preview all" as its default
                    tkey = f"pv_{mid}_{i}"
                    show = c5.toggle(
                        "Show",
                        value=preview_all,
                        key=tkey,
                        help="Toggle to preview the file inline.",
                        disabled=_LOCK
                    )

                    if exists:
                        with fp.open("rb") as fh:
                            c6.download_button(
                                "Download",
                                data=fh.read(),
                                file_name=fp.name,
                                key=f"dl_{mid}_{i}"
                            )
                    else:
                        c6.markdown("—")

                    if show:
                        with st.container():
                            st.markdown(
                                "<div style='margin:8px 0; padding:10px; background:#f7f7f9; border-radius:8px;'>",
                                unsafe_allow_html=True
                            )
                            preview_widget(fp)
                            st.markdown("</div>", unsafe_allow_html=True)

                st.caption(
                    "Tip: Keep evidence in `docs/`, `configs/`, `logs/`, or `data/` so paths stay tidy and portable."
                )

st.markdown("</div>", unsafe_allow_html=True)


# ===== L1 Governance – Compliance Matrix =====
st.markdown("""
<div data-tour-target="l1-governance">
<h3>🏛️ L1 Governance – Compliance Coverage by Framework</h3>
""", unsafe_allow_html=True)

with measure_section("l1_compliance_matrix"):
    l1 = latest.get("L1")
    if l1 and "metrics" in l1:
        fb = l1["metrics"].get("framework_breakdown")
        cov = l1["metrics"].get("coverage_percent")
        if fb and isinstance(fb, list):
            st.caption(
                f"Overall L1 coverage: **{cov:.2f}%** • Each bar shows coverage (weighted clause pass%) for a regulation."
            )
            # Build a tidy dataframe
            frows = []
            for item in fb:
                frows.append({
                    "Framework": item.get("framework"),
                    "Coverage %": float(item.get("coverage_percent", 0.0)),
                    "Passed / Total": f"{item.get('covered', 0)}/{item.get('total', 0)}"
                })
            fdf = pd.DataFrame(frows).sort_values("Coverage %", ascending=True)
            chart = (
                alt.Chart(fdf)
                .mark_bar()
                .encode(
                    x=alt.X("Coverage %:Q", title="Coverage (%)",
                            scale=alt.Scale(domain=[0, 100])),
                    y=alt.Y("Framework:N", sort="-x", title=""),
                    tooltip=["Framework", alt.Tooltip(
                        "Coverage %:Q", format=".2f"), "Passed / Total"]
                )
                .properties(height=220, width="container")
            )
            st.altair_chart(chart, width="stretch")
        else:
            st.info("No per-framework breakdown found in L1 metrics yet.")
    else:
        st.info("Run L1 Governance to populate the compliance matrix.")

    l1 = _normalize_clause_evidence(l1 or {})

st.markdown("</div>", unsafe_allow_html=True)

# ===== L1 Governance – Regulation Traceability Map =====
st.markdown("""
<div data-tour-target="l1-traceability">
<h3>🧭 Governance Traceability (Clause ↔ Metric)</h3>
""", unsafe_allow_html=True)

_trace_path = os.path.join("configs", "trace_map.yaml")
if not os.path.exists(_trace_path):
    st.info("No trace_map.yaml found in configs/. Create it to visualize clause-to-metric links.")
else:
    try:
        with open(_trace_path, "r", encoding="utf-8") as fh:
            _tm = _yaml.safe_load(fh) or {}
        trace_map = (_tm.get("trace_map") or {})
    except Exception as e:
        trace_map = {}
        st.error(f"Failed to load trace_map.yaml: {e}")

    if trace_map:
        fw_names = list(trace_map.keys())
        fw_sel = st.selectbox("Framework", fw_names,
                              index=0, key="trace_fw_sel")

        rows = []
        for clause, metric_list in (trace_map.get(fw_sel, {}) or {}).items():
            for met in (metric_list or []):
                rows.append({"Clause": clause, "Metric": met, "Link": 1})

        if not rows:
            st.info("No mappings in the selected framework yet.")
        else:
            _df = pd.DataFrame(rows)

            # --- Matrix-style heatmap (Clause on Y, Metric on X)
            st.caption(
                "Matrix view: a filled cell indicates a documented linkage between a clause and a model metric."
            )
            pivot = _df.pivot_table(
                index="Clause", columns="Metric", values="Link", fill_value=0
            )
            # Show as table first for clarity
            st.dataframe(pivot, width="stretch")

            try:
                tidy = (
                    _df.groupby(["Clause", "Metric"], as_index=False)["Link"]
                    .sum()
                )
                chart = (
                    alt.Chart(tidy)
                    .mark_rect()
                    .encode(
                        x=alt.X("Metric:N", sort=None, title="Metric"),
                        y=alt.Y("Clause:N", sort=None, title="Clause"),
                        color=alt.Color(
                            "Link:Q",
                            scale=alt.Scale(
                                domain=[0, 1], range=["#f0f0f0", "#1f77b4"]
                            ),
                            legend=None
                        ),
                        tooltip=["Clause", "Metric"]
                    )
                    .properties(height=320, width="container")
                )
                st.altair_chart(chart, theme=None)
            except Exception:
                pass

            with st.expander("Copy metric paths"):
                st.write("Click to copy commonly used metric paths:")
                _uniq_metrics = sorted({r["Metric"] for r in rows})
                for mpath in _uniq_metrics:
                    st.code(mpath)
    else:
        st.info("Trace map is empty. Add entries in configs/trace_map.yaml.")

st.markdown("</div>", unsafe_allow_html=True)


# ===== L1 Governance – What's Missing (Failed Clauses) =====
st.markdown("### ❗ L1 Governance – What’s Missing")
l1 = latest.get("L1")
if l1 and "metrics" in l1 and isinstance(l1["metrics"].get("clauses"), list):
    failed = [c for c in l1["metrics"]["clauses"]
              if not c.get("passed", False)]
    if failed:
        miss_df = pd.DataFrame([
            {
                "Framework": c.get("framework"),
                "Clause": c.get("id"),
                "Description": c.get("description"),
                "Weight": c.get("weight"),
                "Why Failed": c.get("why_failed"),
                "Hint": c.get("hint"),
            }
            for c in failed
        ])
        st.dataframe(miss_df, width="stretch")
        st.download_button(
            "⬇️ Download failed clauses (CSV)",
            data=miss_df.to_csv(index=False).encode("utf-8"),
            file_name="l1_failed_clauses.csv",
            mime="text/csv"
        )
        st.caption(
            "Table shows only failed clauses. Fixing high-weight ones first increases L1 fastest.")
    else:
        st.success("All mapped clauses passed 🎉")
else:
    st.info("No clause-level output from L1 yet. Re-run L1 after updating the evaluator to emit `metrics.clauses`.")
    # === L1 Governance — Evidence Attachments ===
l1 = latest.get("L1")
if l1 and "metrics" in l1:
    failed = [c for c in (l1["metrics"].get("clauses")
                          or []) if not c.get("passed")]
    evidence_files = l1.get("evidence", [])

    with st.expander("📎 L1 Evidence attachments (click to view)", expanded=False):
        if not evidence_files:
            st.info("No evidence files were attached in the L1 config.")
        else:
            from io import BytesIO
            import os

            st.caption(
                "Files below come from **L1.evidence** in your project config.")
            colA, colB = st.columns(2)
            for i, path in enumerate(evidence_files):
                try:
                    with open(path, "rb") as fh:
                        data = fh.read()
                    fname = os.path.basename(path)
                    (colA if i % 2 == 0 else colB).download_button(
                        f"⬇️ {fname}",
                        data=data,
                        file_name=fname,
                        mime="application/octet-stream",
                        key=f"l1_ev_{i}",
                    )
                except Exception:
                    (colA if i % 2 == 0 else colB).warning(f"Missing: {path}")

    # Optional: per-clause quick evidence/context
    if failed:
        st.markdown("#### 🧩 Clause-level context & hints")
        for c in failed:
            title = f"{c.get('framework')} · {c.get('id')} — {c.get('description')}"
            with st.expander(f"❌ {title}"):
                st.write(f"**Why failed:** {c.get('why_failed') or '—'}")
                st.write(f"**Hint:** {c.get('hint') or '—'}")
                if evidence_files:
                    st.caption("Relevant attached evidence:")
                    for p in evidence_files:
                        st.code(p, language="text")
else:
    st.info("Run L1 Governance to see evidence attachments.")


# --- Show active compliance map (Strict / Lenient) ---
with st.expander("🔄 Compliance map profile (Strict / Lenient)"):
    st.write("Run the quick scripts to switch maps, then re-run L1:")
    st.code("use_lenient_map.bat   # or use_strict_map.bat", language="bash")

    # Show which map is currently active by peeking at the pointer file
    try:
        import yaml
        import os
        ptr_path = "configs/compliance_map.yaml"
        if os.path.exists(ptr_path):
            with open(ptr_path, "r", encoding="utf-8") as fh:
                y = yaml.safe_load(fh) or {}
            frameworks = list((y.get("frameworks") or {}).keys())
            st.caption(
                f"✅ Active mapping: `{ptr_path}` • Frameworks loaded: {', '.join(frameworks[:5])}")
        else:
            st.warning(
                "Pointer file configs/compliance_map.yaml not found. Run one of the scripts above.")
    except Exception as e:
        st.warning(f"Could not read active mapping pointer: {e}")

    st.info("After switching, regenerate L1:\n\n"
            "`python -m cli.iraqaf_cli run --module L1 --config configs\\project.example.yaml --out reports`")

    # ===== L1 Governance – Trends Over Time =====
st.markdown("### 🧭 L1 Governance – Trends Over Time")
l1_rows = []
for f in files:
    try:
        with open(f, "r") as fh:
            d = json.load(fh)
        if d.get("module") != "L1":
            continue
        m = d.get("metrics", {}) or {}
        base = os.path.basename(f).replace(".json", "")
        ts_match = re.search(r"(\d{8}-\d{6})", base)
        ts_label = ts_match.group(1) if ts_match else base
        l1_rows.append({
            "Run": base,
            "Label": ts_label,
            "time": os.path.getmtime(f),
            "policy_coverage": (None if m.get("policy_coverage") is None else float(m["policy_coverage"])),
            "audit_completeness": (None if m.get("audit_completeness") is None else float(m["audit_completeness"])),
            "risk_register_age_days": (None if m.get("risk_register_age_days") is None else float(m["risk_register_age_days"])),
            "Score": d.get("score")
        })
    except Exception:
        pass

if not l1_rows:
    st.info("No L1 governance metrics found yet. (Optional keys: policy_coverage, audit_completeness, risk_register_age_days)")
else:
    l1_df = pd.DataFrame(l1_rows).sort_values("time").reset_index(drop=True)
    l1_df["Run #"] = l1_df.index + 1
    st.dataframe(l1_df[["Run #", "Label", "policy_coverage", "audit_completeness", "risk_register_age_days", "Score"]],
                 width="stretch")

    cov_bands = pd.DataFrame({"y0": [0.00, 0.90, 0.95], "y1": [
                             0.90, 0.95, 1.00], "band": ["red", "yellow", "green"]})
    comp_bands = cov_bands.copy()
    age_bands = pd.DataFrame({"y0": [0, 1, 2], "y1": [1, 2, 7], "band": [
                             # 0–1d green, 1–2d yellow, >2d red (demo)
                             "green", "yellow", "red"]})

    def l1_chart(df, metric_col, title, bands_df, y_domain):
        base = alt.Chart(df)
        bands = alt.Chart(bands_df).mark_rect(opacity=0.15).encode(
            y=alt.Y("y0:Q", title=None, scale=alt.Scale(domain=y_domain)),
            y2="y1:Q",
            color=alt.Color("band:N",
                            scale=alt.Scale(domain=["green", "yellow", "red"],
                                            range=["#00c851", "#ffde59", "#ff4b4b"]),
                            legend=None)
        )
        line = base.mark_line(point=True).encode(
            x=alt.X("`Run #`:Q", title="Run # (chronological)"),
            y=alt.Y(f"{metric_col}:Q", title=title,
                    scale=alt.Scale(domain=y_domain)),
            tooltip=["Label:N", alt.Tooltip(f"{metric_col}:Q", format=".3f")]
        ).properties(height=240, width="container")
        return (bands + line)

    c1, c2 = st.columns(2)
    with c1:
        st.altair_chart(l1_chart(l1_df, "policy_coverage", "Policy Coverage (↑ better)", cov_bands, (0, 1.0)),
                        width="stretch")
    with c2:
        st.altair_chart(l1_chart(l1_df, "audit_completeness", "Audit Completeness (↑ better)", comp_bands, (0, 1.0)),
                        width="stretch")

    st.altair_chart(l1_chart(l1_df, "risk_register_age_days", "Risk Register Age (days, ↓ better)", age_bands, (0, 7)),
                    width="stretch")


# ===== L1 – Regulation Traceability Map =====
st.markdown("### 🧭 Regulation Traceability Map (Clause ↔ Metric)")

trace_path = os.path.join("configs", "trace_map.yaml")
if not os.path.exists(trace_path):
    st.info("No trace_map.yaml found in configs/.")
else:
    with open(trace_path, "r", encoding="utf-8") as fh:
        y = yaml.safe_load(fh) or {}
    tm = y.get("trace_map", {})

    rows_tm = []
    for fw, clauses in tm.items():
        for cid, metrics in (clauses or {}).items():
            for mref in (metrics or []):
                rows_tm.append(
                    {"Framework": fw, "Clause": cid, "Metric": str(mref)})

    if not rows_tm:
        st.info("Trace map is empty.")
    else:
        tdf = pd.DataFrame(rows_tm)
        st.caption(
            "Shows how each clause relates to module metrics. Darker = more associations.")

        # Build a matrix: rows=Clause (scoped with Framework), cols=Metric, values=count (1)
        tdf["ClauseFull"] = tdf["Framework"] + " · " + tdf["Clause"]
        mat = (
            tdf.groupby(["ClauseFull", "Metric"]
                        ).size().reset_index(name="count")
        )
        heat = alt.Chart(mat).mark_rect().encode(
            x=alt.X("Metric:N", sort="y", title="Metric"),
            y=alt.Y("ClauseFull:N", sort="-x", title="Clause"),
            color=alt.Color("count:Q", scale=alt.Scale(
                scheme="blues"), legend=None),
            tooltip=["ClauseFull", "Metric"]
        ).properties(height=320, width="container")
        st.altair_chart(heat, theme=None)

# ===== L2 Privacy/Security – Trends Over Time =====
st.markdown("""
<div data-tour-target="l2-privacy">
<h3>🔒 L2 Privacy/Security – Trends Over Time</h3>
""", unsafe_allow_html=True)

l2_rows = []
for f in files:
    try:
        with open(f, "r") as fh:
            d = json.load(fh)
        if d.get("module") != "L2":
            continue
        m = d.get("metrics", {}) or {}
        base = os.path.basename(f).replace(".json", "")
        ts_match = re.search(r"(\d{8}-\d{6})", base)
        ts_label = ts_match.group(1) if ts_match else base
        l2_rows.append({
            "Run": base,
            "Label": ts_label,
            "time": os.path.getmtime(f),
            "encryption_coverage": (None if m.get("encryption_coverage") is None else float(m["encryption_coverage"])),
            "dpia_complete": (None if m.get("dpia_complete") is None else float(m["dpia_complete"])),
            "access_review_age_days": (None if m.get("access_review_age_days") is None else float(m["access_review_age_days"])),
            "incident_rate_per_1k_users": (None if m.get("incident_rate_per_1k_users") is None else float(m["incident_rate_per_1k_users"])),
            "Score": d.get("score")
        })
    except Exception:
        pass

if not l2_rows:
    st.info("No L2 privacy/security metrics found yet. (Optional keys: encryption_coverage, dpia_complete, access_review_age_days, incident_rate_per_1k_users)")
else:
    l2_df = pd.DataFrame(l2_rows).sort_values("time").reset_index(drop=True)
    l2_df["Run #"] = l2_df.index + 1
    st.dataframe(
        l2_df[[
            "Run #", "Label", "encryption_coverage", "dpia_complete",
            "access_review_age_days", "incident_rate_per_1k_users", "Score"
        ]],
        width="stretch"
    )

    cov_bands = pd.DataFrame({"y0": [0.00, 0.90, 0.95], "y1": [0.90, 0.95, 1.00],
                              "band": ["red", "yellow", "green"]})
    dpia_bands = pd.DataFrame({"y0": [0.00, 0.75, 0.90], "y1": [0.75, 0.90, 1.00],
                               "band": ["red", "yellow", "green"]})
    access_age_bands = pd.DataFrame({"y0": [0, 30, 60], "y1": [30, 60, 120],
                                     "band": ["green", "yellow", "red"]})
    incident_bands = pd.DataFrame({"y0": [0, 1, 3], "y1": [1, 3, 10],
                                   "band": ["green", "yellow", "red"]})

    def l2_chart(df, metric_col, title, bands_df, y_domain):
        base = alt.Chart(df)
        bands = alt.Chart(bands_df).mark_rect(opacity=0.15).encode(
            y=alt.Y("y0:Q", title=None, scale=alt.Scale(domain=y_domain)),
            y2="y1:Q",
            color=alt.Color(
                "band:N",
                scale=alt.Scale(domain=["green", "yellow", "red"],
                                range=["#00c851", "#ffde59", "#ff4b4b"]),
                legend=None
            )
        )
        line = base.mark_line(point=True).encode(
            x=alt.X("`Run #`:Q", title="Run # (chronological)"),
            y=alt.Y(f"{metric_col}:Q", title=title,
                    scale=alt.Scale(domain=y_domain)),
            tooltip=["Label:N", alt.Tooltip(f"{metric_col}:Q", format=".3f")]
        ).properties(height=240, width="container")
        return (bands + line)

    c1, c2 = st.columns(2)
    with c1:
        st.altair_chart(
            l2_chart(l2_df, "encryption_coverage",
                     "Encryption Coverage (↑ better)", cov_bands, (0, 1.0)),
            width="stretch"
        )
    with c2:
        st.altair_chart(
            l2_chart(l2_df, "dpia_complete",
                     "DPIA Completion (↑ better)", dpia_bands, (0, 1.0)),
            width="stretch"
        )

    c3, c4 = st.columns(2)
    with c3:
        st.altair_chart(
            l2_chart(l2_df, "access_review_age_days",
                     "Access Review Age (days, ↓ better)", access_age_bands, (0, 120)),
            width="stretch"
        )
    with c4:
        st.altair_chart(
            l2_chart(l2_df, "incident_rate_per_1k_users",
                     "Incident Rate per 1k Users (↓ better)", incident_bands, (0, 10)),
            width="stretch"
        )

st.markdown("</div>", unsafe_allow_html=True)


# ===== L1 – Automated Clause Tagging (lightweight) =====
st.markdown("### 🧠 Automated Clause Tagging (beta)")
desc = st.text_area("Describe the AI system / use case",
                    placeholder="e.g., Clinical decision support using patient EHR + imaging …")
if st.button("Suggest relevant clauses", disabled=_LOCK):
    # simple keyword map → framework.clause_id
    kw_map = {
        "transparency": [("EU_AI_ACT", "Art13_Transparency")],
        "explain": [("EU_AI_ACT", "Art13_Transparency"), ("GDPR", "Art22_Automated_Decisions")],
        "health": [("SaMD", "Quality_Management")],
        "medical": [("SaMD", "Quality_Management")],
        "risk": [("ISO_IEC_42001", "AI_Risk_Management")],
        "logging": [("EU_AI_ACT", "Art15_Robustness"), ("HIPAA", "Security_Integrity")],
        "privacy": [("GDPR", "Art5_Data_Principles")],
        "security": [("GDPR", "Art5_Data_Principles"), ("HIPAA", "Security_Integrity")],
        "fairness": [("EU_AI_ACT", "Art9_Data_Governance")],
        "bias": [("EU_AI_ACT", "Art9_Data_Governance")],
    }
    found = set()
    text = (desc or "").lower()
    for k, targets in kw_map.items():
        if k in text:
            for fw, cid in targets:
                found.add((fw, cid))
    if not found:
        st.info("No suggestions matched. Try adding domain keywords (privacy, logging, medical, transparency, fairness…).")
    else:
        out = [{"Framework": fw, "Clause": cid} for (fw, cid) in sorted(found)]
        df_out = pd.DataFrame(out)
        st.dataframe(df_out, use_container_width=True)
        st.caption(
            "Tip: wire this to your governance process to auto-populate a review list.")

        # NEW: save to configs/suggested_clauses.json
        save_it = st.checkbox(
            "Save to configs/suggested_clauses.json",
            value=True,
            disabled=_LOCK
        )
        if save_it:
            os.makedirs("configs", exist_ok=True)
            payload = {
                "description": (desc or ""),
                "suggestions": [{"framework": r["Framework"], "clause": r["Clause"]} for _, r in df_out.iterrows()],
                "generated_at": pd.Timestamp.now().isoformat(timespec="seconds")
            }
            try:
                with open("configs/suggested_clauses.json", "w", encoding="utf-8") as fh:
                    json.dump(payload, fh, indent=2)
                st.success("Saved → configs/suggested_clauses.json")
            except Exception as e:
                st.error(f"Failed to save: {e}")


# ===== L3 Fairness – Evolution Across Runs (DPG & EOD) =====
st.markdown("""
<div data-tour-target="l3-fairness">
<h3>🧮 L3 Fairness – Evolution Across Runs</h3>
""", unsafe_allow_html=True)

# Make this section collapsible to avoid loading unless needed
show_l3_trends = st.checkbox(
    "📊 Show historical L3 trends", value=False, key="show_l3_trends")

if show_l3_trends:
    with measure_section("l3_fairness_processing"):
        # Collect all L3 reports
        l3_rows = []
        for f in files:
            try:
                with open(f, "r") as fh:
                    d = json.load(fh)
                if d.get("module") != "L3":
                    continue
                m = d.get("metrics", {})
                dpg = m.get("DPG", None)
                eod = m.get("EOD", None)
                auroc = m.get("AUROC", None)
                score = d.get("score", None)

                # label from filename & time (for ordering)
                base = os.path.basename(f).replace(".json", "")
                ts_match = re.search(r"(\d{8}-\d{6})", base)
                ts_label = ts_match.group(1) if ts_match else base
                mtime = os.path.getmtime(f)

                l3_rows.append({
                    "Run": base,
                    "Label": ts_label,
                    "time": mtime,
                    "DPG": None if dpg is None else float(dpg),
                    "EOD": None if eod is None else float(eod),
                    "AUROC": None if auroc in (None, "null") else float(auroc),
                    "Score": score,
                    "rates_by_group": m.get("rates_by_group"),
                    "tpr_by_group": m.get("tpr_by_group"),
                })
            except Exception:
                pass

        if not l3_rows:
            st.info(
                "No L3 fairness reports found yet. Run L3 multiple times to visualize evolution."
            )
        else:
            l3_df = pd.DataFrame(l3_rows).sort_values(
                "time").reset_index(drop=True)
            l3_df["Run #"] = l3_df.index + 1

            # Small, tidy table
            st.dataframe(
                l3_df[["Run #", "Label", "DPG", "EOD", "AUROC", "Score"]],
                width="stretch",
            )

            # Threshold bands (0–0.05 green, 0.05–0.1 yellow, 0.1–0.5 red)
            band_df = pd.DataFrame({
                "y0": [0.00, 0.05, 0.10],
                "y1": [0.05, 0.10, 0.50],
                "band": ["green", "yellow", "red"]
            })

            def fairness_chart(metric_col: str, title: str):
                base = alt.Chart(l3_df)

                bands = alt.Chart(band_df).mark_rect(opacity=0.15).encode(
                    y=alt.Y("y0:Q", title=None,
                            scale=alt.Scale(domain=(0, 0.5))),
                    y2="y1:Q",
                    color=alt.Color(
                        "band:N",
                        scale=alt.Scale(
                            domain=["green", "yellow", "red"],
                            range=["#00c851", "#ffde59", "#ff4b4b"]
                        ),
                        legend=None
                    )
                )

                line = base.mark_line(point=True).encode(
                    x=alt.X("`Run #`:Q", title="Run # (chronological)"),
                    y=alt.Y(
                        f"{metric_col}:Q",
                        title=title,
                        scale=alt.Scale(domain=(0, 0.5))
                    ),
                    tooltip=["Label:N", alt.Tooltip(
                        f"{metric_col}:Q", format=".3f")]
                ).properties(height=260, width="container")

                return (bands + line).resolve_scale(color="independent")

            c1, c2 = st.columns(2)
            with c1:
                st.altair_chart(
                    fairness_chart("DPG", "Demographic Parity Gap (↓ better)"),
                    width="stretch"
                )
            with c2:
                st.altair_chart(
                    fairness_chart(
                        "EOD", "Equal Opportunity Difference (↓ better)"),
                    width="stretch"
                )

            # Latest group diagnostics
            latest_l3 = l3_df.iloc[-1]
            st.markdown("#### 🔎 Latest Run – Group Diagnostics")
            col_a, col_b = st.columns(2)
            with col_a:
                st.write("**Positive Prediction Rate by Group**")
                st.json(latest_l3["rates_by_group"] or {})
            with col_b:
                st.write("**TPR (Sensitivity) by Group**")
                st.json(latest_l3["tpr_by_group"] or {})

st.markdown("</div>", unsafe_allow_html=True)

# ===== L3 Fairness — Group Bias Evolution (per-group DPG/EOD) =====
st.markdown("""
<div data-tour-target="l3-group-bias">
<h3>🧮 L3 Fairness — Group Bias Evolution</h3>
""", unsafe_allow_html=True)

# Make this section collapsible to avoid loading unless needed
show_l3_group_trends = st.checkbox(
    "📊 Show per-group bias evolution",
    value=False,
    key="show_l3_group_trends",
    help="View how fairness metrics evolved across different demographic groups over time"
)

if show_l3_group_trends:
    with measure_section("l3_group_fairness_processing"):
        # Collect all L3 reports efficiently using the helper function
        def extract_l3_group_metrics(report: dict) -> dict:
            """Extract L3 per-group metrics for bias evolution tracking"""
            m = report.get("metrics", {}) or {}
            return {
                "rates_by_group": m.get("rates_by_group") or {},
                "tpr_by_group": m.get("tpr_by_group") or {},
            }

        l3_df = collect_module_history(files, "L3", extract_l3_group_metrics)

        if l3_df.empty:
            st.info(
                "No L3 reports found yet. Run L3 multiple times to visualize per-group bias evolution.")
        else:
            # Longform per-group rates + TPR
            rows_rate, rows_tpr = [], []
            for _, r in l3_df.iterrows():
                run_idx = int(r["Run #"])
                lbl = r["Label"]
                rates = r["rates_by_group"] or {}
                tprs = r["tpr_by_group"] or {}

                # Calculate overall means (used to compute per-group gaps)
                overall_rate = (pd.Series(rates).mean() if rates else None)
                overall_tpr = (pd.Series(tprs).mean() if tprs else None)

                # Process positive prediction rates
                for g, v in (rates.items() if isinstance(rates, dict) else []):
                    if v is not None and overall_rate is not None:
                        rows_rate.append({
                            "Run #": run_idx,
                            "label": lbl,
                            "group": str(g),
                            "Rate": float(v),
                            # per-group parity gap
                            "DPG_group": float(v) - float(overall_rate)
                        })

                # Process TPR values
                for g, v in (tprs.items() if isinstance(tprs, dict) else []):
                    if v is not None and overall_tpr is not None:
                        rows_tpr.append({
                            "Run #": run_idx,
                            "label": lbl,
                            "group": str(g),
                            "TPR": float(v),
                            # per-group equal opp diff
                            "EOD_group": float(v) - float(overall_tpr)
                        })

            st.markdown("#### 👥 Positive Rate & DPG by Group (across runs)")
            if rows_rate:
                rate_df = pd.DataFrame(rows_rate)
                c1, c2 = st.columns(2)

                with c1:
                    ch = alt.Chart(rate_df).mark_line(point=True).encode(
                        x=alt.X("`Run #`:Q", title="Run #"),
                        y=alt.Y("Rate:Q", title="Positive Prediction Rate",
                                scale=alt.Scale(domain=(0, 1))),
                        color=alt.Color("group:N", title="Group"),
                        tooltip=["label:N", "group:N",
                                 alt.Tooltip("Rate:Q", format=".3f")]
                    ).properties(height=260, width="container", title="Positive Rate by Group")
                    st.altair_chart(ch, use_container_width=True)

                with c2:
                    ch = alt.Chart(rate_df).mark_line(point=True).encode(
                        x=alt.X("`Run #`:Q", title="Run #"),
                        y=alt.Y("DPG_group:Q",
                                title="Demographic Parity Gap (group − overall)"),
                        color=alt.Color("group:N", title="Group"),
                        tooltip=["label:N", "group:N", alt.Tooltip(
                            "DPG_group:Q", format="+.3f")]
                    ).properties(height=260, width="container", title="DPG by Group")
                    st.altair_chart(ch, use_container_width=True)
            else:
                st.info("No per-group positive rates available yet.")

            st.markdown("#### 👥 TPR & EOD by Group (across runs)")
            if rows_tpr:
                tpr_df = pd.DataFrame(rows_tpr)
                c3, c4 = st.columns(2)

                with c3:
                    ch = alt.Chart(tpr_df).mark_line(point=True).encode(
                        x=alt.X("`Run #`:Q", title="Run #"),
                        y=alt.Y("TPR:Q", title="True Positive Rate",
                                scale=alt.Scale(domain=(0, 1))),
                        color=alt.Color("group:N", title="Group"),
                        tooltip=["label:N", "group:N",
                                 alt.Tooltip("TPR:Q", format=".3f")]
                    ).properties(height=260, width="container", title="TPR by Group")
                    st.altair_chart(ch, use_container_width=True)

                with c4:
                    # Threshold bands for |EOD_group|
                    band_df = pd.DataFrame({
                        "y0": [-0.50, -0.10, -0.05,  0.05,  0.10],
                        "y1": [-0.10, -0.05,  0.05,  0.10,  0.50],
                        "band": ["red", "yellow", "green", "yellow", "red"]
                    })
                    bands = alt.Chart(band_df).mark_rect(opacity=0.12).encode(
                        y=alt.Y("y0:Q", title=None),
                        y2="y1:Q",
                        color=alt.Color("band:N", scale=alt.Scale(
                            domain=["green", "yellow", "red"],
                            range=["#00c851", "#ffde59", "#ff4b4b"]), legend=None)
                    )
                    line = alt.Chart(tpr_df).mark_line(point=True).encode(
                        x=alt.X("`Run #`:Q", title="Run #"),
                        y=alt.Y("EOD_group:Q",
                                title="Equal Opportunity Diff (group − overall)"),
                        color=alt.Color("group:N", title="Group"),
                        tooltip=["label:N", "group:N", alt.Tooltip(
                            "EOD_group:Q", format="+.3f")]
                    ).properties(height=260, width="container", title="EOD by Group")
                    st.altair_chart((bands + line), use_container_width=True)
            else:
                st.info("No per-group TPR values available yet.")
else:
    st.info("👆 Enable the checkbox above to view per-group fairness trends across multiple evaluation runs.")

st.markdown("</div>", unsafe_allow_html=True)

# ===== Mitigation Recommendations (heuristic) =====
st.markdown("#### 🩹 Fairness Mitigation Recommendations")

latest_l3 = None
for f in sorted([p for p in files if os.path.basename(p).startswith("L3-")]):
    pass
# last one in chronological list:
l3_candidates = sorted([p for p in files if os.path.basename(
    p).startswith("L3-")], key=lambda p: os.path.getmtime(p))
if l3_candidates:
    latest_l3 = load_json(l3_candidates[-1])

if not latest_l3:
    st.info("No latest L3 run detected.")
else:
    m = latest_l3.get("metrics", {}) or {}
    dpg = m.get("DPG")
    eod = m.get("EOD")
    tpr_by_group = m.get("tpr_by_group") or {}
    rates_by_group = m.get("rates_by_group") or {}

    recs = []
    # thresholds matching your L3 scoring targets
    if isinstance(eod, (int, float)) and eod > 0.10:
        recs.append(
            "Equal Opportunity gap is high (>0.10). Increase positive examples for the low-TPR group or adjust threshold per group.")
    elif isinstance(eod, (int, float)) and eod > 0.05:
        recs.append(
            "Equal Opportunity gap is moderate (>0.05). Try threshold tuning or class-balanced loss.")

    if isinstance(dpg, (int, float)) and dpg > 0.05:
        recs.append(
            "Demographic Parity gap >0.05. Consider rebalancing decision thresholds or post-processing to equalize positive rates.")

    # diagnose which group lags by TPR
    try:
        if tpr_by_group:
            g_min = min(tpr_by_group, key=lambda k: (
                999 if tpr_by_group[k] is None else tpr_by_group[k]))
            recs.append(
                f"Group '{g_min}' has the lowest sensitivity (TPR). Add targeted data or feature engineering for this group.")
    except Exception:
        pass

    if not recs:
        recs = ["No major fairness flags detected. Keep monitoring drift across runs."]

    for r in recs:
        st.markdown(f"- {r}")

# ===== Bias Attribution Heatmap (per-group permutation importance) =====
st.markdown("### 🔥 Bias Attribution Heatmap")

default_path = "data/explain_with_group.csv"
opt_col1, opt_col2 = st.columns([2, 1])
with opt_col1:
    st.caption(
        "Provide a CSV with columns: **y**, **group**, and one or more numeric feature columns.")
    user_file = st.file_uploader("Upload CSV (optional override)", type=[
                                 "csv"], disabled=_LOCK, key="l3_attr_upl")
with opt_col2:
    use_default = st.checkbox(f"Use default: {default_path}", value=(
        user_file is None), disabled=_LOCK)

csv_source = None
if user_file is not None and not use_default:
    csv_source = user_file
elif use_default and os.path.exists(default_path):
    csv_source = default_path

if csv_source is None:
    st.info("No CSV available. Upload a file or enable the default path.")
else:
    try:
        from sklearn.linear_model import LogisticRegression
        from sklearn.inspection import permutation_importance

        dfx = pd.read_csv(csv_source).dropna()
        if not {"y", "group"}.issubset(dfx.columns):
            st.warning("CSV must contain 'y' and 'group' plus feature columns.")
        else:
            label_col = "y"
            features = [c for c in dfx.columns if c not in (
                label_col, "group")]
            X0 = dfx[features].apply(pd.to_numeric, errors="coerce")
            dfx = dfx.loc[X0.dropna().index]
            X = dfx[features].astype(float).values
            y = dfx[label_col].values
            groups = dfx["group"].astype(str).values

            clf = LogisticRegression(max_iter=1000)
            clf.fit(X, y)

            rows = []
            for g in sorted(pd.unique(groups)):
                mask = (groups == g)
                if mask.sum() < 5:
                    continue
                perm = permutation_importance(
                    clf, X[mask], y[mask], n_repeats=6, random_state=42)
                imp = perm.importances_mean.clip(min=0.0)
                for j, feat in enumerate(features):
                    rows.append({"group": g, "feature": feat,
                                "importance": float(imp[j])})

            if not rows:
                st.info("Not enough samples per group to compute attribution.")
            else:
                heat = pd.DataFrame(rows)
                heat["norm"] = heat.groupby("group")["importance"].transform(
                    lambda s: (s - s.min()) / (s.max() - s.min() + 1e-9)
                )
                chart = alt.Chart(heat).mark_rect().encode(
                    x=alt.X("feature:N", title="Feature"),
                    y=alt.Y("group:N", title="Group"),
                    color=alt.Color(
                        "norm:Q", title="Relative Importance", scale=alt.Scale(scheme="blues")),
                    tooltip=["group", "feature", alt.Tooltip(
                        "importance:Q", format=".4f")]
                ).properties(height=260, width="container", title="Per-Group Attribution (Permutation Importance)")
                st.altair_chart(chart, use_container_width=True)
    except Exception as e:
        st.warning(f"Heatmap unavailable: {e}")


# ===== Data Provenance & Mitigation Recommendations =====
st.markdown("### 📜 Data Provenance (L3)")
prov = {
    "dataset_version": (latest.get("L3") or {}).get("metrics", {}).get("dataset_version"),
    "last_train_date": (latest.get("L3") or {}).get("metrics", {}).get("last_train_date"),
}
# fallback from project config if missing
try:
    with open(os.path.join("configs", "project.example.yaml"), "r", encoding="utf-8") as fh:
        proj_cfg = yaml.safe_load(fh) or {}
    prov.setdefault("dataset_version", proj_cfg.get(
        "L3", {}).get("dataset_version"))
    prov.setdefault("last_train_date", proj_cfg.get(
        "L3", {}).get("last_train_date"))
except Exception:
    pass

st.json({k: v for k, v in prov.items() if v is not None})
if not any(prov.values()):
    st.caption(
        "Tip: add `dataset_version` and `last_train_date` to L3 metrics or project config for audit trails.")

# Mitigation recommendations
st.markdown("#### 🩹 Mitigation Recommendations")
recos = []
# Simple rules:
latest_l3 = latest.get("L3") or {}
DPG = (latest_l3.get("metrics") or {}).get("DPG")
EOD = (latest_l3.get("metrics") or {}).get("EOD")
if DPG is not None and float(DPG) > 0.1:
    recos.append(
        "High Demographic Parity Gap → consider class rebalancing or group-aware thresholding.")
if EOD is not None and float(EOD) > 0.1:
    recos.append(
        "Equal Opportunity Difference is high → inspect TPR per group and retrain with focal loss or group sampling.")
if not recos:
    st.caption("No high-priority mitigation suggestions at this time.")
else:
    for r in recos:
        st.write(f"• {r}")

# ===== L4 Explainability – Interactive =====
st.markdown("""
<div data-tour-target="l4-explainability">
<h3>🔍 L4 Explainability – Interactive</h3>
""", unsafe_allow_html=True)

l4 = latest.get("L4")

if not (l4 and "metrics" in l4):
    st.info("Run L4 to populate explainability metrics.")
else:
    m = l4["metrics"]
    # importances: prefer SHAP if present and user selects it
    shap_imp = m.get("shap_importance")
    perm_imp = m.get("feature_importance") or {}
    has_shap = bool(shap_imp)

    # --- Health gauge ---
    deletion_drop = float(m.get("deletion_drop", 0.0))
    stability_tau = float(m.get("stability_tau", 0.0))
    l4_score = float(l4.get("score", 0.0))

    with st.container():
        st.caption(
            "Health uses deletion-drop (↑) and stability τ (↑). Targets: deletion_drop ≥ 0.15, τ ≥ 0.85.")
        st.markdown(
            f"<h4 style='text-align:center;'>Explainability Health</h4>", unsafe_allow_html=True)
        st.markdown(
            f"<h2 style='text-align:center;color:#444;'>{'🟢' if l4_score >= 90 else ('🟡' if l4_score >= 75 else '🔴')} {l4_score:.2f}</h2>", unsafe_allow_html=True)
        bar_html = f"""
        <div style='background:#eee;border-radius:10px;height:20px;margin-bottom:6px;'>
          <div style='width:{min(max(l4_score, 0), 100)}%;background:linear-gradient(90deg,#ff4b4b,#ffde59,#00c851);
          height:20px;border-radius:10px;'></div>
        </div>
        """
        st.markdown(bar_html, unsafe_allow_html=True)
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Baseline", f"{m.get('baseline_metric', 0):.3f}")
        c2.metric("Masked", f"{m.get('masked_metric', 0):.3f}")
        c3.metric("Deletion drop",
                  f"{deletion_drop:.3f}", help="Higher is better; ≥ 0.15")
        c4.metric("Stability τ", f"{stability_tau:.3f}",
                  help="Higher is better; ≥ 0.85")

    # --- Importance toggle + Top-N slider ---
    src = st.radio(
        "Importance source",
        ["Permutation", "SHAP (if available)"],
        index=0 if not has_shap else 1,
        horizontal=True,
        help="Choose which importance values to visualize."
    )
    top_n = st.slider("Top-N features", min_value=3, max_value=max(3,
                      len((shap_imp or perm_imp))), value=min(10, len((shap_imp or perm_imp))))

    chosen = shap_imp if (src.startswith("SHAP") and has_shap) else perm_imp
    if not chosen:
        st.info("No importance values found yet.")
    else:
        imp_df = (
            pd.DataFrame({"feature": list(chosen.keys()), "importance": [
                         float(v) for v in chosen.values()]})
            .sort_values("importance", ascending=False)
            .head(top_n)
        )
        total = float(imp_df["importance"].sum()) or 1.0
        imp_df["share_%"] = (imp_df["importance"] / total) * 100.0

        st.caption("Tip: use the slider to focus on the top drivers.")
        chart = alt.Chart(imp_df).mark_bar().encode(
            x=alt.X("importance:Q", title="Importance"),
            y=alt.Y("feature:N", sort="-x", title="Feature"),
            tooltip=["feature", alt.Tooltip(
                "importance:Q", format=".4f"), alt.Tooltip("share_%:Q", format=".2f")]
        ).properties(height=280, width="container")
        st.altair_chart(chart, theme=None)

        st.download_button(
            "⬇️ Download current importances (CSV)",
            data=imp_df.to_csv(index=False).encode("utf-8"),
            file_name=f"L4_importances_{'shap' if (src.startswith('SHAP') and has_shap) else 'perm'}.csv",
            mime="text/csv",
            key="dl_l4_imp",
        )

st.markdown("</div>", unsafe_allow_html=True)

# --- Trends across runs: deletion_drop & stability_tau ---
# Scan all L4 reports and plot over time
l4_rows = []
for f in files:
    try:
        with open(f, "r", encoding="utf-8") as fh:
            d = json.load(fh)
        if d.get("module") != "L4":
            continue
        base = os.path.basename(f).replace(".json", "")
        # label from filename; try to parse timestamp like L4-YYYYMMDD-HHMMSS.json
        ts_match = re.search(r"(\d{8}-\d{6})", base)
        ts_label = ts_match.group(1) if ts_match else base
        mtime = os.path.getmtime(f)
        md = d.get("metrics", {})
        l4_rows.append({
            "Run": base,
            "Label": ts_label,
            "time": mtime,
            "deletion_drop": float(md.get("deletion_drop", 0.0)),
            "stability_tau": float(md.get("stability_tau", 0.0)),
            "score": float(d.get("score", 0.0)),
        })
    except Exception:
        pass

if l4_rows:
    import re
    import os
    l4_df = pd.DataFrame(l4_rows).sort_values(
        "time").reset_index(drop=True)
    l4_df["Run #"] = l4_df.index + 1

    st.markdown("#### 📈 L4 Trends — Deletion drop & Stability τ")
    # background bands for thresholds
    band_df_drop = pd.DataFrame(
        {"y0": [0.00, 0.15], "y1": [0.15, 1.00], "band": ["red", "green"]})
    band_df_tau = pd.DataFrame(
        {"y0": [0.00, 0.85], "y1": [0.85, 1.00], "band": ["red", "green"]})

    def banded(metric, title, domain=(0, 1)):
        bands = alt.Chart(
            band_df_drop if metric == "deletion_drop" else band_df_tau
        ).mark_rect(opacity=0.12).encode(
            y=alt.Y("y0:Q", title=None, scale=alt.Scale(domain=domain)),
            y2="y1:Q",
            color=alt.Color("band:N", scale=alt.Scale(
                domain=["red", "green"],
                range=["#ff4b4b", "#00c851"]
            ), legend=None)
        )
        line = alt.Chart(l4_df).mark_line(point=True).encode(
            x=alt.X("`Run #`:Q", title="Run # (chronological)"),
            y=alt.Y(f"{metric}:Q", title=title,
                    scale=alt.Scale(domain=domain)),
            tooltip=["Label:N", alt.Tooltip(
                f"{metric}:Q", format=".3f"), alt.Tooltip("score:Q", format=".1f")]
        ).properties(height=250, width="container")
        return (bands + line).resolve_scale(color="independent")

    cA, cB = st.columns(2)
    with cA:
        st.altair_chart(banded(
            "deletion_drop", "Deletion drop (≥ 0.15 target)"), width="stretch")
    with cB:
        st.altair_chart(banded(
            "stability_tau", "Stability τ (≥ 0.85 target)"), width="stretch")
else:
    st.info("No historical L4 runs yet to plot trends.")

st.markdown("</div>", unsafe_allow_html=True)


# ===== L4 — SHAP Summary (split view: Beeswarm ▸ left | Insights ▸ right) =====
st.markdown("### 🐝 L4 – SHAP Summary")

# --- Try to auto-pick a CSV from L4 evidence if the user didn't upload one
l4_rep = latest.get("L4") or {}
csv_from_evidence = None
try:
    ev = l4_rep.get("evidence") or []
    if isinstance(ev, list) and ev:
        for p in ev:
            if isinstance(p, str) and p.lower().endswith(".csv") and os.path.exists(p):
                csv_from_evidence = p
                break
except Exception:
    csv_from_evidence = None

# --- Inputs (upload + hints)
c_up, c_help = st.columns([2, 1])
with c_up:
    shap_csv = st.file_uploader(
        "Upload CSV (must include a label column + numeric features)",
        type=["csv"], key="l4_beeswarm_upl", disabled=_LOCK
    )
with c_help:
    st.caption(
        "If no upload, the first CSV found in **L4 evidence** will be used (when available).")

# --- Settings
c_left, c_right = st.columns([2, 1])
with c_left:
    sample_n = st.slider(
        "Sample N for plot (to keep it readable)",
        min_value=100, max_value=2000, value=600, step=100, disabled=_LOCK,
        help="Only affects visualization; SHAP is computed on the test split."
    )
with c_right:
    label_guess = (l4_rep.get("metrics", {}) or {}).get(
        "label_col") or l4_rep.get("label_col") or "y"
    label_col = st.text_input(
        "Label column name", value=label_guess, disabled=_LOCK,
        help="Change if your label column has a different name."
    )

# --- Decide which CSV to use
csv_path_to_use = shap_csv if shap_csv is not None else (
    csv_from_evidence if csv_from_evidence and os.path.exists(
        csv_from_evidence) else None
)

if csv_path_to_use is None:
    st.info(
        "No CSV available. Upload one above or ensure L4 evidence contains a CSV file."
    )
else:
    try:
        import numpy as np
        import pandas as pd
        import altair as alt
        from sklearn.linear_model import LogisticRegression
        from sklearn.model_selection import train_test_split
        from sklearn.metrics import roc_auc_score, accuracy_score

        # Deterministic split so visuals are stable
        SEED = 42
        np.random.seed(SEED)

        df = pd.read_csv(csv_path_to_use)
        if label_col not in df.columns:
            raise ValueError(f"Label column '{label_col}' not found in CSV.")

        feature_cols = [c for c in df.columns if c != label_col]
        for c in feature_cols:
            df[c] = pd.to_numeric(df[c], errors="coerce")
        df = df.dropna(subset=[label_col] + feature_cols)
        if df.empty:
            raise ValueError(
                "No rows left after dropping NaNs; check your CSV."
            )

        X = df[feature_cols].values
        y = df[label_col].values

        X_tr, X_te, y_tr, y_te = train_test_split(
            X,
            y,
            test_size=0.3,
            random_state=SEED,
            stratify=y if len(np.unique(y)) == 2 else None,
        )
        clf = LogisticRegression(max_iter=1000)
        clf.fit(X_tr, y_tr)

        # Compute model metric (for the Insights panel)
        is_binary = len(np.unique(y)) == 2 and hasattr(clf, "predict_proba")
        if is_binary:
            prob = clf.predict_proba(X_te)[:, 1]
            model_metric = roc_auc_score(y_te, prob)
            metric_label = "AUROC"
        else:
            pred = clf.predict(X_te)
            model_metric = accuracy_score(y_te, pred)
            metric_label = "Accuracy"

        # ========== LAYOUT: LEFT (beeswarm) | RIGHT (insights) ==========
        colL, colR = st.columns([7, 5])

        # ---------- LEFT: SHAP Beeswarm (matplotlib if available; Altair fallback) ----------
        with colL:
            computed_shap_vals = None  # Initialize to prevent undefined errors

            try:
                import shap
                import matplotlib.pyplot as plt

                with measure_section("shap_computation"):
                    try:
                        expl = shap.LinearExplainer(
                            clf, X_tr, feature_perturbation="interventional"
                        )
                        shap_vals = expl.shap_values(X_te)
                        computed_shap_vals = shap_vals  # Save for later use

                    except Exception as shap_error:
                        logger.error(
                            f"SHAP computation failed: {shap_error}", exc_info=True
                        )
                        st.error(
                            "⚠️ SHAP computation failed. Falling back to permutation importance."
                        )

                        # Fallback to permutation importance
                        from sklearn.inspection import permutation_importance

                        perm = permutation_importance(
                            clf, X_te, y_te, n_repeats=6, random_state=SEED
                        )

                        # Convert to SHAP-like format (per-sample)
                        shap_vals = np.tile(
                            perm.importances_mean, (X_te.shape[0], 1)
                        )
                        computed_shap_vals = shap_vals

                    # Downsample for plotting aesthetics (only the figure)
                    if X_te.shape[0] > sample_n:
                        idx = np.random.choice(
                            X_te.shape[0], size=sample_n, replace=False
                        )
                        X_plot = X_te[idx]
                        shap_plot_vals = shap_vals[idx]
                    else:
                        X_plot = X_te
                        shap_plot_vals = shap_vals

                st.caption(
                    "Beeswarm shows the distribution of feature impacts across samples (|SHAP|)."
                )
                fig = plt.figure(figsize=(9.5, 5.2))
                plt.title("SHAP Beeswarm", fontsize=12, pad=8)
                shap.summary_plot(
                    shap_plot_vals,
                    features=X_plot,
                    feature_names=feature_cols,
                    show=False,
                    plot_type="dot",
                )
                st.pyplot(fig, clear_figure=True)
                # Keep shap_vals around for the right-side panels
                computed_shap_vals = shap_vals

            except ImportError:
                st.warning(
                    "⚠️ SHAP library not installed. Install with: `pip install shap`")
                st.info("Using permutation importance as fallback...")

                # Fallback visualization using permutation importance
                from sklearn.inspection import permutation_importance
                perm = permutation_importance(
                    clf, X_te, y_te, n_repeats=6, random_state=SEED)

                # Create simple bar chart
                imp_df = pd.DataFrame({
                    "feature": feature_cols,
                    "importance": perm.importances_mean
                }).sort_values("importance", ascending=False).head(10)

                chart = alt.Chart(imp_df).mark_bar().encode(
                    x=alt.X("importance:Q", title="Importance"),
                    y=alt.Y("feature:N", sort="-x", title="Feature"),
                    tooltip=["feature", alt.Tooltip(
                        "importance:Q", format=".4f")]
                ).properties(height=320, width="container")

                st.altair_chart(chart, use_container_width=True)
                computed_shap_vals = None

            except Exception as e:
                show_error_inline(e, "SHAP visualization failed")
                computed_shap_vals = None

            st.caption(
                f"Random seed: {SEED} (deterministic splits for stable visuals)"
            )

        # ---------- RIGHT: Insights panel (metric + mean|SHAP| + PDP-lite) ----------
        with colR:
            st.markdown("#### 📊 Model Insight")
            # Guard against missing SHAP values
            if computed_shap_vals is None:
                st.warning(
                    "⚠️ SHAP values not available. Model insights limited.")
                st.info(
                    "Install SHAP library for full feature importance analysis: `pip install shap matplotlib`")

                # Show basic model metrics only
                c1, c2 = st.columns([1, 1])
                with c1:
                    st.metric(metric_label, f"{model_metric:.3f}")
                with c2:
                    st.caption(f"Test split size: {X_te.shape[0]}")

            else:
                c1, c2 = st.columns([1, 1])
            with c1:
                st.metric(metric_label, f"{model_metric:.3f}")
            with c2:
                st.caption(f"Test split size: {X_te.shape[0]}")

            # Mean absolute SHAP per feature
            import numpy as _np
            mean_abs = _np.mean(_np.abs(computed_shap_vals), axis=0)
            top_df = (
                pd.DataFrame({"Feature": feature_cols,
                              "Mean |SHAP|": mean_abs})
                .sort_values("Mean |SHAP|", ascending=False)
                .reset_index(drop=True)
            )

            top_k = st.slider("Top-K features", 3,
                              min(20, len(top_df)), value=min(10, len(top_df)))
            top_view = top_df.head(top_k)

            bar = alt.Chart(top_view).mark_bar().encode(
                y=alt.Y("Feature:N", sort="-x", title=None),
                x=alt.X("Mean |SHAP|:Q", title="Mean |SHAP|"),
                tooltip=[alt.Tooltip("Feature:N"), alt.Tooltip(
                    "Mean |SHAP|:Q", format=".4f")],
            ).properties(height=max(220, 20 * len(top_view)), width="container")
            st.altair_chart(bar, use_container_width=True)

            # PDP-lite: choose a feature, show average prediction vs binned values
            st.markdown("##### 🔎 Partial Dependence (quick view)")
            pick_feat = st.selectbox(
                "Feature", options=list(top_view["Feature"]), index=0, disabled=_LOCK
            )

            # Build a quick PDP by binning the chosen feature and averaging predictions
            try:
                # Use test split for PDP to reflect what you saw in SHAP
                col_idx = feature_cols.index(pick_feat)
                fvals = X_te[:, col_idx]
                # Quantile bins for robustness
                bins = np.unique(np.quantile(fvals, np.linspace(0, 1, 11)))
                # Guard against degenerate constant feature
                if len(bins) < 3:
                    bins = np.linspace(fvals.min(), fvals.max() + 1e-9, 10)

                cats = pd.cut(fvals, bins=bins, include_lowest=True)
                if is_binary:
                    yhat = prob
                else:
                    # Approx prediction score for multi-class: use decision_function if present else proba argmax
                    try:
                        yhat = clf.decision_function(X_te)
                        if yhat.ndim > 1:
                            yhat = yhat.max(axis=1)
                    except Exception:
                        if hasattr(clf, "predict_proba"):
                            yhat = clf.predict_proba(X_te).max(axis=1)
                        else:
                            yhat = clf.predict(X_te)

                pdp_df = pd.DataFrame({"bin": cats.astype(str), "yhat": yhat})
                pdp_view = pdp_df.groupby("bin", as_index=False)["yhat"].mean()

                line = alt.Chart(pdp_view).mark_line(point=True).encode(
                    x=alt.X("bin:N", title=f"{pick_feat} (binned)"),
                    y=alt.Y("yhat:Q", title="Avg prediction"),
                    tooltip=["bin", alt.Tooltip("yhat:Q", format=".3f")],
                ).properties(height=220, width="container")
                st.altair_chart(line, use_container_width=True)
            except Exception:
                st.caption("PDP not available for this feature.")

            # Download mean |SHAP|
            st.download_button(
                "⬇️ Download mean |SHAP| (CSV)",
                data=top_df.to_csv(index=False).encode("utf-8"),
                file_name="mean_abs_shap.csv",
                mime="text/csv",
                disabled=_LOCK
            )

    except Exception as e:
        # inline, friendly error
        import traceback as _tb
        tb = "".join(_tb.format_exception_only(type(e), e)).strip()
        st.error(f"⚠️ Could not render SHAP summary: {tb}")


# ===== L4 – Compare Models (A vs B) =====
st.markdown("### ⚖️ L4 – Compare Models (A vs B)")

# Deterministic seed for reproducibility (auditable)
SEED = 42
np.random.seed(SEED)
random.seed(SEED)


@st.cache_resource  # cache heavy training for faster re-runs
def _fit_and_measure(csv_or_path, label_col="y", seed: int = SEED):
    """
    Train a simple logistic model deterministically and compute:
      - permutation importances (dict feature->score)
      - deletion_drop (float)
      - base metric (AUROC for binary else accuracy)
    Returns: (importances_dict, deletion_drop, base_metric, seed)
    """
    import pandas as pd
    from sklearn.linear_model import LogisticRegression
    from sklearn.model_selection import train_test_split
    from sklearn.metrics import accuracy_score, roc_auc_score
    from sklearn.inspection import permutation_importance

    # Load CSV
    dfm = pd.read_csv(csv_or_path) if hasattr(
        csv_or_path, "read") else pd.read_csv(csv_or_path)
    feats = [c for c in dfm.columns if c != label_col]
    if not feats:
        raise ValueError(
            "No feature columns found (every column is the label).")

    # Coerce numerics & drop NaNs
    for c in feats:
        dfm[c] = pd.to_numeric(dfm[c], errors="coerce")
    dfm = dfm.dropna(subset=[label_col] + feats)
    if dfm.empty:
        raise ValueError(
            "No rows left after dropping NaNs. Check your CSV formatting.")

    X = dfm[feats].values
    y = dfm[label_col].values

    # Deterministic split/model
    X_tr, X_te, y_tr, y_te = train_test_split(
        X, y, test_size=0.3, random_state=seed,
        stratify=y if len(np.unique(y)) == 2 else None
    )
    clf = LogisticRegression(max_iter=1000, random_state=seed).fit(X_tr, y_tr)

    # Base metric
    binary = len(np.unique(y)) == 2
    prob = clf.predict_proba(X_te)[:, 1] if (
        binary and hasattr(clf, "predict_proba")) else None
    pred = clf.predict(X_te)
    kind = "auroc" if binary else "accuracy"
    base = roc_auc_score(y_te, prob) if (
        kind == "auroc" and prob is not None) else accuracy_score(y_te, pred)

    # Permutation importance
    perm = permutation_importance(
        clf, X_te, y_te, n_repeats=6, random_state=seed)
    imp = perm.importances_mean.clip(min=0.0)

    # Deletion-drop: mask top feature
    top_idx = int(np.argsort(imp)[::-1][0]) if len(imp) > 0 else 0
    train_means = X_tr.mean(axis=0)
    X_mask = X_te.copy()
    if len(imp) > 0:
        X_mask[:, top_idx] = train_means[top_idx]
    prob_m = clf.predict_proba(X_mask)[:, 1] if (
        binary and hasattr(clf, "predict_proba")) else None
    pred_m = clf.predict(X_mask)
    masked = roc_auc_score(y_te, prob_m) if (
        kind == "auroc" and prob_m is not None) else accuracy_score(y_te, pred_m)
    drop = float(base - masked)

    return ({feats[i]: float(imp[i]) for i in range(len(feats))}, drop, float(base), seed)


# Uploads
cA, cB = st.columns(2)
with cA:
    st.write("**Model A CSV**")
    uplA = st.file_uploader("Upload CSV A", type=[
                            "csv"], disabled=_LOCK, key="l4_csv_a")
with cB:
    st.write("**Model B CSV**")
    uplB = st.file_uploader("Upload CSV B", type=[
                            "csv"], disabled=_LOCK, key="l4_csv_b")

if uplA and uplB:
    try:
        impA, dropA, baseA, seedA = cached_fit(uplA)
        impB, dropB, baseB, seedB = cached_fit(uplB)

        st.caption(
            f"Random seed: {SEED} — ensures deterministic and auditable A/B results")

        # Build union of features & deltas
        import pandas as pd
        import altair as alt
        keys = sorted(set(impA.keys()) | set(impB.keys()))
        rows = []
        for k in keys:
            a = float(impA.get(k, 0.0))
            b = float(impB.get(k, 0.0))
            rows.append({"feature": k, "A_importance": a,
                        "B_importance": b, "Δ(B−A)": b - a})
        delta_df = pd.DataFrame(rows).sort_values("Δ(B−A)", ascending=False)

        st.markdown("#### 📊 Top Feature Delta (A vs B)")
        if delta_df.empty:
            st.info("No overlapping or valid features to compare.")
        else:
            st.dataframe(delta_df, use_container_width=True)

            # ---- Deletion-drop / Metric Delta ----
            st.markdown("#### 🧪 Deletion-drop / Metric Delta")
            c1, c2, c3 = st.columns(3)
            c1.metric("A deletion_drop", f"{dropA:.3f}")
            c2.metric("B deletion_drop",
                      f"{dropB:.3f}", f"{(dropB - dropA):+.3f}")
            c3.metric("Δ metric (B−A)", f"{(baseB - baseA):+.3f}")

            # Guard slider bounds even if very few features
            topn_max = max(1, min(15, len(delta_df)))
            topn_default = min(10, topn_max)
            topn = st.slider("Show top-N features", 1, topn_max,
                             topn_default, disabled=_LOCK, key="l4_topn")
            show_df = delta_df.head(topn)

            # ---- Δ(B−A) bar chart ----
            st.markdown("#### ➕ Importance Delta (B − A)")
            ch_delta = alt.Chart(show_df).mark_bar().encode(
                x=alt.X("Δ(B−A):Q", title="Change in importance (B − A)"),
                y=alt.Y("feature:N", sort="-x", title="Feature"),
                tooltip=["feature", alt.Tooltip("Δ(B−A):Q", format="+.4f")]
            ).properties(height=280, width="container")
            st.altair_chart(ch_delta, use_container_width=True)

            # ---- Overlayed importances (A vs B) ----
            st.markdown("#### 🔍 Feature Importances (A vs B)")
            imp_long = pd.melt(
                show_df[["feature", "A_importance", "B_importance"]],
                id_vars="feature",
                var_name="model",
                value_name="importance"
            )
            imp_long["model"] = imp_long["model"].map(
                {"A_importance": "A", "B_importance": "B"})
            bar = alt.Chart(imp_long).mark_bar().encode(
                y=alt.Y("feature:N", sort="-x", title="Feature"),
                x=alt.X("importance:Q", title="Importance"),
                color=alt.Color("model:N", title="Model"),
                tooltip=["feature", "model", alt.Tooltip(
                    "importance:Q", format=".4f")]
            ).properties(height=280, width="container")
            st.altair_chart(bar, use_container_width=True)

            # ---- Deletion-drop A vs B (slope chart) ----
            st.markdown("#### ⛳ Deletion-drop (A vs B)")
            dd_df = pd.DataFrame(
                {"model": ["A", "B"], "deletion_drop": [dropA, dropB]})
            slope = alt.Chart(dd_df).mark_line(point=True).encode(
                x=alt.X("model:N", title="Model"),
                y=alt.Y("deletion_drop:Q", title="Deletion drop"),
                tooltip=["model", alt.Tooltip("deletion_drop:Q", format=".3f")]
            ).properties(height=220, width=380)
            st.altair_chart(slope, use_container_width=False)

    except Exception as e:
        import traceback
        tb = "".join(traceback.format_exception_only(type(e), e)).strip()
        st.error(f"⚠️ L4 comparison failed: {tb}")
else:
    st.caption("Upload two CSVs to compare models’ importances and deletion-drop.")

# ===== L4 – Composite Health Meter =====
st.markdown("### 🧭 L4 – Composite Health Meter")
m = (latest.get("L4") or {}).get("metrics", {}) or {}
dd = float(m.get("deletion_drop", 0.0))
tau = float(m.get("stability_tau", 0.0))
inf = float(m.get("infidelity", 1.0))  # 0 good, 1 bad (proxy)

# normalize components to 0..1 against demo targets
dd_norm = min(max(dd / 0.15, 0), 1)          # ≥0.15 is "1.0"
tau_norm = min(max(tau / 0.85, 0), 1)        # ≥0.85 is "1.0"
inf_norm = 1.0 - min(max(inf, 0), 1)         # lower is better → invert

health = (0.45*dd_norm) + (0.45*tau_norm) + (0.10*inf_norm)
health_pct = 100.0 * health
badge = "🟢" if health_pct >= 90 else ("🟡" if health_pct >= 75 else "🔴")
st.markdown(
    f"<div style='background:var(--metric-bg);border:1px solid var(--border);border-radius:10px;padding:10px;'>"
    f"<b>Composite Health:</b> {badge} {health_pct:.1f}%<br>"
    f"• deletion_drop={dd:.3f} (target ≥0.15) • stability_tau={tau:.3f} (target ≥0.85) • infidelity={inf:.2f}"
    f"</div>", unsafe_allow_html=True
)

# ===== Multi-model Comparison (bulk) =====
st.markdown("### 🧪 Multi-model Comparison (bulk)")
mm_files = st.file_uploader(
    "Upload 2+ CSVs (same schema)",
    type=["csv"], accept_multiple_files=True, disabled=_LOCK, key="l4_multi_csvs"
)
if mm_files:
    import pandas as pd
    rows = []
    for i, fobj in enumerate(mm_files, start=1):
        try:
            imp, drop, base, _seed = _fit_and_measure(fobj)  # unpack 4-tuple
            rows.append(
                {"Model": f"Model_{i}", "Metric": base, "Deletion_drop": drop})
        except Exception as e:
            rows.append(
                {"Model": f"Model_{i}", "Metric": None, "Deletion_drop": None})
    st.dataframe(pd.DataFrame(rows), use_container_width=True)


# ===== Model Card (DOCX) =====
st.markdown("### 🧾 Model Card (DOCX)")


def build_model_card(latest_bundle: dict) -> bytes:
    try:
        from docx import Document  # type: ignore
    except Exception:
        st.error("python-docx not installed. pip install python-docx")
        return b""
    doc = Document()
    doc.add_heading("Model Card — IRAQAF", level=1)
    doc.add_paragraph("Generated from current dashboard state.")
    # L4
    l4r = latest_bundle.get("L4") or {}
    l4m = l4r.get("metrics", {}) or {}
    doc.add_heading("L4 — Explainability", level=2)
    doc.add_paragraph(f"Score: {l4r.get('score')}  |  Band: {l4r.get('band')}")
    doc.add_paragraph(
        f"Deletion drop: {l4m.get('deletion_drop')}, Stability τ: {l4m.get('stability_tau')}, Infidelity: {l4m.get('infidelity')}")
    # L3
    l3r = latest_bundle.get("L3") or {}
    l3m = l3r.get("metrics", {}) or {}
    doc.add_heading("L3 — Fairness", level=2)
    doc.add_paragraph(f"Score: {l3r.get('score')}  |  Band: {l3r.get('band')}")
    doc.add_paragraph(f"DPG: {l3m.get('DPG')}  |  EOD: {l3m.get('EOD')}")
    # L1
    l1r = latest_bundle.get("L1") or {}
    l1m = l1r.get("metrics", {}) or {}
    doc.add_heading("L1 — Governance/Compliance", level=2)
    doc.add_paragraph(f"Score: {l1r.get('score')}  |  Band: {l1r.get('band')}")
    doc.add_paragraph(f"Coverage: {l1m.get('coverage_percent')}%")
    # Evidence
    doc.add_heading("Evidence", level=2)
    for mid in ["L1", "L2", "L3", "L4", "L5"]:
        rep = latest_bundle.get(mid) or {}
        ev = rep.get("evidence") or []
        if ev:
            doc.add_paragraph(f"{mid}:")
            for e in ev:
                doc.add_paragraph(f" - {e}")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


if st.button("Generate Model Card (Word)", disabled=_LOCK):
    payload = build_model_card(latest)
    if payload:
        st.download_button(
            "⬇️ Download Model_Card.docx",
            data=payload,
            file_name="Model_Card.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_model_card"
        )


# ===== L5 Operations – Trends Over Time =====
st.markdown("### 🛠️ L5 Operations – Trends Over Time")
# Collect all L5 reports
l5_rows = []
for f in files:
    try:
        with open(f, "r") as fh:
            d = json.load(fh)
        if d.get("module") != "L5":
            continue
        m = d.get("metrics", {})
        cov = m.get("logging_coverage", None)       # 0–1 (↑ better)
        lat = m.get("alert_latency_h", None)        # hours (↓ better)
        score = d.get("score", None)

        base = os.path.basename(f).replace(".json", "")
        ts_match = re.search(r"(\d{8}-\d{6})", base)
        ts_label = ts_match.group(1) if ts_match else base
        mtime = os.path.getmtime(f)

        l5_rows.append({
            "Run": base,
            "Label": ts_label,
            "time": mtime,
            "logging_coverage": (None if cov is None else float(cov)),
            "alert_latency_h": (None if lat is None else float(lat)),
            "Score": score,
        })
    except Exception:
        pass

if not l5_rows:
    st.info("No L5 operations reports found yet. Run L5 multiple times to visualize ops trends.")
else:
    l5_df = pd.DataFrame(l5_rows).sort_values("time").reset_index(drop=True)
    l5_df["Run #"] = l5_df.index + 1

    st.dataframe(
        l5_df[["Run #", "Label", "logging_coverage", "alert_latency_h", "Score"]],
        width="stretch",
    )

    # Threshold bands:
    # - Coverage (↑ better):    red <0.90, yellow 0.90–0.95, green >=0.95
    cov_bands = pd.DataFrame({
        "y0": [0.00, 0.90, 0.95],
        "y1": [0.90, 0.95, 1.00],
        "band": ["red", "yellow", "green"]
    })
    # - Latency (↓ better):     green <=1h, yellow 1–2h, red >2h (domain up to 4h for display)
    lat_bands = pd.DataFrame({
        "y0": [0.00, 1.00, 2.00],
        "y1": [1.00, 2.00, 4.00],
        "band": ["green", "yellow", "red"]
    })

    def ops_chart(df, metric_col: str, title: str, bands_df: pd.DataFrame, y_domain):
        base = alt.Chart(df)
        bands = alt.Chart(bands_df).mark_rect(opacity=0.15).encode(
            y=alt.Y("y0:Q", title=None, scale=alt.Scale(domain=y_domain)),
            y2="y1:Q",
            color=alt.Color(
                "band:N",
                scale=alt.Scale(
                    domain=["green", "yellow", "red"],
                    range=["#00c851", "#ffde59", "#ff4b4b"]
                ),
                legend=None
            )
        )
        line = base.mark_line(point=True).encode(
            x=alt.X("`Run #`:Q", title="Run # (chronological)"),
            y=alt.Y(f"{metric_col}:Q", title=title,
                    scale=alt.Scale(domain=y_domain)),
            tooltip=["Label:N", alt.Tooltip(f"{metric_col}:Q", format=".3f")]
        ).properties(height=260, width="container")
        return (bands + line).resolve_scale(color="independent")

    c1, c2 = st.columns(2)
    with c1:
        st.altair_chart(
            ops_chart(l5_df, "logging_coverage",
                      "Logging Coverage (↑ better)", cov_bands, (0, 1.0)),
            width="stretch"
        )
    with c2:
        st.altair_chart(
            ops_chart(l5_df, "alert_latency_h",
                      "Alert Latency (hours, ↓ better)", lat_bands, (0, 4.0)),
            width="stretch"
        )

    # Latest snapshot
    latest_l5 = l5_df.iloc[-1]
    st.markdown("#### 🧭 Latest Run – Ops Snapshot")
    st.write(
        f"- **Logging coverage:** {latest_l5['logging_coverage'] if pd.notnull(latest_l5['logging_coverage']) else 'N/A'}"
        f"\n\n- **Alert latency:** {latest_l5['alert_latency_h'] if pd.notnull(latest_l5['alert_latency_h']) else 'N/A'} h"
    )

# ===== L5 — Operations & Live Monitoring =====
st.markdown("""
<div data-tour-target="l5-operations">
<h2>⚙️ L5 — Operations & Live Monitoring</h2>
""", unsafe_allow_html=True)

# --- Source selector + controls
src_col1, src_col2, src_col3, src_col4 = st.columns([1.3, 1.3, 1, 1])
with src_col1:
    source = st.selectbox(
        "Source",
        ["Prometheus", "CSV upload", "JSON upload"],
        index=0,
        disabled=_LOCK,
        help="Pick where to read live metrics from."
    )
with src_col2:
    refresh_secs = st.number_input(
        "Auto-refresh (seconds)",
        min_value=0, max_value=3600, value=0, step=5,
        disabled=_LOCK,
        help="0 = off. Uses streamlit_autorefresh when > 0."
    )
with src_col3:
    latency_slo = st.number_input(
        "Latency SLO (p95, seconds)", min_value=0.0, value=1.0, step=0.1, disabled=_LOCK
    )
with src_col4:
    refresh_now = st.button("🔄 Refresh now", disabled=_LOCK)

# OPTIONAL: light scheduler using streamlit_autorefresh
if refresh_secs and not _LOCK:
    try:
        # pip install streamlit-autorefresh
        from streamlit_autorefresh import st_autorefresh
        st_autorefresh(interval=refresh_secs * 1000, key="l5_autorefresh")
    except Exception:
        st.caption(
            "Tip: install streamlit-autorefresh for timed updates: `pip install streamlit-autorefresh`")

# --- Inputs per source (Live Metrics configuration)
with st.expander("Connection & Sources", expanded=True):
    # Select live data source type
    source = st.selectbox(
        "Live metrics source",
        ["Prometheus API", "Upload CSV", "Upload JSON", "HTTP Endpoint"],
        disabled=_LOCK,
        key="live_source_select"
    )

    # Show current monitoring policy thresholds (loaded from configs/policies.yaml)
    st.caption(
        f"Active policy: Latency ≤ {POLICY['latency_slo']} s • "
        f"Error ≤ {POLICY['error_rate_threshold']} %"
    )

    # Auto-refresh interval (user-adjustable)
    refresh_interval = st.slider(
        "Auto-refresh interval (seconds)",
        5, 300, POLICY.get("refresh_interval_s", 60),
        step=5, disabled=_LOCK,
        key="refresh_interval_slider",
        help="Set how often to refresh live metrics automatically."
    )

    # Auto-refresh trigger (Streamlit built-in)
    st_autorefresh(interval=refresh_interval * 1000, key="l5_refresh")

    # ------------------------------
    # Inputs based on chosen source
    # ------------------------------
    if source == "Prometheus API":
        c1, c2 = st.columns(2)
        with c1:
            prom_latency_url = st.text_input(
                "Prometheus query (latency, seconds)",
                placeholder=(
                    "http://localhost:9090/api/v1/query?"
                    "query=histogram_quantile(0.95, sum(rate(http_request_duration_seconds_bucket[5m])) by (le))"
                ),
                key="prom_latency_url",
                disabled=_LOCK,
            )

            # Validate URL
            if prom_latency_url:
                is_valid, sanitized_url, error = validate_and_sanitize_input(
                    prom_latency_url,
                    input_type="url",
                    max_length=2000
                )
                if not is_valid:
                    st.error(f"⚠️ Invalid URL: {error}")
                    prom_latency_url = None
                else:
                    prom_latency_url = sanitized_url

        with c2:
            prom_error_url = st.text_input(
                "Prometheus query (error rate, %)",
                placeholder=(
                    "http://localhost:9090/api/v1/query?"
                    "query=100*sum(rate(http_requests_total{status=~\"5..\"}[5m]))/"
                    "sum(rate(http_requests_total[5m]))"
                ),
                key="prom_error_url",
                disabled=_LOCK,
            )

            # Validate URL
            if prom_error_url:
                is_valid, sanitized_url, error = validate_and_sanitize_input(
                    prom_error_url,
                    input_type="url",
                    max_length=2000
                )
                if not is_valid:
                    st.error(f"⚠️ Invalid URL: {error}")
                    prom_error_url = None
                else:
                    prom_error_url = sanitized_url

    elif source == "Upload CSV":
        up_csv = st.file_uploader(
            "Upload live metrics CSV (columns: time, latency_s, error_rate_pct)",
            type=["csv"],
            key="l5_live_csv",
            disabled=_LOCK,
        )

    elif source == "Upload JSON":
        up_json = st.file_uploader(
            "Upload live metrics JSON (list of {time, latency_s, error_rate_pct})",
            type=["json"],
            key="l5_live_json",
            disabled=_LOCK,
        )

    elif source == "HTTP Endpoint":
        api_url = st.text_input(
            "HTTP JSON endpoint (GET returning {time, latency_s, error_rate_pct})",
            placeholder="https://example.com/live_metrics",
            key="l5_http_url",
            disabled=_LOCK,
        )

st.markdown("</div>", unsafe_allow_html=True)
# --- Helpers


def _parse_prom_value(resp_json):
    try:
        # Prom API returns data.result[0].value = [ts, value]
        v = float(resp_json["data"]["result"][0]["value"][1])
        t = float(resp_json["data"]["result"][0]["value"][0])
        return pd.to_datetime(t, unit="s"), v
    except Exception:
        return None, None


@st.cache_data(ttl=15, show_spinner=False)
def _fetch_prom(url):
    if not url:
        return None, None

    # Check rate limit
    allowed, error_msg = rate_limiter.is_allowed('prometheus')
    if not allowed:
        logger.warning(f"Prometheus rate limit exceeded: {error_msg}")
        st.warning(f"⚠️ {error_msg}")
        return None, None

    try:
        r = requests.get(url, timeout=10)
        r.raise_for_status()
        return _parse_prom_value(r.json())
    except requests.exceptions.Timeout:
        logger.error(f"Prometheus request timeout: {url}")
        st.error("⚠️ Request timeout. Check Prometheus server.")
        return None, None
    except requests.exceptions.RequestException as e:
        logger.error(f"Prometheus request failed: {e}")
        show_error_inline(e, "Prometheus fetch failed")
        return None, None


def _load_live_df():
    # Priority: uploads → Prometheus → empty
    if source == "CSV upload" and 'up_csv' in locals() and up_csv is not None:
        df0 = pd.read_csv(up_csv)
        # normalize column names
        cols = {c.lower(): c for c in df0.columns}
        # required: time, latency_s, error_rate_pct (case-insensitive)
        tcol = cols.get("time")
        lcol = cols.get("latency_s")
        ecol = cols.get("error_rate_pct")
        if all([tcol, lcol, ecol]):
            try:
                df0[tcol] = pd.to_datetime(df0[tcol])
            except Exception:
                pass
            return (df0
                    .rename(columns={tcol: "time", lcol: "latency_s", ecol: "error_rate_pct"})
                    .sort_values("time"))
        return None

    if source == "JSON upload" and 'up_json' in locals() and up_json is not None:
        try:
            js = json.load(up_json)
            df0 = pd.json_normalize(js)
            df0["time"] = pd.to_datetime(df0["time"])
            return df0[["time", "latency_s", "error_rate_pct"]].sort_values("time")
        except Exception:
            return None

    if source == "Prometheus":
        # If Prometheus endpoints provided, fetch a single point now
        t1, v1 = _fetch_prom(st.session_state.get("prom_latency_url"))
        t2, v2 = _fetch_prom(st.session_state.get("prom_error_url"))
        if t1 or t2:
            rows = [{
                "time": t1 or t2 or pd.Timestamp.utcnow(),
                "latency_s": (v1 if v1 is not None else math.nan),
                "error_rate_pct": (v2 if v2 is not None else math.nan),
            }]
            return pd.DataFrame(rows)

    return None


# Force refresh on click (clear cached Prometheus calls)
if refresh_now:
    try:
        _fetch_prom.clear()
    except Exception:
        pass

live_df = _load_live_df()

# --- UI rendering for snapshot/charts/incidents (unchanged logic)
if live_df is None or live_df.empty:
    st.info("No live data yet. Provide Prometheus queries or upload a CSV/JSON.")
else:
    st.subheader("📡 Current Live Snapshot")
    latest_row = live_df.iloc[-1]
    c1, c2, c3 = st.columns(3)
    with c1:
        st.metric("p95 latency (s)",
                  f"{float(latest_row['latency_s']):.3f}", help="Lower is better")
    with c2:
        st.metric("Error rate (%)",
                  f"{float(latest_row['error_rate_pct']):.2f}", help="Lower is better")
    with c3:
        ok = float(latest_row['latency_s']) <= latency_slo
        st.metric(
            "SLO status",
            "✅ OK" if ok else "❌ Breach",
            help=f"Threshold p95 ≤ {latency_slo:.2f}s"
        )

    # Charts (Altair)
    st.markdown("#### Trends")
    if "time" in live_df.columns:
        base = alt.Chart(live_df).encode(x=alt.X("time:T", title="Time"))
        lat_chart = base.mark_line().encode(
            y=alt.Y("latency_s:Q", title="Latency (s)")
        ).properties(height=220)
        err_chart = base.mark_line().encode(
            y=alt.Y("error_rate_pct:Q", title="Error rate (%)")
        ).properties(height=220)
        st.altair_chart(lat_chart, theme=None)
        st.altair_chart(err_chart, theme=None)

    # Generate incident if thresholds exceeded (write to reports/incidents.json)
    # Simple policy: latency > SLO OR error_rate > 5% → incident
    inc = None
    try:
        lat_ok = float(latest_row["latency_s"]) <= latency_slo
        err_ok = float(latest_row["error_rate_pct"]) <= 5.0
        if (not lat_ok) or (not err_ok):
            inc = {
                "time": datetime.utcnow().isoformat(timespec="seconds"),
                "kind": "SLO breach",
                "severity": "high" if float(latest_row["error_rate_pct"]) > 10.0 else "medium",
                "latency_s": float(latest_row["latency_s"]),
                "error_rate_pct": float(latest_row["error_rate_pct"]),
                "note": "Auto-detected breach from live metrics panel"
            }
    except Exception as e:
        logger.error(f"Incident detection failed: {e}", exc_info=True)
        show_error_inline(e, "Incident detection failed")

    # Only attempt to log if an incident was created
    if inc is not None:
        if audit_locked():
            st.warning(
                "Audit Mode is ON – automatic incident logging disabled.")
            logger.warning(
                "Incident detected but not logged (audit mode active)")
        else:
            os.makedirs("reports", exist_ok=True)
            inc_path = os.path.join("reports", "incidents.json")

            try:
                # Thread-safe file operations
                import threading
                _incident_lock = st.session_state.get('_incident_lock')
                if _incident_lock is None:
                    _incident_lock = threading.Lock()
                    st.session_state._incident_lock = _incident_lock

                with _incident_lock:
                    data = []
                    if os.path.exists(inc_path):
                        with open(inc_path, "r", encoding="utf-8") as fh:
                            data = json.load(fh) or []

                    data.append(inc)

                    # Atomic write using temp file
                    temp_path = f"{inc_path}.tmp"
                    with open(temp_path, "w", encoding="utf-8") as fh:
                        json.dump(data, fh, indent=2)

                    # Atomic rename (overwrites target)
                    os.replace(temp_path, inc_path)

                logger.info(
                    f"Incident recorded: {inc['kind']} severity={inc['severity']}")
                st.success("Incident recorded.")
            except Exception as e:
                logger.error(f"Failed to record incident: {e}", exc_info=True)
                show_error_inline(e, "Incident logging failed")


# --- Incident timeline + scheduler tips
with st.expander("⚙️ Run Scheduler Tips", expanded=False):
    st.markdown(
        "- **Auto-refresh:** uses `streamlit-autorefresh` for UI updates.\n"
        "- **Batch runs:** schedule your CLI (e.g., `python -m cli.iraqaf_cli run --module ALL --out reports`) with:\n"
        "  - **Windows Task Scheduler** (Triggers → Daily/On logon; Action → Start a program)\n"
        "  - **cron** on Linux/macOS (e.g., `0 * * * * /usr/bin/python /path/cli.py`)\n"
        "  - **GitHub Actions/CI** for nightly runs"
    )

# ===== L5 — Incident Timeline =====
st.markdown("### 🧯 Incident Timeline")

inc_path = os.path.join("reports", "incidents.json")
inc_rows = []
if os.path.exists(inc_path):
    try:
        with open(inc_path, "r", encoding="utf-8") as fh:
            inc_rows = json.load(fh) or []
    except Exception:
        inc_rows = []

inc_df = None
if not inc_rows:
    st.info("No incidents logged yet.")
else:
    inc_df = pd.json_normalize(inc_rows)
    # normalize & sort
    try:
        inc_df["time"] = pd.to_datetime(inc_df["time"])
    except Exception:
        pass
    inc_df = inc_df.sort_values("time")
    st.dataframe(
        inc_df[["time", "kind", "severity", "latency_s", "error_rate_pct", "note"]],
        use_container_width=True
    )

    # Altair timeline (points across severity bands)
    chart = (
        alt.Chart(inc_df)
        .mark_point(size=120)
        .encode(
            x=alt.X("time:T", title="Time"),
            y=alt.Y("severity:N", title="Severity", sort=[
                    "low", "medium", "high", "critical"]),
            color=alt.Color("kind:N", title="Incident type"),
            tooltip=["time", "kind", "severity",
                     "latency_s", "error_rate_pct", "note"]
        )
        .properties(height=220)
    )
    st.altair_chart(chart, theme=None, use_container_width=True)

# ---------- Demo helper: add a synthetic incident (for testing the timeline) ----------

INC_PATH = Path("reports/incidents.json")
INC_PATH.parent.mkdir(parents=True, exist_ok=True)


def _read_incidents(path: Path = INC_PATH) -> list[dict]:
    """Return a list[dict] from incidents.json (robust to legacy dict payloads)."""
    if not path.exists():
        return []
    try:
        obj = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(obj, list):
            return [x for x in obj if isinstance(x, dict)]
        if isinstance(obj, dict):
            return [obj]
        return []
    except Exception as e:
        show_error_inline(e, "Read incidents.json failed")
        return []


def _write_incidents(items: list[dict], path: Path = INC_PATH) -> bool:
    """Write incidents list back to disk (guarded by Audit Mode)."""
    if _LOCK:  # respect Audit Mode
        st.warning("Audit Mode is ON – writes disabled.")
        return False
    try:
        path.write_text(json.dumps(items, indent=2), encoding="utf-8")
        return True
    except Exception as e:
        show_error_inline(e, "Write incidents.json failed")
        return False


demo_col1, demo_col2 = st.columns([1, 1])

with demo_col1:
    if st.button("➕ Add demo incident (medium)", disabled=_LOCK, help="Adds a synthetic SLO breach now"):
        demo_inc = {
            "time": pd.Timestamp.utcnow().isoformat(timespec="seconds"),
            "kind": "SLO breach",
            "severity": "medium",
            "latency_s": float((live_df.iloc[-1]["latency_s"] if live_df is not None and not live_df.empty else 1.5)),
            "error_rate_pct": float((live_df.iloc[-1]["error_rate_pct"] if live_df is not None and not live_df.empty else 6.0)),
            "note": "Demo incident"
        }
        incidents_list = _read_incidents(INC_PATH)
        incidents_list.append(demo_inc)
        if _write_incidents(incidents_list, INC_PATH):
            st.success("Demo incident added.")

with demo_col2:
    # Use current table if you already built inc_df above; otherwise rebuild quickly
    if 'inc_df' in locals() and inc_df is not None and not inc_df.empty:
        export_df = inc_df
    else:
        _rows = _read_incidents(INC_PATH)
        export_df = pd.json_normalize(_rows) if _rows else None

    if export_df is not None and not export_df.empty:
        st.download_button(
            "⬇️ Download incidents.csv",
            data=export_df.to_csv(index=False).encode("utf-8"),
            file_name="incidents.csv",
            mime="text/csv",
            disabled=_LOCK
        )
    else:
        st.caption("No incidents to export yet.")

# ===== Quality Drift Alerts (GQAS) =====
st.markdown("### 🚨 Quality Drift Alerts")

# find last two AGG reports
agg_paths = [p for p in files if os.path.basename(p).startswith("AGG-")]
agg_paths = sorted(agg_paths, key=lambda p: os.path.getmtime(p))
if len(agg_paths) < 2:
    st.info("Not enough AGG runs to compute drift. Run the audit at least twice.")
else:
    prev = load_json(agg_paths[-2])
    curr = load_json(agg_paths[-1])
    g0 = (prev or {}).get("gqas")
    g1 = (curr or {}).get("gqas")
    if g0 is None or g1 is None:
        st.info("Missing GQAS in AGG reports.")
    else:
        delta = float(g1) - float(g0)
        st.write(
            f"Previous: **{g0:.2f}** → Current: **{g1:.2f}** (Δ {delta:+.2f})")
        breach = (delta <= -5.0)
        st.markdown(
            f"Status: {'🟥 Drift > 5 points' if breach else '🟩 Stable'}")

        with st.expander("Notify (Slack / Email)"):
            st.caption(
                "Send a one-time alert to your team. Fields are locked in Audit Mode.")

            # one atomic form = no duplicate element IDs + cleaner UX
            with st.form("drift_notify_form"):
                slack_url = st.text_input(
                    "Slack webhook URL",
                    placeholder="https://hooks.slack.com/services/XXX/YYY/ZZZ",
                    type="password",
                    key="drift_slack_url_input",
                    disabled=audit_mode,
                    help="Stored only for this session, not persisted"
                )
                email_to = st.text_input(
                    "Email to (comma-separated)",
                    placeholder="ops@example.com",
                    key="drift_email_to_input",
                    disabled=audit_mode
                )
                smtp_host = st.text_input(
                    "SMTP host",
                    placeholder="smtp.example.com",
                    key="drift_smtp_host_input",
                    disabled=audit_mode
                )
                smtp_user = st.text_input(
                    "SMTP user",
                    key="drift_smtp_user_input",
                    disabled=audit_mode
                )
                smtp_pass = st.text_input(
                    "SMTP password",
                    type="password",
                    key="drift_smtp_pass_input",
                    disabled=audit_mode,
                    help="⚠️ Not stored - enter each time you send alerts"
                )
                st.caption(
                    "🔒 Credentials are used only for this notification and are NOT saved")

                submit = st.form_submit_button(
                    "Send Notifications",
                    disabled=audit_mode
                )

            # Only run network calls if not in audit mode and user submitted
            if submit and not audit_mode:
                alert_msg = (
                    f":rotating_light: IRAQAF GQAS drift detected: "
                    f"{g0:.2f} → {g1:.2f} (Δ {delta:+.2f}). "
                    f"Check recent incidents and module scores."
                )

                def send_slack(msg, url):
                    if not url:
                        st.warning("Slack webhook URL missing.")
                        return False
                    try:
                        import requests
                        r = requests.post(url, json={"text": msg}, timeout=10)
                        if 200 <= r.status_code < 300:
                            return True
                        st.error(
                            f"Slack error: {r.status_code} {r.text[:200]}")
                    except Exception as e:
                        st.error(f"Slack exception: {e}")
                    return False

                def send_email(subject, body, to_csv, host, user, pwd):
                    try:
                        import smtplib
                        from email.mime.text import MIMEText
                        tos = [t.strip()
                               for t in (to_csv or "").split(",") if t.strip()]
                        if not (tos and host):
                            st.warning(
                                "Email: recipient(s) or SMTP host missing.")
                            return False
                        msg = MIMEText(body)
                        msg["Subject"] = subject
                        msg["From"] = user or "iraqaf@local"
                        msg["To"] = ", ".join(tos)
                        with smtplib.SMTP(host, 587, timeout=15) as s:
                            s.starttls()
                            if user and pwd:
                                s.login(user, pwd)
                            s.send_message(msg)
                        return True
                    except Exception as e:
                        st.error(f"Email exception: {e}")
                        return False

                # Fire
                sent_any = False
                if slack_url:
                    sent_any = send_slack(alert_msg, slack_url) or sent_any
                if email_to:
                    sent_any = send_email("IRAQAF GQAS Drift Alert", alert_msg,
                                          email_to, smtp_host, smtp_user, smtp_pass) or sent_any

                if sent_any:
                    st.success("Notification(s) sent.")
                else:
                    st.info("No notifications sent (fill Slack URL and/or Email).")

            # Handle submit outside the form context for clarity
            if submit:
                # persist values
                st.session_state["drift_slack_url"] = slack_url
                st.session_state["drift_email_to"] = email_to
                st.session_state["drift_smtp_host"] = smtp_host
                st.session_state["drift_smtp_user"] = smtp_user
                logger.info(
                    f"Drift alert triggered: GQAS {g0:.2f} → {g1:.2f} (Δ {delta:+.2f})")

                sent_any = False
                if slack_url:
                    if send_slack(alert_msg, slack_url):
                        st.success("Slack alert sent.")
                        sent_any = True
                if email_to and smtp_host:
                    if send_email("IRAQAF GQAS Drift Alert", alert_msg, email_to, smtp_host, smtp_user, smtp_pass):
                        st.success("Email sent.")
                        sent_any = True

                if not sent_any:
                    st.info(
                        "Nothing sent. Provide a Slack webhook and/or email + SMTP settings.")


# ========== AGGREGATE SECTION ==========
st.markdown("""
<div data-tour-target="gqas-aggregate">
<h3>⚙️ Aggregate Global QA Score (GQAS)</h3>
""", unsafe_allow_html=True)

with measure_section("aggregate_gqas_render"):
    agg = latest["AGG"]

    def gqas_band(gqas: float):
        if gqas >= 92:
            return "green", "🟢"
        elif gqas >= 88:
            return "yellow", "🟡"
        return "red", "🔴"

# ===== Risk profile toggle =====
risk = st.radio("Risk profile", ["High", "Medium"], horizontal=True, index=0)


def floors_check(module_reports: dict):
    if risk == "High":
        floors = {"L1": 90, "L2": 90, "L3": 95, "L4": 90, "L5": 85}
    else:
        floors = {"L1": 85, "L2": 85, "L3": 90, "L4": 85, "L5": 80}

    results = {}
    for mid, floor in floors.items():
        rep = module_reports.get(mid)
        if not rep:
            # status, emoji, floor, score
            results[mid] = ("missing", "❔", floor, None)
            continue

        score = rep["score"]
        passed = score >= floor
        sym = "🟢" if passed else "🔴"

        results[mid] = (
            "pass" if passed else "fail",
            sym,
            floor,
            score,
        )
    return results


if agg:
    gqas = agg.get("gqas", 0.0)
    band, emoji = gqas_band(gqas)

    # Centered GQAS card (unchanged)
    st.markdown(
        f"<h3 style='text-align:center;'>⚙️ <b>Aggregate Global QA Score (GQAS)</b></h3>"
        "<p style='text-align:center;color:gray;'>Weighted overall quality from all five modules</p>",
        unsafe_allow_html=True,
    )
    st.markdown(
        f"<h2 style='text-align:center;color:#444;'>{emoji} {gqas:.2f}</h2>", unsafe_allow_html=True)

    progress_html = f"""
    <div style='background:#eee;border-radius:10px;height:25px;margin-bottom:15px;'>
      <div style='width:{gqas}%;background:linear-gradient(90deg,#ff4b4b,#ffde59,#00c851);
      height:25px;border-radius:10px;'></div>
    </div>
    """
    st.markdown(progress_html, unsafe_allow_html=True)

    st.markdown("#### 🧱 Shipping Floors")

    floors = floors_check(latest)
    cols = st.columns(5)
    order = ["L1", "L2", "L3", "L4", "L5"]

    for i, mid in enumerate(order):
        status, sym, floor, score = floors[mid]
        label = name_for(mid)

        if status == "missing":
            cols[i].markdown(
                "<div style='padding:10px;border-radius:10px;background:#fdf5d4;"
                "text-align:center;min-height:110px;display:flex;flex-direction:column;"
                "justify-content:center;'>"
                f"<div>❔ <b>{label}</b></div>"
                "<div><i>Missing report</i></div>"
                "</div>",
                unsafe_allow_html=True,
            )
        else:
            bg = "#d4edda" if status == "pass" else "#f8d7da"
            cols[i].markdown(
                f"""
    <div style="
        background:{bg};
        padding:12px;
        border-radius:14px;
        text-align:center;
        height:150px;                 /* 🔥 HARD-FIXED HEIGHT */
        display:flex;
        flex-direction:column;
        justify-content:center;       /* vertically center */
        align-items:center;
    ">
        <div style="max-width:180px;text-align:center;">
            <b>{sym} {label}</b>
        </div>
        <div>Floor <b>{floor}</b> | Score <b>{score:.1f}</b></div>
        <div>{'✅ Meets' if status == 'pass' else '❌ Fails'}</div>
    </div>
    """,
                unsafe_allow_html=True
            )

    # Verdict badge (unchanged)
    floors_met = all(v[0] == "pass" for v in floors.values()
                     if v[0] != "missing")
    color = "#d4edda" if floors_met else "#f8d7da"
    msg = "✅ Floors met" if floors_met else "❌ Floors not met"
    st.markdown(
        f"<div style='background:{color};padding:10px;margin-top:15px;border-radius:8px;"
        "box-shadow:0 2px 6px rgba(0,0,0,0.1);text-align:center;'>"
        f"<b>{msg}</b></div>",
        unsafe_allow_html=True
    )

# Module scores table (risk-aware: only green/red)
st.markdown("#### 📊 Module Scores (for record)")
score_rows = []
for mid in order:
    rep = latest.get(mid)
    status, sym, floor, score = floors.get(mid, ("missing", "❔", None, None))

    if status == "missing" or rep is None:
        score_rows.append({
            "Module": name_for(mid),
            "Score": None,
            "Band": None,
        })
    else:
        # derive band from floor pass/fail → only green/red
        band = "green" if status == "pass" else "red"
        score_rows.append({
            "Module": name_for(mid),
            "Score": score,
            "Band": band,
        })

st.dataframe(pd.DataFrame(score_rows), width="stretch")


# ===== Data Provenance Summary =====
st.markdown("### 🧾 Data Provenance")

try:
    with open("configs/project.example.yaml", "r", encoding="utf-8") as fh:
        proj_cfg = yaml.safe_load(fh) or {}
    meta = proj_cfg.get("META") or {}
    dv = meta.get("dataset_version", "—")
    ltd = meta.get("last_train_date", "—")
    owner = meta.get("data_owner", "—")
    st.markdown(
        f"""
        <div style='background:#f8f9fa;border:1px solid #eee;border-radius:10px;padding:10px;margin-bottom:10px;'>
          <b>Dataset Version:</b> {dv}<br>
          <b>Last Train Date:</b> {ltd}<br>
          <b>Data Owner:</b> {owner}
        </div>
        """, unsafe_allow_html=True
    )
except Exception:
    st.info("Provenance metadata not found in configs/project.example.yaml → META block (optional).")

    # ===== Compare Runs & Trend Visualization =====
st.markdown("### 📈 Compare Runs & Trend Over Time")

# Find all available AGG reports
agg_files = [f for f in files if f.endswith(".json") and "AGG" in f]
run_options = []
for f in agg_files:
    try:
        with open(f, "r") as fh:
            d = json.load(fh)
        if d.get("module") == "AGG":
            label = os.path.basename(f).replace(".json", "")
            d["_label"] = label
            d["_mtime"] = os.path.getmtime(f)  # for chronological sort
            run_options.append(d)
    except Exception:
        pass

if not run_options:
    st.info("No AGG reports found to compare yet. Run the framework multiple times to generate more reports.")
else:
    # sort all runs by time
    run_options = sorted(run_options, key=lambda r: r["_mtime"])

    labels = [r["_label"] for r in run_options]
    selection = st.multiselect(
        "Select runs to compare",
        labels,
        default=(labels[-2:] if len(labels) >= 2 else labels),
    )

    if selection:
        sel_data = [r for r in run_options if r["_label"] in selection]
        # keep selected data chrono-sorted
        sel_data = sorted(sel_data, key=lambda r: r["_mtime"])

        cmp_rows = []
        for r in sel_data:
            row = {
                "Run": r["_label"],
                "GQAS": r.get("gqas"),
                "Floors Met": r.get("floors_met"),
            }
            scores = r.get("scores", {})  # e.g., CRS/SAI/FI/TS/OPS
            for mid, val in scores.items():
                row[mid] = val
            cmp_rows.append(row)

        cmp_df = pd.DataFrame(cmp_rows)
        st.dataframe(cmp_df, width="stretch")

        # Trend visualization (chronological)
        trend_df = cmp_df.copy()
        trend_df["Run Index"] = range(1, len(trend_df) + 1)
        trend_chart = (
            alt.Chart(trend_df)
            .mark_line(point=True)
            .encode(
                x=alt.X("Run Index:Q", title="Run # (chronological)"),
                y=alt.Y("GQAS:Q", title="Aggregate Quality Score"),
                tooltip=["Run", alt.Tooltip("GQAS:Q", format=".2f")],
            )
            .properties(height=250, width="container", title="GQAS Trend Over Time")
        )
        st.altair_chart(trend_chart, width="stretch")

        st.caption(
            "Shows GQAS change across multiple evaluation runs. Each run = one complete quality audit.")
    else:
        st.info("Select at least one run to visualize comparisons.")

        # ===== AGG Snapshot Compare (two runs) =====
st.markdown("### 🧾 Snapshot Compare – Aggregate (AGG)")
agg_paths = [f for f in files if f.endswith(
    ".json") and ("AGG" in os.path.basename(f))]
if len(agg_paths) < 2:
    st.info("Create at least two AGG reports to compare snapshots.")
else:
    agg_labels = {p: _label_from_path(p) for p in agg_paths}
    left, right = st.columns(2)
    with left:
        sel_a = st.selectbox("Baseline run (A)", options=agg_paths,
                             format_func=lambda p: agg_labels[p], key="aggA")
    with right:
        sel_b = st.selectbox("Compare to (B)", options=[
                             p for p in agg_paths if p != sel_a], format_func=lambda p: agg_labels[p], key="aggB")

    A = load_json(sel_a)
    B = load_json(sel_b)
    if not (A and B and A.get("module") == "AGG" and B.get("module") == "AGG"):
        st.warning("Could not load both AGG reports.")
    else:
        a_scores = A.get("scores", {})
        b_scores = B.get("scores", {})
        rows = []
        for k in ["CRS", "SAI", "FI", "TS", "OPS"]:
            rows.append({
                "Metric": k,
                f"{agg_labels[sel_a]}": a_scores.get(k),
                f"{agg_labels[sel_b]}": b_scores.get(k),
                "Δ": (None if (k not in a_scores or k not in b_scores) else (float(b_scores[k]) - float(a_scores[k]))),
            })
        rows.append({
            "Metric": "GQAS",
            f"{agg_labels[sel_a]}": A.get("gqas"),
            f"{agg_labels[sel_b]}": B.get("gqas"),
            "Δ": (None if (A.get("gqas") is None or B.get("gqas") is None) else float(B["gqas"]) - float(A["gqas"]))
        })
        cmp_df = pd.DataFrame(rows)

        st.dataframe(cmp_df, width="stretch")

        # Delta bars (positive good)
        dplot = cmp_df.dropna(subset=["Δ"]).copy()
        dplot = dplot[dplot["Metric"].isin(
            ["CRS", "SAI", "FI", "TS", "OPS", "GQAS"])]
        chart = alt.Chart(dplot).mark_bar().encode(
            x=alt.X("Δ:Q", title="Delta (B - A)", scale=alt.Scale(domain=(min(-5,
                    float(dplot["Δ"].min())-1), max(5, float(dplot["Δ"].max())+1)))),
            y=alt.Y("Metric:N", sort=None),
            tooltip=["Metric", alt.Tooltip("Δ:Q", format="+.2f")]
        ).properties(height=220, width="container", title="Change vs Baseline")
        st.altair_chart(chart, width="stretch")


st.markdown("</div>", unsafe_allow_html=True)
# ===== L1 Governance – Clause Drill-through (revised) =====
st.markdown("### 🏛️ L1 Governance – Clause Drill-through")

l1_paths = [f for f in files if f.endswith(
    ".json") and os.path.basename(f).startswith("L1-")]
if len(l1_paths) < 1:
    st.info("No L1 reports found. Run L1 to enable clause drill-through.")
else:
    pre_a = _nearest_by_time(l1_paths, agg_paths[0]) if agg_paths else None
    pre_b = _nearest_by_time(l1_paths, agg_paths[-1]) if agg_paths else None

    disabled = get_audit_mode()

    c1, c2 = st.columns(2)
    with c1:
        l1_a = st.selectbox(
            "Run A (baseline)",
            options=l1_paths,
            index=(l1_paths.index(pre_a) if pre_a in l1_paths else 0),
            format_func=lambda p: _label_from_path(p),
            key="l1A",
            disabled=_LOCK
        )
    with c2:
        l1_b = st.selectbox(
            "Run B (compare)",
            options=[p for p in l1_paths if p != l1_a],
            format_func=lambda p: _label_from_path(p),
            key="l1B",
            disabled=_LOCK
        )

    L1A = load_json(l1_a) or {}
    L1B = load_json(l1_b) or {}

    if not (L1A and L1B):
        st.warning("Could not load selected L1 reports.")
    else:
        # normalize so each clause has evidence_links
        L1A = _normalize_clause_evidence(L1A)
        L1B = _normalize_clause_evidence(L1B)

        ca = {(c.get("framework"), c.get("id")): c for c in (
            L1A.get("metrics", {}).get("clauses") or [])}
        cb = {(c.get("framework"), c.get("id")): c for c in (
            L1B.get("metrics", {}).get("clauses") or [])}
        keys = sorted(set(ca.keys()) | set(cb.keys()))

        diff_rows = []
        for k in keys:
            Acl, Bcl = ca.get(k), cb.get(k)
            Apass = (Acl.get("passed") if Acl else None)
            Bpass = (Bcl.get("passed") if Bcl else None)
            change = "↑ pass" if (Apass is False and Bpass is True) else (
                "↓ fail" if (Apass is True and Bpass is False) else "—")
            diff_rows.append({
                "Framework": k[0],
                "Clause": k[1],
                "A Passed": Apass,
                "B Passed": Bpass,
                "Change": change,
                "Evidence (A)": "📎" if (Acl and Acl.get("evidence_links")) else "—",
                "Evidence (B)": "📎" if (Bcl and Bcl.get("evidence_links")) else "—",
            })

        if not diff_rows:
            st.info(
                "No clause-level details were found in either run (metrics.clauses missing).")
        else:
            def _fmt(v): return "🟢 Pass" if v is True else (
                "🔴 Fail" if v is False else "—")
            view = pd.DataFrame(diff_rows)
            if "A Passed" in view:
                view["A Passed"] = view["A Passed"].map(_fmt)
            if "B Passed" in view:
                view["B Passed"] = view["B Passed"].map(_fmt)
            view["Change"] = view["Change"].fillna("—")
            st.dataframe(view, use_container_width=True)

            # Per-clause evidence picker + preview
            st.markdown("#### 📎 Open Evidence (per clause)")
            pick = st.selectbox(
                "Select a clause to view evidence",
                options=[f"{fw} · {cid}" for (fw, cid) in keys],
                disabled=_LOCK
            )

            if pick:
                fw, cid = pick.split(" · ", 1)
                clA, clB = ca.get((fw, cid)), cb.get((fw, cid))

                def _render_evidence(tag: str, clause: dict):
                    with st.expander(f"Run {tag} — Evidence for {fw} · {cid}", expanded=False):
                        links = (clause or {}).get("evidence_links") or []
                        if not links:
                            st.caption("No evidence links for this clause.")
                            return
                        for i, p in enumerate(links, 1):
                            if re.match(r"^https?://", str(p), flags=re.I):
                                st.markdown(f"- 🌐 [{p}]({p})")
                            else:
                                pp = Path(p)
                                if pp.exists():
                                    st.markdown(
                                        f"- 📄 `{pp.name}`  \n<small style='color:#888'>{pp.as_posix()}</small>",
                                        unsafe_allow_html=True
                                    )
                                    with pp.open("rb") as fh:
                                        st.download_button("Download", data=fh.read(
                                        ), file_name=pp.name, key=f"dl_evd_{tag}_{i}")
                                    _preview_widget(pp)
                                else:
                                    st.markdown(f"- ❌ Missing: `{p}`")

                _render_evidence("A", clA)
                _render_evidence("B", clB)


# ===== Executive Summary (Auto) =====
st.markdown("### 🧠 Executive Summary (auto)")


def _band_emoji(band: str) -> str:
    return {"green": "🟢", "yellow": "🟡", "red": "🔴"}.get((band or "").lower(), "⚪")


def _floor_table(latest_reports: dict, risk_profile: str):
    # same floors as dashboard logic
    if risk_profile == "High":
        floors = {"L1": 90, "L2": 90, "L3": 95, "L4": 90, "L5": 85}
    else:
        floors = {"L1": 85, "L2": 85, "L3": 90, "L4": 85, "L5": 80}
    rows = []
    for mid, floor in floors.items():
        rep = latest_reports.get(mid)
        score = rep.get("score") if rep else None
        band = rep.get("band") if rep else None
        status = (score is not None) and (score >= floor)
        rows.append({
            "Module": mid,
            "Name": {
                "L1": "Governance & Regulatory", "L2": "Privacy & Security",
                "L3": "Fairness & Ethics", "L4": "Explainability & Transparency",
                "L5": "Operations & Monitoring",
            }[mid],
            "Score": (None if score is None else f"{score:.1f}"),
            "Band": (None if band is None else f"{_band_emoji(band)} {band.capitalize()}"),
            "Floor": floor,
            "Pass": ("✅" if status else ("❌" if score is not None else "—"))
        })
    return pd.DataFrame(rows)


def _top_drivers(latest_reports: dict):
    items = []
    for mid in ["L1", "L2", "L3", "L4", "L5"]:
        rep = latest_reports.get(mid)
        if rep and isinstance(rep.get("score"), (int, float)):
            items.append((mid, rep["score"], rep.get("band")))
    if not items:
        return None, None
    items_sorted = sorted(items, key=lambda x: x[1], reverse=True)
    best = items_sorted[0]
    worst = items_sorted[-1]
    return best, worst


def _failed_clauses_summary(l1_report: dict, max_items: int = 6):
    fails = []
    try:
        for c in (l1_report.get("metrics", {}).get("clauses") or []):
            if c.get("passed") is False:
                fails.append({
                    "framework": c.get("framework"),
                    "id": c.get("id"),
                    "desc": c.get("description"),
                    "hint": c.get("hint") or "",
                    "why": c.get("why_failed") or ""
                })
    except Exception:
        return []
    # prioritize by presence of hint/why, then by name
    fails = sorted(fails, key=lambda x: (x["framework"], x["id"]))
    return fails[:max_items]


# Gather data
agg = latest.get("AGG", {})
gqas = agg.get("gqas")
gqas_band = None
if gqas is not None:
    if gqas >= 92:
        gqas_band = "green"
    elif gqas >= 88:
        gqas_band = "yellow"
    else:
        gqas_band = "red"

# Risk profile (you already define it earlier; we read it if exists)
# If you named it `risk`, reuse; else default to "High".
risk_profile = globals().get("risk", "High")

# Build narrative
lines = []
lines.append("**Overview**")
if gqas is not None:
    lines.append(
        f"- Aggregate quality (GQAS): **{gqas:.2f}** {_band_emoji(gqas_band)} ({gqas_band.capitalize() if gqas_band else ''})")
else:
    lines.append("- Aggregate quality (GQAS): not computed yet.")

best, worst = _top_drivers(latest)
if best:
    lines.append(
        f"- Strongest module: **{best[0]}** ({best[1]:.1f}) {_band_emoji(best[2]) if best[2] else ''}")
if worst:
    lines.append(
        f"- Weakest module: **{worst[0]}** ({worst[1]:.1f}) {_band_emoji(worst[2]) if worst[2] else ''}")

# Floors
st.markdown("\n".join(lines))
st.markdown("---")
st.markdown(f"**Shipping Floors ({risk_profile} risk)**")
floor_df = _floor_table(latest, risk_profile)
st.dataframe(floor_df, width="stretch")

# L1 gaps
l1 = latest.get("L1")
if l1:
    failures = _failed_clauses_summary(l1, max_items=6)
    st.markdown(
        "**Governance gaps (top):**" if failures else "**Governance gaps:** None detected in current mapping.")
    if failures:
        for f in failures:
            hint = f["hint"].strip()
            why = f["why"].strip()
            bullets = []
            if why:
                bullets.append(f"_Why:_ {why}")
            if hint:
                bullets.append(f"_Hint:_ {hint}")
            tail = (" — " + " ".join(bullets)) if bullets else ""
            st.markdown(
                f"- **{f['framework']} – {f['id']}**: {f['desc']}{tail}")

# Next actions (simple rules)
actions = []
if gqas is not None and gqas < 92:
    actions.append(
        "Raise overall GQAS ≥ 92 by addressing the weakest module and any failed floors.")
if latest.get("L3") and latest["L3"]["score"] < 95:
    actions.append(
        "Reduce fairness gaps (DPG/EOD) or improve L3 score ≥ 95 for high-risk products.")
if latest.get("L4") and latest["L4"]["score"] < 90:
    actions.append(
        "Improve explainability: ensure deletion_drop ≥ 0.15 and stability τ ≥ 0.85.")
if latest.get("L5") and latest["L5"]["score"] < (85 if risk_profile == "High" else 80):
    actions.append(
        "Increase logging coverage ≥ 0.95 and keep alert latency ≤ 1h.")

st.markdown("---")
st.markdown(
    "**Next actions:**" if actions else "**Next actions:** none — thresholds currently met.")
for a in actions:
    st.markdown(f"- {a}")

# Compose a plain-text export
summary_text = []
summary_text.append("IRAQAF – Executive Summary")
summary_text.append("="*28)
if gqas is not None:
    summary_text.append(f"GQAS: {gqas:.2f} ({gqas_band})")
summary_text.append(f"Risk profile: {risk_profile}")
summary_text.append("")
summary_text.append("Floors:")
for _, row in floor_df.iterrows():
    summary_text.append(
        f"- {row['Module']} {row['Name']}: Score {row['Score']}, Floor {row['Floor']} → {row['Pass']}")
summary_text.append("")
if l1 and (fails := _failed_clauses_summary(l1, max_items=12)):
    summary_text.append("Governance gaps:")
    for f in fails:
        line = f"- {f['framework']} – {f['id']}: {f['desc']}"
        if f["why"]:
            line += f" | Why: {f['why']}"
        if f["hint"]:
            line += f" | Hint: {f['hint']}"
        summary_text.append(line)
summary_text.append("")
if actions:
    summary_text.append("Next actions:")
    for a in actions:
        summary_text.append(f"- {a}")

summary_blob = "\n".join(summary_text).encode("utf-8")
st.download_button("⬇️ Download Executive Summary (.txt)", data=summary_blob,
                   file_name="IRAQAF_Executive_Summary.txt", mime="text/plain")

# Optional: Word export using python-docx (reuses same content)
try:
    from io import BytesIO
    doc = Document()
    doc.add_heading("IRAQAF – Executive Summary", level=1)
    if gqas is not None:
        doc.add_paragraph(f"GQAS: {gqas:.2f} ({gqas_band})")
    doc.add_paragraph(f"Risk profile: {risk_profile}")

    doc.add_heading("Floors", level=2)
    tbl = doc.add_table(rows=1, cols=5)
    hdr = tbl.rows[0].cells
    hdr[0].text = "Module"
    hdr[1].text = "Name"
    hdr[2].text = "Score"
    hdr[3].text = "Floor"
    hdr[4].text = "Pass"
    for _, r in floor_df.iterrows():
        row = tbl.add_row().cells
        row[0].text = str(r["Module"])
        row[1].text = str(r["Name"])
        row[2].text = str(r["Score"])
        row[3].text = str(r["Floor"])
        row[4].text = str(r["Pass"])

    if l1 and (fails := _failed_clauses_summary(l1, max_items=12)):
        doc.add_heading("Governance gaps", level=2)
        for f in fails:
            p = doc.add_paragraph()
            p.add_run(f"{f['framework']} – {f['id']}: ").bold = True
            p.add_run(f"{f['desc']}")
            if f["why"]:
                doc.add_paragraph(f"Why: {f['why']}")
            if f["hint"]:
                doc.add_paragraph(f"Hint: {f['hint']}")

    if actions:
        doc.add_heading("Next actions", level=2)
        for a in actions:
            doc.add_paragraph(a, style="List Bullet")

    bio = BytesIO()
    doc.save(bio)
    st.download_button("⬇️ Download Executive Summary (.docx)", data=bio.getvalue(),
                       file_name="IRAQAF_Executive_Summary.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
except Exception:
    st.caption("Tip: `pip install python-docx` to enable Word export here.")

st.markdown("### 🤖 AI-Generated Executive Summary (optional)")
use_llm = st.checkbox(
    "Generate with LLM (requires OPENAI_API_KEY)", value=False, disabled=_LOCK)

if use_llm and not _LOCK:
    try:
        import os
        import textwrap
        from datetime import datetime
        from openai import OpenAI  # pip install openai

        if not os.getenv("OPENAI_API_KEY"):
            st.warning("Set OPENAI_API_KEY env variable to enable LLM.")
        else:
            client = OpenAI()
            prompt = f"""
You are an AI auditor. Write a crisp, actionable executive summary for stakeholders.
Use these artifacts (JSON-like): latest={json.dumps(latest)[:50000]}
Emphasize: GQAS, floors (pass/fail), biggest gaps, top 3 actions, risks, and time-sensitive notes.
Date: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}
"""
            with st.spinner("Generating summary…"):
                resp = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.3,
                )
            summary = resp.choices[0].message.content.strip()
            st.markdown(summary)
            st.download_button("⬇️ Download AI Summary (txt)", summary.encode("utf-8"),
                               file_name="AI_Executive_Summary.txt", mime="text/plain", disabled=_LOCK)
    except Exception as e:
        st.warning(f"LLM summary unavailable: {e}")
else:
    st.caption("Toggle on if you have an API key.")


# =========================
# Bottom Panels (polished)
# Evidence Tray & Pins • Sync to YAML • Exports • Radar
# =========================

# ---------- tiny style helpers ----------
st.markdown("""
<style>
/* Section headers */
.block-title {font-weight:700;font-size:1.15rem;margin:0 0 6px}
.block-subtle {color:#6b7280;font-size:0.9rem;margin-top:-2px}

/* Generic "card" */
.iraqaf-card{
  background:#ffffff;
  border:1px solid #e5e7eb;
  border-radius:12px;
  padding:14px 14px 10px;
  margin-bottom:12px;
  box-shadow:0 1px 2px rgba(0,0,0,0.04);
}

/* Pill buttons feel */
button[kind="secondary"]{border-radius:999px}

/* Make expanders look card-like */
.iraqaf-expander > div{border:1px solid #e5e7eb;border-radius:12px}
</style>
""", unsafe_allow_html=True)


def card(title: str, subtitle: str | None = None):
    st.markdown(f"<div class='iraqaf-card'><div class='block-title'>{title}</div>" +
                (f"<div class='block-subtle'>{subtitle}</div>" if subtitle else ""),
                unsafe_allow_html=True)


def close_card():
    st.markdown("</div>", unsafe_allow_html=True)


# =========================
# 📎 Evidence Tray & File Pins
# =========================
card("📎 Evidence Tray & Pins",
     "Drop files, select module(s), and Save → evidence/. Then Sync to YAML to add paths under each module’s `evidence:` list.")

EVID_DIR = Path("evidence")
EVID_DIR.mkdir(exist_ok=True)
INDEX_PATH = Path("configs/evidence_index.json")
INDEX_PATH.parent.mkdir(exist_ok=True)

MODULE_LABELS = {
    "L1": "Governance & Regulatory",
    "L2": "Privacy & Security",
    "L3": "Fairness & Ethics",
    "L4": "Explainability & Transparency",
    "L5": "Operations & Monitoring",
}


def _load_index() -> dict:
    if INDEX_PATH.exists():
        try:
            with open(INDEX_PATH, "r", encoding="utf-8") as fh:
                return json.load(fh)
        except Exception:
            return {}
    return {}


def _save_index(ix: dict):
    with open(INDEX_PATH, "w", encoding="utf-8") as fh:
        json.dump(ix, fh, indent=2)


def _sanitize_name(name: str) -> str:
    keep = "._-abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789"
    return "".join(c if c in keep else "_" for c in name)
# ============================================================================
# UNDO/REDO FUNCTIONALITY FOR EVIDENCE OPERATIONS
# ============================================================================


def init_evidence_history():
    """Initialize undo/redo stacks for evidence operations."""
    if 'evidence_history' not in st.session_state:
        st.session_state.evidence_history = {
            'undo_stack': [],
            'redo_stack': [],
            'max_history': 20
        }


def record_evidence_operation(operation_type: str, data: dict):
    """
    Record an evidence operation for undo/redo.

    Args:
        operation_type: Type of operation (add, remove, update)
        data: Operation data (file paths, module IDs, etc.)
    """
    init_evidence_history()

    operation = {
        'type': operation_type,
        'data': data,
        'timestamp': datetime.now().isoformat()
    }

    # Add to undo stack
    st.session_state.evidence_history['undo_stack'].append(operation)

    # Clear redo stack (new operation invalidates redo)
    st.session_state.evidence_history['redo_stack'] = []

    # Limit history size
    max_history = st.session_state.evidence_history['max_history']
    if len(st.session_state.evidence_history['undo_stack']) > max_history:
        st.session_state.evidence_history['undo_stack'].pop(0)


def undo_evidence_operation() -> bool:
    """Undo the last evidence operation. Returns True if successful."""
    init_evidence_history()

    undo_stack = st.session_state.evidence_history['undo_stack']
    if not undo_stack:
        return False

    operation = undo_stack.pop()
    st.session_state.evidence_history['redo_stack'].append(operation)

    try:
        if operation['type'] == 'add':
            # Remove the files that were added
            for module_id, file_path in operation['data']['files']:
                ix = _load_index()
                if module_id in ix and file_path in ix[module_id]:
                    ix[module_id].remove(file_path)
                    _save_index(ix)

                # Delete physical file
                if Path(file_path).exists():
                    Path(file_path).unlink()

        elif operation['type'] == 'remove':
            # Restore the files that were removed
            for module_id, file_path in operation['data']['files']:
                ix = _load_index()
                if module_id not in ix:
                    ix[module_id] = []
                if file_path not in ix[module_id]:
                    ix[module_id].append(file_path)
                    _save_index(ix)

        logger.info(f"Undo operation: {operation['type']}")
        return True

    except Exception as e:
        logger.error(f"Undo failed: {e}", exc_info=True)
        return False


def redo_evidence_operation() -> bool:
    """Redo the last undone operation. Returns True if successful."""
    init_evidence_history()

    redo_stack = st.session_state.evidence_history['redo_stack']
    if not redo_stack:
        return False

    operation = redo_stack.pop()
    st.session_state.evidence_history['undo_stack'].append(operation)

    try:
        if operation['type'] == 'add':
            # Re-add the files
            for module_id, file_path in operation['data']['files']:
                ix = _load_index()
                if module_id not in ix:
                    ix[module_id] = []
                if file_path not in ix[module_id]:
                    ix[module_id].append(file_path)
                    _save_index(ix)

        elif operation['type'] == 'remove':
            # Re-remove the files
            for module_id, file_path in operation['data']['files']:
                ix = _load_index()
                if module_id in ix and file_path in ix[module_id]:
                    ix[module_id].remove(file_path)
                    _save_index(ix)

        logger.info(f"Redo operation: {operation['type']}")
        return True

    except Exception as e:
        logger.error(f"Redo failed: {e}", exc_info=True)
        return False


def batch_process_evidence_files(files_list: list, modules: list, base_dir: Path) -> tuple[list, list]:
    """
    Process multiple evidence files in batch with enhanced security.

    Args:
        files_list: List of uploaded file objects
        modules: List of module IDs to pin to
        base_dir: Base directory for evidence storage

    Returns:
        Tuple of (successful_files, failed_files)
    """
    # Security constants
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    ALLOWED_EXTENSIONS = {'.pdf', '.csv', '.png', '.jpg',
                          '.jpeg', '.txt', '.docx', '.xlsx', '.json', '.yaml', '.yml'}
    DANGEROUS_EXTENSIONS = {'.exe', '.bat', '.sh', '.ps1',
                            '.cmd', '.com', '.scr', '.vbs', '.js', '.jar'}

    successful = []
    failed = []

    for uf in files_list:
        try:
            # Read file data
            data = uf.read()

            # 1. Check file size
            if len(data) > MAX_FILE_SIZE:
                failed.append(
                    (uf.name, f"File too large ({len(data) / 1024 / 1024:.1f}MB > 50MB limit)"))
                continue

            # 2. Check for empty files
            if len(data) == 0:
                failed.append((uf.name, "Empty file"))
                continue

            # 3. Validate filename for path traversal
            if ".." in uf.name or "/" in uf.name or "\\" in uf.name:
                logger.warning(f"Path traversal attempt blocked: {uf.name}")
                failed.append(
                    (uf.name, "Invalid filename (path traversal attempt)"))
                continue

            # 4. Check file extension
            file_ext = Path(uf.name).suffix.lower()

            # Block dangerous extensions
            if file_ext in DANGEROUS_EXTENSIONS:
                logger.warning(f"Dangerous file type blocked: {uf.name}")
                failed.append((uf.name, f"File type not allowed: {file_ext}"))
                continue

            # Warn on unknown extensions (but allow with warning)
            if file_ext and file_ext not in ALLOWED_EXTENSIONS:
                logger.warning(f"Unusual file extension: {uf.name}")
                # Continue processing but log the warning

            # 5. Generate safe filename with hash
            stem = _sanitize_name(Path(uf.name).stem)
            ext = _sanitize_name(file_ext) or ".bin"
            h = _hash_bytes(data)
            fname = f"{stem}-{h}{ext}"

            # 6. Validate generated filename
            if ".." in fname or "/" in fname or "\\" in fname:
                logger.warning(f"Generated suspicious filename: {fname}")
                failed.append((uf.name, "Sanitization failed"))
                continue

            # 7. Check for duplicate files (by hash)
            duplicate_found = False
            for mid in modules:
                target_dir = base_dir / mid
                if target_dir.exists():
                    for existing_file in target_dir.glob(f"*-{h}{ext}"):
                        if existing_file.exists():
                            duplicate_found = True
                            logger.info(
                                f"Duplicate file detected (same hash): {uf.name}")
                            break
                if duplicate_found:
                    break

            if duplicate_found:
                failed.append(
                    (uf.name, "Duplicate file (same content already exists)"))
                continue

            # 8. Save to each module with atomic operations
            saved_paths = []
            for mid in modules:
                target_dir = base_dir / mid
                target_dir.mkdir(parents=True, exist_ok=True)

                # Resolve and validate path is within target directory
                outp = (target_dir / fname).resolve()
                if not str(outp).startswith(str(target_dir.resolve())):
                    logger.error(f"Path traversal attempt blocked: {outp}")
                    failed.append(
                        (uf.name, "Security: Path traversal blocked"))
                    continue

                # Atomic write using temp file
                temp_path = outp.with_suffix(outp.suffix + '.tmp')
                try:
                    with open(temp_path, "wb") as fh:
                        fh.write(data)

                    # Verify written file
                    if temp_path.stat().st_size != len(data):
                        temp_path.unlink()
                        raise IOError("File write verification failed")

                    # Atomic rename
                    temp_path.replace(outp)
                    saved_paths.append((mid, str(outp.as_posix())))

                except Exception as write_error:
                    if temp_path.exists():
                        temp_path.unlink()
                    raise write_error

            if saved_paths:
                successful.append((fname, saved_paths))
                logger.info(
                    f"Saved evidence: {fname} to {len(modules)} module(s), size={len(data)} bytes, hash={h}")
            else:
                failed.append((uf.name, "Failed to save to any module"))

        except Exception as e:
            logger.error(f"Failed to save {uf.name}: {e}", exc_info=True)
            failed.append((uf.name, f"Error: {str(e)[:100]}"))

    return successful, failed
# =============================================================================
# UNDO/REDO SYSTEM (Already implemented in your code, but add UI polish)
# =============================================================================

# LOCATION: Add after line ~1450 (after evidence upload section)


# Enhanced undo/redo UI with visual feedback
if 'evidence_history' in st.session_state:
    history = st.session_state.evidence_history
    undo_count = len(history['undo_stack'])
    redo_count = len(history['redo_stack'])

    if undo_count > 0 or redo_count > 0:
        st.markdown("""
        <div style='
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            padding: 12px 16px;
            border-radius: 10px;
            margin: 16px 0;
            color: white;
        '>
            <h4 style='margin: 0 0 8px 0; color: white;'>🕐 History</h4>
        """, unsafe_allow_html=True)

        undo_col, redo_col, info_col = st.columns([1, 1, 2])

        with undo_col:
            undo_disabled = undo_count == 0 or _LOCK
            if st.button(
                f"↶ Undo ({undo_count})",
                disabled=undo_disabled,
                use_container_width=True,
                key="evidence_undo_btn_ui",
                help="Undo last evidence operation"
            ):
                if undo_evidence_operation():
                    st.success("✅ Undone!")
                    time.sleep(0.5)
                    st.rerun()
                else:
                    st.error("❌ Undo failed")

        with redo_col:
            redo_disabled = redo_count == 0 or _LOCK
            if st.button(
                f"↷ Redo ({redo_count})",
                disabled=redo_disabled,
                use_container_width=True,
                key="evidence_redo_btn_ui",
                help="Redo last undone operation"
            ):
                if redo_evidence_operation():
                    st.success("✅ Redone!")
                    time.sleep(0.5)
                    st.rerun()
                else:
                    st.error("❌ Redo failed")

        with info_col:
            if undo_count > 0:
                last_op = history['undo_stack'][-1]
                st.caption(
                    f"Last: {last_op['type']} at {last_op['timestamp'][:19]}")

        st.markdown("</div>", unsafe_allow_html=True)


def _hash_bytes(b: bytes) -> str:
    return hashlib.sha1(b).hexdigest()[:10]


ix = _load_index()

u1, u2 = st.columns([2, 1])
with u1:
    uploads = st.file_uploader(
        "Upload evidence files (PDF, CSV, PNG, TXT, DOCX)",
        type=["pdf", "csv", "png", "txt", "docx"],
        accept_multiple_files=True,
        disabled=_LOCK
    )
with u2:
    modules_to_pin = st.multiselect(
        "Pin uploaded files to module(s)",
        options=list(MODULE_LABELS.keys()),
        default=["L1"],
        format_func=lambda m: f"{m} – {MODULE_LABELS[m]}",
        help="Each selected module receives a reference to every uploaded file.",
        disabled=_LOCK
    )

b1, b2 = st.columns([1, 1])

with colA:
    if st.button("💾 Save to evidence/ and pin", disabled=_LOCK):
        if not uploads:
            st.warning("No files uploaded.")
        elif not modules_to_pin:
            st.warning("Pick at least one module to pin.")
        else:
            with progress_tracker("Saving evidence files", total=len(uploads)) as tracker:
                successful, failed = batch_process_evidence_files(
                    uploads, modules_to_pin, EVID_DIR)

                # Update index for successful saves
                for fname, paths in successful:
                    for mid, path in paths:
                        ix.setdefault(mid, [])
                        if path not in ix[mid]:
                            ix[mid].append(path)

                # Save updated index
                if successful:
                    _save_index(ix)
                    st.success(
                        f"✅ Saved {len(successful)} file(s): {', '.join(f[0] for f in successful)}")
# Save updated index
                if successful:
                    _save_index(ix)

                    # Record operation for undo/redo
                    all_saved_paths = []
                    for fname, paths in successful:
                        all_saved_paths.extend(paths)

                    record_evidence_operation('add', {
                        'files': all_saved_paths,
                        'timestamp': datetime.now().isoformat()
                    })

                    st.success(
                        f"✅ Saved {len(successful)} file(s): {', '.join(f[0] for f in successful)}")
                # Show failures if any
                if failed:
                    with st.expander(f"⚠️ {len(failed)} file(s) failed", expanded=True):
                        for fname, error in failed:
                            st.error(f"❌ {fname}: {error}")

with b2:
    if st.button("🧭 Open evidence folder (path)", use_container_width=True):
        st.info(f"Evidence folder: `{EVID_DIR.resolve()}`")
        # Undo/Redo controls
init_evidence_history()
undo_count = len(st.session_state.evidence_history['undo_stack'])
redo_count = len(st.session_state.evidence_history['redo_stack'])

if undo_count > 0 or redo_count > 0:
    st.markdown("**History**")
    undo_col, redo_col = st.columns(2)

    with undo_col:
        undo_disabled = undo_count == 0 or _LOCK
        if st.button(
            f"↶ Undo ({undo_count})",
            disabled=undo_disabled,
            use_container_width=True,
            key="evidence_undo_btn",
            help="Undo last evidence operation"
        ):
            if undo_evidence_operation():
                st.success("✅ Undone!")
                st.rerun()
            else:
                st.error("❌ Undo failed")

    with redo_col:
        redo_disabled = redo_count == 0 or _LOCK
        if st.button(
            f"↷ Redo ({redo_count})",
            disabled=redo_disabled,
            use_container_width=True,
            key="evidence_redo_btn",
            help="Redo last undone operation"
        ):
            if redo_evidence_operation():
                st.success("✅ Redone!")
                st.rerun()
            else:
                st.error("❌ Redo failed")

with st.expander("📚 Current evidence index (per module)", expanded=False):
    st.json(ix or {})


close_card()

# =========================
# 🔄 Sync to YAML + Export
# =========================
card("🔄 Sync to `configs/project.example.yaml`",
     "Merges pins into each module’s `evidence: [...]`. Creates keys if missing and de-duplicates.")


def _sync_to_yaml(yaml_path="configs/project.example.yaml"):
    yaml_path = Path(yaml_path)
    if not yaml_path.exists():
        st.error(f"Config not found: {yaml_path}")
        return False
    try:
        import yaml
    except Exception:
        st.error("Missing PyYAML. `pip install pyyaml`")
        return False

    with open(yaml_path, "r", encoding="utf-8") as fh:
        try:
            cfg = yaml.safe_load(fh) or {}
        except Exception as e:
            st.error(f"YAML parse error: {e}")
            return False

    for mid in MODULE_LABELS.keys():
        cfg.setdefault(mid, {})
        ev = cfg[mid].get("evidence", [])
        if ev is None:
            ev = []
        if not isinstance(ev, list):
            ev = [str(ev)]
        for rel in ix.get(mid, []):
            if rel not in ev:
                ev.append(rel)
        cfg[mid]["evidence"] = sorted(ev)

    backup = yaml_path.with_suffix(f".backup.{int(time.time())}.yaml")
    shutil.copyfile(yaml_path, backup)
    with open(yaml_path, "w", encoding="utf-8") as fh:
        yaml.safe_dump(cfg, fh, sort_keys=False)
    st.success(f"Synced to {yaml_path}  •  Backup: {backup.name}")
    return True


s1, s2 = st.columns([1, 1])
with c1:
    if st.button("🔁 Sync now", disabled=_LOCK):
        try:
            _sync_to_yaml()
        except Exception as e:
            show_error_inline(e, "YAML sync failed")
        else:
            _sync_to_yaml()
with s2:
    st.caption(
        "Tip: after syncing, re-run L1/L2 so reports include the new `evidence` references.")

# Inline export card


# =============================================================================
# EXPORT PRESETS (Enhanced version already in your code - add visual polish)
# =============================================================================

st.markdown("""
<div style='
    background: linear-gradient(135deg, #10b981 0%, #059669 100%);
    padding: 16px;
    border-radius: 12px;
    margin-bottom: 16px;
    color: white;
'>
    <h3 style='margin: 0; color: white;'>💾 Quick Export Presets</h3>
    <p style='margin: 4px 0 0 0; opacity: 0.9; font-size: 0.9rem;'>
        One-click exports for different audiences
    </p>
</div>
""", unsafe_allow_html=True)

preset_col1, preset_col2, preset_col3 = st.columns(3)

with preset_col1:
    st.markdown("""
    <div style='text-align: center; margin-bottom: 8px;'>
        <div style='font-size: 2rem;'>📊</div>
        <div style='font-weight: 600; color: #1f2937;'>Executive Summary</div>
        <div style='font-size: 0.875rem; color: #6b7280;'>High-level metrics</div>
    </div>
    """, unsafe_allow_html=True)

    if st.button("Export Executive", use_container_width=True, key="export_exec_preset", type="primary"):
        export_data = {
            "GQAS": agg.get("gqas"),
            "Floors": {
                "L1": latest.get("L1", {}).get("score"),
                "L2": latest.get("L2", {}).get("score"),
                "L3": latest.get("L3", {}).get("score"),
                "L4": latest.get("L4", {}).get("score"),
                "L5": latest.get("L5", {}).get("score"),
            },
            "Generated": datetime.now().isoformat(),
            "Risk_Profile": risk_profile if 'risk_profile' in locals() else "High"
        }
        st.download_button(
            "⬇️ Download JSON",
            data=json.dumps(export_data, indent=2).encode("utf-8"),
            file_name=f"executive_summary_{datetime.now().strftime('%Y%m%d')}.json",
            mime="application/json",
            key="dl_exec_summary",
            use_container_width=True
        )

with preset_col2:
    st.markdown("""
    <div style='text-align: center; margin-bottom: 8px;'>
        <div style='font-size: 2rem;'>🔬</div>
        <div style='font-weight: 600; color: #1f2937;'>Technical Deep Dive</div>
        <div style='font-size: 0.875rem; color: #6b7280;'>Full details + metrics</div>
    </div>
    """, unsafe_allow_html=True)

    if st.button("Export Technical", use_container_width=True, key="export_tech_preset", type="primary"):
        st.download_button(
            "⬇️ Download JSON",
            data=json.dumps(latest, indent=2).encode("utf-8"),
            file_name=f"technical_report_{datetime.now().strftime('%Y%m%d')}.json",
            mime="application/json",
            key="dl_tech_report",
            use_container_width=True
        )

with preset_col3:
    st.markdown("""
    <div style='text-align: center; margin-bottom: 8px;'>
        <div style='font-size: 2rem;'>📋</div>
        <div style='font-weight: 600; color: #1f2937;'>Compliance Only</div>
        <div style='font-size: 0.875rem; color: #6b7280;'>L1 governance data</div>
    </div>
    """, unsafe_allow_html=True)

    if st.button("Export Compliance", use_container_width=True, key="export_compliance_preset", type="primary"):
        compliance_data = {
            "L1_Governance": latest.get("L1", {}),
            "Generated": datetime.now().isoformat()
        }
        st.download_button(
            "⬇️ Download JSON",
            data=json.dumps(compliance_data, indent=2).encode("utf-8"),
            file_name=f"compliance_report_{datetime.now().strftime('%Y%m%d')}.json",
            mime="application/json",
            key="dl_compliance_report",
            use_container_width=True
        )

st.markdown("---")

# Original export buttons
agg_json = json.dumps(agg, indent=2).encode("utf-8")
csv_df = pd.DataFrame(score_rows)
x1, x2 = st.columns(2)
close_card()

# =========================
# 🧾 Auto-Report (Word)
# =========================
card("🧾 Auto-Report", "Generate a Word report with the latest bundle.")


def build_docx(report_bundle: dict) -> bytes:
    try:
        from docx import Document  # type: ignore
    except Exception:
        st.error("`python-docx` not installed. `pip install python-docx`")
        return b""
    doc = Document()
    doc.add_heading("IRAQAF QA Report", level=1)
    doc.add_paragraph("Automatically generated from the latest IRAQAF run.")

    agg_rep = report_bundle.get("AGG")
    if agg_rep:
        doc.add_heading("Aggregate (GQAS)", level=2)
        doc.add_paragraph(f"GQAS: {agg_rep.get('gqas')}")
        doc.add_paragraph(f"Floors met: {agg_rep.get('floors_met')}")

    doc.add_heading("Modules", level=2)
    for m in ["L1", "L2", "L3", "L4", "L5"]:
        rep_m = report_bundle.get(m)
        if not rep_m:
            continue
        doc.add_heading(NAMES[m], level=3)
        doc.add_paragraph(f"Score: {rep_m['score']}  |  Band: {rep_m['band']}")
        doc.add_paragraph("Metrics:")
        doc.add_paragraph(json.dumps(rep_m.get("metrics", {}), indent=2))

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


dl_slot = st.empty()
if audit_mode:
    st.caption("🔒 Disabled in Audit Mode.")
    st.button("Generate Word Report", key="gen_docx", disabled=_LOCK)
else:
    if st.button("Generate Word Report", key="gen_docx", use_container_width=True):
        payload = build_docx(latest)
        if payload:
            with dl_slot:
                st.download_button("⬇️ Download IRAQAF_Report.docx",
                                   data=payload,
                                   file_name="IRAQAF_Report.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   key="dl_docx",
                                   use_container_width=True)
close_card()

if not agg:
    st.info("Run all modules to create an aggregate report.")

# =========================
# 📄 HTML/PDF Report Export
# =========================
card("📄 HTML/PDF Report Export")


def _render_html_report(bundle: dict) -> str:
    gqas = (bundle.get("AGG") or {}).get("gqas")
    rows = []
    for mid in ["L1", "L2", "L3", "L4", "L5"]:
        rep = bundle.get(mid) or {}
        rows.append(
            f"<tr><td>{mid}</td><td>{rep.get('score', '')}</td><td>{rep.get('band', '')}</td></tr>")
    table = "\n".join(rows)
    return f"""<!doctype html>
<html><head><meta charset="utf-8"><title>IRAQAF Report</title>
<style>
body{{font-family:Inter,Arial,sans-serif;margin:0;background:#f8fafc}}
.wrap{{max-width:960px;margin:24px auto;padding:0 16px}}
.card{{background:#fff;border:1px solid #e5e7eb;border-radius:12px;padding:18px 16px;margin-bottom:12px}}
h1{{margin:0 0 8px}} h2{{margin:18px 0 8px}}
table{{border-collapse:collapse;width:100%}}
td,th{{border:1px solid #e5e7eb;padding:8px}} th{{background:#f6f9ff;text-align:left}}
.small{{color:#6b7280}}
</style></head><body>
<div class="wrap">
  <div class="card">
    <h1>IRAQAF Report</h1>
    <div class="small">Generated {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}</div>
    <h2>Aggregate</h2>
    <p><b>GQAS:</b> {gqas}</p>
    <h2>Modules</h2>
    <table><tr><th>Module</th><th>Score</th><th>Band</th></tr>
    {table}
    </table>
  </div>
</div>
</body></html>"""


html_blob = _render_html_report(latest)
h1, h2 = st.columns([1, 1])
with h1:
    st.download_button("⬇️ Download HTML report",
                       data=html_blob.encode("utf-8"),
                       file_name="IRAQAF_Report.html",
                       mime="text/html",
                       use_container_width=True,
                       disabled=_LOCK)
with h2:
    if st.checkbox("Also generate PDF (requires pdfkit + wkhtmltopdf)", value=False, disabled=_LOCK):
        try:
            import pdfkit
            pdf_bytes = pdfkit.from_string(html_blob, False)
            st.download_button("⬇️ Download PDF report", data=pdf_bytes,
                               file_name="IRAQAF_Report.pdf", mime="application/pdf", disabled=_LOCK)
        except Exception as e:
            show_error_inline(e, "PDF export unavailable")

close_card()

# =========================
# 🧭 Maturity Radar (L1–L5)
# =========================
card("🧭 Maturity Radar (L1–L5)",
     "Polar plot of module scores (0–100). Hover for values.")

radar_rep = {m: (latest.get(m) or {}).get("score")
             for m in ["L1", "L2", "L3", "L4", "L5"]}
if not any(isinstance(v, (int, float)) for v in radar_rep.values()):
    st.info("No module scores available for radar.")
else:
    # Prepare a closed polygon in polar coordinates, then project to XY.
    names = ["L1", "L2", "L3", "L4", "L5"]
    scores = [float(radar_rep.get(m) or 0.0) for m in names]
    names.append(names[0])
    scores.append(scores[0])  # close
    n = len(names)
    rad_df = pd.DataFrame({"module": names, "score": scores})
    rad_df["idx"] = range(n)
    rad_df["angle"] = rad_df["idx"] * (2*math.pi/(n-1))
    rad_df["x"] = rad_df["score"] * np.cos(rad_df["angle"])
    rad_df["y"] = rad_df["score"] * np.sin(rad_df["angle"])

    # Soft background grid/rings
    rings = pd.DataFrame({"r": [25, 50, 75, 100]})
    ring_df = pd.concat(
        [pd.DataFrame({"r": [r]*361, "deg": np.arange(361)})
         for r in rings["r"]],
        ignore_index=True
    )
    ring_df["rad"] = np.deg2rad(ring_df["deg"])
    ring_df["x"] = ring_df["r"]*np.cos(ring_df["rad"])
    ring_df["y"] = ring_df["r"]*np.sin(ring_df["rad"])

    grid = alt.Chart(ring_df).mark_line(opacity=0.15).encode(x="x:Q", y="y:Q")
    poly = alt.Chart(rad_df).mark_area(opacity=0.15).encode(
        x="x:Q", y="y:Q")
    outline = alt.Chart(rad_df).mark_line(point=True).encode(
        x="x:Q", y="y:Q", tooltip=["module", "score"])

    # spokes with labels
    spoke_base = pd.DataFrame(
        {"module": names[:-1], "idx": range(len(names)-1)})
    spoke_base["angle"] = spoke_base["idx"] * (2*math.pi/(len(names)-1))
    spoke_base["x"] = 110*np.cos(spoke_base["angle"])
    spoke_base["y"] = 110*np.sin(spoke_base["angle"])
    label_layer = alt.Chart(spoke_base).mark_text(fontSize=12).encode(
        x="x:Q", y="y:Q", text="module:N")

    chart = (grid + poly + outline +
             label_layer).properties(width=420, height=420)
    st.altair_chart(chart, use_container_width=False)

# ========================================
# Performance Summary (Debug Mode)
# ========================================

if st.sidebar.checkbox("🔧 Show Debug Info", value=False, key="debug_mode"):
    st.markdown("---")

    # Styled container for debug info
    st.markdown("""
    <div style='
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 20px;
        border-radius: 12px;
        margin-bottom: 20px;
        color: white;
    '>
        <h2 style='margin: 0; color: white;'>🔧 Debug & Performance</h2>
        <p style='margin: 5px 0 0 0; opacity: 0.9;'>System diagnostics and monitoring</p>
    </div>
    """, unsafe_allow_html=True)

    # Performance Metrics Card
    with st.container():
        st.markdown("""
        <div style='
            background: white;
            border: 1px solid #e5e7eb;
            border-radius: 10px;
            padding: 16px;
            margin-bottom: 16px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        '>
            <h3 style='margin: 0 0 12px 0; color: #1f2937;'>⚡ Performance Metrics</h3>
        """, unsafe_allow_html=True)

        if PERFORMANCE_MONITORING:
            monitor = get_monitor()
            all_stats = monitor.get_all_stats()

            if all_stats:
                # Show top 5 slowest operations
                sorted_ops = sorted(
                    all_stats.items(),
                    key=lambda x: x[1].get("mean", 0),
                    reverse=True
                )[:5]

                for op_name, stats in sorted_ops:
                    mean = stats['mean']
                    color = "#10b981" if mean < 0.5 else (
                        "#f59e0b" if mean < 2.0 else "#ef4444")
                    emoji = "🟢" if mean < 0.5 else ("🟡" if mean < 2.0 else "🔴")

                    st.markdown(f"""
                    <div style='
                        display: flex;
                        justify-content: space-between;
                        padding: 8px;
                        background: #f9fafb;
                        border-radius: 6px;
                        margin-bottom: 6px;
                    '>
                        <span style='color: #6b7280;'>{emoji} {op_name}</span>
                        <span style='color: {color}; font-weight: 600;'>{mean:.3f}s</span>
                    </div>
                    """, unsafe_allow_html=True)
            else:
                st.info("No performance data collected yet")
        else:
            st.warning("Performance monitoring is disabled")

        st.markdown("</div>", unsafe_allow_html=True)

    # System Info Card
    with st.container():
        st.markdown("""
        <div style='
            background: white;
            border: 1px solid #e5e7eb;
            border-radius: 10px;
            padding: 16px;
            margin-bottom: 16px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        '>
            <h3 style='margin: 0 0 12px 0; color: #1f2937;'>💾 System Info</h3>
        """, unsafe_allow_html=True)

        mem_usage = get_memory_usage()
        if mem_usage:
            rss_mb = mem_usage.get('rss_mb', 0)
            mem_pct = mem_usage.get('percent', 0)

            # Memory usage bar
            mem_color = "#10b981" if mem_pct < 50 else (
                "#f59e0b" if mem_pct < 80 else "#ef4444")
            st.markdown(f"""
            <div style='margin-bottom: 12px;'>
                <div style='display: flex; justify-content: space-between; margin-bottom: 4px;'>
                    <span style='color: #6b7280; font-size: 0.875rem;'>Memory Usage</span>
                    <span style='color: {mem_color}; font-weight: 600;'>{rss_mb:.1f} MB ({mem_pct:.1f}%)</span>
                </div>
                <div style='
                    width: 100%;
                    height: 8px;
                    background: #e5e7eb;
                    border-radius: 4px;
                    overflow: hidden;
                '>
                    <div style='
                        width: {mem_pct}%;
                        height: 100%;
                        background: {mem_color};
                        transition: width 0.3s ease;
                    '></div>
                </div>
            </div>
            """, unsafe_allow_html=True)

        # Reports loaded
        st.markdown(f"""
        <div style='
            display: flex;
            justify-content: space-between;
            padding: 8px;
            background: #f9fafb;
            border-radius: 6px;
            margin-bottom: 6px;
        '>
            <span style='color: #6b7280;'>📊 Reports loaded</span>
            <span style='color: #1f2937; font-weight: 600;'>{len(files)}</span>
        </div>
        """, unsafe_allow_html=True)

        # Modules found
        modules_count = sum(1 for v in latest.values() if v)
        st.markdown(f"""
        <div style='
            display: flex;
            justify-content: space-between;
            padding: 8px;
            background: #f9fafb;
            border-radius: 6px;
            margin-bottom: 6px;
        '>
            <span style='color: #6b7280;'>📦 Modules found</span>
            <span style='color: #1f2937; font-weight: 600;'>{modules_count}/6</span>
        </div>
        """, unsafe_allow_html=True)

        # Cache info
        try:
            cache_info = st.cache_data.get_stats()
            hit_rate = 0
            if len(cache_info) > 0:
                total_calls = sum(info.get("cache_misses", 0) + info.get("cache_hits", 0)
                                  for info in cache_info if isinstance(info, dict))
                total_hits = sum(info.get("cache_hits", 0)
                                 for info in cache_info if isinstance(info, dict))
                if total_calls > 0:
                    hit_rate = (total_hits / total_calls) * 100

            cache_color = "#10b981" if hit_rate > 50 else (
                "#f59e0b" if hit_rate > 20 else "#ef4444")
            st.markdown(f"""
            <div style='
                display: flex;
                justify-content: space-between;
                padding: 8px;
                background: #f9fafb;
                border-radius: 6px;
            '>
                <span style='color: #6b7280;'>💿 Cache hit rate</span>
                <span style='color: {cache_color}; font-weight: 600;'>{hit_rate:.1f}%</span>
            </div>
            """, unsafe_allow_html=True)
        except:
            pass

        st.markdown("</div>", unsafe_allow_html=True)

    # Actions Card
    with st.container():
        st.markdown("""
        <div style='
            background: white;
            border: 1px solid #e5e7eb;
            border-radius: 10px;
            padding: 16px;
            margin-bottom: 16px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        '>
            <h3 style='margin: 0 0 12px 0; color: #1f2937;'>🔧 Actions</h3>
        """, unsafe_allow_html=True)

        col1, col2 = st.columns(2)

        with col1:
            if st.button("🗑️ Clear All Caches", use_container_width=True, key="debug_clear_cache"):
                st.cache_data.clear()
                st.success("✅ Caches cleared!")
                st.rerun()

        with col2:
            if st.button("📊 Export Performance Log", use_container_width=True, key="debug_export_perf"):
                if PERFORMANCE_MONITORING:
                    monitor = get_monitor()
                    perf_data = json.dumps(monitor.get_all_stats(), indent=2)
                    st.download_button(
                        "⬇️ Download performance.json",
                        data=perf_data.encode("utf-8"),
                        file_name=f"performance_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                        mime="application/json",
                        key="debug_download_perf"
                    )
                else:
                    st.warning("Performance monitoring not enabled")

        st.markdown("</div>", unsafe_allow_html=True)

    # Security Card
    with st.container():
        st.markdown("""
        <div style='
            background: white;
            border: 1px solid #e5e7eb;
            border-radius: 10px;
            padding: 16px;
            margin-bottom: 16px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        '>
            <h3 style='margin: 0 0 12px 0; color: #1f2937;'>🔒 Security</h3>
        """, unsafe_allow_html=True)

        # Security status checks
        security_checks = {
            "Rate limiting": ("✅", "Active"),
            "Input validation": ("✅", "Active"),
            "File upload limits": ("✅", "50MB max"),
            "Audit mode": ("✅" if get_audit_mode() else "⚪", "Enabled" if get_audit_mode() else "Disabled"),
        }

        for check, (emoji, status) in security_checks.items():
            status_color = "#10b981" if emoji == "✅" else "#9ca3af"
            st.markdown(f"""
            <div style='
                display: flex;
                justify-content: space-between;
                padding: 8px;
                background: #f9fafb;
                border-radius: 6px;
                margin-bottom: 6px;
            '>
                <span style='color: #6b7280;'>{check}</span>
                <span style='color: {status_color}; font-weight: 600;'>{emoji} {status}</span>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("</div>", unsafe_allow_html=True)

    # Last refresh timestamp
    st.markdown(f"""
    <div style='
        text-align: center;
        color: #9ca3af;
        font-size: 0.875rem;
        padding: 12px;
    '>
        🕐 Last refresh: {st.session_state.get('last_refresh', 'Never')}
    </div>
    """, unsafe_allow_html=True)

close_card()
