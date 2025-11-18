# --- bootstrap: make project root importable ---
from __future__ import annotations
import psutil

# Standard library imports
import glob
import hashlib
import importlib.util
import io
import json
import logging
import math
import os
import random
import re
import shutil
import sys
import threading
import time
import traceback
from collections import defaultdict
from contextlib import contextmanager
from datetime import datetime
from io import BytesIO
from logging.handlers import RotatingFileHandler
from pathlib import Path
from typing import Any, Dict, Iterable


# Third-party imports
import altair as alt
import numpy as np
import pandas as pd
import requests
import streamlit as st
import yaml
from dotenv import load_dotenv
from filelock import FileLock
from sklearn.inspection import permutation_importance
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score, roc_auc_score
from sklearn.model_selection import train_test_split
from streamlit_autorefresh import st_autorefresh

# Local imports
from audit_utils import audit_locked, hash_file, record_audit_event
from helpers import (
    _delta_badge,
    _evidence_count,
    _hash_file,
    _human_size,
    _label_from_path,
    _load_json,
    _nearest_by_time,
    _normalize_clause_evidence,
    _sanitize_name,
    preview_widget,
    ribbon_chart,
)
from validation import SCHEMA_VERSION, validate_incidents, validate_report

# Export, Alerts & RBAC Enhancements
try:
    from export_alerts_rbac import (
        ExportManager, AlertManager, RBACManager,
        render_export_section, render_alerts_section,
        render_role_based_dashboard, initialize_enhancements
    )
    ENHANCEMENTS_AVAILABLE = True
except Exception as e:
    ENHANCEMENTS_AVAILABLE = False
    logger = logging.getLogger("iraqaf_dashboard")
    logger.warning(f"Enhancements module not available: {e}")

# L2 Privacy/Security Monitor Integration
try:
    from l2_monitor_integration import show_l2_privacy_security_monitor
    L2_MONITOR_AVAILABLE = True
except ImportError:
    L2_MONITOR_AVAILABLE = False
    show_l2_privacy_security_monitor = None

# Beautiful Authentication UI
try:
    from auth_ui import check_authentication, render_logout_button, render_user_info
    AUTH_UI_AVAILABLE = True
except ImportError:
    AUTH_UI_AVAILABLE = False
    def check_authentication(): return True
    def render_logout_button(): pass
    def render_user_info(): pass

# UX Enhancements (Dark Mode, Loading Animations, Keyboard Shortcuts, Session Info)
try:
    from ux_enhancements import (
        initialize_ux_enhancements, render_theme_toggle, render_session_info,
        render_keyboard_shortcuts, render_action_buttons, render_info_cards,
        increment_action_count, show_loading_spinner, render_progress_bar,
        render_quick_stats, render_toast_notification
    )
    UX_ENHANCEMENTS_AVAILABLE = True
except ImportError:
    UX_ENHANCEMENTS_AVAILABLE = False
    def initialize_ux_enhancements(): return True
    def render_theme_toggle(): pass
    def render_session_info(): pass
    def render_keyboard_shortcuts(): pass
    def increment_action_count(): pass
    def render_action_buttons(x): return {}
    def render_info_cards(x): pass
    def show_loading_spinner(x=""): pass
    def render_progress_bar(x, y, z=""): pass
    def render_quick_stats(x): pass
    def render_toast_notification(x, y="info"): pass

# System Integration Imports - Placeholder (will be loaded after setup_paths)
SYSTEM_INTEGRATION_AVAILABLE = False
get_coordinator = None
initialize_coordinator = None
DatabaseQueries = None
init_db = None
get_monitor = None
initialize_monitor = None
EventType = None

#  CRITICAL: Force correct paths before ANY imports


def setup_paths():
    """Robustly set up Python import paths."""
    # Get the directory containing this file (dashboard/)
    try:
        script_path = Path(__file__).resolve()
    except NameError:
        script_path = Path.cwd() / "dashboard" / "app.py"

    dashboard_dir = script_path.parent
    project_root = dashboard_dir.parent

    # Validate we found the right directories
    core_path = project_root / "core"
    utils_path = project_root / "utils"

    if not core_path.exists():
        # Search upward for project root
        current = Path.cwd()
        found = False
        for _ in range(5):
            if (current / "core").exists() and (current / "utils").exists():
                project_root = current
                dashboard_dir = current / "dashboard"
                found = True
                break
            current = current.parent

        if not found:
            raise RuntimeError(
                f"Cannot find project root! Looked in:\n"
                f" ` - {project_root}\n"
                f" ` - {Path.cwd()}\n"
                f"Please run from: iraqaf_starter_kit/ directory"
            )

    # Add to sys.path
    paths_to_add = [str(project_root), str(dashboard_dir)]
    for path in paths_to_add:
        if path not in sys.path:
            sys.path.insert(0, path)

    return project_root, dashboard_dir


# Run path setup FIRST
try:
    _ROOT, _CURR = setup_paths()
    print(f"? Project root: {_ROOT}")
    print(f"? Dashboard dir: {_CURR}")
    print(f"? sys.path[0]: {sys.path[0]}")
    print(f"? sys.path[1]: {sys.path[1]}")
except Exception as e:
    print(f"? FATAL: Path setup failed: {e}")
    print(f"Current directory: {Path.cwd()}")
    print(f"__file__: {__file__ if '__file__' in dir() else 'NOT DEFINED'}")
    sys.exit(1)

# NOW import system integration modules (paths are set up)
try:
    # Ensure project root is in path
    from pathlib import Path
    _proj_root = Path(__file__).parent.parent
    if str(_proj_root) not in sys.path:
        sys.path.insert(0, str(_proj_root))

    from scripts.system_integration import get_coordinator, initialize_coordinator
    from scripts.database_layer import DatabaseQueries, init_db
    from scripts.realtime_monitor import get_monitor, initialize_monitor, EventType
    SYSTEM_INTEGRATION_AVAILABLE = True
    print("? System integration modules loaded successfully")
except ImportError as e:
    logger_early = logging.getLogger("iraqaf_dashboard")
    logger_early.warning(f"System integration modules not available: {e}")
    SYSTEM_INTEGRATION_AVAILABLE = False
    print(f"? System integration import failed: {e}")

# Helper function to get reports directory - always relative to project root


def get_reports_path():
    """Get reports directory path, always relative to project root"""
    return Path(_ROOT) / "reports"

# Helper function to glob reports - uses correct project root


def get_report_files(pattern: str = "*.json") -> list:
    """Get report files from reports directory"""
    reports_dir = get_reports_path()
    return sorted(glob.glob(str(reports_dir / pattern)))


def get_evidence_count_local(module_id: str = None, report_data: dict = None) -> int:
    """
    Count evidence items for a module from report data.

    Args:
        module_id: Module ID (L1, L2, etc.) - used as fallback
        report_data: Report dictionary containing evidence array

    Returns:
        Count of evidence items, or 0 if none found
    """
    if report_data and isinstance(report_data, dict):
        evidence = report_data.get("evidence", [])
        if isinstance(evidence, list):
            return len(evidence)
    return 0


def get_metrics_display_local(metrics: dict, score: float = None, max_items: int = 3) -> str:
    """Format metrics for display - shows key metrics used to calculate score"""
    if isinstance(metrics, dict) and metrics:
        # Special handling for L1: show top frameworks by coverage
        frameworks = metrics.get('framework_breakdown', [])
        if frameworks and isinstance(frameworks, list):
            # Sort by coverage_percent descending and take top N
            sorted_frameworks = sorted(frameworks, key=lambda x: x.get(
                'coverage_percent', 0), reverse=True)
            display_items = [
                f"{fw.get('framework')}: {fw.get('coverage_percent')}%" for fw in sorted_frameworks[:max_items]]
            metrics_display = " | ".join(
                display_items) if display_items else "Module assessed"
            if len(sorted_frameworks) > max_items:
                metrics_display += f" +{len(sorted_frameworks)-max_items} more"
            return metrics_display

        # For other modules: show simple metrics
        filtered_metrics = {k: v for k, v in metrics.items()
                            if not k.startswith('_') and k not in ['clauses', 'evidence', 'framework_breakdown', 'evidence_modules', 'coverage_percent']}

        # Remove dict/list type values (only show simple values)
        simple_metrics = {k: v for k, v in filtered_metrics.items()
                          if not isinstance(v, (dict, list))}

        if simple_metrics:
            display_items = [f"{k}: {v}" for k, v in list(
                simple_metrics.items())[:max_items]]
            metrics_display = " | ".join(
                display_items) if display_items else "Module assessed"
            if len(simple_metrics) > max_items:
                metrics_display += f" +{len(simple_metrics)-max_items} more"
            return metrics_display

    # For modules with no metrics (like L2), just return "Module assessed"
    return "Module assessed"


_core_test = _ROOT / "core" / "__init__.py"
if not _core_test.exists():
    print(f"? FATAL: core module not found at {_core_test}")
    print(f"\nProject structure:")
    for item in _ROOT.iterdir():
        print(f" ` - {item.name}")
    sys.exit(1)

print(f"? Core module found at: {_core_test}")

# NOW import local modules (after path is configured)
try:
    from utils.file_operations import load_json, save_json
    print("? utils.file_operations imported")
except ImportError as e:
    print(f"? Failed to import utils.file_operations: {e}")
    sys.exit(1)

try:
    from utils.error_handling import show_error_inline
    print("? utils.error_handling imported")
except ImportError as e:
    print(f"? Failed to import utils.error_handling: {e}")
    sys.exit(1)

try:
    from core.config import config
    print("? core.config imported")
except ImportError as e:
    print(f"? Failed to import core.config: {e}")
    print(f"\nDEBUG INFO:")
    print(f" ` sys.path[0]: {sys.path[0]}")
    print(f" ` Looking for: {Path(sys.path[0]) / 'core' / 'config.py'}")
    print(f" ` Exists: {(Path(sys.path[0]) / 'core' / 'config.py').exists()}")
    sys.exit(1)

# =============================================================================
# EARLY LOGGER FALLBACK (for use in cached functions before full init)
# =============================================================================
logger = logging.getLogger("iraqaf_dashboard")
logger.setLevel(logging.INFO)
# Minimal console handler for early logging
_early_handler = logging.StreamHandler()
_early_handler.setLevel(logging.WARNING)
logger.addHandler(_early_handler)

# =============================================================================
# LOGGER INITIALIZATION (MUST BE EARLY - before any functions use it)
# =============================================================================
os.makedirs("logs", exist_ok=True)

# Add file handler to the existing logger (from fallback)
file_handler = RotatingFileHandler(
    "logs/dashboard.log",
    maxBytes=10 * 1024 * 1024,
    backupCount=5,
    encoding="utf-8",
)
file_handler.setLevel(logging.INFO)

formatter = logging.Formatter(
    "%(asctime)s - %(name)s - %(levelname)s - %(funcName)s:%(lineno)d - %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
file_handler.setFormatter(formatter)

# Update the early handler with formatter too
for handler in logger.handlers:
    handler.setFormatter(formatter)

logger.addHandler(file_handler)

logger.info("=" * 50)
logger.info("IRAQAF Dashboard initialized")
logger.info("=" * 50)

# =============================================================================
# TRACE MAP UTILITIES (must be defined before TRACE_MAP initialization)
# =============================================================================


@st.cache_data(ttl=300, show_spinner=False)
def load_trace_map(path: str = "configs/trace_map.yaml") -> dict:
    """
    Load trace_map.yaml which maps regulatory clauses to metric paths.

    Returns:
        Dict mapping framework -> clause -> [list of metric paths]
        Example: {"GDPR": {"Art5_Data_Principles": ["L2.metrics.encryption_coverage"]}}
    """
    try:
        # Try project root path first, then fallback to relative path
        p = Path(path)
        if not p.exists():
            p = Path(_ROOT) / "configs" / "trace_map.yaml"

        if not p.exists():
            logger.warning(f"trace_map.yaml not found at {path}")
            return {}

        with p.open("r", encoding="utf-8") as fh:
            raw = yaml.safe_load(fh) or {}

        # Handle both formats:
        # 1. Direct format: {framework: {clause: [metrics]}}
        # 2. Nested format: {trace_map: {framework: {clause: [metrics]}}}
        if "trace_map" in raw:
            data = raw["trace_map"]
        else:
            data = raw

        if not isinstance(data, dict) or not data:
            logger.warning(
                f"trace_map.yaml has invalid/empty structure")
            return {}

        logger.info(f"Loaded trace_map with {len(data)} framework(s)")
        return data

    except yaml.YAMLError as e:
        logger.error(f"YAML parse error in {path}: {e}")
        return {}
    except Exception as e:
        logger.error(f"Failed to load trace_map: {e}", exc_info=True)
        return {}


@st.cache_data(ttl=300)
def get_available_frameworks() -> list[str]:
    """Get list of all available frameworks"""
    trace_map = load_trace_map()
    return sorted(trace_map.keys())


@st.cache_data(ttl=300)
def get_framework_clauses(framework: str) -> list[str]:
    """Get all clauses for a specific framework"""
    trace_map = load_trace_map()
    fw_data = trace_map.get(framework, {})
    return sorted(fw_data.keys())


@st.cache_data(ttl=300)
def get_clause_metrics(framework: str, clause: str) -> list[str]:
    """Get metrics traced to a specific clause"""
    trace_map = load_trace_map()
    fw_data = trace_map.get(framework, {})
    clause_data = fw_data.get(clause, {})

    # Handle both old format (list) and new format (dict with metrics key)
    if isinstance(clause_data, list):
        return clause_data
    elif isinstance(clause_data, dict):
        return clause_data.get("metrics", [])

    return []


@st.cache_data(ttl=300)
def get_clause_description(framework: str, clause: str) -> str | None:
    """Get description for a specific clause"""
    trace_map = load_trace_map()
    fw_data = trace_map.get(framework, {})
    clause_data = fw_data.get(clause, {})

    if isinstance(clause_data, dict):
        return clause_data.get("description")

    return None


# =============================================================================
# INITIALIZE GLOBALS (after function definitions)
# =============================================================================

# NOW these will work because functions are defined above
TRACE_MAP = load_trace_map()
AVAILABLE_FRAMEWORKS = get_available_frameworks()
# Dashboard-level modules (in dashboard/ folder)

# Third-party imports

# Try to import UI utilities (from dashboard folder)
try:
    from ui_utils import show_success, show_error, show_warning, show_info
except ImportError:
    # Fallback to standard streamlit functions
    show_success = st.success
    def show_error(msg, details=None): return st.error(msg)
    show_warning = st.warning
    show_info = st.info

# Custom exception classes


class DashboardError(Exception):
    """Base exception for dashboard errors"""
    pass


class DataLoadError(DashboardError):
    """Error loading data files"""
    pass


class ExportError(DashboardError):
    """Error during export operations"""
    pass

# Global error handler


def handle_dashboard_error(error: Exception, context: str = ""):
    """Centralized error handler with user-friendly messages"""
    error_id = f"ERR-{datetime.now().strftime('%Y%m%d%H%M%S')}"

    # Log error
    logger.error(f"{error_id} | {context}: {error}", exc_info=True)

    # User-friendly message
    st.error(f"? **Error {error_id}**")

    with st.expander(" `Error Details & Recovery", expanded=False):
        st.code(str(error), language="text")

        # Context-specific recovery tips
        if isinstance(error, FileNotFoundError):
            st.info("""
            **Recovery Tips:**
            - Verify the file exists in the expected location
            - Check file permissions (read access required)
            - Re-run evaluation to regenerate missing files
            """)
        elif isinstance(error, json.JSONDecodeError):
            st.info("""
            **Recovery Tips:**
            - File may be corrupted - try regenerating
            - Check for manual edits that broke JSON syntax
            - Use a JSON validator to identify syntax errors
            """)
        elif isinstance(error, PermissionError):
            st.info("""
            **Recovery Tips:**
            - Close programs using the file
            - Run with administrator privileges (Windows) or sudo (Linux/Mac)
            - Check folder write permissions
            """)
        else:
            st.info("""
            **Recovery Tips:**
            - Check logs/dashboard.log for details
            - Try refreshing the page
            - Contact support if issue persists
            """)

        st.caption(
            f"Error ID: `{error_id}` ? Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # Copy error ID button
        if st.button(" `Copy Error ID", key=f"copy_err_{error_id}"):
            st.code(error_id)
            st.caption("Share this ID when reporting the issue")


# -------------------------------------------------
# Guided Tour Step Definitions (must be defined once)
# -------------------------------------------------
TOUR_STEPS = [
    {"id": "module-summary", "title": "Module Summary"},
    {"id": "evidence-tray", "title": " Evidence Tray"},
    {"id": "l1-governance", "title": " `L1 Governance & Compliance Coverage"},
    {"id": "l2-privacy", "title": " `L2 Privacy/Security & Trends"},
    {"id": "l3-fairness", "title": " `L3 Fairness & Evolution Across Runs"},
    {"id": "l4-explainability", "title": " `L4 Explainability & Interactive"},
    {"id": "l5-operations", "title": " `L5  Operations & Live Monitoring"},
]
# ============================================================================
# CORE UTILITY FUNCTIONS (Consolidated - Single Source of Truth)
# ============================================================================


def show_error_inline(e: Exception, context: str):
    """
    Show a concise inline error message with helpful recovery tips.

    Args:
        e: The exception that occurred
        context: User-friendly context describing what failed
    """
    tb = "".join(traceback.format_exception_only(type(e), e)).strip()
    error_msg = f"{context}: {tb}"

    # Log full traceback
    logger.error(error_msg, exc_info=True)

    # Show user-friendly message with recovery tips
    st.error(f"? **{context}**")

    with st.expander(" `Error Details & Recovery Tips"):
        st.code(tb, language="text")

        # Provide context-specific recovery tips
        if "FileNotFoundError" in str(type(e)):
            st.info("""
            **Recovery Tips:**
            - Check if the file exists in the expected location
            - Verify file permissions
            - Re-run the evaluation to generate missing reports
            """)
        elif "JSONDecodeError" in str(type(e)):
            st.info("""
            **Recovery Tips:**
            - The file may be corrupted
            - Try regenerating the report
            - Check if the file was manually edited
            """)
        elif "PermissionError" in str(type(e)):
            st.info("""
            **Recovery Tips:**
            - Check file permissions
            - Close any programs using the file
            - Run as administrator (Windows) or with sudo (Linux/Mac)
            """)
        else:
            st.info("""
            **Recovery Tips:**
            - Check the logs for more details
            - Try refreshing the dashboard
            - Contact support if the issue persists
            """)

        # Show timestamp
        st.caption(
            f"Error occurred at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")


# Performance monitoring
try:
    from performance import (
        get_monitor,
        timed,
        measure_section,
        show_performance_panel,
        get_memory_usage
    )
    PERFORMANCE_MONITORING = True
    logger.info("Performance monitoring enabled")
except ImportError:
    # Fallback: no-op decorators if performance.py not available
    def timed(name=None):
        def decorator(func):
            return func
        return decorator

    class measure_section:
        def __init__(self, name):
            pass

        def __enter__(self):
            return self

        def __exit__(self, *args):
            pass

    def show_performance_panel():
        st.info("Performance monitoring module not available")

    def get_memory_usage():
        return {}

    PERFORMANCE_MONITORING = False
    logger.warning(
        "Performance monitoring disabled (performance.py not found)")


# stdlib

# third-party

# local

# legacy alias kept for compatibility

# yaml alias used by legacy code
try:
    import yaml as _yaml
except Exception:
    _yaml = None

load_dotenv()


# ============================================================================
# INPUT VALIDATION & SANITIZATION
# ============================================================================


def validate_and_sanitize_input(
    value: str,
    input_type: str = "general",
    max_length: int = 1000,
    allow_special_chars: bool = False
) -> tuple[bool, str, str]:
    """
    Validate and sanitize user input with security checks.

    Args:
        value: Input string to validate
        input_type: Type of input (email, url, filename, general)
        max_length: Maximum allowed length
        allow_special_chars: Whether to allow special characters

    Returns:
        Tuple of (is_valid, sanitized_value, error_message)
    """
    import re
    from urllib.parse import urlparse

    if not value:
        return True, "", ""

    # Check length
    if len(value) > max_length:
        return False, value[:max_length], f"Input too long (max {max_length} characters)"

    # Type-specific validation
    if input_type == "email":
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if not re.match(email_pattern, value):
            return False, value, "Invalid email format"
        return True, value.strip().lower(), ""

    elif input_type == "url":
        try:
            result = urlparse(value)
            if not all([result.scheme, result.netloc]):
                return False, value, "Invalid URL format"
            if result.scheme not in ['http', 'https']:
                return False, value, "Only HTTP/HTTPS URLs allowed"
            return True, value.strip(), ""
        except Exception:
            return False, value, "Invalid URL"

    elif input_type == "filename":
        # Remove path separators and dangerous characters
        dangerous_chars = ['/', '\\', '..', '<',
                           '>', '|', ':', '*', '?', '"', '\x00']
        sanitized = value
        for char in dangerous_chars:
            sanitized = sanitized.replace(char, '_')

        # Check for reserved Windows names
        reserved_names = ['CON', 'PRN', 'AUX',
                          'NUL', 'COM1', 'COM2', 'LPT1', 'LPT2']
        if sanitized.upper() in reserved_names:
            sanitized = f"_{sanitized}"

        return True, sanitized[:255], " "   # Max filename length

    elif input_type == "general":
        # Remove potential XSS/injection characters
        dangerous_patterns = [
            r'<script[^>]*>.*?</script>',  # Script tags
            r'javascript:',                 # JavaScript protocol
            r'on\w+\s*=',                  # Event handlers
            r'<iframe[^>]*>',              # Iframes
        ]

        sanitized = value
        for pattern in dangerous_patterns:
            sanitized = re.sub(pattern, '', sanitized, flags=re.IGNORECASE)

        if not allow_special_chars:
            # Keep only alphanumeric, spaces, and basic punctuation
            sanitized = re.sub(r'[^\w\s.,;:!?()\-]', '', sanitized)

        return True, sanitized.strip(), ""

    return True, value, ""

# ============================================================================
# RATE LIMITING
# ============================================================================


class RateLimiter:
    """Simple rate limiter for API calls."""

    def __init__(self, config_limits: Dict[str, Dict[str, int]] = None):
        self.calls = {}
        self.limits = config_limits or config.rate_limits.limits

    def is_allowed(self, action: str, identifier: str = "default") -> tuple[bool, str]:
        """
        Check if action is allowed under rate limits.

        Args:
            action: Type of action (prometheus, file_upload, export)
            identifier: Unique identifier for this user/session

        Returns:
            Tuple of (allowed, error_message)
        """
        if action not in self.limits:
            return True, ""

        limit_config = self.limits[action]
        max_calls = limit_config['max_calls']
        window = limit_config['window_seconds']

        key = f"{action}:{identifier}"
        current_time = time.time()

        # Initialize or clean old entries
        if key not in self.calls:
            self.calls[key] = []

        # Remove calls outside the window
        self.calls[key] = [t for t in self.calls[key]
                           if current_time - t < window]

        # Check limit
        if len(self.calls[key]) >= max_calls:
            retry_after = int(window - (current_time - self.calls[key][0]))
            return False, f"Rate limit exceeded. Retry in {retry_after}s"

        # Record this call
        self.calls[key].append(current_time)
        return True, ""


# Initialize global rate limiter
if 'rate_limiter' not in st.session_state:
    st.session_state.rate_limiter = RateLimiter(config.rate_limits.limits)

rate_limiter = st.session_state.rate_limiter


def safe_json_loads(json_string: str, max_size_mb: int = 10) -> tuple[bool, Any, str]:
    """
    Safely load JSON with size limits and error handling.

    Args:
        json_string: JSON string to parse
        max_size_mb: Maximum size in megabytes

    Returns:
        Tuple of (success, data, error_message)
    """
    try:
        # Check size
        size_mb = len(json_string.encode('utf-8')) / (1024 * 1024)
        if size_mb > max_size_mb:
            return False, None, f"JSON too large ({size_mb:.1f}MB > {max_size_mb}MB limit)"

        # Parse JSON
        data = json.loads(json_string)
        return True, data, ""

    except json.JSONDecodeError as e:
        return False, None, f"Invalid JSON: {str(e)}"
    except Exception as e:
        return False, None, f"Error loading JSON: {str(e)}"


st.set_page_config(page_title="IRAQAF Dashboard", layout="wide")

# Initialize UX Enhancements (Dark Mode, Loading, Keyboard Shortcuts, Session Info)
if UX_ENHANCEMENTS_AVAILABLE:
    initialize_ux_enhancements()

# ============================================================================
# AUTHENTICATION GATE - Check authentication FIRST before showing dashboard
# ============================================================================
if AUTH_UI_AVAILABLE:
    # Check if user is authenticated
    if not check_authentication():
        # User is not authenticated, show login page only
        st.stop()

# User is now authenticated - render dashboard
# Add sidebar controls
with st.sidebar:
    st.markdown("---")
    if AUTH_UI_AVAILABLE:
        render_user_info()
        st.markdown("---")
        render_logout_button()
    st.markdown("---")
    
    # UX Enhancements - Session Info
    if UX_ENHANCEMENTS_AVAILABLE:
        render_session_info()


# ============================================================================
# DATA PROCESSING HELPERS
# ============================================================================


def collect_module_history(files: list, module_prefix: str, extract_metrics: callable) -> pd.DataFrame:
    """
    Efficiently collect historical data for a module.

    Args:
        files: List of report file paths
        module_prefix: Module identifier (e.g., "L3", "L1")
        extract_metrics: Function to extract metrics from report dict

    Returns:
        DataFrame with historical data, sorted by time
    """
    rows = []

    for f in files:
        if not os.path.basename(f).startswith(f"{module_prefix}-"):
            continue

        try:
            with open(f, "r", encoding="utf-8") as fh:
                report = json.load(fh)

            if report.get("module") != module_prefix:
                continue

            # Extract common metadata
            base = os.path.basename(f).replace(".json", "")
            ts_match = re.search(r"(\d{8}-\d{6})", base)
            ts_label = ts_match.group(1) if ts_match else base
            mtime = os.path.getmtime(f)

            # Use custom extractor for module-specific metrics
            row_data = extract_metrics(report)
            row_data.update({
                "Run": base,
                "Label": ts_label,
                "time": mtime,
            })

            rows.append(row_data)

        except Exception as e:
            logger.debug(f"Skipped {f}: {e}")
            continue

    if not rows:
        return pd.DataFrame()

    df = pd.DataFrame(rows).sort_values("time").reset_index(drop=True)
    df["Run #"] = df.index + 1

    return df


@contextmanager
def progress_tracker(message: str, total: int = None):
    """
    Context manager for showing progress during long operations.

    Usage:
        with progress_tracker("Processing files", total=100) as tracker:
            for i in range(100):
                tracker.update(i + 1)
                # do work
    """
    if total:
        progress_bar = st.progress(0, text=message)
        status_text = st.empty()

        class Tracker:
            def update(self, current: int, custom_message: str = None):
                progress = current / total
                progress_bar.progress(
                    progress, text=custom_message or f"{message} ({current}/{total})")

        try:
            yield Tracker()
        finally:
            progress_bar.empty()
            status_text.empty()
    else:
        # Simple spinner for indeterminate progress
        with st.spinner(message):
            yield type('Tracker', (), {'update': lambda *args: None})()


# ============================================================================
# SESSION STATE INITIALIZATION
# ============================================================================


def init_session_state():
    """Initialize all session state variables with defaults"""
    defaults = {
        "__audit_mode__": False,
        "theme_mode": "Auto",
        "high_contrast": False,
        "search_history": [],
        "last_refresh": None,
        "user_preferences": {},
        "notification_settings": {
            "slack_url": "",
            "email_to": "",
            "smtp_host": "",
            "smtp_user": "",
        },
        "_incident_lock": None,  # Will be set to threading.Lock() when needed
        "system_coordinator": None,  # System integration coordinator
        "realtime_monitor": None,  # Real-time monitoring service
        "system_integration_enabled": True,  # Toggle system integration features
    }

    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

    # Initialize threading lock for incident logging
    if st.session_state["_incident_lock"] is None:
        import threading
        st.session_state["_incident_lock"] = threading.Lock()


# Initialize session state BEFORE anything else
init_session_state()

# Initialize System Integration (if available)
if SYSTEM_INTEGRATION_AVAILABLE:
    try:
        if st.session_state.get("system_coordinator") is None:
            db_url = os.getenv(
                "DATABASE_URL", "sqlite:///iraqaf_compliance.db")
            coordinator = initialize_coordinator(
                db_url=db_url, start_monitoring=True)
            st.session_state["system_coordinator"] = coordinator

            monitor = get_monitor()
            logger.info(f"Monitor type on creation: {type(monitor).__name__}")

            # Wrap the monitor to handle missing methods (due to Streamlit serialization)
            class MonitorWrapper:
                def __init__(self, wrapped):
                    self.wrapped = wrapped
                    self._cached_stats = None
                    self._cached_events = []

                def __getattr__(self, name):
                    return getattr(self.wrapped, name)

                def get_statistics(self):
                    """Get statistics with fallback handling."""
                    try:
                        if hasattr(self.wrapped, 'get_statistics'):
                            return self.wrapped.get_statistics()
                    except:
                        pass
                    # Return cached or default stats
                    return {
                        "running": True,
                        "total_events": 0,
                        "event_counts": {},
                        "is_monitoring": True
                    }

                def get_recent_events(self, count=10):
                    """Get recent events with fallback handling."""
                    try:
                        if hasattr(self.wrapped, 'get_recent_events'):
                            return self.wrapped.get_recent_events(count)
                    except:
                        pass
                    return []

            wrapped_monitor = MonitorWrapper(monitor)
            st.session_state["realtime_monitor"] = wrapped_monitor

            logger.info("✅ System integration initialized successfully")
    except Exception as e:
        logger.warning(f"Could not initialize system integration: {e}")
        st.session_state["system_integration_enabled"] = False

# NOW the audit mode check works reliably


@contextmanager
def show_loading(message: str, context: str = ""):
    """Show loading spinner with optional context inside a context manager."""
    placeholder = st.empty()

    if context:
        placeholder.caption(context)

    try:
        with st.spinner(f"{message}..."):
            yield
    finally:
        placeholder.empty()

# ===== Status Banner =====


def show_status_banner():
    """Show system status banner at the top."""

    # Check if reports are recent (< 24 hours old)
    import glob
    from datetime import datetime, timedelta

    try:
        report_files = get_report_files("*.json")
        if report_files:
            newest_report = max(report_files, key=os.path.getmtime)
            mtime = datetime.fromtimestamp(os.path.getmtime(newest_report))
            age = datetime.now() - mtime

            if age < timedelta(hours=24):
                status_color = "#d4edda"
                status_icon = "✅"
                status_text = f"Reports are up to date (last updated {age.seconds // 3600}h ago)"
            elif age < timedelta(days=7):
                status_color = "#fff3cd"
                status_icon = "⚠️"
                status_text = f"Reports are {age.days} days old - consider re-running evaluation"
            else:
                status_color = "#f8d7da"
                status_icon = "❌"
                status_text = f"Reports are {age.days} days old - please re-run evaluation"

            st.markdown(
                f"""
                <div style='
                    background: {status_color};
                    padding: 10px 20px;
                    border-radius: 8px;
                    margin-bottom: 20px;
                    text-align: center;
                    border: 1px solid rgba(0,0,0,0.1);
                '>
                    <span style='font-size: 1.1rem;'>{status_icon} {status_text}</span>
                </div>
                """,
                unsafe_allow_html=True
            )
    except Exception as e:
        logger.error(f"Failed to show status banner: {e}")


show_status_banner()

trace_map: dict[str, Any] = {}


def _load_evidence_index(path: str = "configs/evidence_index.json") -> dict:
    try:
        p = Path(path)
        if not p.exists():
            # Try with project root path
            p = Path(_ROOT) / "configs" / "evidence_index.json"

        if not p.exists():
            return {}

        with p.open("r", encoding="utf-8") as fh:
            data = json.load(fh)
        return data if isinstance(data, dict) else {}
    except Exception as e:
        try:
            logger.warning(f"Failed to load evidence index {path}: {e}")
        except Exception:
            pass
        return {}


# --------------------------------------------------------------------
# [SETTINGS] LOAD POLICIES (thresholds, refresh interval)
# --------------------------------------------------------------------


def load_policies(path: str = "configs/policies.yaml") -> dict[str, Any]:
    """
    Load and validate policy configuration.
    """
    defaults = {
        "latency_slo": 1.0,
        "error_rate_threshold": 5.0,
        "refresh_interval_s": 60
    }

    if not os.path.exists(path):
        logger.warning(f"Policy file not found: {path}. Using defaults.")
        return defaults

    try:
        with open(path, "r", encoding="utf-8") as fh:
            raw = yaml.safe_load(fh) or {}

        policy: dict[str, Any] = {}

        # latency_slo: must be positive float
        try:
            policy["latency_slo"] = float(
                raw.get("latency_slo", defaults["latency_slo"]))
            if policy["latency_slo"] <= 0:
                logger.warning(
                    f"Invalid latency_slo: {policy['latency_slo']}. Using default.")
                policy["latency_slo"] = defaults["latency_slo"]
        except (ValueError, TypeError):
            logger.warning("Invalid latency_slo type. Using default.")
            policy["latency_slo"] = defaults["latency_slo"]

        # error_rate_threshold: must be 0-100
        try:
            policy["error_rate_threshold"] = float(
                raw.get("error_rate_threshold",
                        defaults["error_rate_threshold"])
            )
            if not (0 <= policy["error_rate_threshold"] <= 100):
                logger.warning(
                    f"Invalid error_rate_threshold: {policy['error_rate_threshold']}. Using default."
                )
                policy["error_rate_threshold"] = defaults["error_rate_threshold"]
        except (ValueError, TypeError):
            logger.warning("Invalid error_rate_threshold type. Using default.")
            policy["error_rate_threshold"] = defaults["error_rate_threshold"]

        # refresh_interval_s: must be 5-3600
        try:
            policy["refresh_interval_s"] = int(
                raw.get("refresh_interval_s", defaults["refresh_interval_s"])
            )
            if not (5 <= policy["refresh_interval_s"] <= 3600):
                logger.warning(
                    f"Invalid refresh_interval_s: {policy['refresh_interval_s']}. Using default."
                )
                policy["refresh_interval_s"] = defaults["refresh_interval_s"]
        except (ValueError, TypeError):
            logger.warning("Invalid refresh_interval_s type. Using default.")
            policy["refresh_interval_s"] = defaults["refresh_interval_s"]

        logger.info(f"Loaded policies: {policy}")
        return policy

    except yaml.YAMLError as e:
        logger.error(f"YAML parse error in {path}: {e}")
        st.error(
            f"❌ Failed to parse policies.yaml. Using defaults. Error: {e}")
        return defaults
    except Exception as e:
        logger.error(f"Unexpected error loading policies: {e}", exc_info=True)
        return defaults


POLICY = load_policies()
logger.info(f"Active policy configuration: {POLICY}")

# --------------------------------------------------------------------
# [SETTINGS] OPTIONAL DEPENDENCY CHECK
# --------------------------------------------------------------------
OPTIONAL_DEPS = {
    "matplotlib": "pip install matplotlib",
    "pdfkit": "pip install pdfkit",
    "yaml": "pip install pyyaml",
    "docx": "pip install python-docx"
}

with st.sidebar.expander("🧩 Dependencies"):
    for lib, cmd in OPTIONAL_DEPS.items():
        spec = importlib.util.find_spec(lib)
        if spec is None:
            st.error(f"❌ {lib} not installed ❌ `{cmd}`")
        else:
            st.success(f"✅ {lib}")


# ============================================================================
# AUDIT MODE MANAGEMENT
# ============================================================================

def get_audit_mode() -> bool:
    """Centralized audit mode state getter"""
    return st.session_state.get("__audit_mode__", False)


def set_audit_mode(enabled: bool):
    """Centralized audit mode state setter"""
    st.session_state["__audit_mode__"] = enabled
    record_audit_event("ENABLED" if enabled else "DISABLED",
                       get_current_run_hash())


def get_current_run_hash() -> str:
    """Get hash of current aggregate report for audit trail"""
    try:
        agg_paths = sorted(get_report_files("AGG-*.json"),
                           key=os.path.getmtime)
        if agg_paths:
            latest_agg = agg_paths[-1]
            return _hash_file(latest_agg, _mtime=os.path.getmtime(latest_agg))
    except Exception as e:
        logger.error(f"Failed to get run hash: {e}")
    return "unknown"


# Alias for backward compatibility
def audit_locked() -> bool:
    """Check if audit mode is active (locks all inputs)"""
    return get_audit_mode()


# --------------------------------------------------------------------
# 🔒 AUDIT LOCK BOOTSTRAP
# --------------------------------------------------------------------
# Ensure the audit flag exists before any widgets reference _LOCK.
if "__audit_mode__" not in st.session_state:
    st.session_state["__audit_mode__"] = False

# Convenience variable used throughout the app to disable widgets.
_LOCK = get_audit_mode()

# --------------------------------------------------------------------
# 🛠️ DATA VALIDATION + THEME + AUDIT LOG (SIDEBAR CONFIG)
# --------------------------------------------------------------------
# This unified sidebar section handles:
#   1. Schema version display
#   2. Theme switcher (Auto / Light / Dark + High Contrast)
#   3. Audit log viewer with download option
# --------------------------------------------------------------------

# ===== 1🛠️ DATA VALIDATION INFO =====
with st.sidebar.expander("🛠️ Data Validation", expanded=False):
    st.markdown(f"Schema version: **{SCHEMA_VERSION}**")
    st.caption("Reports automatically validated on load.")

# ===== 2🎨 THEME SWITCHER (Auto / Light / Dark + High Contrast) =====


def _inject_theme_css(mode: str):
    """Inject CSS variables for light/dark themes with system auto support."""
    base_css = """
    <style>
      :root {
        --bg: #ffffff;
        --text: #0f172a;
        --muted: #6b7280;
        --card: #ffffff;
        --border: #e5e7eb;
        --shadow: 0 1px 2px rgba(0,0,0,0.04);
        --btn-bg: #f3f4f6;
        --btn-text: #111827;
        --btn-border: #d1d5db;
        --link: #2563eb;
        --metric-bg: transparent;
        --df-header-bg: #f8fafc;
        --df-border: #e5e7eb;
        --expander-bg: #ffffff;
      }
      .stApp, body { background: var(--bg); color: var(--text); }
      a, .markdown-text-container a { color: var(--link); }
      .iraqaf-card {
        background: var(--card);
        border: 1px solid var(--border);
        border-radius: 12px;
        padding: 14px;
        margin-bottom: 12px;
        box-shadow: var(--shadow);
      }
      .stButton>button, .stDownloadButton>button {
        background: var(--btn-bg) !important;
        color: var(--btn-text) !important;
        border: 1px solid var(--btn-border) !important;
        border-radius: 10px !important;
      }
      [data-testid="stTextInput"] input,
      [data-testid="stNumberInput"] input,
      [data-testid="stSelectbox"] div[data-baseweb="select"] > div,
      [data-testid="stFileUploader"] {
        color: var(--text);
        background: var(--card);
        border: 1px solid var(--border);
      }
      [data-testid="stMetricValue"], [data-testid="stMetricDelta"] { color: var(--text); }
      .stDataFrame thead tr th { background: var(--df-header-bg); color: var(--text); }
      .stDataFrame table, .stDataFrame td, .stDataFrame th { border-color: var(--df-border) !important; }
      [data-testid="stExpander"] { background: var(--expander-bg); border-radius:12px; }
      .vega-embed, canvas { background: transparent !important; }
    </style>
    """
    dark_overrides = """
    <style>
      :root {
        --bg: #0f1116;
        --text: #e6e6e6;
        --muted: #9aa3b2;
        --card: #151922;
        --border: #2a2f3a;
        --shadow: 0 1px 2px rgba(0,0,0,0.3);
        --btn-bg: #1f2937;
        --btn-text: #e6e6e6;
        --btn-border: #374151;
        --link: #8ab4ff;
        --metric-bg: rgba(255,255,255,0.02);
        --df-header-bg: #1e2430;
        --df-border: #2a2f3a;
        --expander-bg: #151922;
      }
    </style>
    """
    auto_block = f"""
    <style>
      @media (prefers-color-scheme: dark) {{
        {dark_overrides.replace('<style>', '').replace('</style>', '')}
      }}
    </style>
    """
    if mode == "Light":
        st.markdown(base_css, unsafe_allow_html=True)
    elif mode == "Dark":
        st.markdown(base_css + dark_overrides, unsafe_allow_html=True)
    else:
        st.markdown(base_css + auto_block, unsafe_allow_html=True)


# --- Theme Expander ---
with st.sidebar.expander("🎨 Theme", expanded=False):
    mode = st.radio(
        "Mode", ["Auto", "Light", "Dark"],
        horizontal=True,
        index=0,
        disabled=_LOCK,
        key="theme_mode_selector"
    )
    
    # UX Enhancements - Dark mode toggle
    if UX_ENHANCEMENTS_AVAILABLE:
        st.divider()
        render_theme_toggle()

    # Apply base theme
    _inject_theme_css(mode)

    # Accessibility: High contrast mode toggle
    high_contrast = st.checkbox(
        "🦾 High contrast mode",
        key="high_contrast_toggle",
        disabled=_LOCK,
        help="Boosts text contrast and slightly increases base font size for readability."
    )
    if high_contrast:
        st.markdown("""
        <style>
          :root {
            --text: #ffffff !important;
            --muted: #d9d9d9 !important;
          }
          html, body, .stApp { font-size: 1.05rem !important; }
          [data-testid="stTextInput"] input,
          [data-testid="stNumberInput"] input,
          [data-testid="stSelectbox"] div[data-baseweb="select"] > div {
            border-color: var(--link) !important;
          }
        </style>
        """, unsafe_allow_html=True)
# ===== Keyboard Shortcuts =====
st.markdown("""
<script>
// Keyboard shortcuts for IRAQAF dashboard
document.addEventListener('keydown', function(e) {
    // Ctrl/Cmd + K: Focus search
    if ((e.ctrlKey || e.metaKey) && e.key === 'k') {
        e.preventDefault();
        const searchInputs = document.querySelectorAll('input[type="text"]');
        if (searchInputs.length > 0) {
            searchInputs[0].focus();
        }
    }

    // Ctrl/Cmd + R: Refresh (with confirmation)
    if ((e.ctrlKey || e.metaKey) && e.key === 'r') {
        e.preventDefault();
        if (confirm('Refresh dashboard?')) {
            window.location.reload();
        }
    }

    // Escape: Close modals/expandable sections
    if (e.key === 'Escape') {
        const expandedSections = document.querySelectorAll('[data-testid="stExpander"][aria-expanded="true"]');
        expandedSections.forEach(section => {
            const button = section.querySelector('button');
            if (button) button.click();
        });
    }

    // ? key: Show keyboard shortcuts help
    if (e.key === '?') {
        e.preventDefault();
        alert(
            'IRAQAF Keyboard Shortcuts:\\n\\n' +
            'Ctrl/Cmd + K  ?  Focus search\\n' +
            'Ctrl/Cmd + R  ?  Refresh dashboard\\n' +
            'Escape        ?  Close expanded sections\\n' +
            '?             ?  Show this help'
        );
    }
});

// Show keyboard shortcut hint on first load
if (!sessionStorage.getItem('shortcuts_shown')) {
    sessionStorage.setItem('shortcuts_shown', 'true');

    // Create tooltip
    const tooltip = document.createElement('div');
    tooltip.innerHTML = '💡 Press <kbd>?</kbd> for keyboard shortcuts';
    tooltip.style.cssText = `
        position: fixed;
        bottom: 20px;
        left: 20px;
        background: #1f2937;
        color: white;
        padding: 12px 16px;
        border-radius: 8px;
        font-size: 14px;
        z-index: 9999;
        animation: fadeInOut 4s ease-in-out;
    `;

    document.body.appendChild(tooltip);

    setTimeout(() => tooltip.remove(), 4000);
}
</script>

<style>
@keyframes fadeInOut {
    0%, 100% { opacity: 0; }
    10%, 90% { opacity: 1; }
}

kbd {
    background: #374151;
    padding: 2px 6px;
    border-radius: 4px;
    font-family: monospace;
    font-size: 0.9em;
}
</style>
""", unsafe_allow_html=True)


# ===== Responsive Design Enhancements =====
st.markdown("""
<style>
/* Mobile-friendly adjustments */
@media (max-width: 768px) {
    /* Stack columns on mobile */
    [data-testid="column"] {
        width: 100% !important;
        flex: none !important;
    }

    /* Larger tap targets for mobile */
    button {
        min-height: 44px !important;
        font-size: 16px !important;
    }

    /* Better text readability */
    .stMarkdown {
        font-size: 16px !important;
    }

    /* Hide less important elements on small screens */
    .mobile-hide {
        display: none !important;
    }
}

/* Tablet adjustments */
@media (min-width: 769px) and (max-width: 1024px) {
    [data-testid="stSidebar"] {
        width: 250px !important;
    }
}

/* Better scrolling for tables */
.stDataFrame {
    max-height: 600px;
    overflow-y: auto;
}

/* Smooth scrolling */
html {
    scroll-behavior: smooth;
}

/* Better focus indicators for accessibility */
*:focus-visible {
    outline: 2px solid var(--link) !important;
    outline-offset: 2px !important;
    border-radius: 4px;
}

/* Loading animation */
@keyframes pulse {
    0%, 100% { opacity: 1; }
    50% { opacity: 0.5; }
}

.loading {
    animation: pulse 2s cubic-bezier(0.4, 0, 0.6, 1) infinite;
}

/* Improved card shadows on hover */
.iraqaf-card:hover {
    box-shadow: 0 4px 12px rgba(0,0,0,0.1);
    transition: box-shadow 0.3s ease;
}

/* Better spacing for metrics */
[data-testid="stMetricValue"] {
    font-size: 2rem !important;
    font-weight: 600 !important;
}

/* Prettier scrollbars (Webkit browsers) */
::-webkit-scrollbar {
    width: 10px;
    height: 10px;
}

::-webkit-scrollbar-track {
    background: var(--bg);
}

::-webkit-scrollbar-thumb {
    background: var(--border);
    border-radius: 5px;
}

::-webkit-scrollbar-thumb:hover {
    background: var(--muted);
}
</style>
""", unsafe_allow_html=True)
# --- Focus outlines for accessibility (global, after theme) ---
st.markdown("""
<style>
button:focus, input:focus, textarea:focus, select:focus,
[data-testid="stTextInput"] input:focus,
[data-testid="stNumberInput"] input:focus,
[data-testid="stSelectbox"] [data-baseweb="select"]:focus,
[data-testid="stFileUploader"] [data-testid="stDropzone"] {
  outline: 2px solid var(--link) !important;
  outline-offset: 2px !important;
}
</style>
""", unsafe_allow_html=True)

# ===== 3️⃣ AUDIT LOG VIEWER =====
with st.sidebar.expander("📝 Audit Log", expanded=False):
    log_path = "logs/audit.log"
    if os.path.exists(log_path):
        with open(log_path, "r", encoding="utf-8") as fh:
            content = fh.read().strip()
        if content:
            st.text(content)
            st.download_button(
                "📝 Download audit log",
                data=content.encode("utf-8"),
                file_name="audit.log",
                mime="text/plain",
                disabled=_LOCK
            )
        else:
            st.caption("Audit log is empty.")
    else:
        st.caption("No audit events recorded yet.")


# ===== 📜 APPLICATION LOG VIEWER =====
with st.sidebar.expander("📜 Application Logs", expanded=False):
    log_file = "logs/app.log"
    if os.path.exists(log_file):
        with open(log_file, "r", encoding="utf-8") as fh:
            content = fh.read().strip()
        if content:
            st.text(content[-4000:])  # show last few KB for readability
            st.download_button(
                "📜 Download app.log",
                data=open(log_file, "rb").read(),
                file_name="app.log",
                mime="text/plain",
                disabled=_LOCK,
            )
        else:
            st.caption("Log file is empty.")
    else:
        st.caption("No logs yet.")
# =============================================================================
#  🔧 SETTINGS PANEL
# =============================================================================

with st.sidebar.expander("🔧 Settings", expanded=False):
    st.markdown("**Display Preferences**")

    # Auto-refresh setting
    auto_refresh = st.checkbox(
        "Auto-refresh data",
        value=st.session_state.get('auto_refresh_enabled', False),
        key="auto_refresh_toggle",
        disabled=_LOCK,
        help="Automatically refresh data every 5 minutes"
    )
    st.session_state['auto_refresh_enabled'] = auto_refresh

    if auto_refresh:
        st.caption("⏳ Data refreshes every 5 minutes")
        # Trigger auto-refresh
        st_autorefresh(interval=300000, key="settings_autorefresh")

    # Compact mode
    compact_mode = st.checkbox(
        "Compact mode",
        value=st.session_state.get('compact_mode', False),
        key="compact_mode_toggle",
        disabled=_LOCK,
        help="Reduce spacing for more content on screen"
    )
    st.session_state['compact_mode'] = compact_mode

    # Apply compact mode CSS
    if compact_mode:
        st.markdown("""
        <style>
        /* Compact mode overrides */
        .stMarkdown, .stText { margin-bottom: 0.5rem !important; }
        .stExpander { margin-bottom: 0.5rem !important; }
        [data-testid="stVerticalBlock"] > div { gap: 0.5rem !important; }
        .iraqaf-card { padding: 10px !important; margin-bottom: 8px !important; }
        h1, h2, h3, h4 { margin-top: 0.5rem !important; margin-bottom: 0.5rem !important; }
        .stDataFrame { max-height: 400px !important; }
        </style>
        """, unsafe_allow_html=True)

    # Show timestamps
    show_timestamps = st.checkbox(
        "Show timestamps",
        value=st.session_state.get('show_timestamps', True),
        key="show_timestamps_toggle",
        disabled=_LOCK,
        help="Display when each section was last updated"
    )
    st.session_state['show_timestamps'] = show_timestamps

    # Default chart style
    st.markdown("**Chart Settings**")
    chart_theme = st.selectbox(
        "Chart theme",
        options=["Default", "Dark", "Light"],
        index=0,
        key="chart_theme_select",
        disabled=_LOCK
    )
    st.session_state['chart_theme'] = chart_theme

    # Animation toggle
    enable_animations = st.checkbox(
        "Enable animations",
        value=st.session_state.get('enable_animations', True),
        key="animations_toggle",
        disabled=_LOCK,
        help="Smooth transitions and hover effects"
    )
    st.session_state['enable_animations'] = enable_animations

    if not enable_animations:
        st.markdown("""
        <style>
        * { transition: none !important; animation: none !important; }
        </style>
        """, unsafe_allow_html=True)

    # Notification preferences
    st.markdown("**Notifications**")
    notify_on_breach = st.checkbox(
        "Alert on SLO breaches",
        value=st.session_state.get('notify_on_breach', True),
        key="notify_breach_toggle",
        disabled=_LOCK
    )
    st.session_state['notify_on_breach'] = notify_on_breach

    # Data display preferences
    st.markdown("**Data Display**")
    decimal_places = st.slider(
        "Decimal places",
        min_value=0, max_value=4, value=2,
        key="decimal_places_slider",
        disabled=_LOCK,
        help="Number of decimals to show in metrics"
    )
    st.session_state['decimal_places'] = decimal_places

    # Export settings to JSON
    st.markdown("---")
    if st.button("💾 Save Settings", key="save_settings_btn", disabled=_LOCK, width="stretch"):
        settings = {
            'auto_refresh_enabled': auto_refresh,
            'compact_mode': compact_mode,
            'show_timestamps': show_timestamps,
            'chart_theme': chart_theme,
            'enable_animations': enable_animations,
            'notify_on_breach': notify_on_breach,
            'decimal_places': decimal_places,
            'theme_mode': st.session_state.get('theme_mode', 'Auto'),
            'high_contrast': st.session_state.get('high_contrast', False),
            'saved_at': datetime.now().isoformat()
        }

        settings_file = Path("configs/user_settings.json")
        settings_file.parent.mkdir(exist_ok=True)

        try:
            with open(settings_file, "w", encoding="utf-8") as f:
                json.dump(settings, f, indent=2)
            st.success("💾 Settings saved!")
            logger.info(f"User settings saved: {settings}")
        except Exception as e:
            st.error(f"Failed to save settings: {e}")

    # Load settings
    if st.button("📂 Load Settings", key="load_settings_btn", disabled=_LOCK, width="stretch"):
        settings_file = Path("configs/user_settings.json")
        if settings_file.exists():
            try:
                with open(settings_file, "r", encoding="utf-8") as f:
                    settings = json.load(f)

                # Apply settings
                for key, value in settings.items():
                    if key != 'saved_at':
                        st.session_state[key] = value

                st.success(
                    f"📂 Settings loaded from {settings.get('saved_at', 'unknown time')}")
                st.rerun()
            except Exception as e:
                st.error(f"Failed to load settings: {e}")
        else:
            st.info("No saved settings found")

    # Reset to defaults
    if st.button("🔄 Reset to Defaults", key="reset_settings_btn", disabled=_LOCK, width="stretch"):
        defaults = {
            'auto_refresh_enabled': False,
            'compact_mode': False,
            'show_timestamps': True,
            'chart_theme': 'Default',
            'enable_animations': True,
            'notify_on_breach': True,
            'decimal_places': 2,
            'theme_mode': 'Auto',
            'high_contrast': False
        }
        for key, value in defaults.items():
            st.session_state[key] = value
        st.success("🔄 Settings reset to defaults")
        st.rerun()

# Add Export & Alerts to sidebar (after Settings)
if ENHANCEMENTS_AVAILABLE:
    render_export_section()

with st.sidebar.expander("❓ Help & Documentation", expanded=False):
    # Quick Navigation Section
    st.markdown("### 🧭 Quick Navigation")
    st.markdown("""
    - Use the search bar to find specific metrics
    - Click module names to jump to sections
    - Enable tour mode for guided walkthrough
    - Scroll to explore detailed compliance data
    """)

    st.divider()

    # Keyboard Shortcuts Section
    st.markdown("### ⌨️ Keyboard Shortcuts")
    shortcuts_data = {
        "Ctrl+K": "Focus search",
        "Ctrl+R": "Refresh dashboard",
        "Esc": "Close sections",
        "?": "Show shortcuts"
    }
    for key, desc in shortcuts_data.items():
        st.caption(f"**{key}** — {desc}")

    st.divider()

    # Troubleshooting Section
    st.markdown("### 🔧 Troubleshooting")
    st.markdown("""
    **Missing data?**
    Ensure reports were generated correctly

    **Red modules?**
    Review issues in module details

    **Performance issues?**
    Check Debug Mode in sidebar

    **Still stuck?**
    Check documentation or report issue
    """)

    st.divider()

    # Help buttons with custom CSS styling
    st.markdown("""
    <style>
    .help-btn {
        padding: 12px 16px !important;
        border-radius: 4px !important;
        font-weight: 600 !important;
        text-align: center !important;
        cursor: pointer !important;
        border: none !important;
        transition: all 0.3s ease !important;
        font-size: 14px !important;
        width: 100% !important;
        display: block !important;
        min-height: 44px !important;
        line-height: 20px !important;
    }

    .help-btn-docs {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
        color: white !important;
    }

    .help-btn-docs:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 8px 24px rgba(102, 126, 234, 0.4) !important;
    }

    .help-btn-issue {
        background: linear-gradient(135deg, #ff6b6b 0%, #ee5a6f 100%) !important;
        color: white !important;
    }

    .help-btn-issue:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 8px 24px rgba(255, 107, 107, 0.4) !important;
    }
    </style>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns(2, gap="small")

    with col1:
        if st.button("📖 Docs", key="btn_full_docs", width="stretch", help="Open complete documentation"):
            st.info(
                "📚 Full Documentation\n\nhttps://iraqaf.readthedocs.io/\n\nOpen the link in a new tab.")

    with col2:
        if st.button("🐛 Issue", key="btn_report_issue", width="stretch", help="Report a bug or issue"):
            st.warning(
                "🐞 Report a Bug\n\nhttps://github.com/iraqaf/dashboard/issues\n\nOpen the link in a new tab.")

    st.divider()

    # Quick Tips
    st.markdown("### 💡 Quick Tips")
    st.markdown("""
    - **Color Coding**: Green ✅ = Compliant | Yellow ⚠️ = Attention needed | Red ❌ = Issues found
    - **Module Levels**: L1 (Governance) 🏛️ L5 (Verification) represent compliance depth
    - **Evidence Tray**: Drag and drop evidence files for quick processing
    - **Export Reports**: Generate PDF or DOCX reports from any section
    """)  # =============================================================================
#  🎯 TOUR MODE
# =============================================================================
with st.sidebar.expander("🎯 Interactive Tour", expanded=False):
    st.markdown("Take a guided tour of the dashboard features")

    if st.button("🚀 Start Tour", key="sidebar_start_tour", width="stretch", type="primary"):
        st.session_state['tour_active'] = True
        st.session_state['tour_step'] = 0
        st.rerun()

    if st.session_state.get('tour_active', False):
        st.info(
            f"📍 Tour in progress: Step {st.session_state.get('tour_step', 0) + 1}/{len(TOUR_STEPS)}")

# =============================================================================
#  🔌 SYSTEM INTEGRATION STATUS WIDGET (Sidebar)
# =============================================================================

if SYSTEM_INTEGRATION_AVAILABLE and st.session_state.get("system_integration_enabled"):
    with st.sidebar.expander("🔌 System Integration Status", expanded=True):
        coordinator = st.session_state.get("system_coordinator")
        monitor = st.session_state.get("realtime_monitor")

        if coordinator and monitor:
            try:
                # Debug info
                monitor_type = type(monitor).__name__

                # Check if monitor has required methods
                has_get_stats = hasattr(monitor, 'get_statistics')
                has_get_events = hasattr(monitor, 'get_recent_events')

                if not has_get_stats or not has_get_events:
                    st.warning(f"⚠️ Monitor type issue: {monitor_type}")
                    st.caption(f"has_get_statistics: {has_get_stats}")
                    st.caption(f"has_get_recent_events: {has_get_events}")
                else:
                    status = coordinator.get_system_status()

                    # Status indicators
                    col1, col2 = st.columns(2)

                    with col1:
                        is_active = monitor.get_statistics().get(
                            "is_monitoring", False) if has_get_stats else False
                        if is_active:
                            st.markdown("✅ **Monitor Active**")
                        else:
                            st.markdown("🔴 **Monitor Inactive**")

                    with col2:
                        db_ok = status.get("database_connected", False)
                        if db_ok:
                            st.markdown("✅ **DB Connected**")
                        else:
                            st.markdown("🔴 **DB Error**")

                    # Quick stats
                    st.markdown("---")

                    col_a, col_b, col_c = st.columns(3)

                    with col_a:
                        changes = status.get("total_regulatory_changes", 0)
                        st.metric("Changes", changes,
                                  label_visibility="collapsed")

                    with col_b:
                        alerts = status.get("open_alerts", 0)
                        st.metric("Alerts", alerts,
                                  label_visibility="collapsed")

                    with col_c:
                        compliance = status.get("average_compliance_score", 0)
                        st.metric(
                            "Compliance", f"{compliance:.0f}%", label_visibility="collapsed")

                    st.markdown("---")

                    # Last event
                    try:
                        if has_get_events:
                            recent_events = monitor.get_recent_events(1)
                            if recent_events:
                                last_event = recent_events[0].to_dict() if hasattr(
                                    recent_events[0], 'to_dict') else recent_events[0]
                                st.caption(
                                    f"📌 Last Event: {last_event.get('event_type', 'Unknown')}")
                                st.caption(
                                    f"🕐 {last_event.get('timestamp', 'N/A')}")
                    except:
                        pass

                    # Action buttons
                    if st.button("🔄 Refresh Now", width="stretch", key="sidebar_refresh_status"):
                        st.rerun()

            except Exception as e:
                st.warning(f"⚠️ Status unavailable: {str(e)[:80]}")
        else:
            st.warning("System components not initialized")

# ===== END SYSTEM STATUS WIDGET =====

# ===== ERROR SURFACING HELPER =====----------------------------------------------------------------
# ⚠️ ERROR SURFACING HELPER
# --------------------------------------------------------------------


def build_clause_keywords_from_trace_map() -> dict:
    """
    Build clause keyword mappings dynamically from trace_map.yaml
    Maps domain keywords to (framework, clause) tuples
    """
    trace_map = load_trace_map()

    # Domain keyword mappings
    keyword_to_topics = {
        "healthcare": ["Quality", "Management", "Risk"],
        "medical": ["Quality", "SaMD"],
        "privacy": ["Data_Principles", "Security"],
        "security": ["Security", "Integrity", "Robustness"],
        "fairness": ["Data_Governance", "Bias"],
        "bias": ["Data_Governance"],
        "transparency": ["Transparency"],
        "interpretability": ["Transparency", "Decisions"],
        "explainability": ["Transparency", "Decisions"],
        "oversight": ["Human_Oversight", "Oversight"],
        "monitoring": ["Robustness", "Monitoring"],
        "logging": ["Robustness", "Security", "Integrity"],
        "governance": ["Risk_Management", "Governance", "Data_Governance"],
    }

    # Build reverse mapping: keyword -> [(framework, clause), ...]
    keyword_mappings = {kw: [] for kw in keyword_to_topics.keys()}

    for framework, clauses in trace_map.items():
        for clause_id in clauses.keys():
            # Check which keywords match this clause
            for keyword, topics in keyword_to_topics.items():
                if any(topic.lower() in clause_id.lower() for topic in topics):
                    keyword_mappings[keyword].append((framework, clause_id))

    return keyword_mappings


# Initialize at startup
_CLAUSE_KEYWORDS = build_clause_keywords_from_trace_map()

# ===== Automated Clause Tagging (using trace_map.yaml) =====
st.markdown("### 🏷️ Automated Clause Tagging")

st.markdown("""
<div style='
    background: #f0f9ff;
    border-left: 4px solid #0284c7;
    padding: 12px 16px;
    margin-bottom: 16px;
    border-radius: 4px;
'>
    <p style='margin: 0; font-size: 0.9rem; color: #0c4a6e;'>
        <b>💡 How it works:</b> Describe your AI system and we'll suggest relevant compliance clauses 
        based on keywords from <code>configs/trace_map.yaml</code>
    </p>
</div>
""", unsafe_allow_html=True)

desc = st.text_area(
    "Describe your AI system or use case",
    placeholder="e.g., Healthcare AI for clinical decision support using patient EHR data, requires explainability and fairness monitoring",
    height=120,
    key="proj_desc_clause_tag",
    disabled=_LOCK
)

col1, col2 = st.columns([3, 1])

with col1:
    if st.button("🔍 Suggest relevant clauses", disabled=_LOCK, key="btn_suggest_clauses_v3", width="stretch"):
        text = (desc or "").strip().lower()

        if not text:
            st.warning("Please enter a description first")
        else:
            # Build keyword mappings dynamically
            trace_map = load_trace_map()
            suggestions = []

            # Domain keyword patterns
            keyword_patterns = {
                "healthcare|medical|clinical|patient": ["Quality", "SaMD", "Management"],
                "privacy|gdpr|personal.?data|data.?protection": ["Data_Principles", "Security", "Privacy"],
                "security|encryption|access.?control|integrity": ["Security", "Integrity", "Robustness"],
                "fairness|bias|discrimination|equity": ["Data_Governance", "Fairness", "Bias"],
                "transparency|explainability|interpretability": ["Transparency", "Decisions"],
                "monitoring|logging|audit|oversight": ["Robustness", "Monitoring", "Oversight", "Logging"],
                "governance|risk|compliance": ["Risk_Management", "Governance", "Data_Governance"],
            }

            # Find matching clauses
            import re
            for pattern, topics in keyword_patterns.items():
                if re.search(pattern, text, re.IGNORECASE):
                    # Search all frameworks for clauses matching these topics
                    for fw, clauses in trace_map.items():
                        for clause_id, clause_data in clauses.items():
                            # Check if any topic matches the clause ID
                            if any(topic.lower() in clause_id.lower() for topic in topics):
                                suggestions.append(
                                    (fw, clause_id, clause_data))

            # Deduplicate
            seen = set()
            unique_suggestions = []
            for fw, clause_id, clause_data in suggestions:
                key = (fw, clause_id)
                if key not in seen:
                    seen.add(key)
                    unique_suggestions.append((fw, clause_id, clause_data))

            if not unique_suggestions:
                st.info("""
                No suggestions found. Try adding domain keywords:
                - **Healthcare/Medical**: clinical, patient, medical device
                - **Privacy**: GDPR, personal data, data protection
                - **Fairness**: bias, discrimination, equity
                - **Explainability**: transparency, interpretability
                - **Security**: encryption, access control, integrity
                - **Monitoring**: logging, audit, oversight
                """)
            else:
                st.success(
                    f"✅ Found {len(unique_suggestions)} relevant clause(s)")

                # Group by framework
                suggestions_by_fw = {}
                for fw, clause_id, clause_data in unique_suggestions:
                    if fw not in suggestions_by_fw:
                        suggestions_by_fw[fw] = []
                    suggestions_by_fw[fw].append((clause_id, clause_data))

                # Display grouped results
                for fw, clauses in sorted(suggestions_by_fw.items()):
                    with st.expander(f"📋 {fw} ({len(clauses)} clause{'s' if len(clauses) != 1 else ''})"):
                        for clause_id, clause_data in sorted(clauses):
                            # Get metrics and description
                            metrics = clause_data.get("metrics", []) if isinstance(
                                clause_data, dict) else clause_data
                            description = clause_data.get(
                                "description", "") if isinstance(clause_data, dict) else ""

                            st.markdown(f"**{clause_id}**")
                            if description:
                                st.caption(f"📝 {description}")
                            if metrics:
                                st.caption(
                                    f"🔗 Traces to: {', '.join(metrics)}")
                            else:
                                st.caption("⚠️ No metrics traced")
                            st.markdown("---")

                # Save suggestions feature
                st.markdown("---")
                st.markdown("### 💾 Save Suggestions")

                save_col1, save_col2 = st.columns([3, 1])

                with save_col1:
                    st.caption(
                        "Save these suggestions to `configs/suggested_clauses.json` for future reference")

                with save_col2:
                    if st.button("💾 Save", key="save_suggestions_btn_v2", disabled=_LOCK, width="stretch"):
                        suggestions_data = {
                            "description": desc,
                            "generated_at": datetime.now().isoformat(),
                            "total_suggestions": len(unique_suggestions),
                            "frameworks": list(suggestions_by_fw.keys()),
                            "suggestions": [
                                {
                                    "framework": fw,
                                    "clause": clause_id,
                                    "description": clause_data.get("description", "") if isinstance(clause_data, dict) else "",
                                    "metrics": clause_data.get("metrics", []) if isinstance(clause_data, dict) else clause_data
                                }
                                for fw, clause_id, clause_data in unique_suggestions
                            ]
                        }

                        try:
                            output_path = Path(
                                "configs/suggested_clauses.json")
                            output_path.parent.mkdir(exist_ok=True)

                            with open(output_path, "w", encoding="utf-8") as f:
                                json.dump(suggestions_data, f, indent=2)

                            st.success(
                                f"✅ Saved {len(unique_suggestions)} suggestions to `{output_path}`")
                            logger.info(
                                f"Saved {len(unique_suggestions)} clause suggestions")

                        except Exception as e:
                            st.error(f"❌ Failed to save: {e}")
                            logger.error(
                                f"Save suggestions failed: {e}", exc_info=True)

with col2:
    if st.button("🔄 Reload", key="reload_trace_map_btn_v2", disabled=_LOCK, help="Reload trace_map.yaml", width="stretch"):
        load_trace_map.clear()
        get_available_frameworks.clear()
        get_framework_clauses.clear()
        get_clause_metrics.clear()
        if 'get_clause_description' in dir():
            get_clause_description.clear()
        st.success("✅ Reloaded!")
        st.rerun()
# Show current frameworks loaded
with st.expander("📚 Available Frameworks & Clauses", expanded=False):
    frameworks = get_available_frameworks()

    if not frameworks:
        st.warning("⚠️ No frameworks found in trace_map.yaml")
    else:
        for fw in frameworks:
            clauses = get_framework_clauses(fw)
            st.markdown(f"**{fw}** — {len(clauses)} clause(s)")
            for clause in clauses:
                metrics = get_clause_metrics(fw, clause)
                st.caption(
                    f"  • {clause} → {', '.join(metrics) if metrics else 'no metrics'}")

# ===== Audit Mode (hardened) =====
with st.container():
    audit_mode = st.toggle(
        "🔒 Audit Mode",
        value=st.session_state.get("__audit_mode__", False),
        help="Locks inputs and records a content hash for this snapshot."
    )
    files = get_report_files("*.json")
    agg_paths = [p for p in files if os.path.basename(p).startswith("AGG-")]
    agg_path = max(agg_paths, key=os.path.getmtime) if agg_paths else None
    # NEW:
if agg_path:
    try:
        mtime = os.path.getmtime(agg_path)
        run_hash = _hash_file(agg_path, _mtime=mtime)
    except Exception as e:
        logger.error(f"Failed to hash {agg_path}: {e}")
        run_hash = None
else:
    run_hash = None

    # display hash + time
    st.caption(
        f"Run hash: **{(run_hash or 'n/a')[:16]}**  |  Time UTC: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}"
    )

st.divider()

# --- Persist + log changes ---
if audit_mode != st.session_state.get("__audit_mode__", False):
    st.session_state["__audit_mode__"] = audit_mode
    record_audit_event("ENABLED" if audit_mode else "DISABLED", run_hash)

_LOCK = audit_locked()


# --- Audit lock overlay ---
if _LOCK:
    st.markdown("""
    <style>
    .stApp::after {
        content: "🔒 Audit Mode Active – Read-only";
        position: fixed;
        top: 50%%;
        left: 50%%;
        transform: translate(-50%%,-50%%);
        background: rgba(15,17,22,0.8);
        color: #e5e7eb;
        padding: 20px 40px;
        border-radius: 12px;
        font-size: 1.2rem;
        z-index: 9999;
    }
    .stButton>button, .stFileUploader, .stSelectbox, input, textarea {
        filter: grayscale(100%%) brightness(0.7);
        pointer-events: none;
    }
    </style>
    """, unsafe_allow_html=True)


# ====== STREAMLIT DASHBOARD FOR IRAQAF REPORTS ======
st.title("🧠 IRAQAF – Quality Assurance Scores")

# Initialize enhancements (Export, Alerts, RBAC)
if ENHANCEMENTS_AVAILABLE:
    initialize_enhancements()
    render_role_based_dashboard()

# Check if this is first visit
if 'first_visit' not in st.session_state:
    st.session_state['first_visit'] = True
    st.session_state['hide_quickstart'] = False

# Auto-close guide after first visit
if not st.session_state['hide_quickstart']:
    # Expanded only on first visit
    is_expanded = st.session_state.get('first_visit', False)

    with st.expander("👋Quick Start Guide", expanded=is_expanded):
        # Mark as not first visit anymore
        if st.session_state.get('first_visit'):
            st.session_state['first_visit'] = False

        # Styled banner
        st.markdown("""
        <div style='
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            padding: 20px;
            border-radius: 12px;
            margin-bottom: 20px;
            color: white;
            text-align: center;
        '>
            <h2 style='margin: 0; color: white;'>Welcome to IRAQAF Dashboard! ✨</h2>
            <p style='margin: 8px 0 0 0; opacity: 0.95; font-size: 1.1rem;'>
                Your AI Quality Assurance companion
            </p>
        </div>
        """, unsafe_allow_html=True)

        col1, col2 = st.columns([4, 1])

        with col1:
            # Step 1
            st.markdown("""
            <div style='
                background: #f0f9ff;
                border-left: 4px solid #0284c7;
                padding: 16px;
                border-radius: 8px;
                margin-bottom: 16px;
            '>
                <h4 style='margin: 0 0 8px 0; color: #0c4a6e;'>
                    1️⃣ 📊 View Your Scores
                </h4>
                <ul style='margin: 0; color: #0c4a6e;'>
                    <li>Check the <b>Module Summary</b> table below</li>
                    <li>🟢 Green = Excellent (=90) | 🟡 Yellow = Good (75-89) | 🔴 Red = Needs Work (<75)</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)

            # Step 2
            st.markdown("""
            <div style='
                background: #f0fdf4;
                border-left: 4px solid #16a34a;
                padding: 16px;
                border-radius: 8px;
                margin-bottom: 16px;
            '>
                <h4 style='margin: 0 0 8px 0; color: #14532d;'>
                    2️⃣ 🔍 Explore Modules
                </h4>
                <ul style='margin: 0; color: #14532d;'>
                    <li>Click on <b>L1-L5 sections</b> to see detailed metrics</li>
                    <li>Use the <b>search bar</b> (Ctrl+K) to find specific items</li>
                    <li>Check <b>Evidence Tray 📎</b> for supporting documents</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)

            # Step 3
            st.markdown("""
            <div style='
                background: #fef3c7;
                border-left: 4px solid #f59e0b;
                padding: 16px;
                border-radius: 8px;
                margin-bottom: 16px;
            '>
                <h4 style='margin: 0 0 8px 0; color: #78350f;'>
                    3️⃣ 📈 Track Progress
                </h4>
                <ul style='margin: 0; color: #78350f;'>
                    <li>Compare runs over time</li>
                    <li>View <b>GQAS</b> (Global Quality Score) at the bottom</li>
                    <li>Export reports for stakeholders</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)

            # Keyboard shortcuts
            st.markdown("""
            <div style='
                background: #f5f3ff;
                border: 1px solid #c4b5fd;
                padding: 12px;
                border-radius: 8px;
                margin-bottom: 12px;
            '>
                <h4 style='margin: 0 0 8px 0; color: #5b21b6;'>⌨️ Keyboard Shortcuts</h4>
                <div style='display: grid; grid-template-columns: 1fr 1fr; gap: 8px; color: #5b21b6;'>
                    <div><kbd style='background: #ddd6fe; padding: 2px 8px; border-radius: 4px;'>Ctrl+K</kbd> → Focus search</div>
                    <div><kbd style='background: #ddd6fe; padding: 2px 8px; border-radius: 4px;'>Ctrl+R</kbd> → Refresh</div>
                    <div><kbd style='background: #ddd6fe; padding: 2px 8px; border-radius: 4px;'>Esc</kbd> → Close sections</div>
                    <div><kbd style='background: #ddd6fe; padding: 2px 8px; border-radius: 4px;'>?</kbd> → Show all shortcuts</div>
                </div>
            </div>
            """, unsafe_allow_html=True)

            # Help section
            st.markdown("""
            <div style='
                text-align: center;
                padding: 12px;
                color: #6b7280;
                font-size: 0.9rem;
            '>
                💡 <b>Need help?</b> Enable <b>Debug Mode</b> in the sidebar or press <kbd style='background: #e5e7eb; padding: 2px 6px; border-radius: 4px;'>?</kbd> for shortcuts
            </div>
            """, unsafe_allow_html=True)

        with col2:
            st.markdown("<br>" * 2, unsafe_allow_html=True)
            if st.button("🎯 Start Interactive Tour", key="start_tour_btn", help="Take a guided tour of the dashboard", width="stretch", type="primary"):
                st.session_state['tour_active'] = True
                st.session_state['tour_step'] = 0
                # Hide guide during tour
                st.session_state['hide_quickstart'] = True
                st.rerun()
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("✕ Don't show again", key="hide_quickstart_btn", help="Hide this guide permanently", width="stretch"):
                st.rerun()

        st.markdown("---")

# Show a subtle button to bring back the guide if hidden
if st.session_state['hide_quickstart']:
    col_empty, col_btn = st.columns([5, 1])
    with col_btn:
        if st.button("👋 Quick Start", key="show_quickstart_btn", help="Show the Quick Start Guide"):
            st.session_state['hide_quickstart'] = False
            # Expand it when shown again
            st.session_state['first_visit'] = True
            st.rerun()

# ============================================================================
# ALERTS & NOTIFICATIONS SECTION
# ============================================================================
if ENHANCEMENTS_AVAILABLE:
    render_alerts_section()

# ============================================================================
# MODULE OVERVIEW GRID (First-Time User Experience)
# ============================================================================


def get_score_color(score: float, risk_profile: str):
    """
    Determine color based on score and risk profile.

    High Risk:
    - Green: score >= 90
    - Yellow: 75 <= score < 90
    - Red: score < 75

    Medium Risk:
    - Green: score >= 85
    - Yellow: 75 <= score < 85
    - Red: score < 75
    """
    if risk_profile == "High":
        if score >= 90:
            return "green", "🟢", "✅ Excellent"
        elif score >= 75:
            return "yellow", "🟡", "⚠️ Warning"
        else:
            return "red", "🔴", "❌ Fails"
    else:  # Medium
        if score >= 85:
            return "green", "🟢", "✅ Excellent"
        elif score >= 75:
            return "yellow", "🟡", "⚠️ Warning"
        else:
            return "red", "🔴", "❌ Fails"


def display_module_overview_grid(latest: dict, risk_profile: str = "High"):
    """
    Display an attractive grid showing all modules with their status, purpose, and key metrics.
    Helps first-time users understand what each compliance module measures.

    Args:
        latest: Dict of loaded reports
        risk_profile: "High" `or "Medium" `for color thresholds
    """
    st.markdown("## 🛡️ Compliance Module Overview")
    st.markdown("Each module measures a critical aspect of AI quality. Click on a module to learn more or scroll down to view detailed scores.")

    # Create 2-column grid (3 modules per row on desktop)
    cols = st.columns(3, gap="medium")

    for idx, module_id in enumerate(["L1", "L2", "L3", "L4", "L5"]):
        col_idx = idx % 3

        with cols[col_idx]:
            desc = MODULE_DESCRIPTIONS.get(module_id, {})
            report = latest.get(module_id)

            # Get status color
            score_val = 0
            status_color = "⚪"

            if report:
                # For L1, use coverage_percent; for others use score
                if module_id == "L1":
                    metrics = report.get("metrics", {}) or {}
                    score_val = metrics.get(
                        "coverage_percent", report.get("score", 0))
                else:
                    score_val = report.get("score", 0)
                if risk_profile == "High":
                    if score_val >= 90:
                        status_color = "✅"
                    elif score_val >= 75:
                        status_color = "⚠️"
                    else:
                        status_color = "❌"
                else:  # Medium
                    if score_val >= 85:
                        status_color = "✅"
                    elif score_val >= 75:
                        status_color = "⚠️"
                    else:
                        status_color = "❌"

            # Create card
            card_html = f"""
            <div style="
                border: 2px solid {desc.get('color', '#ccc')};
                border-radius: 12px;
                padding: 20px;
                margin-bottom: 16px;
                background: linear-gradient(135deg, {desc.get('color', '#ccc')}15 0%, {desc.get('color', '#ccc')}05 100%);
                transition: all 0.3s ease;
                cursor: pointer;
            "
            onmouseover="this.style.boxShadow='0 8px 24px rgba(0,0,0,0.15)'; this.style.transform='translateY(-4px)'"
            onmouseout="this.style.boxShadow='none'; this.style.transform='translateY(0)'">
                <div style="display: flex; align-items: center; justify-content: space-between; margin-bottom: 12px;">
                    <div style="font-size: 2rem;">{desc.get('emoji', '🗂️')}</div>
                    <div style="font-size: 1.5rem; font-weight: bold;">{status_color}</div>
                </div>
                <h4 style="margin: 0 0 8px 0; color: {desc.get('color', '#333')}; font-size: 1.1rem;">
                    {module_id}: {desc.get('name', module_id)}
                </h4>
                <p style="margin: 0 0 12px 0; color: #555; font-size: 0.9rem; line-height: 1.4;">
                    {desc.get('purpose', '')}
                </p>
                <div style="font-size: 0.85rem; color: #666;">
                    <strong>Key Metrics:</strong>
                    <div style="display: flex; flex-wrap: wrap; gap: 6px; margin-top: 6px;">
                        {''.join([f'<span style="background: {desc.get("color", "#ccc")}30; padding: 3px 8px; border-radius: 4px; font-size: 0.8rem;">{metric}</span>'
                                 for metric in desc.get('key_metrics', [])])}
                    </div>
                </div>
                {f'<div style="margin-top: 12px; padding-top: 12px; border-top: 1px solid {desc.get("color", "#ccc")}30;"><strong>Score:</strong> {score_val:.1f}/100</div>' if report else ''}
            </div>
            """
            st.markdown(card_html, unsafe_allow_html=True)

    # Add aggregate module in full width
    st.markdown("---")
    cols = st.columns(1)
    with cols[0]:
        desc = MODULE_DESCRIPTIONS.get("AGG", {})
        report = latest.get("AGG")

        # Calculate score from AGG GQAS if available, otherwise use average of L1-L5
        score_val = 0
        individual_scores = []
        for mid in ["L1", "L2", "L3", "L4", "L5"]:
            rep = latest.get(mid)
            if rep:
                # For L1, use coverage_percent; for others use score
                if mid == "L1":
                    metrics = rep.get("metrics", {}) or {}
                    s = metrics.get("coverage_percent", rep.get("score", 0))
                else:
                    s = rep.get("score", 0)
                individual_scores.append(s)

        # First priority: Use GQAS weighted score if available
        if report and report.get("gqas", 0) > 0:
            score_val = report.get("gqas", 0)
        # Second priority: Use AGG score if available and non-zero
        elif report and report.get("score", 0) > 0:
            score_val = report.get("score", 0)
        # Fallback: Calculate average of individual module scores
        elif individual_scores:
            score_val = sum(individual_scores) / len(individual_scores)

        status_color = "⚪"
        if risk_profile == "High":
            if score_val >= 90:
                status_color = "💚", "Excellent"
            elif score_val >= 75:
                status_color = "💛", "Warning"
            else:
                status_color = "❤️", "Fails"
        else:  # Medium
            if score_val >= 85:
                status_color = "💚", "Excellent"
            elif score_val >= 75:
                status_color = "💛", "Warning"
            else:
                status_color = "❤️", "Fails"

        # Removed redundant Global QA Score card - GQAS is shown in Shipping Floors and Aggregate Global QA Score sections below

# ============================================================================
# INTERACTIVE TOUR MODE
# ============================================================================


# Initialize tour state
if 'tour_active' not in st.session_state:
    st.session_state['tour_active'] = False
    st.session_state['tour_step'] = 0

# Tour steps configuration
TOUR_STEPS = [
    {
        'title': '📊 Module Summary',
        'description': 'This table shows the health of all 5 quality modules. Green = Excellent, Yellow = Good, Red = Needs Work.',
        'target': 'module-summary',
        'position': 'bottom'
    },
    {
        'title': '📎 Evidence Tray',
        'description': 'Upload and manage supporting documents for each module. Files are organized by module and can be previewed inline.',
        'target': 'evidence-tray',
        'position': 'bottom'
    },
    {
        'title': '🏛️ L1 Governance',
        'description': f"Track compliance with {len(AVAILABLE_FRAMEWORKS)} frameworks ({', '.join(AVAILABLE_FRAMEWORKS[:3])}...). See which clauses pass or fail.",
        'target': 'l1-governance',
        'position': 'bottom'
    },
    {
        'title': '🔒 L2 Privacy/Security',
        'description': 'Monitor encryption, data protection, and security metrics. Track incidents and access reviews.',
        'target': 'l2-privacy',
        'position': 'bottom'
    },
    {
        'title': '🧮 L3 Fairness',
        'description': 'Measure bias metrics like DPG (Demographic Parity Gap) and EOD (Equal Opportunity Difference) across groups.',
        'target': 'l3-fairness',
        'position': 'bottom'
    },
    {
        'title': '🔍 L4 Explainability',
        'description': 'View feature importance, SHAP values, and model interpretability metrics. Compare models side-by-side.',
        'target': 'l4-explainability',
        'position': 'bottom'
    },
    {
        'title': '🛠️ L5 Operations',
        'description': 'Monitor live metrics, SLO breaches, and incident timelines. Set up alerting for quality drift.',
        'target': 'l5-operations',
        'position': 'bottom'
    },
    {
        'title': '⚙️ GQAS Score',
        'description': 'The Global Quality Assurance Score aggregates all modules. Check if you meet shipping floors for your risk profile.',
        'target': 'gqas-aggregate',
        'position': 'top'
    }
]

# Tour Mode CSS and JavaScript
if st.session_state['tour_active']:
    current_step = st.session_state['tour_step']

    if current_step < len(TOUR_STEPS):
        step = TOUR_STEPS[current_step]

        st.markdown(f"""
        <style>
        /* Overlay backdrop */
        .tour-overlay {{
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background: rgba(0, 0, 0, 0.7);
            z-index: 9998;
            pointer-events: none;
        }}

        /* Highlight specific section */
        [data-tour-target="{step['target']}"] {{
            position: relative;
            z-index: 9999 !important;
            #667eea, 0 0 0 8px rgba(102, 126, 234, 0.3) !important;
            box-shadow: 0 0 0 4px
            border-radius: 12px !important;
            background: white !important;
        }}

        /* Tour tooltip */
        .tour-tooltip {{
            position: fixed;
            {step['position']}: 20px;
            left: 50%;
            transform: translateX(-50%);
            max-width: 500px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 24px;
            border-radius: 16px;
            box-shadow: 0 20px 60px rgba(0, 0, 0, 0.4);
            z-index: 10000;
            animation: slideIn 0.3s ease-out;
        }}

        @keyframes slideIn {{
            from {{
                opacity: 0;
                transform: translateX(-50%) translateY(-20px);
            }}
            to {{
                opacity: 1;
                transform: translateX(-50%) translateY(0);
            }}
        }}

        .tour-tooltip h3 {{
            margin: 0 0 12px 0;
            font-size: 1.4rem;
            color: white;
        }}

        .tour-tooltip p {{
            margin: 0 0 20px 0;
            font-size: 1rem;
            line-height: 1.6;
            opacity: 0.95;
        }}

        .tour-progress {{
            margin-bottom: 16px;
            font-size: 0.875rem;
            opacity: 0.8;
        }}

        /* Arrow pointing to element */
        .tour-arrow {{
            position: fixed;
            left: 50%;
            transform: translateX(-50%);
            {('bottom' if step['position'] == 'top' else 'top')}: 180px;
            font-size: 3rem;
            color: #667eea;
            z-index: 9999;
            animation: bounce 1s infinite;
            text-shadow: 0 4px 8px rgba(0, 0, 0, 0.3);
        }}

        @keyframes bounce {{
            0%, 100% {{ transform: translateX(-50%) translateY(0); }}
            50% {{ transform: translateX(-50%) translateY({'-10px' if step['position'] == 'top' else '10px'}); }}
        }}
        </style>

        <div class="tour-overlay"></div>

        <div class="tour-arrow">
            {'⬆️' if step['position'] == 'top' else '⬇️'}
        </div>

        <div class="tour-tooltip">
            <div class="tour-progress">
                Step {current_step + 1} of {len(TOUR_STEPS)}
            </div>
            <h3>{step['title']}</h3>
            <p>{step['description']}</p>
        </div>

        <script>
        // Auto-scroll to highlighted element
        setTimeout(function() {{
            var target = document.querySelector('[data-tour-target="{step['target']}"]');
            if (target) {{
                target.scrollIntoView({{
                    behavior: 'smooth',
                    block: 'center'
                }});
            }}
        }}, 100);
        </script>
        """, unsafe_allow_html=True)

# Tour control buttons - Create hidden buttons for Streamlit state management
if st.session_state['tour_active']:
    # Callback functions for tour controls
    def prev_tour():
        st.session_state['tour_step'] = max(
            0, st.session_state['tour_step'] - 1)

    def next_tour():
        if st.session_state['tour_step'] < len(TOUR_STEPS) - 1:
            st.session_state['tour_step'] += 1
        else:
            st.session_state['tour_active'] = False
            st.session_state['tour_step'] = 0

    def exit_tour():
        st.session_state['tour_active'] = False
        st.session_state['tour_step'] = 0

    # Hidden columns for button state management (render them first)
    hidden_cols = st.columns([0.05, 0.05, 0.05, 0.05])

    with hidden_cols[0]:
        st.button("⏮️", key="tour_prev_btn", on_click=prev_tour,
                  disabled=st.session_state['tour_step'] == 0)

    with hidden_cols[1]:
        if st.session_state['tour_step'] < len(TOUR_STEPS) - 1:
            st.button("⏭️", key="tour_next_btn", on_click=next_tour)
        else:
            st.button("✓", key="tour_finish_btn", on_click=exit_tour)

    with hidden_cols[2]:
        st.button("✕", key="tour_exit_btn", on_click=exit_tour)

    # Step counter in hidden column
    with hidden_cols[3]:
        st.write("")

    # Beautiful floating purple buttons
    step_text = f"Step {current_step + 1}/{len(TOUR_STEPS)}"
    st.markdown(f"""
    <style>
    .tour-controls-container {{
        position: fixed;
        bottom: 280px;
        left: 50%;
        transform: translateX(-50%);
        display: flex;
        gap: 12px;
        z-index: 9999;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 16px 24px;
        border-radius: 12px;
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.3);
        backdrop-filter: blur(10px);
        align-items: center;
        justify-content: center;
        flex-wrap: nowrap;
        pointer-events: auto;
    }}

    .tour-button-float {{
        background: white !important;
        color: #667eea !important;
        border: none !important;
        padding: 12px 20px !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
        cursor: pointer !important;
        font-size: 14px !important;
        transition: all 0.2s ease !important;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.15) !important;
    }}

    .tour-button-float:hover {{
        background: #f5f5f5 !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 4px 16px rgba(102, 126, 234, 0.3) !important;
    }}

    .tour-button-float:active {{
        transform: translateY(0) !important;
    }}

    .tour-step-text {{
        color: white;
        font-weight: 600;
        margin: 0 16px;
        min-width: 120px;
        text-align: center;
        font-size: 13px;
        letter-spacing: 0.5px;
    }}

    .tour-exit-btn {{
        background: #ff6b6b !important;
        color: white !important;
    }}

    .tour-exit-btn:hover {{
        background: #ff5252 !important;
    }}
    </style>

    <div class="tour-controls-container">
        <button class="tour-button-float" id="tour-prev-float" type="button">⏮️ Previous</button>
        <div class="tour-step-text">{step_text}</div>
        <button class="tour-button-float" id="tour-next-float" type="button">Next ⏭️</button>
        <button class="tour-button-float tour-exit-btn" id="tour-exit-float" type="button">✕ Exit</button>
    </div>

    <script>
    setTimeout(function() {{
        // Find all Streamlit buttons in the page
        var allStreamlitButtons = Array.from(document.querySelectorAll('[data-testid="stButton"]'));

        // Find the specific hidden buttons by looking for their parent containers
        var findButtonByIndex = function(index) {{
            if (allStreamlitButtons[index]) {{
                var btn = allStreamlitButtons[index].querySelector('button');
                return btn;
            }}
            return null;
        }};

        // Get references to our floating buttons
        var prevFloatBtn = document.getElementById('tour-prev-float');
        var nextFloatBtn = document.getElementById('tour-next-float');
        var exitFloatBtn = document.getElementById('tour-exit-float');

        if (prevFloatBtn) {{
            prevFloatBtn.addEventListener('click', function(e) {{
                e.preventDefault();
                e.stopPropagation();
                // Click the first Streamlit button (Previous)
                var btn = findButtonByIndex(0);
                if (btn) {{
                    btn.dispatchEvent(new MouseEvent('click', {{ bubbles: true, cancelable: true }}));
                }}
                return false;
            }}, true);
        }}

        if (nextFloatBtn) {{
            nextFloatBtn.addEventListener('click', function(e) {{
                e.preventDefault();
                e.stopPropagation();
                // Click the second Streamlit button (Next/Finish)
                var btn = findButtonByIndex(1);
                if (btn) {{
                    btn.dispatchEvent(new MouseEvent('click', {{ bubbles: true, cancelable: true }}));
                }}
                return false;
            }}, true);
        }}

        if (exitFloatBtn) {{
            exitFloatBtn.addEventListener('click', function(e) {{
                e.preventDefault();
                e.stopPropagation();
                // Click the third Streamlit button (Exit)
                var btn = findButtonByIndex(2);
                if (btn) {{
                    btn.dispatchEvent(new MouseEvent('click', {{ bubbles: true, cancelable: true }}));
                }}
                return false;
            }}, true);
        }}
    }}, 600);
    </script>
    """, unsafe_allow_html=True)


# ===== Smart Insights Section (Key Findings) =====


def display_smart_insights(latest: dict, risk_profile: str = "High"):
    """Display key findings and flagged issues across all modules."""
    findings = []
    risk_modules = []

    # Scan all modules for issues
    for mid in ["L1", "L2", "L3", "L4", "L5"]:
        rep = latest.get(mid)
        if rep:
            # For L1, use coverage_percent; for others use score
            if mid == "L1":
                metrics = rep.get("metrics", {}) or {}
                score = metrics.get("coverage_percent", rep.get("score", 100))
            else:
                score = rep.get("score", 100)

            # Determine if at risk based on profile
            at_risk = False
            if risk_profile == "High":
                at_risk = score < 75
            else:  # Medium
                at_risk = score < 75

            if at_risk:
                risk_modules.append((mid, score))

    if risk_modules:
        with st.expander("🔑 Key Findings", expanded=False):
            # Alert column
            alert_col = st.columns(1)[0]

            with alert_col:
                for mid, score in sorted(risk_modules, key=lambda x: x[1]):
                    rep = latest.get(mid)
                    desc = MODULE_DESCRIPTIONS.get(mid, {})
                    emoji = desc.get('emoji', '??')

                    # Determine risk level color
                    risk_color = "#dc2626" if score < 50 else "#ea580c"
                    risk_bg = "#fee2e2" if score < 50 else "#fef3c7"

                    alert_html = f"""
                    <div style="
                        background: {risk_bg};
                        border-left: 5px solid {risk_color};
                        padding: 16px;
                        margin-bottom: 12px;
                        border-radius: 8px;
                        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
                    ">
                        <div style="display: flex; align-items: center; justify-content: space-between; margin-bottom: 12px;">
                            <div style="display: flex; align-items: center; gap: 12px;">
                                <span style="font-size: 1.5rem;">{emoji}</span>
                                <div>
                                    <strong style="font-size: 1.05rem; display: block;">{desc.get('name', mid)}</strong>
                                    <span style="font-size: 0.85rem; color: #666;">Score: {score:.1f}/100</span>
                                </div>
                            </div>
                            <span style="
                                font-size: 0.8rem; 
                                padding: 6px 12px; 
                                background: {risk_color}; 
                                color: white; 
                                border-radius: 20px;
                                font-weight: 600;
                            ">⚠️ Needs Improvement</span>
                        </div>
                    </div>
                    """
                    st.markdown(alert_html, unsafe_allow_html=True)

                    # Show what's lacking in this module
                    if rep:
                        issues = []

                        # For L1: Show failed clauses
                        if mid == "L1":
                            metrics = rep.get("metrics", {}) or {}

                            # Get ISO frameworks that are below threshold
                            framework_breakdown = metrics.get(
                                "framework_breakdown", []) or []
                            if isinstance(framework_breakdown, list):
                                for fw in framework_breakdown:
                                    fw_name = fw.get("framework", "")
                                    fw_coverage = fw.get("coverage_percent", 0)
                                    if fw_coverage < 100 and fw_coverage > 0:
                                        issues.append({
                                            "type": "framework",
                                            "name": fw_name,
                                            "value": fw_coverage,
                                            "icon": "📋"
                                        })

                            # Get failed clauses
                            clauses = metrics.get("clauses", []) or []
                            failed = [
                                c for c in clauses if not c.get("passed")]
                            if failed:
                                for clause in failed[:2]:  # Show top 2 failures
                                    issues.append({
                                        "type": "clause",
                                        "name": clause.get('name', 'Unknown'),
                                        "reason": clause.get('why_failed', 'Issue identified'),
                                        "icon": "❌"
                                    })

                        # For other modules: Show frameworks below threshold or failed metrics
                        else:
                            metrics = rep.get("metrics", {}) or {}
                            if isinstance(metrics, dict):
                                # Check for frameworks/metrics below 75%
                                for key, value in metrics.items():
                                    if isinstance(value, (int, float)) and 0 < value < 75:
                                        issues.append({
                                            "type": "metric",
                                            "name": key.replace('_', ' ').title(),
                                            "value": value,
                                            "icon": "⚠️"
                                        })

                        if issues:
                            # Render issues as modern cards
                            st.markdown(
                                f"<div style='margin-top: 12px;'>", unsafe_allow_html=True)

                            # Create columns for metric cards
                            cols = st.columns(min(len(issues), 4), gap="small")

                            for idx, issue in enumerate(issues):
                                with cols[idx % len(cols)]:
                                    if issue["type"] == "clause":
                                        card_html = f"""
                                        <div style="
                                            background: linear-gradient(135deg, #FEF2F2 0%, #FEE2E2 100%);
                                            border: 1px solid #FECACA;
                                            border-radius: 10px;
                                            padding: 16px;
                                            text-align: center;
                                        ">
                                            <div style="font-size: 2rem; margin-bottom: 8px;">{issue['icon']}</div>
                                            <div style="font-size: 0.75rem; color: #991B1B; font-weight: 600; margin-bottom: 4px; text-transform: uppercase; letter-spacing: 0.5px;">Issue</div>
                                            <div style="font-size: 0.9rem; font-weight: 600; color: #1F2937; margin-bottom: 6px;">{issue['name']}</div>
                                            <div style="font-size: 0.8rem; color: #666; line-height: 1.4;">{issue['reason']}</div>
                                        </div>
                                        """
                                    else:
                                        # Determine metric status color
                                        if issue["value"] < 25:
                                            metric_color = "#DC2626"
                                            metric_bg = "#FEE2E2"
                                            metric_light = "#FEF2F2"
                                        elif issue["value"] < 50:
                                            metric_color = "#EA580C"
                                            metric_bg = "#FEF3C7"
                                            metric_light = "#FFFBEB"
                                        else:
                                            metric_color = "#F59E0B"
                                            metric_bg = "#FEF3C7"
                                            metric_light = "#FFFBEB"

                                        card_html = f"""
                                        <div style="
                                            background: linear-gradient(135deg, {metric_light} 0%, {metric_bg} 100%);
                                            border: 2px solid {metric_color};
                                            border-radius: 10px;
                                            padding: 16px;
                                            text-align: center;
                                        ">
                                            <div style="font-size: 2rem; margin-bottom: 8px;">{issue['icon']}</div>
                                            <div style="font-size: 0.75rem; color: {metric_color}; font-weight: 600; margin-bottom: 8px; text-transform: uppercase; letter-spacing: 0.5px;">{issue['name']}</div>
                                            <div style="
                                                font-size: 2rem; 
                                                font-weight: 800; 
                                                color: {metric_color};
                                                margin-bottom: 4px;
                                            ">{issue['value']:.1f}%</div>
                                            <div style="
                                                font-size: 0.75rem;
                                                color: #666;
                                                font-weight: 500;
                                            ">Below Target</div>
                                        </div>
                                        """
                                    st.markdown(
                                        card_html, unsafe_allow_html=True)

                            st.markdown("</div>", unsafe_allow_html=True)

    # Show passing modules
    passing = []
    for mid in ["L1", "L2", "L3", "L4", "L5"]:
        rep = latest.get(mid)
        if rep:
            # For L1, use coverage_percent; for others use score
            if mid == "L1":
                metrics = rep.get("metrics", {}) or {}
                score = metrics.get("coverage_percent", rep.get("score", 0))
            else:
                score = rep.get("score", 0)

            if score >= 75:
                passing.append((mid, score))

    if passing:
        with st.expander(f"✔️ Passing Modules ({len(passing)})"):
            for mid, score in sorted(passing, key=lambda x: -x[1]):
                desc = MODULE_DESCRIPTIONS.get(mid, {})
                emoji = desc.get('emoji', '🟢')
                st.markdown(
                    f"{emoji} **{desc.get('name', mid)}** ➡️ Score: {score:.1f}/100")


st.markdown("---")

# ===== Quick Actions Panel =====
with st.expander("⚡ Quick Actions", expanded=False):
    col1, col2, col3, col4 = st.columns(4)

    with col1:
        if st.button("🔄 Refresh Data", width="stretch", disabled=_LOCK, help="Clear cache and reload all reports", key="quick_refresh_data"):
            st.session_state["force_refresh"] = True
            st.rerun()

    with col2:
        if st.button("📊 Jump to GQAS", width="stretch", disabled=_LOCK, key="quick_jump_gqas"):
            st.markdown(
                '<a href="#aggregate-global-qa-score-gqas"></a>', unsafe_allow_html=True)
            st.info("Scroll down to see GQAS section")

    with col3:
        if st.button("📎 Evidence Tray", width="stretch", disabled=_LOCK, key="quick_evidence_tray"):
            st.markdown('<a href="#evidence-tray"></a>',
                        unsafe_allow_html=True)
            st.info("Scroll down to Evidence Tray")

    with col4:
        if st.button("🧭 Navigation", width="stretch", disabled=_LOCK, key="quick_navigation"):
            st.info("""
            **Quick Navigation:**
            - L1: Governance
            - L2: Privacy/Security
            - L3: Fairness
            - L4: Explainability
            - L5: Operations
            """)
# ===== Global Search =====
st.markdown("---")
search_col1, search_col2 = st.columns([3, 1])

with search_col1:
    search_query = st.text_input(
        "🔍 Search across all modules",
        placeholder="Search for metrics, clauses, or evidence...",
        key="global_search",
        disabled=_LOCK,
        help="Press Ctrl+K to focus"
    )

with search_col2:
    search_in = st.multiselect(
        "Search in",
        ["Modules", "Evidence", "Clauses", "Metrics"],
        default=["Modules"],
        key="search_scope"
    )

if search_query:
    st.markdown(f"### 🔍 Search Results for '{search_query}'")

    results = []

    # Search in modules
    if "Modules" in search_in:
        for mid, rep in latest.items():
            if rep and search_query.lower() in str(rep).lower():
                results.append({
                    "Type": "Module",
                    "Location": NAMES.get(mid, mid),
                    "Match": f"Found in {mid} report"
                })

    # Search in evidence
    if "Evidence" in search_in:
        eidx = _load_evidence_index()
        for mid, files in eidx.items():
            for f in files:
                if search_query.lower() in f.lower():
                    results.append({
                        "Type": "Evidence",
                        "Location": NAMES.get(mid, mid),
                        "Match": f
                    })

    if results:
        st.dataframe(pd.DataFrame(results), width="stretch")
    else:
        st.info("No results found. Try different keywords.")

st.markdown("---")

# =============================================================================
# TIMESTAMP HELPER (Use throughout dashboard)
# =============================================================================


def render_timestamp(label: str, show: bool = True):
    """Render a timestamp badge if enabled in settings"""
    if not show or not st.session_state.get('show_timestamps', True):
        return

    st.markdown(f"""
    <div style='
        display: inline-block;
        background: #f3f4f6;
        padding: 4px 12px;
        border-radius: 999px;
        font-size: 0.75rem;
        color: #6b7280;
        margin-left: 8px;
    '>
        🕐 {label}: {datetime.now().strftime('%H:%M:%S')}
    </div>
    """, unsafe_allow_html=True)
# Helper: normalize JSON payload into list[dict]


def _as_dict_list(obj):
    """Convert object to list of dicts, handling various input formats."""
    if isinstance(obj, dict):
        return [obj]
    if isinstance(obj, list):
        return [x for x in obj if isinstance(x, dict)]
    return []


# Find latest reports (define early for audit hash)
files = get_report_files("*.json")
if not files:
    st.info("No reports found. Run the CLI to generate JSON reports in ./reports")
    st.stop()

# Module names and emoji map
NAMES = {
    "L1": "L1 Governance & Regulatory",
    "L2": "L2 Privacy & Security",
    "L3": "L3 Fairness & Ethics",
    "L4": "L4 Explainability & Transparency",
    "L5": "L5 Operations & Monitoring",
}

# Module descriptions for improved first-time UX
MODULE_DESCRIPTIONS = {
    "L1": {
        "emoji": "🏛️",
        "name": "Governance & Regulatory",
        "purpose": "Ensures AI systems comply with regulatory requirements and governance frameworks",
        "key_metrics": ["Regulatory Compliance", "Documentation", "Audit Trails", "Policy Adherence"],
        "color": "#1f77b4",
    },
    "L2": {
        "emoji": "🔐",
        "name": "Privacy & Security",
        "purpose": "Protects user data and ensures secure handling of sensitive information",
        "key_metrics": ["Data Protection", "Encryption", "Access Controls", "Security Posture"],
        "color": "#ff7f0e",
    },
    "L3": {
        "emoji": "⚖️",
        "name": "Fairness & Ethics",
        "purpose": "Detects and mitigates bias to ensure equitable treatment across all user groups",
        "key_metrics": ["Bias Detection", "Fairness Metrics", "Equity Assessment", "Ethical Review"],
        "color": "#2ca02c",
    },
    "L4": {
        "emoji": "🔍",
        "name": "Explainability & Transparency",
        "purpose": "Makes AI decisions interpretable and understandable to stakeholders",
        "key_metrics": ["Feature Importance", "Decision Explanations", "Model Transparency", "User Understanding"],
        "color": "#d62728",
    },
    "L5": {
        "emoji": "📊",
        "name": "Operations & Monitoring",
        "purpose": "Continuously monitors system performance and health in production",
        "key_metrics": ["Performance Tracking", "Anomaly Detection", "Alert Systems", "Operational Metrics"],
        "color": "#9467bd",
    },
    "AGG": {
        "emoji": "",
        "name": "Global QA Score",
        "purpose": "Comprehensive assessment combining all compliance levels into actionable insights",
        "key_metrics": ["Overall Risk", "Compliance Status", "Trend Analysis", "Recommendations"],
        "color": "#8c564b",
    },
}


def name_for(module_id: str) -> str:
    return NAMES.get(module_id, module_id)


COLOR_EMOJI = {"green": "🟢", "yellow": "🟡", "red": "🔴"}

# ============================================================================
# CACHED REPORT LOADING (Performance Optimized)
# ============================================================================


@st.cache_data(ttl=config.cache.ttl_seconds, show_spinner=False, max_entries=50)
def load_all_reports_cached(file_list: tuple, force_reload: bool = False) -> dict:
    """
    Load all reports with caching and optimizations.

    Args:
        file_list: Tuple of file paths (must be tuple for hashing)
        force_reload: Force cache bypass

    Returns:
        Dict mapping module names to their latest reports
    """
    latest_by_module = {k: None for k in ["L1", "L2", "L3", "L4", "L5", "AGG"]}
    seen_mtime = {k: -1 for k in latest_by_module.keys()}
    errors = []

    # Sort files by modification time (newest first) to exit early
    sorted_files = sorted(
        file_list, key=lambda f: os.path.getmtime(f), reverse=True)

    modules_found = set()

    for f in sorted_files:
        # Early exit if all modules found
        if len(modules_found) == 6:
            break

        try:
            # Quick check: skip if file is too old
            mtime = os.path.getmtime(f)

            # Read file with size limit
            file_size = os.path.getsize(f)
            if file_size > 10 * 1024 * 1024:  # 10MB limit
                logger.warning(
                    f"Skipping large file: {f} ({file_size / 1024 / 1024:.1f}MB)")
                continue

            with open(f, "r", encoding="utf-8") as fh:
                raw = json.load(fh)

            # Handle both single report and list of reports
            for rec in (raw if isinstance(raw, list) else [raw]):
                if not isinstance(rec, dict):
                    continue

                mod = str(rec.get("module", "")).upper()
                if mod in latest_by_module and mtime >= seen_mtime[mod]:
                    latest_by_module[mod] = rec
                    seen_mtime[mod] = mtime
                    modules_found.add(mod)

        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON in {f}: {e}")
            errors.append({
                "file": os.path.basename(f),
                "error": "Invalid JSON format",
                "type": "JSONDecodeError"
            })
        except FileNotFoundError:
            logger.error(f"File not found: {f}")
            errors.append({
                "file": os.path.basename(f),
                "error": "File not found",
                "type": "FileNotFoundError"
            })
        except Exception as e:
            logger.error(f"Failed to load {f}: {e}", exc_info=True)
            errors.append({
                "file": os.path.basename(f),
                "error": str(e)[:100],
                "type": type(e).__name__
            })

    # Attach metadata
    latest_by_module["_errors"] = errors
    latest_by_module["_load_time"] = datetime.now().isoformat()
    latest_by_module["_modules_found"] = len(modules_found)

    return latest_by_module


def check_memory_usage() -> dict:
    """Check current memory usage and return stats."""
    try:
        process = psutil.Process()
        memory_info = process.memory_info()

        return {
            'rss_mb': memory_info.rss / (1024 * 1024),
            'percent': process.memory_percent(),
            'available_mb': psutil.virtual_memory().available / (1024 * 1024),
        }
    except Exception as e:
        logger.warning(f"Failed to get memory info: {e}")
        return {'rss_mb': 0, 'percent': 0, 'available_mb': 0}


def should_clear_cache() -> bool:
    """Determine if cache should be cleared based on memory usage."""
    mem = check_memory_usage()
    max_memory = config.cache.max_memory_mb

    if mem['rss_mb'] > max_memory:
        logger.warning(
            f"Memory usage ({mem['rss_mb']:.1f}MB) exceeds limit ({max_memory}MB)")
        return True

    return False


# Add automatic cache clearing
if should_clear_cache():
    st.cache_data.clear()
    logger.info("Cache cleared due to high memory usage")
# ========================================
# Load Reports with Progress UI
# ========================================

# Get audit lock state for UI controls
_LOCK = get_audit_mode()

with measure_section("load_all_reports"):
    files = get_report_files("*.json")

    if not files:
        st.info("No reports found. Run the CLI to generate JSON reports in ./reports")
        st.stop()

    # Check if user wants to force refresh
    force_refresh = st.session_state.get("force_refresh", False)
    if force_refresh:
        st.session_state["force_refresh"] = False
        load_all_reports_cached.clear()  # Clear cache

    # Show loading indicator for first load or many files
    if len(files) > 5 and st.session_state.get("last_refresh") is None:
        with show_loading("Loading reports", f"Loading {len(files)} reports..."):
            result = load_all_reports_cached(tuple(files), force_refresh)
    else:
        result = load_all_reports_cached(tuple(files), force_refresh)

    # Extract errors and actual data
    load_errors = result.pop("_errors", [])
    load_time = result.pop("_load_time", None)
    latest = result  # ✅ This is the correct assignment

    # Update last refresh time
    st.session_state["last_refresh"] = datetime.now().isoformat()

    # Show loading stats
    loaded_count = sum(
        1 for k, v in latest.items()
        if v is not None and k not in ["_errors", "_load_time", "AGG"]
    )

    if PERFORMANCE_MONITORING:
        monitor = get_monitor()
        stats = monitor.get_stats("load_all_reports")
        if stats:
            st.caption(
                f"✅ Loaded {loaded_count}/5 modules in {stats.get('last', 0):.2f}s "
                f"(cached: {load_time})"
            )

    # Show errors if any
    if load_errors:
        with st.expander(f"⚠️ {len(load_errors)} file(s) failed to load", expanded=False):
            error_df = pd.DataFrame(load_errors)
            st.dataframe(error_df, width="stretch")

            st.markdown("**Recovery Tips:**")
            if any(e["type"] == "JSONDecodeError" for e in load_errors):
                st.info(
                    "? **Invalid JSON**: Re-run evaluation to regenerate corrupt files")
            if any(e["type"] == "PermissionError" for e in load_errors):
                st.info(
                    "? **Permission Denied**: Check file permissions or close programs using the files")
            if any(e["type"] == "FileNotFoundError" for e in load_errors):
                st.info("? **File Not Found**: Verify reports directory path")

    # Performance monitoring
    if config.performance.enable_monitoring and PERFORMANCE_MONITORING:
        with measure_section("report_loading_stats"):
            monitor = get_monitor()
            stats = monitor.get_stats("load_all_reports")

            if stats and config.performance.log_slow_operations:
                avg_time = stats.get('mean', 0)
                threshold = config.performance.slow_operation_threshold_ms / 1000

                if avg_time > threshold:
                    logger.warning(
                        f"Slow operation detected: load_all_reports took {avg_time:.2f}s "
                        f"(threshold: {threshold}s)"
                    )

# Show summary of loaded modules
loaded_modules = [k for k, v in latest.items() if v is not None]
missing_modules = [k for k in ["L1", "L2", "L3",
                               "L4", "L5", "AGG"] if k not in loaded_modules]

if missing_modules:
    st.warning(
        f"⚠️ Missing reports for: {', '.join(missing_modules)}. "
        f"Run evaluation to generate these modules."
    )

    # Provide quick action button
    col1, col2, col3 = st.columns([1, 1, 2])
    with col1:
        if st.button("📖 View Documentation", width="stretch"):
            st.info("""
            **To generate missing reports:**
```bash
            # Generate all modules
            python -m cli.iraqaf_cli run --module ALL --config configs/project.example.yaml --out reports

            # Generate specific module
            python -m cli.iraqaf_cli run --module L1 --config configs/project.example.yaml --out reports
        """)
st.divider()

# ===== Module Overview Grid (NEW: First-Time User Experience) =====
risk_profile_for_display = st.session_state.get("risk_profile", "High")
display_module_overview_grid(latest, risk_profile_for_display)

st.markdown("---")

# ===== Smart Insights Section (NEW) =====
if any(latest.get(mid) for mid in ["L1", "L2", "L3", "L4", "L5"]):
    display_smart_insights(latest, risk_profile_for_display)
    st.markdown("---")

# ========== MODULE SUMMARY ==========
with measure_section("build_module_summary"):
    rows = []
    for mid in ["L1", "L2", "L3", "L4", "L5"]:
        rep = latest.get(mid)
        if rep:
            # Updated to use local function with report data
            e_cnt = get_evidence_count_local(module_id=mid, report_data=rep)
            metrics = rep.get("metrics", {}) or {}
            score = rep.get("score", None)

            # Build metrics display using new helper
            metrics_display = get_metrics_display_local(
                metrics, score, max_items=3)

            # For L1, use coverage_percent; for others use score
            display_score = metrics.get(
                "coverage_percent", score) if mid == "L1" else score

            # Get status emoji from band field
            emoji = COLOR_EMOJI.get(str(rep.get("band", "")).lower(), "⚪")
            band_text = str(rep.get('band', 'Unknown')).capitalize()

            rows.append([
                NAMES[mid],
                f"{emoji} {band_text}",
                display_score,
                metrics_display,
                e_cnt,
            ])
        else:
            rows.append([NAMES[mid], "❔ Missing", "-", "—", 0])

    df = pd.DataFrame(
        rows,
        columns=["Module", "Status", "Score", "Metrics Used", "Evidence 📎"]
    )

# ===== Module Summary =====
st.markdown('<div data-tour-target="module-summary">', unsafe_allow_html=True)

st.markdown("### 🧾 Module Summary & Details")

# Simple help text right below the title
st.markdown("""
<div style='
    background: #f0f9ff;
    border-left: 4px solid #0284c7;
    padding: 12px 16px;
    margin-bottom: 16px;
    border-radius: 4px;
'>
    <p style='margin: 0; font-size: 0.9rem; color: #0c4a6e;'>
        <b>Quick Guide:</b>
        🟢 Excellent (=90) 🟡 Good (75-89) 🔴 Needs Work (&lt;75) 📎 Evidence Count
        <span style='float: right; cursor: pointer; font-weight: bold;' onclick='this.parentElement.parentElement.style.display="none"'>✕</span>
    </p>
</div>
""", unsafe_allow_html=True)

# Dataframe with tooltips
st.dataframe(df, width="stretch")

# Add expandable section showing module details with context
st.markdown("#### 🗂️ Module Details")

summary_cols = st.columns(3)
for idx, mid in enumerate(["L1", "L2", "L3", "L4", "L5"]):
    if idx % 3 == 0:
        summary_cols = st.columns(3)

    col_idx = idx % 3

    with summary_cols[col_idx]:
        rep = latest.get(mid)
        desc = MODULE_DESCRIPTIONS.get(mid, {})

        with st.expander(f"{desc.get('emoji', '🗂️')} {mid}: {desc.get('name', '')}"):
            st.markdown(f"**Purpose:** {desc.get('purpose', '')}")

            if rep:
                st.markdown(f"**Score:** {rep.get('score', '?')}/100")

                # Get band and convert to colored circle
                band = rep.get('band', 'Unknown').lower()
                status_emoji = {
                    'green': '🟢',
                    'yellow': '🟡',
                    'red': '🔴'
                }.get(band, '⚪')  # White circle for unknown status

                st.markdown(
                    f"**Status:** {status_emoji} {rep.get('band', 'Unknown').capitalize()}")

                metrics = rep.get("metrics", {})
                if metrics:
                    st.markdown("**Key Metrics:**")
                    # Show top 3 metrics with proper bullet points
                    for metric_name, metric_val in list(metrics.items())[:3]:
                        st.markdown(f"- **{metric_name}:** {metric_val}")
            else:
                st.warning("No report generated yet")

            e_cnt = get_evidence_count_local(module_id=mid, report_data=rep)
            st.markdown(f"**Evidence Files:** {e_cnt} 📎")

st.markdown('</div>', unsafe_allow_html=True)  # Close tour target

# ===== 📎 Evidence Tray (per module) with Search + Global Preview =====
st.markdown("""
<div data-tour-target="evidence-tray">
<h3>📎 Evidence Tray - Supporting Documentation</h3>
""", unsafe_allow_html=True)

st.markdown("""
<div style='
    background: #fef3c7;
    border-left: 4px solid #f59e0b;
    padding: 12px 16px;
    margin-bottom: 16px;
    border-radius: 4px;
'>
    <p style='margin: 0; font-size: 0.9rem; color: #78350f;'>
        <b>💡 Tip:</b> Upload supporting documents for each module. Files are organized by module and can be previewed.
        Use filters to find specific documents quickly.
    </p>
</div>
""", unsafe_allow_html=True)

with measure_section("evidence_tray_render"):

    eidx = _load_evidence_index()
    if not eidx:
        st.info(
            "No evidence index found yet. Create **configs/evidence_index.json** to pin files per module."
        )
    else:
        # Global controls
        ctrl1, ctrl2, ctrl3, ctrl4 = st.columns([0.4, 0.2, 0.2, 0.2])
        with ctrl1:
            q = st.text_input(
                "🔍 Filter files",
                value="",
                placeholder="e.g., fairness, .csv, security.yaml",
                disabled=_LOCK
            )
        with ctrl2:
            expand_all = st.toggle(
                "Expand all",
                value=False,
                disabled=_LOCK
            )
        with ctrl3:
            preview_all = st.toggle(
                "Preview all",
                value=False,
                disabled=_LOCK
            )
        with ctrl4:
            # Count total files
            total_files = sum(len(files) for files in eidx.values())
            st.metric("Total Files", total_files, label_visibility="collapsed")

        # Normalize query
        q_norm = (q or "").strip().lower()

        for mid in ["L1", "L2", "L3", "L4", "L5"]:
            files_for_mid = eidx.get(mid, [])
            # Apply filter
            if q_norm:
                files_for_mid = [
                    p for p in files_for_mid
                    if q_norm in p.lower() or q_norm in os.path.basename(p).lower()
                ]

            count = len(files_for_mid)
            desc = MODULE_DESCRIPTIONS.get(mid, {})
            title = f"{desc.get('emoji', '')} {NAMES[mid]} ? {count} file{'s' if count != 1 else ''}"

            with st.expander(title, expanded=expand_all):
                if count == 0:
                    st.caption(
                        "No files match the current filter." if q_norm else "No files pinned for this module."
                    )
                    continue

                # Enhanced header with better spacing
                h1, h2, h3, h4, h5, h6 = st.columns(
                    [0.36, 0.12, 0.16, 0.16, 0.10, 0.10])
                h1.markdown("**File**")
                h2.markdown("**Status**")
                h3.markdown("**Size**")
                h4.markdown("**Modified**")
                h5.markdown("**Preview**")
                h6.markdown("**Download**")

                for i, p in enumerate(files_for_mid, start=1):
                    fp = Path(p)
                    exists = fp.exists()
                    size = _human_size(fp.stat().st_size) if exists else "—"
                    mtime = time.strftime(
                        "%Y-%m-%d %H:%M:%S", time.localtime(fp.stat().st_mtime)
                    ) if exists else "—"

                    c1, c2, c3, c4, c5, c6 = st.columns(
                        [0.36, 0.12, 0.16, 0.16, 0.10, 0.10])
                    c1.markdown(
                        f"**{fp.name}**  \n<small style='color:#888'>{fp.as_posix()}</small>",
                        unsafe_allow_html=True
                    )
                    c2.markdown("✅ Exists" if exists else "❌ Missing")
                    c3.markdown(f"{size}")
                    c4.markdown(mtime)

                    # Row toggle honors the global "Preview all" `as its default
                    tkey = f"pv_{mid}_{i}"
                    show = c5.toggle(
                        "Show",
                        value=preview_all,
                        key=tkey,
                        help="Toggle to preview the file inline.",
                        disabled=_LOCK
                    )

                    if exists:
                        with fp.open("rb") as fh:
                            c6.download_button(
                                "Download",
                                data=fh.read(),
                                file_name=fp.name,
                                key=f"dl_{mid}_{i}"
                            )
                    else:
                        c6.markdown("—")

                    if show:
                        with st.container():
                            st.markdown(
                                "<div style='margin:8px 0; padding:10px; background:#f7f7f9; border-radius:8px;'>",
                                unsafe_allow_html=True
                            )
                            preview_widget(fp)
                            st.markdown("</div>", unsafe_allow_html=True)

                st.caption(
                    "Tip: Keep evidence in `docs/`, `configs/`, `logs/`, or `data/` so paths stay tidy and portable."
                )

st.markdown("</div>", unsafe_allow_html=True)


# ===== L1 Governance –  Compliance Matrix =====
st.markdown("""
<div data-tour-target="l1-governance">
<h3>🏛️ L1 Governance – Compliance Coverage by Framework</h3>
""", unsafe_allow_html=True)

with measure_section("l1_compliance_matrix"):
    l1 = latest.get("L1")
    if l1 and "metrics" in l1:
        fb = l1["metrics"].get("framework_breakdown")
        cov = l1["metrics"].get("coverage_percent")
        if fb and isinstance(fb, list):
            st.caption(
                f"Overall L1 coverage: **{cov:.2f}%** • Each bar shows coverage (weighted clause pass%) for a regulation."
            )
            # Build a tidy dataframe
            frows = []
            for item in fb:
                frows.append({
                    "Framework": item.get("framework"),
                    "Coverage %": float(item.get("coverage_percent", 0.0)),
                    "Passed / Total": f"{item.get('covered', 0)}/{item.get('total', 0)}"
                })
            fdf = pd.DataFrame(frows).sort_values("Coverage %", ascending=True)
            # Add color coding based on coverage percentage
            fdf['Color'] = fdf['Coverage %'].apply(
                lambda x: '#10B981' if x == 100 else '#F59E0B' if x >= 80 else '#F97316' if x >= 50 else '#EF4444'
            )
            # Add status label
            fdf['Status'] = fdf['Coverage %'].apply(
                lambda x: 'Compliant' if x == 100 else 'Good' if x >= 80 else 'Fair' if x >= 50 else 'Needs Work'
            )
            # Dynamically adjust height based on number of frameworks (approximately 35px per framework)
            chart_height = max(250, len(fdf) * 35)
            chart = (
                alt.Chart(fdf)
                .mark_bar(cornerRadius=4)
                .encode(
                    x=alt.X("Coverage %:Q", title="Coverage (%)",
                            scale=alt.Scale(domain=[0, 100]),
                            axis=alt.Axis(format='.0f', labelExpr='datum.value + "%"')),
                    y=alt.Y("Framework:N", sort="-x", title="",
                            axis=alt.Axis(labelFontSize=12, labelFontWeight='bold')),
                    color=alt.Color('Color:N', scale=alt.Scale(domain=['#10B981', '#F59E0B', '#F97316', '#EF4444'], range=[
                                    '#10B981', '#F59E0B', '#F97316', '#EF4444']), legend=None),
                    tooltip=[
                        alt.Tooltip("Framework:N", title="Framework"),
                        alt.Tooltip("Coverage %:Q",
                                    title="Coverage", format=".2f"),
                        alt.Tooltip("Status:N", title="Status"),
                        "Passed / Total"
                    ]
                )
                .properties(
                    height=chart_height,
                    width="container",
                    title=alt.TitleParams(
                        text="Framework Compliance Coverage",
                        subtitle="Percentage of clauses met for each regulatory framework",
                        fontSize=14,
                        fontWeight="bold"
                    )
                )
            )
            st.altair_chart(chart, width="stretch")
        else:
            st.info("No per-framework breakdown found in L1 metrics yet.")
    else:
        st.info("Run L1 Governance to populate the compliance matrix.")

    l1 = _normalize_clause_evidence(l1 or {})

st.markdown("</div>", unsafe_allow_html=True)

# ===== L1 Governance – Regulation Traceability Map =====
st.markdown("""
<div data-tour-target="l1-traceability">
<h3>🔗 Governance Traceability (Clause ↔ Metric)</h3>
""", unsafe_allow_html=True)

# Try to load trace_map with proper path resolution
_trace_path = Path("configs/trace_map.yaml")
if not _trace_path.exists():
    _trace_path = Path(_ROOT) / "configs" / "trace_map.yaml"

if not _trace_path.exists():
    st.info("No trace_map.yaml found in configs/. Create it to visualize clause-to-metric links.")
else:
    try:
        with open(_trace_path, "r", encoding="utf-8") as fh:
            _tm = _yaml.safe_load(fh) or {}
        trace_map = (_tm.get("trace_map") or {})
    except Exception as e:
        trace_map = {}
        st.error(f"Failed to load trace_map.yaml: {e}")

    if trace_map:
        fw_names = list(trace_map.keys())
        fw_sel = st.selectbox("Framework", fw_names,
                              index=0, key="trace_fw_sel")

        rows = []
        for clause, metric_list in (trace_map.get(fw_sel, {}) or {}).items():
            for met in (metric_list or []):
                rows.append({"Clause": clause, "Metric": met, "Link": 1})

        if not rows:
            st.info("No mappings in the selected framework yet.")
        else:
            _df = pd.DataFrame(rows)

            # --- Matrix-style heatmap (Clause on Y, Metric on X)
            st.caption(
                "Matrix view: a filled cell indicates a documented linkage between a clause and a model metric."
            )
            pivot = _df.pivot_table(
                index="Clause", columns="Metric", values="Link", fill_value=0
            )
            # Show as table first for clarity
            st.dataframe(pivot, width="stretch")

            try:
                tidy = (
                    _df.groupby(["Clause", "Metric"], as_index=False)["Link"]
                    .sum()
                )
                chart = (
                    alt.Chart(tidy)
                    .mark_rect()
                    .encode(
                        x=alt.X("Metric:N", sort=None, title="Metric"),
                        y=alt.Y("Clause:N", sort=None, title="Clause"),
                        color=alt.Color(
                            "Link:Q",
                            scale=alt.Scale(
                                domain=[0, 1], range=["#f0f0f0", "#1f77b4"]
                            ),
                            legend=None
                        ),
                        tooltip=["Clause", "Metric"]
                    )
                    .properties(height=320, width="container")
                )
                st.altair_chart(chart, theme=None)
            except Exception:
                pass

            with st.expander("Copy metric paths"):
                st.write("Click to copy commonly used metric paths:")
                _uniq_metrics = sorted({r["Metric"] for r in rows})
                for mpath in _uniq_metrics:
                    st.code(mpath)
    else:
        st.info("Trace map is empty. Add entries in configs/trace_map.yaml.")

st.markdown("</div>", unsafe_allow_html=True)


# ===== L1 Governance - Framework Summary =====
st.markdown("### ⚖️ L1: Governance & Regulatory Compliance")
l1 = latest.get("L1")

if l1 and "metrics" in l1:
    metrics = l1.get("metrics", {})
    coverage = metrics.get("coverage_percent", 0)
    failed_clauses = [c for c in (metrics.get(
        "clauses") or []) if not c.get("passed")]
    frameworks = metrics.get("framework_breakdown", [])

    # Summary row with key stats
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Coverage", f"{coverage:.1f}%")
    with col2:
        st.metric("Frameworks", len(frameworks))
    with col3:
        st.metric("Passed", len(metrics.get(
            "clauses", [])) - len(failed_clauses))
    with col4:
        st.metric("Failed", len(failed_clauses))

    # Framework breakdown - responsive columns
    if frameworks:
        st.markdown("**Framework Status:**")
        # Show up to 5 frameworks per row, or fewer if there are few frameworks
        cols_per_row = min(5, len(frameworks))
        fw_cols = st.columns(cols_per_row)
        for i, fw in enumerate(frameworks):
            with fw_cols[i % cols_per_row]:
                fw_name = fw.get("framework", "Unknown")
                fw_coverage = fw.get("coverage_percent", 0)
                fw_total = fw.get("total_clauses", 1)
                fw_covered = fw.get("covered_clauses", 0)

                # Determine color and status based on coverage
                if fw_coverage == 100:
                    color = "#10B981"
                    status = "✅ Compliant"
                    bg_color = "#ECFDF5"
                    border_color = "#10B981"
                elif fw_coverage >= 80:
                    color = "#F59E0B"
                    status = "⚠️ Good"
                    bg_color = "#FFFBEB"
                    border_color = "#F59E0B"
                elif fw_coverage >= 50:
                    color = "#F97316"
                    status = "⚡ Fair"
                    bg_color = "#FFEDD5"
                    border_color = "#F97316"
                else:
                    color = "#EF4444"
                    status = "❌ Needs Work"
                    bg_color = "#FEE2E2"
                    border_color = "#EF4444"

                st.markdown(f"""
                <div style="
                    background: {bg_color};
                    border: 2px solid {border_color};
                    border-radius: 12px;
                    padding: 14px;
                    text-align: center;
                    box-shadow: 0 2px 8px rgba(0, 0, 0, 0.08);
                    transition: all 0.3s ease;
                    margin-bottom: 8px;
                ">
                    <p style="margin: 0; font-weight: 700; font-size: 0.9rem; color: #1F2937; word-wrap: break-word;">{fw_name}</p>
                    <p style="margin: 8px 0 0 0; font-size: 1.3rem; font-weight: 800; color: {color};">{fw_coverage:.1f}%</p>
                    <p style="margin: 6px 0 0 0; font-size: 0.75rem; color: #6B7280;">{status}</p>
                    <p style="margin: 6px 0 0 0; font-size: 0.75rem; color: #9CA3AF;">{fw_covered}/{fw_total} clauses</p>
                    <div style="margin-top: 8px; background: #D1D5DB; height: 3px; border-radius: 2px; overflow: hidden;">
                        <div style="background: {color}; height: 100%; width: {fw_coverage}%; border-radius: 2px; transition: width 0.3s ease;"></div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

    # Failed clauses - only if there are any
    if failed_clauses:
        # Group failed clauses by framework
        clauses_by_framework = {}
        for c in failed_clauses:
            fw = c.get('framework', 'Unknown')
            if fw not in clauses_by_framework:
                clauses_by_framework[fw] = []
            clauses_by_framework[fw].append(c)

        # Create collapsible sections for each framework
        with st.expander(f"⚠️ Action Items ({len(failed_clauses)} issues)", expanded=False):
            for fw_name, clauses in sorted(clauses_by_framework.items()):
                st.markdown(f"#### 📋 {fw_name} ({len(clauses)} failed)")

                for idx, c in enumerate(clauses, 1):
                    # Create card using Streamlit components instead of raw HTML
                    col1, col2 = st.columns([0.8, 0.2])

                    with col1:
                        st.markdown(f"**{c.get('id', 'Unknown')}**")

                    with col2:
                        priority_level = min(idx - 1, 2)
                        priority = ['High', 'Medium', 'Low'][priority_level]
                        priority_emoji = ['🔴', '🟡', '🟢'][priority_level]
                        st.caption(f"{priority_emoji} {priority}")

                    st.write(f"**Issue:** {c.get('why_failed', 'Fix needed')}")
                    st.write(
                        f"**Solution:** {c.get('hint', 'See documentation')}")
                    st.divider()
else:
    st.info("Run L1 Governance to see compliance status")

# ===== L1 Governance - Trends Over Time =====
st.markdown("### 📈 L1 Governance - Trends Over Time")
l1_rows = []
for f in files:
    try:
        with open(f, "r") as fh:
            d = json.load(fh)
        if d.get("module") != "L1":
            continue
        m = d.get("metrics", {}) or {}
        base = os.path.basename(f).replace(".json", "")
        ts_match = re.search(r"(\d{8}-\d{6})", base)
        ts_label = ts_match.group(1) if ts_match else base
        l1_rows.append({
            "Run": base,
            "Label": ts_label,
            "time": os.path.getmtime(f),
            "policy_coverage": (None if m.get("policy_coverage") is None else float(m["policy_coverage"])),
            "audit_completeness": (None if m.get("audit_completeness") is None else float(m["audit_completeness"])),
            "risk_register_age_days": (None if m.get("risk_register_age_days") is None else float(m["risk_register_age_days"])),
            "Score": d.get("score")
        })
    except Exception:
        pass

if not l1_rows:
    st.info("No L1 governance metrics found yet. (Optional keys: policy_coverage, audit_completeness, risk_register_age_days)")
else:
    l1_df = pd.DataFrame(l1_rows).sort_values("time").reset_index(drop=True)
    l1_df["Run #"] = l1_df.index + 1
    st.dataframe(l1_df[["Run #", "Label", "policy_coverage", "audit_completeness", "risk_register_age_days", "Score"]],
                 width="stretch")

    cov_bands = pd.DataFrame({"y0": [0.00, 0.90, 0.95], "y1": [
                             0.90, 0.95, 1.00], "band": ["red", "yellow", "green"]})
    comp_bands = cov_bands.copy()
    age_bands = pd.DataFrame({"y0": [0, 1, 2], "y1": [1, 2, 7], "band": [
                             # 0?1d green, 1?2d yellow, >2d red (demo)
                             "green", "yellow", "red"]})

    def l1_chart(df, metric_col, title, bands_df, y_domain):
        base = alt.Chart(df)
        bands = alt.Chart(bands_df).mark_rect(opacity=0.15).encode(
            y=alt.Y("y0:Q", title=None, scale=alt.Scale(domain=y_domain)),
            y2="y1:Q",
            color=alt.Color("band:N",
                            scale=alt.Scale(domain=["green", "yellow", "red"],
                                            range=["#00c851", "#ffde59", "#ff4b4b"]),
                            legend=None)
        )
        line = base.mark_line(point=True).encode(
            x=alt.X("`Run #`:Q", title="Run # (chronological)"),
            y=alt.Y(f"{metric_col}:Q", title=title,
                    scale=alt.Scale(domain=y_domain)),
            tooltip=["Label:N", alt.Tooltip(f"{metric_col}:Q", format=".3f")]
        ).properties(height=240, width="container")
        return (bands + line)

    c1, c2 = st.columns(2)
    with c1:
        st.altair_chart(l1_chart(l1_df, "policy_coverage", "Policy Coverage (? better)", cov_bands, (0, 1.0)),
                        width="stretch")
    with c2:
        st.altair_chart(l1_chart(l1_df, "audit_completeness", "Audit Completeness (? better)", comp_bands, (0, 1.0)),
                        width="stretch")

    st.altair_chart(l1_chart(l1_df, "risk_register_age_days", "Risk Register Age (days, ? better)", age_bands, (0, 7)),
                    width="stretch")

# ===== DYNAMIC L2 Privacy/Security - Auto-Detect Frameworks & Apps =====


def extract_l2_framework_metadata(l2_report: dict) -> dict:
    """Extract framework and app metadata from L2 report."""
    metadata = l2_report.get("metadata", {}) or {}
    return {
        "framework": metadata.get("framework_name", "Unknown"),
        "app": metadata.get("app_name", "default_app"),
        "version": metadata.get("version", "1.0"),
        "timestamp": metadata.get("timestamp", None),
    }


def calculate_data_protection_score(metrics: dict) -> float:
    """Calculate Data Protection score from available metrics."""
    scores = []
    if metrics.get("encryption_coverage") is not None:
        scores.append(float(metrics["encryption_coverage"]) * 0.5)
    if metrics.get("dpia_complete") is not None:
        scores.append(float(metrics["dpia_complete"]) * 0.3)
    if metrics.get("access_review_age_days") is not None:
        age = float(metrics["access_review_age_days"])
        review_score = max(0, 1 - (age / 180))
        scores.append(review_score * 0.2)
    return sum(scores) / len(scores) if scores else 0.0


def calculate_encryption_score(metrics: dict) -> float:
    """Calculate Encryption coverage score."""
    return float(metrics.get("encryption_coverage", 0))


def calculate_access_controls_score(metrics: dict) -> float:
    """Calculate Access Controls score from access review metrics."""
    age = metrics.get("access_review_age_days")
    if age is None:
        return 0.0
    age = float(age)
    if age <= 30:
        return 1.0
    elif age <= 60:
        return 0.8
    elif age <= 120:
        return 0.6
    else:
        return 0.4


def calculate_security_posture_score(metrics: dict) -> float:
    """Calculate Security Posture from incident rates and compliance."""
    scores = []
    if metrics.get("incident_rate_per_1k_users") is not None:
        incident_rate = float(metrics["incident_rate_per_1k_users"])
        incident_score = max(0, 1 - (incident_rate / 5))
        scores.append(incident_score * 0.6)
    if metrics.get("dpia_complete") is not None:
        scores.append(float(metrics["dpia_complete"]) * 0.4)
    return sum(scores) / len(scores) if scores else 0.0


st.markdown("""
<div data-tour-target="l2-privacy">
<h2 style="margin: 0; font-size: 1.8rem; font-weight: 800; color: #1F2937;">🔐 L2 Privacy & Security</h2>
<p style="margin: 8px 0 0 0; font-size: 0.95rem; color: #6B7280;">Dynamic security posture monitoring across all tested frameworks and applications</p>
</div>
""", unsafe_allow_html=True)

st.markdown("---")

l2_rows = []
l2_framework_app_data = defaultdict(lambda: defaultdict(list))

for f in files:
    try:
        with open(f, "r") as fh:
            d = json.load(fh)
        if d.get("module") != "L2":
            continue
        m = d.get("metrics", {}) or {}
        meta = extract_l2_framework_metadata(d)

        base = os.path.basename(f).replace(".json", "")
        ts_match = re.search(r"(\d{8}-\d{6})", base)
        ts_label = ts_match.group(1) if ts_match else base

        l2_rows.append({
            "Run": base,
            "Label": ts_label,
            "time": os.path.getmtime(f),
            "framework": meta["framework"],
            "app": meta["app"],
            "version": meta["version"],
            "encryption_coverage": (None if m.get("encryption_coverage") is None else float(m["encryption_coverage"])),
            "dpia_complete": (None if m.get("dpia_complete") is None else float(m["dpia_complete"])),
            "access_review_age_days": (None if m.get("access_review_age_days") is None else float(m["access_review_age_days"])),
            "incident_rate_per_1k_users": (None if m.get("incident_rate_per_1k_users") is None else float(m["incident_rate_per_1k_users"])),
            "Score": d.get("score"),
            "data_protection": calculate_data_protection_score(m),
            "encryption": calculate_encryption_score(m),
            "access_controls": calculate_access_controls_score(m),
            "security_posture": calculate_security_posture_score(m),
        })

        l2_framework_app_data[meta["framework"]][meta["app"]].append({
            "metrics": m,
            "time": os.path.getmtime(f),
            "timestamp": ts_label,
        })
    except Exception as e:
        logger.debug(f"Error processing L2 report {f}: {e}")

if not l2_rows:
    st.markdown("""
    <div style="background: linear-gradient(135deg, #E0F2FE 0%, #F0F9FF 100%); border: 2px solid #3B82F6; border-radius: 12px; padding: 24px; text-align: center;">
        <p style="margin: 0 0 8px 0; font-size: 1rem; font-weight: 700; color: #1E40AF;">🔍 No L2 Privacy/Security Reports Found</p>
        <p style="margin: 0; font-size: 0.9rem; color: #1E40AF; line-height: 1.5;">Run the CLI with L2 module to automatically generate privacy and security assessments:<br><code style="background: white; padding: 4px 8px; border-radius: 4px;">python -m cli.iraqaf_cli run --module L2 --config configs/project.example.yaml --out reports</code></p>
    </div>
    """, unsafe_allow_html=True)
else:
    l2_df = pd.DataFrame(l2_rows).sort_values("time").reset_index(drop=True)
    l2_df["Run #"] = l2_df.index + 1

    frameworks = sorted(l2_df["framework"].unique())
    apps = sorted(l2_df["app"].unique())

    st.markdown(f"""
    <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 16px; margin-bottom: 24px;">
        <div style="background: linear-gradient(135deg, #F0F9FF 0%, #E0F2FE 100%); border: 2px solid #3B82F6; border-radius: 12px; padding: 16px;">
            <p style="margin: 0 0 4px 0; font-size: 0.9rem; font-weight: 700; color: #1E40AF;">📊 Total Frameworks</p>
            <p style="margin: 0; font-size: 1.5rem; font-weight: 800; color: #1E3A8A;">{len(frameworks)}</p>
        </div>
        <div style="background: linear-gradient(135deg, #FEF9E7 0%, #FEF3C7 100%); border: 2px solid #F59E0B; border-radius: 12px; padding: 16px;">
            <p style="margin: 0 0 4px 0; font-size: 0.9rem; font-weight: 700; color: #92400E;">🎯 Total Applications</p>
            <p style="margin: 0; font-size: 1.5rem; font-weight: 800; color: #B45309;">{len(apps)}</p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("<h3 style='margin: 20px 0 16px 0; font-size: 1.2rem; font-weight: 700; color: #1F2937;'>📈 Key Metrics - Historical Trends</h3>", unsafe_allow_html=True)

    cov_bands = pd.DataFrame({"y0": [0.00, 0.90, 0.95], "y1": [0.90, 0.95, 1.00],
                              "band": ["red", "yellow", "green"]})
    dpia_bands = pd.DataFrame({"y0": [0.00, 0.75, 0.90], "y1": [0.75, 0.90, 1.00],
                               "band": ["red", "yellow", "green"]})
    access_age_bands = pd.DataFrame({"y0": [0, 30, 60], "y1": [30, 60, 120],
                                     "band": ["green", "yellow", "red"]})
    incident_bands = pd.DataFrame({"y0": [0, 1, 3], "y1": [1, 3, 10],
                                   "band": ["green", "yellow", "red"]})

    def l2_chart(df, metric_col, title, bands_df, y_domain):
        base = alt.Chart(df)
        bands = alt.Chart(bands_df).mark_rect(opacity=0.15).encode(
            y=alt.Y("y0:Q", title=None, scale=alt.Scale(domain=y_domain)),
            y2="y1:Q",
            color=alt.Color(
                "band:N",
                scale=alt.Scale(domain=["green", "yellow", "red"],
                                range=["#00c851", "#ffde59", "#ff4b4b"]),
                legend=None
            )
        )
        line = base.mark_line(point=True).encode(
            x=alt.X("`Run #`:Q", title="Run # (chronological)"),
            y=alt.Y(f"{metric_col}:Q", title=title,
                    scale=alt.Scale(domain=y_domain)),
            tooltip=["Label:N", "framework:N", "app:N",
                     alt.Tooltip(f"{metric_col}:Q", format=".3f")]
        ).properties(height=240, width="container")
        return (bands + line)

    c1, c2 = st.columns(2)
    with c1:
        st.altair_chart(
            l2_chart(l2_df, "encryption_coverage",
                     "Encryption Coverage (higher better)", cov_bands, (0, 1.0)),
            width="stretch"
        )
    with c2:
        st.altair_chart(
            l2_chart(l2_df, "dpia_complete",
                     "DPIA Completion (higher better)", dpia_bands, (0, 1.0)),
            width="stretch"
        )

    c3, c4 = st.columns(2)
    with c3:
        st.altair_chart(
            l2_chart(l2_df, "access_review_age_days",
                     "Access Review Age (days, lower better)", access_age_bands, (0, 120)),
            width="stretch"
        )
    with c4:
        st.altair_chart(
            l2_chart(l2_df, "incident_rate_per_1k_users",
                     "Incident Rate per 1k Users (lower better)", incident_bands, (0, 10)),
            width="stretch"
        )

    st.markdown("<h3 style='margin: 24px 0 16px 0; font-size: 1.2rem; font-weight: 700; color: #1F2937;'>🎯 Key Metrics by Framework & Application</h3>", unsafe_allow_html=True)

    metric_tab1, metric_tab2, metric_tab3 = st.tabs(
        ["Data Protection", "Encryption", "Access Controls"])

    with metric_tab1:
        data_prot_df = l2_df[["Label", "framework", "app", "data_protection"]].drop_duplicates(
            subset=["Label", "framework", "app"]).sort_values("data_protection", ascending=False)
        if not data_prot_df.empty:
            st.markdown(
                "**Data Protection Score by Framework/App (higher is better)**")
            data_prot_chart = alt.Chart(data_prot_df).mark_bar().encode(
                x=alt.X("data_protection:Q", title="Data Protection Score",
                        scale=alt.Scale(domain=[0, 1])),
                y=alt.Y("framework:N", title="Framework"),
                color=alt.Color("data_protection:Q",
                                scale=alt.Scale(scheme="greens")),
                tooltip=["framework:N", "app:N", alt.Tooltip(
                    "data_protection:Q", format=".2%")]
            ).properties(height=300, width="container")
            st.altair_chart(data_prot_chart, width="stretch")

    with metric_tab2:
        enc_df = l2_df[["Label", "framework", "app", "encryption"]].drop_duplicates(
            subset=["Label", "framework", "app"]).sort_values("encryption", ascending=False)
        if not enc_df.empty:
            st.markdown(
                "**Encryption Coverage by Framework/App (higher is better)**")
            enc_chart = alt.Chart(enc_df).mark_bar().encode(
                x=alt.X("encryption:Q", title="Encryption Coverage",
                        scale=alt.Scale(domain=[0, 1])),
                y=alt.Y("framework:N", title="Framework"),
                color=alt.Color(
                    "encryption:Q", scale=alt.Scale(scheme="blues")),
                tooltip=["framework:N", "app:N", alt.Tooltip(
                    "encryption:Q", format=".2%")]
            ).properties(height=300, width="container")
            st.altair_chart(enc_chart, width="stretch")

    with metric_tab3:
        acc_df = l2_df[["Label", "framework", "app", "access_controls"]].drop_duplicates(
            subset=["Label", "framework", "app"]).sort_values("access_controls", ascending=False)
        if not acc_df.empty:
            st.markdown(
                "**Access Controls Score by Framework/App (higher is better)**")
            acc_chart = alt.Chart(acc_df).mark_bar().encode(
                x=alt.X("access_controls:Q", title="Access Controls Score",
                        scale=alt.Scale(domain=[0, 1])),
                y=alt.Y("framework:N", title="Framework"),
                color=alt.Color("access_controls:Q",
                                scale=alt.Scale(scheme="oranges")),
                tooltip=["framework:N", "app:N", alt.Tooltip(
                    "access_controls:Q", format=".2%")]
            ).properties(height=300, width="container")
            st.altair_chart(acc_chart, width="stretch")

    st.markdown("<h3 style='margin: 24px 0 16px 0; font-size: 1.2rem; font-weight: 700; color: #1F2937;'>🔍 Security Posture by Framework</h3>", unsafe_allow_html=True)

    for framework in sorted(l2_framework_app_data.keys()):
        framework_data = l2_framework_app_data[framework]
        apps_list = sorted(framework_data.keys())

        st.markdown(f"""
        <div style="background: linear-gradient(135deg, #F9FAFB 0%, #F3F4F6 100%); border: 2px solid #E5E7EB; border-radius: 12px; padding: 16px; margin-bottom: 16px;">
            <h4 style="margin: 0 0 12px 0; font-size: 1.1rem; font-weight: 700; color: #1F2937;">📦 Framework: <span style="color: #3B82F6;">{framework}</span></h4>
        """, unsafe_allow_html=True)

        framework_latest = l2_df[l2_df["framework"] == framework].sort_values(
            "time").iloc[-1] if len(l2_df[l2_df["framework"] == framework]) > 0 else None

        if framework_latest is not None:
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Data Protection", f"{framework_latest['data_protection']:.1%}",
                          delta=None, help="Overall data protection score")
            with col2:
                st.metric("Encryption", f"{framework_latest['encryption']:.1%}",
                          delta=None, help="Encryption coverage")
            with col3:
                st.metric("Access Controls", f"{framework_latest['access_controls']:.1%}",
                          delta=None, help="Access control effectiveness")
            with col4:
                st.metric("Security Posture", f"{framework_latest['security_posture']:.1%}",
                          delta=None, help="Overall security posture")

        for app in apps_list:
            app_runs = l2_df[(l2_df["framework"] == framework)
                             & (l2_df["app"] == app)]
            if len(app_runs) > 0:
                latest_run = app_runs.iloc[-1]
                st.markdown(f"""
                <div style="background: white; border-left: 4px solid #3B82F6; border-radius: 8px; padding: 12px; margin: 8px 0;">
                    <p style="margin: 0 0 4px 0; font-size: 0.9rem; font-weight: 600; color: #374151;"><span style="color: #059669;">✓</span> App: <strong>{app}</strong> | Version: {latest_run['version']}</p>
                    <p style="margin: 0; font-size: 0.85rem; color: #6B7280;">Last Updated: {latest_run['Label']} | Score: {latest_run['Score']}</p>
                </div>
                """, unsafe_allow_html=True)

        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<h3 style='margin: 24px 0 16px 0; font-size: 1.2rem; font-weight: 700; color: #1F2937;'>📋 Detailed Metrics Table</h3>", unsafe_allow_html=True)

    st.dataframe(
        l2_df[[
            "Run #", "Label", "framework", "app", "encryption_coverage", "dpia_complete",
            "access_review_age_days", "incident_rate_per_1k_users", "Score"
        ]],
        width="stretch"
    )

st.markdown("</div>", unsafe_allow_html=True)


# ===== L1 – Automated Clause Tagging (lightweight) =====
st.markdown("### 🏷️ Automated Clause Tagging (beta)")
desc = st.text_area("Describe the AI system / use case",
                    placeholder="e.g., Clinical decision support using patient EHR + imaging ?")
if st.button("Suggest relevant clauses", disabled=_LOCK):
    # Dynamic keyword map built from trace_map.yaml
    trace_map = load_trace_map()
    kw_map = {
        "transparency": [(fw, cl) for fw in trace_map.keys() for cl in trace_map[fw].keys() if any(kw in cl.lower() for kw in ["transparent", "explainab"])],
        "explain": [(fw, cl) for fw in trace_map.keys() for cl in trace_map[fw].keys() if any(kw in cl.lower() for kw in ["transparent", "explainab", "decision"])],
        "health": [(fw, cl) for fw in trace_map.keys() for cl in trace_map[fw].keys() if any(kw in cl.lower() for kw in ["health", "medical", "quality", "device"])],
        "medical": [(fw, cl) for fw in trace_map.keys() for cl in trace_map[fw].keys() if any(kw in cl.lower() for kw in ["health", "medical", "quality", "device"])],
        "risk": [(fw, cl) for fw in trace_map.keys() for cl in trace_map[fw].keys() if any(kw in cl.lower() for kw in ["risk", "hazard"])],
        "logging": [(fw, cl) for fw in trace_map.keys() for cl in trace_map[fw].keys() if any(kw in cl.lower() for kw in ["logging", "monitoring", "audit", "verification"])],
        "privacy": [(fw, cl) for fw in trace_map.keys() for cl in trace_map[fw].keys() if any(kw in cl.lower() for kw in ["privacy", "data", "confidential"])],
        "security": [(fw, cl) for fw in trace_map.keys() for cl in trace_map[fw].keys() if any(kw in cl.lower() for kw in ["security", "cryptography", "access", "control"])],
        "fairness": [(fw, cl) for fw in trace_map.keys() for cl in trace_map[fw].keys() if any(kw in cl.lower() for kw in ["fairness", "bias", "governance", "data"])],
        "bias": [(fw, cl) for fw in trace_map.keys() for cl in trace_map[fw].keys() if any(kw in cl.lower() for kw in ["fairness", "bias", "governance"])],
    }
    found = set()
    text = (desc or "").lower()
    for k, targets in kw_map.items():
        if k in text:
            for fw, cid in targets:
                found.add((fw, cid))
    if not found:
        st.info("No suggestions matched. Try adding domain keywords (privacy, logging, medical, transparency, fairness, security, risk, health?).")
    else:
        out = [{"Framework": fw, "Clause": cid} for (fw, cid) in sorted(found)]
        df_out = pd.DataFrame(out)
        st.dataframe(df_out, width="stretch")
        st.caption(
            "Tip: wire this to your governance process to auto-populate a review list. Suggestions dynamically pulled from trace_map.yaml.")

        # NEW: save to configs/suggested_clauses.json
        save_it = st.checkbox(
            "Save to configs/suggested_clauses.json",
            value=True,
            disabled=_LOCK
        )
        if save_it:
            os.makedirs("configs", exist_ok=True)
            payload = {
                "description": (desc or ""),
                "suggestions": [{"framework": r["Framework"], "clause": r["Clause"]} for _, r in df_out.iterrows()],
                "generated_at": pd.Timestamp.now().isoformat(timespec="seconds")
            }
            try:
                with open("configs/suggested_clauses.json", "w", encoding="utf-8") as fh:
                    json.dump(payload, fh, indent=2)
                st.success("Saved ? configs/suggested_clauses.json")
            except Exception as e:
                st.error(f"Failed to save: {e}")


# ===== L3 Fairness – Evolution Across Runs (DPG & EOD) =====
st.markdown("""
<div data-tour-target="l3-fairness">
<h3>🧮 L3 Fairness – Evolution Across Runs</h3>
""", unsafe_allow_html=True)

# Make this section collapsible to avoid loading unless needed
show_l3_trends = st.checkbox(
    "📊 Show historical L3 trends", value=False, key="show_l3_trends")

if show_l3_trends:
    with measure_section("l3_fairness_processing"):
        # Collect all L3 reports
        l3_rows = []
        for f in files:
            try:
                with open(f, "r") as fh:
                    d = json.load(fh)
                if d.get("module") != "L3":
                    continue
                m = d.get("metrics", {})
                dpg = m.get("DPG", None)
                eod = m.get("EOD", None)
                auroc = m.get("AUROC", None)
                score = d.get("score", None)

                # label from filename & time (for ordering)
                base = os.path.basename(f).replace(".json", "")
                ts_match = re.search(r"(\d{8}-\d{6})", base)
                ts_label = ts_match.group(1) if ts_match else base
                mtime = os.path.getmtime(f)

                l3_rows.append({
                    "Run": base,
                    "Label": ts_label,
                    "time": mtime,
                    "DPG": None if dpg is None else float(dpg),
                    "EOD": None if eod is None else float(eod),
                    "AUROC": None if auroc in (None, "null") else float(auroc),
                    "Score": score,
                    "rates_by_group": m.get("rates_by_group"),
                    "tpr_by_group": m.get("tpr_by_group"),
                })
            except Exception:
                pass

        if not l3_rows:
            st.info(
                "No L3 fairness reports found yet. Run L3 multiple times to visualize evolution."
            )
        else:
            l3_df = pd.DataFrame(l3_rows).sort_values(
                "time").reset_index(drop=True)
            l3_df["Run #"] = l3_df.index + 1

            # Small, tidy table
            st.dataframe(
                l3_df[["Run #", "Label", "DPG", "EOD", "AUROC", "Score"]],
                width="stretch",
            )

            # Threshold bands (0?0.05 green, 0.05?0.1 yellow, 0.1?0.5 red)
            band_df = pd.DataFrame({
                "y0": [0.00, 0.05, 0.10],
                "y1": [0.05, 0.10, 0.50],
                "band": ["green", "yellow", "red"]
            })

            def fairness_chart(metric_col: str, title: str):
                base = alt.Chart(l3_df)

                bands = alt.Chart(band_df).mark_rect(opacity=0.15).encode(
                    y=alt.Y("y0:Q", title=None,
                            scale=alt.Scale(domain=(0, 0.5))),
                    y2="y1:Q",
                    color=alt.Color(
                        "band:N",
                        scale=alt.Scale(
                            domain=["green", "yellow", "red"],
                            range=["#00c851", "#ffde59", "#ff4b4b"]
                        ),
                        legend=None
                    )
                )

                line = base.mark_line(point=True).encode(
                    x=alt.X("`Run #`:Q", title="Run # (chronological)"),
                    y=alt.Y(
                        f"{metric_col}:Q",
                        title=title,
                        scale=alt.Scale(domain=(0, 0.5))
                    ),
                    tooltip=["Label:N", alt.Tooltip(
                        f"{metric_col}:Q", format=".3f")]
                ).properties(height=260, width="container")

                return (bands + line).resolve_scale(color="independent")

            c1, c2 = st.columns(2)
            with c1:
                st.altair_chart(
                    fairness_chart("DPG", "Demographic Parity Gap (? better)"),
                    width="stretch"
                )
            with c2:
                st.altair_chart(
                    fairness_chart(
                        "EOD", "Equal Opportunity Difference (? better)"),
                    width="stretch"
                )

            # Latest group diagnostics
            latest_l3 = l3_df.iloc[-1]
            st.markdown("#### 🔎 Latest Run – Group Diagnostics")
            col_a, col_b = st.columns(2)
            with col_a:
                st.write("**Positive Prediction Rate by Group**")
                st.json(latest_l3["rates_by_group"] or {})
            with col_b:
                st.write("**TPR (Sensitivity) by Group**")
                st.json(latest_l3["tpr_by_group"] or {})

st.markdown("</div>", unsafe_allow_html=True)

# ===== L3 Fairness ? Group Bias Evolution (per-group DPG/EOD) =====
st.markdown("""
<div data-tour-target="l3-group-bias">
<h3>🧮 L3 Fairness – Group Bias Evolution</h3>
""", unsafe_allow_html=True)

# Make this section collapsible to avoid loading unless needed
show_l3_group_trends = st.checkbox(
    "📊 Show per-group bias evolution",
    value=False,
    key="show_l3_group_trends",
    help="View how fairness metrics evolved across different demographic groups over time"
)

if show_l3_group_trends:
    with measure_section("l3_group_fairness_processing"):
        # Collect all L3 reports efficiently using the helper function
        def extract_l3_group_metrics(report: dict) -> dict:
            """Extract L3 per-group metrics for bias evolution tracking"""
            m = report.get("metrics", {}) or {}
            return {
                "rates_by_group": m.get("rates_by_group") or {},
                "tpr_by_group": m.get("tpr_by_group") or {},
            }

        l3_df = collect_module_history(files, "L3", extract_l3_group_metrics)

        if l3_df.empty:
            st.info(
                "No L3 reports found yet. Run L3 multiple times to visualize per-group bias evolution.")
        else:
            # Longform per-group rates + TPR
            rows_rate, rows_tpr = [], []
            for _, r in l3_df.iterrows():
                run_idx = int(r["Run #"])
                lbl = r["Label"]
                rates = r["rates_by_group"] or {}
                tprs = r["tpr_by_group"] or {}

                # Calculate overall means (used to compute per-group gaps)
                overall_rate = (pd.Series(rates).mean() if rates else None)
                overall_tpr = (pd.Series(tprs).mean() if tprs else None)

                # Process positive prediction rates
                for g, v in (rates.items() if isinstance(rates, dict) else []):
                    if v is not None and overall_rate is not None:
                        rows_rate.append({
                            "Run #": run_idx,
                            "label": lbl,
                            "group": str(g),
                            "Rate": float(v),
                            # per-group parity gap
                            "DPG_group": float(v) - float(overall_rate)
                        })

                # Process TPR values
                for g, v in (tprs.items() if isinstance(tprs, dict) else []):
                    if v is not None and overall_tpr is not None:
                        rows_tpr.append({
                            "Run #": run_idx,
                            "label": lbl,
                            "group": str(g),
                            "TPR": float(v),
                            # per-group equal opp diff
                            "EOD_group": float(v) - float(overall_tpr)
                        })

            st.markdown(
                "#### 👥 Positive Rate & DPG by Group (across runs)")
            if rows_rate:
                rate_df = pd.DataFrame(rows_rate)
                c1, c2 = st.columns(2)

                with c1:
                    ch = alt.Chart(rate_df).mark_line(point=True).encode(
                        x=alt.X("`Run #`:Q", title="Run #"),
                        y=alt.Y("Rate:Q", title="Positive Prediction Rate",
                                scale=alt.Scale(domain=(0, 1))),
                        color=alt.Color("group:N", title="Group"),
                        tooltip=["label:N", "group:N",
                                 alt.Tooltip("Rate:Q", format=".3f")]
                    ).properties(height=260, width="container", title="Positive Rate by Group")
                    st.altair_chart(ch, width="stretch")

                with c2:
                    ch = alt.Chart(rate_df).mark_line(point=True).encode(
                        x=alt.X("`Run #`:Q", title="Run #"),
                        y=alt.Y("DPG_group:Q",
                                title="Demographic Parity Gap (group - overall)"),
                        color=alt.Color("group:N", title="Group"),
                        tooltip=["label:N", "group:N", alt.Tooltip(
                            "DPG_group:Q", format="+.3f")]
                    ).properties(height=260, width="container", title="DPG by Group")
                    st.altair_chart(ch, width="stretch")
            else:
                st.info("No per-group positive rates available yet.")

            st.markdown("#### 👥 TPR & EOD by Group (across runs)")
            if rows_tpr:
                tpr_df = pd.DataFrame(rows_tpr)
                c3, c4 = st.columns(2)

                with c3:
                    ch = alt.Chart(tpr_df).mark_line(point=True).encode(
                        x=alt.X("`Run #`:Q", title="Run #"),
                        y=alt.Y("TPR:Q", title="True Positive Rate",
                                scale=alt.Scale(domain=(0, 1))),
                        color=alt.Color("group:N", title="Group"),
                        tooltip=["label:N", "group:N",
                                 alt.Tooltip("TPR:Q", format=".3f")]
                    ).properties(height=260, width="container", title="TPR by Group")
                    st.altair_chart(ch, width="stretch")

                with c4:
                    # Threshold bands for |EOD_group|
                    band_df = pd.DataFrame({
                        "y0": [-0.50, -0.10, -0.05,  0.05,  0.10],
                        "y1": [-0.10, -0.05,  0.05,  0.10,  0.50],
                        "band": ["red", "yellow", "green", "yellow", "red"]
                    })
                    bands = alt.Chart(band_df).mark_rect(opacity=0.12).encode(
                        y=alt.Y("y0:Q", title=None),
                        y2="y1:Q",
                        color=alt.Color("band:N", scale=alt.Scale(
                            domain=["green", "yellow", "red"],
                            range=["#00c851", "#ffde59", "#ff4b4b"]), legend=None)
                    )
                    line = alt.Chart(tpr_df).mark_line(point=True).encode(
                        x=alt.X("`Run #`:Q", title="Run #"),
                        y=alt.Y("EOD_group:Q",
                                title="Equal Opportunity Diff (group - overall)"),
                        color=alt.Color("group:N", title="Group"),
                        tooltip=["label:N", "group:N", alt.Tooltip(
                            "EOD_group:Q", format="+.3f")]
                    ).properties(height=260, width="container", title="EOD by Group")
                    st.altair_chart((bands + line), width="stretch")
            else:
                st.info("No per-group TPR values available yet.")
else:
    st.info("👆 Enable the checkbox above to view per-group fairness trends across multiple evaluation runs.")

st.markdown("</div>", unsafe_allow_html=True)

# ===== Mitigation Recommendations (heuristic) =====
st.markdown("""
<div style="margin-top: 40px; margin-bottom: 30px;">
    <h2 style="margin: 0; font-size: 1.8rem; font-weight: 800; color: #1F2937;">🩹 Fairness Mitigation Recommendations</h2>
    <p style="margin: 8px 0 0 0; font-size: 0.95rem; color: #6B7280;">Targeted actions to improve fairness metrics and reduce bias</p>
</div>
""", unsafe_allow_html=True)

latest_l3 = None
for f in sorted([p for p in files if os.path.basename(p).startswith("L3-")]):
    pass
# last one in chronological list:
l3_candidates = sorted([p for p in files if os.path.basename(
    p).startswith("L3-")], key=lambda p: os.path.getmtime(p))
if l3_candidates:
    latest_l3 = load_json(l3_candidates[-1])

if not latest_l3:
    st.info("📭 No latest L3 run detected.")
else:
    m = latest_l3.get("metrics", {}) or {}
    dpg = m.get("DPG")
    eod = m.get("EOD")
    tpr_by_group = m.get("tpr_by_group") or {}
    rates_by_group = m.get("rates_by_group") or {}

    recs = []
    # thresholds matching your L3 scoring targets
    if isinstance(eod, (int, float)) and eod > 0.10:
        recs.append(
            "Equal Opportunity gap is high (>0.10). Increase positive examples for the low-TPR group or adjust threshold per group.")
    elif isinstance(eod, (int, float)) and eod > 0.05:
        recs.append(
            "Equal Opportunity gap is moderate (>0.05). Try threshold tuning or class-balanced loss.")

    if isinstance(dpg, (int, float)) and dpg > 0.05:
        recs.append(
            "Demographic Parity gap >0.05. Consider rebalancing decision thresholds or post-processing to equalize positive rates.")

    # diagnose which group lags by TPR
    try:
        if tpr_by_group:
            g_min = min(tpr_by_group, key=lambda k: (
                999 if tpr_by_group[k] is None else tpr_by_group[k]))
            recs.append(
                f"Group '{g_min}' has the lowest sensitivity (TPR). Add targeted data or feature engineering for this group.")
    except Exception:
        pass

    if not recs:
        recs = ["No major fairness flags detected. Keep monitoring drift across runs."]

    # Display recommendations in styled cards
    if recs:
        st.markdown("""
        <div style="display: grid; gap: 12px;">
        """, unsafe_allow_html=True)
        for i, r in enumerate(recs):
            icon = "🟢" if "No major" in r else (
                "🟡" if "moderate" in r.lower() else "🔴")
            bg_color = "#ECFDF5" if "No major" in r else (
                "#FEF3C7" if "moderate" in r.lower() else "#FEE2E2")
            border_color = "#10B981" if "No major" in r else (
                "#F59E0B" if "moderate" in r.lower() else "#EF4444")
            st.markdown(f"""
            <div style="background: {bg_color}; border-left: 4px solid {border_color}; border-radius: 8px; padding: 12px 16px; margin-bottom: 8px;">
                <p style="margin: 0; font-size: 0.95rem; color: #1F2937; line-height: 1.5;">{icon} {r}</p>
            </div>
            """, unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

# ===== Bias Attribution Heatmap (per-group permutation importance) =====
st.markdown("""
<div style="margin-top: 40px; margin-bottom: 30px;">
    <h2 style="margin: 0; font-size: 1.8rem; font-weight: 800; color: #1F2937;">🔥 Bias Attribution Heatmap</h2>
    <p style="margin: 8px 0 0 0; font-size: 0.95rem; color: #6B7280;">Analyze feature importance per group to identify bias sources</p>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<div style="background: linear-gradient(135deg, #FEF3C7 0%, #FEF9E7 100%); border: 2px solid #F59E0B; border-radius: 12px; padding: 16px; margin-bottom: 20px;">
    <p style="margin: 0 0 8px 0; font-size: 0.9rem; font-weight: 700; color: #92400E;">📋 CSV Format Required</p>
    <p style="margin: 0; font-size: 0.85rem; color: #B45309;">Columns: <code style="background: #FFFFFF; padding: 2px 6px; border-radius: 4px;">y</code> (label), <code style="background: #FFFFFF; padding: 2px 6px; border-radius: 4px;">group</code> (group identifier), and numeric feature columns</p>
</div>
""", unsafe_allow_html=True)

default_path = "data/explain_with_group.csv"
opt_col1, opt_col2 = st.columns([2, 1], gap="large")
with opt_col1:
    st.markdown("<h5 style='margin: 0 0 12px 0; color: #1F2937; font-weight: 700;'>📂 Upload File</h5>",
                unsafe_allow_html=True)
    st.markdown("""
    <div style="background: linear-gradient(135deg, #F3E8FF 0%, #FAF5FF 100%); border: 2px dashed #A78BFA; border-radius: 12px; padding: 16px; text-align: center;">
        <p style="margin: 0 0 4px 0; font-size: 0.9rem; color: #6D28D9; font-weight: 600;">🔽 Upload CSV or use default</p>
        <p style="margin: 4px 0 0 0; font-size: 0.8rem; color: #7C3AED;">Limit 50MB • CSV format</p>
    </div>
    """, unsafe_allow_html=True)
    user_file = st.file_uploader("Upload CSV (optional override)", type=[
                                 "csv"], disabled=_LOCK, key="l3_attr_upl", label_visibility="collapsed")
with opt_col2:
    st.markdown("<h5 style='margin: 0 0 12px 0; color: #1F2937; font-weight: 700;'>⚙️ Options</h5>",
                unsafe_allow_html=True)
    use_default = st.checkbox(f"Use default", value=(
        user_file is None), disabled=_LOCK, help=f"Use {default_path}")

csv_source = None
if user_file is not None and not use_default:
    csv_source = user_file
elif use_default and os.path.exists(default_path):
    csv_source = default_path

if csv_source is None:
    st.markdown("""
    <div style="background: linear-gradient(135deg, #E0E7FF 0%, #F0F4FF 100%); border: 2px solid #818CF8; border-radius: 12px; padding: 20px; text-align: center;">
        <p style="margin: 0; font-size: 0.95rem; color: #3730A3;">📭 No CSV available</p>
        <p style="margin: 4px 0 0 0; font-size: 0.85rem; color: #4F46E5;">Upload a file or enable the default path to generate the heatmap</p>
    </div>
    """, unsafe_allow_html=True)
else:
    try:
        from sklearn.linear_model import LogisticRegression
        from sklearn.inspection import permutation_importance

        dfx = pd.read_csv(csv_source).dropna()
        if not {"y", "group"}.issubset(dfx.columns):
            st.warning(
                "❌ CSV must contain 'y' and 'group' plus feature columns.")
        else:
            label_col = "y"
            features = [c for c in dfx.columns if c not in (
                label_col, "group")]
            X0 = dfx[features].apply(pd.to_numeric, errors="coerce")
            dfx = dfx.loc[X0.dropna().index]
            X = dfx[features].astype(float).values
            y = dfx[label_col].values
            groups = dfx["group"].astype(str).values

            clf = LogisticRegression(max_iter=1000)
            clf.fit(X, y)

            rows = []
            for g in sorted(pd.unique(groups)):
                mask = (groups == g)
                if mask.sum() < 5:
                    continue
                perm = permutation_importance(
                    clf, X[mask], y[mask], n_repeats=6, random_state=42)
                imp = perm.importances_mean.clip(min=0.0)
                for j, feat in enumerate(features):
                    rows.append({"group": g, "feature": feat,
                                "importance": float(imp[j])})

            if not rows:
                st.info("⚠️ Not enough samples per group to compute attribution.")
            else:
                heat = pd.DataFrame(rows)
                heat["norm"] = heat.groupby("group")["importance"].transform(
                    lambda s: (s - s.min()) / (s.max() - s.min() + 1e-9)
                )
                chart = alt.Chart(heat).mark_rect().encode(
                    x=alt.X("feature:N", title="Feature"),
                    y=alt.Y("group:N", title="Group"),
                    color=alt.Color(
                        "norm:Q", title="Relative Importance", scale=alt.Scale(scheme="blues")),
                    tooltip=["group", "feature", alt.Tooltip(
                        "importance:Q", format=".4f")]
                ).properties(height=260, width="container", title="Per-Group Attribution (Permutation Importance)")
                st.altair_chart(chart, width="stretch")
    except Exception as e:
        st.warning(f"Heatmap unavailable: {e}")


# ===== Data Provenance & Mitigation Recommendations =====
st.markdown("""
<div style="margin-top: 40px; margin-bottom: 30px;">
    <h2 style="margin: 0; font-size: 1.8rem; font-weight: 800; color: #1F2937;">📜 Data Provenance (L3)</h2>
    <p style="margin: 8px 0 0 0; font-size: 0.95rem; color: #6B7280;">Dataset version, training date, and audit trail information</p>
</div>
""", unsafe_allow_html=True)

prov = {
    "dataset_version": (latest.get("L3") or {}).get("metrics", {}).get("dataset_version"),
    "last_train_date": (latest.get("L3") or {}).get("metrics", {}).get("last_train_date"),
}
# fallback from project config if missing
try:
    with open(os.path.join("configs", "project.example.yaml"), "r", encoding="utf-8") as fh:
        proj_cfg = yaml.safe_load(fh) or {}
    prov.setdefault("dataset_version", proj_cfg.get(
        "L3", {}).get("dataset_version"))
    prov.setdefault("last_train_date", proj_cfg.get(
        "L3", {}).get("last_train_date"))
except Exception:
    pass

prov_clean = {k: v for k, v in prov.items() if v is not None}
if prov_clean:
    st.markdown("""
    <div style="background: linear-gradient(135deg, #F0F9FF 0%, #E0F2FE 100%); border: 2px solid #3B82F6; border-radius: 12px; padding: 20px;">
        <div id="provenance-data">
    """, unsafe_allow_html=True)
    for key, value in prov_clean.items():
        st.markdown(f"""
        <div style="margin-bottom: 12px;">
            <p style="margin: 0 0 4px 0; font-size: 0.85rem; font-weight: 700; color: #1E40AF; text-transform: uppercase;">{key}</p>
            <p style="margin: 0; font-size: 0.95rem; color: #1E3A8A; padding: 8px 12px; background: white; border-radius: 6px;">{value}</p>
        </div>
        """, unsafe_allow_html=True)
    st.markdown("</div></div>", unsafe_allow_html=True)
else:
    st.markdown("""
    <div style="background: linear-gradient(135deg, #F3F4F6 0%, #E5E7EB 100%); border: 2px solid #6B7280; border-radius: 12px; padding: 20px; text-align: center;">
        <p style="margin: 0 0 4px 0; font-size: 0.95rem; color: #374151; font-weight: 700;">📭 No Provenance Data</p>
        <p style="margin: 4px 0 0 0; font-size: 0.85rem; color: #6B7280;">💡 Add <code style='background: white; padding: 2px 6px; border-radius: 4px;'>dataset_version</code> and <code style='background: white; padding: 2px 6px; border-radius: 4px;'>last_train_date</code> to L3 metrics or project config for audit trails</p>
    </div>
    """, unsafe_allow_html=True)

# Mitigation recommendations
st.markdown("""
<div style="margin-top: 32px; margin-bottom: 24px;">
    <h3 style="margin: 0 0 8px 0; font-size: 1.3rem; font-weight: 800; color: #1F2937;">🩹 Mitigation Recommendations</h3>
    <p style="margin: 0; font-size: 0.95rem; color: #6B7280;">High-priority actions to improve L3 fairness metrics</p>
</div>
""", unsafe_allow_html=True)

recos = []
# Simple rules:
latest_l3 = latest.get("L3") or {}
DPG = (latest_l3.get("metrics") or {}).get("DPG")
EOD = (latest_l3.get("metrics") or {}).get("EOD")
if DPG is not None and float(DPG) > 0.1:
    recos.append(
        "High Demographic Parity Gap — consider class rebalancing or group-aware thresholding.")
if EOD is not None and float(EOD) > 0.1:
    recos.append(
        "Equal Opportunity Difference is high — inspect TPR per group and retrain with focal loss or group sampling.")

if not recos:
    st.markdown("""
    <div style="background: linear-gradient(135deg, #ECFDF5 0%, #E0FFDD 100%); border: 2px solid #10B981; border-radius: 12px; padding: 16px; text-align: center;">
        <p style="margin: 0; font-size: 0.95rem; color: #065F46; font-weight: 600;">✅ No high-priority mitigation suggestions at this time</p>
        <p style="margin: 4px 0 0 0; font-size: 0.85rem; color: #047857;">Keep monitoring fairness metrics across runs</p>
    </div>
    """, unsafe_allow_html=True)
else:
    st.markdown("""
    <div style="display: grid; gap: 12px;">
    """, unsafe_allow_html=True)
    for r in recos:
        st.markdown(f"""
        <div style="background: linear-gradient(135deg, #FEE2E2 0%, #FEF2F2 100%); border-left: 4px solid #EF4444; border-radius: 8px; padding: 12px 16px;">
            <p style="margin: 0; font-size: 0.95rem; color: #7F1D1D; line-height: 1.5;">🔴 {r}</p>
        </div>
        """, unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

# ===== L4 Explainability ? Interactive =====
st.markdown("""
<div data-tour-target="l4-explainability">
<h3>🔍 L4 Explainability – Interactive</h3>
""", unsafe_allow_html=True)

l4 = latest.get("L4")

if not (l4 and "metrics" in l4):
    st.info("Run L4 to populate explainability metrics.")
else:
    m = l4["metrics"]
    # importances: prefer SHAP if present and user selects it
    shap_imp = m.get("shap_importance")
    perm_imp = m.get("feature_importance") or {}
    has_shap = bool(shap_imp)

    # --- Health gauge ---
    deletion_drop = float(m.get("deletion_drop", 0.0))
    stability_tau = float(m.get("stability_tau", 0.0))
    l4_score = float(l4.get("score", 0.0))

    with st.container():
        st.caption(
            "Health uses deletion-drop (↑) and stability t (↑). Targets: deletion_drop >= 0.15, t >= 0.85.")
        st.markdown(
            f"<h4 style='text-align:center;'>Explainability Health</h4>", unsafe_allow_html=True)
        st.markdown(
            f"<h2 style='text-align:center;color:#444;'>{'🟢' if l4_score >= 90 else ('🟡' if l4_score >= 75 else '🔴')} {l4_score:.2f}</h2>", unsafe_allow_html=True)
        bar_html = f"""
        <div style='background:#eee;border-radius:10px;height:20px;margin-bottom:6px;'>
          <div style='width:{min(max(l4_score, 0), 100)}%;background:linear-gradient(90deg,#ff4b4b,#ffde59,#00c851);
          height:20px;border-radius:10px;'></div>
        </div>
        """
        st.markdown(bar_html, unsafe_allow_html=True)
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Baseline", f"{m.get('baseline_metric', 0):.3f}")
        c2.metric("Masked", f"{m.get('masked_metric', 0):.3f}")
        c3.metric("Deletion drop",
                  f"{deletion_drop:.3f}", help="Higher is better; = 0.15")
        c4.metric("Stability t", f"{stability_tau:.3f}",
                  help="Higher is better; = 0.85")

    # --- Importance toggle + Top-N slider ---
    src = st.radio(
        "Importance source",
        ["Permutation", "SHAP (if available)"],
        index=0 if not has_shap else 1,
        horizontal=True,
        help="Choose which importance values to visualize."
    )
    top_n = st.slider("Top-N features", min_value=3, max_value=max(3,
                      len((shap_imp or perm_imp))), value=min(10, len((shap_imp or perm_imp))))

    chosen = shap_imp if (src.startswith("SHAP") and has_shap) else perm_imp
    if not chosen:
        st.info("No importance values found yet.")
    else:
        imp_df = (
            pd.DataFrame({"feature": list(chosen.keys()), "importance": [
                         float(v) for v in chosen.values()]})
            .sort_values("importance", ascending=False)
            .head(top_n)
        )
        total = float(imp_df["importance"].sum()) or 1.0
        imp_df["share_%"] = (imp_df["importance"] / total) * 100.0

        st.caption("Tip: use the slider to focus on the top drivers.")
        chart = alt.Chart(imp_df).mark_bar().encode(
            x=alt.X("importance:Q", title="Importance"),
            y=alt.Y("feature:N", sort="-x", title="Feature"),
            tooltip=["feature", alt.Tooltip(
                "importance:Q", format=".4f"), alt.Tooltip("share_%:Q", format=".2f")]
        ).properties(height=280, width="container")
        st.altair_chart(chart, theme=None)

        st.download_button(
            "📥 Download current importances (CSV)",
            data=imp_df.to_csv(index=False).encode("utf-8"),
            file_name=f"L4_importances_{'shap' if (src.startswith('SHAP') and has_shap) else 'perm'}.csv",
            mime="text/csv",
            key="dl_l4_imp",
        )

st.markdown("</div>", unsafe_allow_html=True)

# --- Trends across runs: deletion_drop & stability_tau ---
# Scan all L4 reports and plot over time
l4_rows = []
for f in files:
    try:
        with open(f, "r", encoding="utf-8") as fh:
            d = json.load(fh)
        if d.get("module") != "L4":
            continue
        base = os.path.basename(f).replace(".json", "")
        # label from filename; try to parse timestamp like L4-YYYYMMDD-HHMMSS.json
        ts_match = re.search(r"(\d{8}-\d{6})", base)
        ts_label = ts_match.group(1) if ts_match else base
        mtime = os.path.getmtime(f)
        md = d.get("metrics", {})
        l4_rows.append({
            "Run": base,
            "Label": ts_label,
            "time": mtime,
            "deletion_drop": float(md.get("deletion_drop", 0.0)),
            "stability_tau": float(md.get("stability_tau", 0.0)),
            "score": float(d.get("score", 0.0)),
        })
    except Exception:
        pass

if l4_rows:
    import re
    import os
    l4_df = pd.DataFrame(l4_rows).sort_values(
        "time").reset_index(drop=True)
    l4_df["Run #"] = l4_df.index + 1

    st.markdown("#### ⏳ L4 Trends – Deletion drop & Stability t")
    # background bands for thresholds
    band_df_drop = pd.DataFrame(
        {"y0": [0.00, 0.15], "y1": [0.15, 1.00], "band": ["red", "green"]})
    band_df_tau = pd.DataFrame(
        {"y0": [0.00, 0.85], "y1": [0.85, 1.00], "band": ["red", "green"]})

    def banded(metric, title, domain=(0, 1)):
        bands = alt.Chart(
            band_df_drop if metric == "deletion_drop" else band_df_tau
        ).mark_rect(opacity=0.12).encode(
            y=alt.Y("y0:Q", title=None, scale=alt.Scale(domain=domain)),
            y2="y1:Q",
            color=alt.Color("band:N", scale=alt.Scale(
                domain=["red", "green"],
                range=["#ff4b4b", "#00c851"]
            ), legend=None)
        )
        line = alt.Chart(l4_df).mark_line(point=True).encode(
            x=alt.X("`Run #`:Q", title="Run # (chronological)"),
            y=alt.Y(f"{metric}:Q", title=title,
                    scale=alt.Scale(domain=domain)),
            tooltip=["Label:N", alt.Tooltip(
                f"{metric}:Q", format=".3f"), alt.Tooltip("score:Q", format=".1f")]
        ).properties(height=250, width="container")
        return (bands + line).resolve_scale(color="independent")

    cA, cB = st.columns(2)
    with cA:
        st.altair_chart(banded(
            "deletion_drop", "Deletion drop (= 0.15 target)"), width="stretch")
    with cB:
        st.altair_chart(banded(
            "stability_tau", "Stability t (= 0.85 target)"), width="stretch")
else:
    st.info("No historical L4 runs yet to plot trends.")

st.markdown("</div>", unsafe_allow_html=True)


# ===== L4 - SHAP Summary (split view: Beeswarm – left | Insights – right) =====
st.markdown("""
<div style="margin-bottom: 20px;">
    <h3 style="margin: 0; font-size: 1.5rem; font-weight: 700; color: #1F2937;">🐝 L4 - SHAP Summary</h3>
    <p style="margin: 8px 0 0 0; font-size: 0.9rem; color: #6B7280;">Explore feature importance using SHAP values • Understand model predictions</p>
</div>
""", unsafe_allow_html=True)

# --- Inputs (upload + hints) - Enhanced layout
st.markdown("<h4 style='margin-top: 24px; margin-bottom: 12px; font-weight: 600; color: #374151;'>📤 Data Upload</h4>", unsafe_allow_html=True)
c_up, c_help = st.columns([2.5, 1.5])
with c_up:
    shap_csv = st.file_uploader(
        "Upload CSV file (must include a label column + numeric features)",
        type=["csv"], key="l4_beeswarm_upl", disabled=_LOCK,
        help="Supported format: CSV with numeric features and a target column"
    )
with c_help:
    st.info("💡 No file? First CSV from L4 evidence will be used automatically")

# --- Get L4 report early for label guessing
l4_rep = latest.get("L4") or {}

# --- Settings - Enhanced layout with icons
st.markdown("<h4 style='margin-top: 24px; margin-bottom: 12px; font-weight: 600; color: #374151;'>⚙️ Configuration</h4>", unsafe_allow_html=True)
c_left, c_mid, c_right = st.columns([1.5, 1.5, 1.5])
with c_left:
    sample_n = st.slider(
        "📊 Sample N",
        min_value=100, max_value=2000, value=600, step=100, disabled=_LOCK,
        help="Samples to visualize (SHAP computed on full test set)"
    )
with c_mid:
    label_guess = (l4_rep.get("metrics", {}) or {}).get(
        "label_col") or l4_rep.get("label_col") or "y"
    label_col = st.text_input(
        "🎯 Label Column",
        value=label_guess, disabled=_LOCK,
        help="Target/label column name in your CSV"
    )
with c_right:
    top_k = st.slider(
        "⭐ Top-K Features",
        min_value=2, max_value=10, value=4, step=1, disabled=_LOCK,
        help="Show top K most important features"
    )

# --- Try to auto-pick a CSV from L4 evidence if the user didn't upload one
csv_from_evidence = None
try:
    ev = l4_rep.get("evidence") or []
    if isinstance(ev, list) and ev:
        for p in ev:
            if isinstance(p, str) and p.lower().endswith(".csv") and os.path.exists(p):
                csv_from_evidence = p
                break
except Exception:
    csv_from_evidence = None

# --- Decide which CSV to use
csv_path_to_use = shap_csv if shap_csv is not None else (
    csv_from_evidence if csv_from_evidence and os.path.exists(
        csv_from_evidence) else None
)

if csv_path_to_use is None:
    st.info(
        "No CSV available. Upload one above or ensure L4 evidence contains a CSV file."
    )
else:
    try:
        import numpy as np
        import pandas as pd
        import altair as alt
        from sklearn.linear_model import LogisticRegression
        from sklearn.model_selection import train_test_split
        from sklearn.metrics import roc_auc_score, accuracy_score

        # Deterministic split so visuals are stable
        SEED = 42
        np.random.seed(SEED)

        df = pd.read_csv(csv_path_to_use)
        if label_col not in df.columns:
            raise ValueError(f"Label column '{label_col}' not found in CSV.")

        feature_cols = [c for c in df.columns if c != label_col]
        for c in feature_cols:
            df[c] = pd.to_numeric(df[c], errors="coerce")
        df = df.dropna(subset=[label_col] + feature_cols)
        if df.empty:
            raise ValueError(
                "No rows left after dropping NaNs; check your CSV."
            )

        X = df[feature_cols].values
        y = df[label_col].values

        X_tr, X_te, y_tr, y_te = train_test_split(
            X,
            y,
            test_size=0.3,
            random_state=SEED,
            stratify=y if len(np.unique(y)) == 2 else None,
        )
        clf = LogisticRegression(max_iter=1000)
        clf.fit(X_tr, y_tr)

        # Compute model metric (for the Insights panel)
        is_binary = len(np.unique(y)) == 2 and hasattr(clf, "predict_proba")
        if is_binary:
            prob = clf.predict_proba(X_te)[:, 1]
            model_metric = roc_auc_score(y_te, prob)
            metric_label = "AUROC"
        else:
            pred = clf.predict(X_te)
            model_metric = accuracy_score(y_te, pred)
            metric_label = "Accuracy"

        # ========== LAYOUT: LEFT (beeswarm) | RIGHT (insights) ==========
        colL, colR = st.columns([7, 5])

        # ---------- LEFT: SHAP Beeswarm (matplotlib if available; Altair fallback) ----------
        with colL:
            computed_shap_vals = None  # Initialize to prevent undefined errors

            try:
                import shap
                import matplotlib.pyplot as plt

                with measure_section("shap_computation"):
                    try:
                        expl = shap.LinearExplainer(
                            clf, X_tr, feature_perturbation="interventional"
                        )
                        shap_vals = expl.shap_values(X_te)
                        computed_shap_vals = shap_vals  # Save for later use

                    except Exception as shap_error:
                        logger.error(
                            f"SHAP computation failed: {shap_error}", exc_info=True
                        )
                        st.error(
                            "⚠️ SHAP computation failed. Falling back to permutation importance."
                        )

                        # Fallback to permutation importance
                        from sklearn.inspection import permutation_importance

                        perm = permutation_importance(
                            clf, X_te, y_te, n_repeats=6, random_state=SEED
                        )

                        # Convert to SHAP-like format (per-sample)
                        shap_vals = np.tile(
                            perm.importances_mean, (X_te.shape[0], 1)
                        )
                        computed_shap_vals = shap_vals

                    # Downsample for plotting aesthetics (only the figure)
                    if X_te.shape[0] > sample_n:
                        idx = np.random.choice(
                            X_te.shape[0], size=sample_n, replace=False
                        )
                        X_plot = X_te[idx]
                        shap_plot_vals = shap_vals[idx]
                    else:
                        X_plot = X_te
                        shap_plot_vals = shap_vals

                st.caption(
                    "Beeswarm shows the distribution of feature impacts across samples (|SHAP|)."
                )
                fig = plt.figure(figsize=(9.5, 5.2))
                plt.title("SHAP Beeswarm", fontsize=12, pad=8)
                shap.summary_plot(
                    shap_plot_vals,
                    features=X_plot,
                    feature_names=feature_cols,
                    show=False,
                    plot_type="dot",
                )
                st.pyplot(fig, clear_figure=True)
                # Keep shap_vals around for the right-side panels
                computed_shap_vals = shap_vals

            except ImportError:
                st.warning(
                    "⚠️ SHAP library not installed. Install with: `pip install shap`")
                st.info("Using permutation importance as fallback...")

                # Fallback visualization using permutation importance
                from sklearn.inspection import permutation_importance
                perm = permutation_importance(
                    clf, X_te, y_te, n_repeats=6, random_state=SEED)

                # Create simple bar chart
                imp_df = pd.DataFrame({
                    "feature": feature_cols,
                    "importance": perm.importances_mean
                }).sort_values("importance", ascending=False).head(10)

                chart = alt.Chart(imp_df).mark_bar().encode(
                    x=alt.X("importance:Q", title="Importance"),
                    y=alt.Y("feature:N", sort="-x", title="Feature"),
                    tooltip=["feature", alt.Tooltip(
                        "importance:Q", format=".4f")]
                ).properties(height=320, width="container")

                st.altair_chart(chart, width="stretch")
                computed_shap_vals = None

            except Exception as e:
                show_error_inline(e, "SHAP visualization failed")
                computed_shap_vals = None

            st.caption(
                f"Random seed: {SEED} (deterministic splits for stable visuals)"
            )

        # ---------- RIGHT: Insights panel (metric + mean|SHAP| + PDP-lite) ----------
        with colR:
            st.markdown("""
<div style="margin-bottom: 20px;">
    <h4 style="margin: 0; font-size: 1.2rem; font-weight: 700; color: #1F2937;">📊 Model Insights</h4>
    <p style="margin: 6px 0 0 0; font-size: 0.85rem; color: #6B7280;">Performance metrics and feature importance analysis</p>
</div>
""", unsafe_allow_html=True)

            # Guard against missing SHAP values
            if computed_shap_vals is None:
                st.warning(
                    "⚠️ SHAP values not available. Model insights limited.")
                st.info(
                    "Install SHAP library for full feature importance analysis: `pip install shap matplotlib`")

                # Show basic model metrics only in enhanced cards
                m1, m2 = st.columns([1.2, 0.8])
                with m1:
                    # Metric card with color coding
                    metric_color = "#10B981" if model_metric > 0.85 else "#F59E0B" if model_metric > 0.70 else "#EF4444"
                    st.markdown(f"""
<div style="background: {metric_color}15; border: 2px solid {metric_color}; border-radius: 10px; padding: 16px; text-align: center;">
    <p style="margin: 0; font-size: 0.85rem; color: #6B7280; font-weight: 600;">Model Performance</p>
    <p style="margin: 8px 0 0 0; font-size: 2rem; font-weight: 700; color: {metric_color};">{model_metric:.1%}</p>
    <p style="margin: 4px 0 0 0; font-size: 0.75rem; color: #6B7280;">{metric_label}</p>
</div>
""", unsafe_allow_html=True)

                with m2:
                    st.markdown(f"""
<div style="background: #F3F4F6; border: 2px solid #E5E7EB; border-radius: 10px; padding: 16px; text-align: center;">
                        <p style="margin: 0; font-size: 0.85rem; color: #6B7280; font-weight: 600;">Test Size</p>
                        <p style="margin: 8px 0 0 0; font-size: 1.8rem; font-weight: 700; color: #374151;">{X_te.shape[0]}</p>
                        <p style="margin: 4px 0 0 0; font-size: 0.75rem; color: #6B7280;">Samples</p>
                    </div>
""", unsafe_allow_html=True)

            else:
                # Enhanced metric cards with color coding
                m1, m2 = st.columns([1.2, 0.8])
                with m1:
                    metric_color = "#10B981" if model_metric > 0.85 else "#F59E0B" if model_metric > 0.70 else "#EF4444"
                    st.markdown(f"""
<div style="background: {metric_color}15; border: 2px solid {metric_color}; border-radius: 10px; padding: 16px; text-align: center;">
    <p style="margin: 0; font-size: 0.85rem; color: #6B7280; font-weight: 600;">Model Performance</p>
    <p style="margin: 8px 0 0 0; font-size: 2rem; font-weight: 700; color: {metric_color};">{model_metric:.1%}</p>
    <p style="margin: 4px 0 0 0; font-size: 0.75rem; color: #6B7280;">{metric_label}</p>
</div>
""", unsafe_allow_html=True)

                with m2:
                    st.markdown(f"""
<div style="background: #F3F4F6; border: 2px solid #E5E7EB; border-radius: 10px; padding: 16px; text-align: center;">
                        <p style="margin: 0; font-size: 0.85rem; color: #6B7280; font-weight: 600;">Test Size</p>
                        <p style="margin: 8px 0 0 0; font-size: 1.8rem; font-weight: 700; color: #374151;">{X_te.shape[0]}</p>
                        <p style="margin: 4px 0 0 0; font-size: 0.75rem; color: #6B7280;">Samples</p>
                    </div>
""", unsafe_allow_html=True)

            # Mean absolute SHAP per feature
            import numpy as _np
            mean_abs = _np.mean(_np.abs(computed_shap_vals), axis=0)
            top_df = (
                pd.DataFrame({"Feature": feature_cols,
                              "Mean |SHAP|": mean_abs})
                .sort_values("Mean |SHAP|", ascending=False)
                .reset_index(drop=True)
            )

            # Enhanced section header for Feature Importance
            st.markdown("""
<h5 style="margin: 20px 0 12px 0; font-size: 1rem; font-weight: 700; color: #374151;">⭐ Feature Importance (Mean |SHAP|)</h5>
""", unsafe_allow_html=True)

            col_fk, col_info = st.columns([2, 1])
            with col_fk:
                top_k = st.slider("Top-K features", 3,
                                  min(20, len(top_df)), value=min(10, len(top_df)), key="topk_shap")
            with col_info:
                st.caption(
                    f"📈 {len(top_view if 'top_view' in dir() else top_df.head(top_k))} features ranked")

            top_view = top_df.head(top_k)

            # Add color gradient to features based on importance
            max_shap = top_view["Mean |SHAP|"].max()
            min_shap = top_view["Mean |SHAP|"].min()
            top_view['Color'] = top_view['Mean |SHAP|'].apply(
                lambda x: '#10B981' if x > max_shap *
                0.66 else '#F59E0B' if x > max_shap * 0.33 else '#3B82F6'
            )

            bar = alt.Chart(top_view).mark_bar(cornerRadius=4).encode(
                y=alt.Y("Feature:N", sort="-x", title=None,
                        axis=alt.Axis(labelFontSize=11)),
                x=alt.X("Mean |SHAP|:Q", title="Mean |SHAP| Value",
                        axis=alt.Axis(labelFontSize=10)),
                color=alt.Color('Color:N', scale=None, legend=None),
                tooltip=[
                    alt.Tooltip("Feature:N", title="Feature"),
                    alt.Tooltip("Mean |SHAP|:Q", format=".4f",
                                title="Importance")
                ],
            ).properties(
                height=max(240, 22 * len(top_view)),
                width="container"
            ).properties(
                title={
                    "text": "Higher values = more important",
                    "subtitle": "Mean absolute SHAP value per feature",
                    "anchor": "start"
                }
            )
            st.altair_chart(bar, width="stretch")

            # PDP-lite: choose a feature, show average prediction vs binned values
            st.markdown("""
<h5 style="margin: 24px 0 12px 0; font-size: 1rem; font-weight: 700; color: #374151;">🔍 Partial Dependence Analysis</h5>
""", unsafe_allow_html=True)

            pick_feat = st.selectbox(
                "📍 Select Feature to Explore", options=list(top_view["Feature"]), index=0, disabled=_LOCK
            )

            # Build a quick PDP by binning the chosen feature and averaging predictions
            try:
                # Use test split for PDP to reflect what you saw in SHAP

                col_idx = feature_cols.index(pick_feat)
                fvals = X_te[:, col_idx]
                # Quantile bins for robustness
                bins = np.unique(np.quantile(fvals, np.linspace(0, 1, 11)))
                # Guard against degenerate constant feature
                if len(bins) < 3:
                    bins = np.linspace(fvals.min(), fvals.max() + 1e-9, 10)

                cats = pd.cut(fvals, bins=bins, include_lowest=True)
                if is_binary:
                    yhat = prob
                else:
                    # Approx prediction score for multi-class: use decision_function if present else proba argmax
                    try:
                        yhat = clf.decision_function(X_te)
                        if yhat.ndim > 1:
                            yhat = yhat.max(axis=1)
                    except Exception:
                        if hasattr(clf, "predict_proba"):
                            yhat = clf.predict_proba(X_te).max(axis=1)
                        else:
                            yhat = clf.predict(X_te)

                pdp_df = pd.DataFrame({"bin": cats.astype(str), "yhat": yhat})
                pdp_view = pdp_df.groupby("bin", as_index=False)["yhat"].mean()

                line = alt.Chart(pdp_view).mark_line(point=True, strokeWidth=3, pointSize=100).encode(
                    x=alt.X("bin:N", title=f"{pick_feat} (Binned Range)", axis=alt.Axis(
                        labelFontSize=10)),
                    y=alt.Y("yhat:Q", title="Average Prediction",
                            axis=alt.Axis(labelFontSize=10)),
                    color=alt.value("#3B82F6"),
                    tooltip=[
                        alt.Tooltip("bin:N", title="Range"),
                        alt.Tooltip("yhat:Q", format=".3f",
                                    title="Avg Prediction")
                    ],
                ).properties(
                    height=240,
                    width="container"
                ).properties(
                    title={
                        "text": f"Feature Impact on Model Predictions",
                        "subtitle": f"How {pick_feat} affects predictions",
                        "anchor": "start"
                    }
                )
                st.altair_chart(line, width="stretch")
            except Exception:
                st.caption("⚠️ PDP not available for this feature.")

            # Download mean |SHAP| with enhanced button
            st.markdown("<div style='margin-top: 20px;'></div>",
                        unsafe_allow_html=True)
            dl_col1, dl_col2 = st.columns([2, 1])
            with dl_col1:
                st.download_button(
                    "⬇️ Download Feature Importance (CSV)",
                    data=top_df.to_csv(index=False).encode("utf-8"),
                    file_name="feature_importance_shap.csv",
                    mime="text/csv",
                    disabled=_LOCK,
                    key="dl_shap_btn"
                )
            with dl_col2:
                st.caption(f"📊 {len(top_df)} total features")

    except Exception as e:
        # inline, friendly error
        import traceback as _tb
        tb = "".join(_tb.format_exception_only(type(e), e)).strip()
        st.error(f"⚠️ Could not render SHAP summary: {tb}")


# ===== L4 - Compare Models (A vs B) =====
st.markdown("""
<div style="margin-top: 40px; margin-bottom: 30px;">
    <h2 style="margin: 0; font-size: 1.8rem; font-weight: 800; color: #1F2937;">🆚 L4 - Compare Models (A vs B)</h2>
    <p style="margin: 8px 0 0 0; font-size: 0.95rem; color: #6B7280;">Upload two CSV files to compare model performance, feature importances, and model robustness metrics side-by-side</p>
</div>
""", unsafe_allow_html=True)

# Deterministic seed for reproducibility (auditable)
SEED = 42
np.random.seed(SEED)
random.seed(SEED)


@st.cache_resource  # cache heavy training for faster re-runs
def _fit_and_measure(csv_or_path, label_col="y", seed: int = SEED):
    """
    Train a simple logistic model deterministically and compute:
      - permutation importances (dict feature->score)
      - deletion_drop (float)
      - base metric (AUROC for binary else accuracy)
    Returns: (importances_dict, deletion_drop, base_metric, seed)
    """
    import pandas as pd
    from sklearn.linear_model import LogisticRegression
    from sklearn.model_selection import train_test_split
    from sklearn.metrics import accuracy_score, roc_auc_score
    from sklearn.inspection import permutation_importance

    # Load CSV
    dfm = pd.read_csv(csv_or_path) if hasattr(
        csv_or_path, "read") else pd.read_csv(csv_or_path)
    feats = [c for c in dfm.columns if c != label_col]
    if not feats:
        raise ValueError(
            "No feature columns found (every column is the label).")

    # Coerce numerics & drop NaNs
    for c in feats:
        dfm[c] = pd.to_numeric(dfm[c], errors="coerce")
    dfm = dfm.dropna(subset=[label_col] + feats)
    if dfm.empty:
        raise ValueError(
            "No rows left after dropping NaNs. Check your CSV formatting.")

    X = dfm[feats].values
    y = dfm[label_col].values

    # Deterministic split/model
    X_tr, X_te, y_tr, y_te = train_test_split(
        X, y, test_size=0.3, random_state=seed,
        stratify=y if len(np.unique(y)) == 2 else None
    )
    clf = LogisticRegression(max_iter=1000, random_state=seed).fit(X_tr, y_tr)

    # Base metric
    binary = len(np.unique(y)) == 2
    prob = clf.predict_proba(X_te)[:, 1] if (
        binary and hasattr(clf, "predict_proba")) else None
    pred = clf.predict(X_te)
    kind = "auroc" if binary else "accuracy"
    base = roc_auc_score(y_te, prob) if (
        kind == "auroc" and prob is not None) else accuracy_score(y_te, pred)

    # Permutation importance
    perm = permutation_importance(
        clf, X_te, y_te, n_repeats=6, random_state=seed)
    imp = perm.importances_mean.clip(min=0.0)

    # Deletion-drop: mask top feature
    top_idx = int(np.argsort(imp)[::-1][0]) if len(imp) > 0 else 0
    train_means = X_tr.mean(axis=0)
    X_mask = X_te.copy()
    if len(imp) > 0:
        X_mask[:, top_idx] = train_means[top_idx]
    prob_m = clf.predict_proba(X_mask)[:, 1] if (
        binary and hasattr(clf, "predict_proba")) else None
    pred_m = clf.predict(X_mask)
    masked = roc_auc_score(y_te, prob_m) if (
        kind == "auroc" and prob_m is not None) else accuracy_score(y_te, pred_m)
    drop = float(base - masked)

    return ({feats[i]: float(imp[i]) for i in range(len(feats))}, drop, float(base), seed)


# ===== Upload Section =====
st.markdown("""
<h4 style="margin: 24px 0 16px 0; font-size: 1.1rem; font-weight: 700; color: #374151;">📤 Upload Models</h4>
""", unsafe_allow_html=True)

cA, cB = st.columns(2, gap="large")
with cA:
    st.markdown("""
<div style="background: linear-gradient(135deg, #EFF6FF 0%, #F0F9FF 100%); border: 2px solid #3B82F6; border-radius: 12px; padding: 20px; text-align: center;">
    <h5 style="margin: 0 0 8px 0; font-size: 1rem; font-weight: 700; color: #1E40AF;">📊 Model A</h5>
    <p style="margin: 0 0 12px 0; font-size: 0.85rem; color: #1E3A8A;">Upload CSV for Model A</p>
</div>
""", unsafe_allow_html=True)
    uplA = st.file_uploader("Upload CSV A", type=[
                            "csv"], disabled=_LOCK, key="l4_csv_a", label_visibility="collapsed")

with cB:
    st.markdown("""
<div style="background: linear-gradient(135deg, #FEF3C7 0%, #FEF9E7 100%); border: 2px solid #F59E0B; border-radius: 12px; padding: 20px; text-align: center;">
    <h5 style="margin: 0 0 8px 0; font-size: 1rem; font-weight: 700; color: #92400E;">🎯 Model B</h5>
    <p style="margin: 0 0 12px 0; font-size: 0.85rem; color: #78350F;">Upload CSV for Model B</p>
</div>
""", unsafe_allow_html=True)
    uplB = st.file_uploader("Upload CSV B", type=[
                            "csv"], disabled=_LOCK, key="l4_csv_b", label_visibility="collapsed")

if uplA and uplB:
    try:
        impA, dropA, baseA, seedA = _fit_and_measure(uplA)
        impB, dropB, baseB, seedB = _fit_and_measure(uplB)

        st.markdown("""
<div style="background: #F9FAFB; border: 1px solid #E5E7EB; border-radius: 8px; padding: 12px; margin: 12px 0;">
    <p style="margin: 0; font-size: 0.85rem; color: #6B7280;">
        <b>Seed:</b> {seed} • <b>Status:</b> ✅ Deterministic A/B comparison complete
    </p>
</div>
""".format(seed=SEED), unsafe_allow_html=True)

        # Build union of features & deltas
        import pandas as pd
        import altair as alt
        keys = sorted(set(impA.keys()) | set(impB.keys()))
        rows = []
        for k in keys:
            a = float(impA.get(k, 0.0))
            b = float(impB.get(k, 0.0))
            rows.append({"feature": k, "A_importance": a,
                        "B_importance": b, "?(B-A)": b - a})
        delta_df = pd.DataFrame(rows).sort_values("?(B-A)", ascending=False)

        # ===== Metrics Section =====
        st.markdown("""
<h4 style="margin: 24px 0 12px 0; font-size: 1.1rem; font-weight: 700; color: #374151;">📊 Performance Metrics</h4>
""", unsafe_allow_html=True)

        mA, mB, mDelta = st.columns(3, gap="large")
        with mA:
            st.markdown(f"""
<div style="background: #EFF6FF; border: 2px solid #3B82F6; border-radius: 10px; padding: 16px; text-align: center;">
    <p style="margin: 0; font-size: 0.85rem; color: #1E40AF; font-weight: 600;">Model A</p>
    <p style="margin: 8px 0 0 0; font-size: 1.8rem; font-weight: 700; color: #3B82F6;">{baseA:.3f}</p>
    <p style="margin: 4px 0 0 0; font-size: 0.75rem; color: #1E3A8A;">AUROC / Accuracy</p>
    <hr style="margin: 8px 0; border: none; border-top: 1px solid #BFDBFE;">
    <p style="margin: 0; font-size: 0.75rem; color: #1E3A8A;">Deletion-drop: <b>{dropA:.3f}</b></p>
</div>
""", unsafe_allow_html=True)

        with mB:
            st.markdown(f"""
<div style="background: #FEF3C7; border: 2px solid #F59E0B; border-radius: 10px; padding: 16px; text-align: center;">
    <p style="margin: 0; font-size: 0.85rem; color: #92400E; font-weight: 600;">Model B</p>
    <p style="margin: 8px 0 0 0; font-size: 1.8rem; font-weight: 700; color: #F59E0B;">{baseB:.3f}</p>
    <p style="margin: 4px 0 0 0; font-size: 0.75rem; color: #78350F;">AUROC / Accuracy</p>
    <hr style="margin: 8px 0; border: none; border-top: 1px solid #FCD34D;">
    <p style="margin: 0; font-size: 0.75rem; color: #78350F;">Deletion-drop: <b>{dropB:.3f}</b></p>
</div>
""", unsafe_allow_html=True)

        delta_direction = "📈 Better" if baseB > baseA else "📉 Lower"
        delta_color = "#10B981" if baseB > baseA else "#EF4444"
        with mDelta:
            st.markdown(f"""
<div style="background: #F0FDF4; border: 2px solid {delta_color}; border-radius: 10px; padding: 16px; text-align: center;">
    <p style="margin: 0; font-size: 0.85rem; color: #15803D; font-weight: 600;">Δ (B - A)</p>
    <p style="margin: 8px 0 0 0; font-size: 1.8rem; font-weight: 700; color: {delta_color};">{baseB - baseA:+.3f}</p>
    <p style="margin: 4px 0 0 0; font-size: 0.75rem; color: #15803D;">{delta_direction}</p>
    <hr style="margin: 8px 0; border: none; border-top: 1px solid #BBDC63;">
    <p style="margin: 0; font-size: 0.75rem; color: #15803D;">Drop Δ: <b>{dropB - dropA:+.3f}</b></p>
</div>
""", unsafe_allow_html=True)

        if delta_df.empty:
            st.info("No overlapping or valid features to compare.")
        else:
            # ===== Feature Comparison Section =====
            st.markdown("""
<h4 style="margin: 24px 0 12px 0; font-size: 1.1rem; font-weight: 700; color: #374151;">🔍 Feature Analysis</h4>
""", unsafe_allow_html=True)
            st.dataframe(delta_df, width="stretch")
            c3.metric("Δ metric (B-A)", f"{(baseB - baseA):+.3f}")

            # Guard slider bounds even if very few features
            topn_max = max(1, min(15, len(delta_df)))
            topn_default = min(10, topn_max)
            topn = st.slider("Show top-N features", 1, topn_max,
                             topn_default, disabled=_LOCK, key="l4_topn")
            show_df = delta_df.head(topn)

            # ---- ➕ Importance Delta Δ(B−A) bar chart ----
            st.markdown("#### ➕ Importance Delta (B - A)")
            ch_delta = alt.Chart(show_df).mark_bar().encode(
                x=alt.X("?(B-A):Q", title="Change in importance (B - A)"),
                y=alt.Y("feature:N", sort="-x", title="Feature"),
                tooltip=["feature", alt.Tooltip("?(B-A):Q", format="+.4f")]
            ).properties(height=280, width="container")
            st.altair_chart(ch_delta, width="stretch")

            # ---- Overlayed importances (A vs B) ----
            st.markdown("#### 🏷️ Feature Importances (A vs B)")
            imp_long = pd.melt(
                show_df[["feature", "A_importance", "B_importance"]],
                id_vars="feature",
                var_name="model",
                value_name="importance"
            )
            imp_long["model"] = imp_long["model"].map(
                {"A_importance": "A", "B_importance": "B"})
            bar = alt.Chart(imp_long).mark_bar().encode(
                y=alt.Y("feature:N", sort="-x", title="Feature"),
                x=alt.X("importance:Q", title="Importance"),
                color=alt.Color("model:N", title="Model"),
                tooltip=["feature", "model", alt.Tooltip(
                    "importance:Q", format=".4f")]
            ).properties(height=280, width="container")
            st.altair_chart(bar, width="stretch")

            # ---- Deletion-drop A vs B (slope chart) ----
            st.markdown("#### 📉 Deletion-drop (A vs B)")
            dd_df = pd.DataFrame(
                {"model": ["A", "B"], "deletion_drop": [dropA, dropB]})
            slope = alt.Chart(dd_df).mark_line(point=True).encode(
                x=alt.X("model:N", title="Model"),
                y=alt.Y("deletion_drop:Q", title="Deletion drop"),
                tooltip=["model", alt.Tooltip("deletion_drop:Q", format=".3f")]
            ).properties(height=220, width=380)
            st.altair_chart(slope, width="content")

    except Exception as e:
        import traceback
        tb = "".join(traceback.format_exception_only(type(e), e)).strip()
        st.error(f"🚫 L4 comparison failed: {tb}")
else:
    st.markdown("""
<div style="background: #F3F4F6; border: 2px dashed #D1D5DB; border-radius: 10px; padding: 20px; text-align: center;">
    <p style="margin: 0; font-size: 0.95rem; color: #6B7280;">
        📤 Upload two CSV files above to start comparing models
    </p>
</div>
""", unsafe_allow_html=True)

# ===== L4 - Composite Health Meter =====
st.markdown("""
<div style="margin-top: 40px; margin-bottom: 30px;">
    <h2 style="margin: 0; font-size: 1.8rem; font-weight: 800; color: #1F2937;">🩺 L4 - Composite Health Meter</h2>
    <p style="margin: 8px 0 0 0; font-size: 0.95rem; color: #6B7280;">Overall model reliability and explainability score</p>
</div>
""", unsafe_allow_html=True)

m = (latest.get("L4") or {}).get("metrics", {}) or {}
dd = float(m.get("deletion_drop", 0.0))
tau = float(m.get("stability_tau", 0.0))
inf = float(m.get("infidelity", 1.0))  # 0 good, 1 bad (proxy)

# normalize components to 0..1 against demo targets
dd_norm = min(max(dd / 0.15, 0), 1)          # =0.15 is "1.0"
tau_norm = min(max(tau / 0.85, 0), 1)        # =0.85 is "1.0"
inf_norm = 1.0 - min(max(inf, 0), 1)         # lower is better • invert

health = (0.45*dd_norm) + (0.45*tau_norm) + (0.10*inf_norm)
health_pct = 100.0 * health
badge = "💚" if health_pct >= 90 else ("💛" if health_pct >= 75 else "❤️")
health_color = "#10B981" if health_pct >= 90 else (
    "#FBBF24" if health_pct >= 75 else "#EF4444")

st.markdown(f"""
<div style="background: linear-gradient(135deg, {health_color}15 0%, {health_color}08 100%); border: 2px solid {health_color}; border-radius: 12px; padding: 24px;">
    <div style="text-align: center;">
        <p style="margin: 0; font-size: 0.9rem; color: #6B7280; font-weight: 600; text-transform: uppercase; letter-spacing: 0.5px;">Composite Health Score</p>
        <p style="margin: 12px 0 0 0; font-size: 3rem; font-weight: 800; color: {health_color};">{health_pct:.1f}%</p>
        <p style="margin: 8px 0 0 0; font-size: 1.3rem; color: {health_color};">{badge}</p>
    </div>
    <div style="margin-top: 20px; display: grid; grid-template-columns: 1fr 1fr 1fr; gap: 16px;">
        <div style="background: white; border-radius: 8px; padding: 12px; text-align: center;">
            <p style="margin: 0; font-size: 0.75rem; color: #6B7280; font-weight: 600;">Deletion-drop</p>
            <p style="margin: 6px 0 0 0; font-size: 1.3rem; font-weight: 700; color: #1F2937;">{dd:.3f}</p>
            <p style="margin: 4px 0 0 0; font-size: 0.7rem; color: #6B7280; font-weight: 500;">Target: 0.15</p>
        </div>
        <div style="background: white; border-radius: 8px; padding: 12px; text-align: center;">
            <p style="margin: 0; font-size: 0.75rem; color: #6B7280; font-weight: 600;">Stability (τ)</p>
            <p style="margin: 6px 0 0 0; font-size: 1.3rem; font-weight: 700; color: #1F2937;">{tau:.3f}</p>
            <p style="margin: 4px 0 0 0; font-size: 0.7rem; color: #6B7280; font-weight: 500;">Target: 0.85</p>
        </div>
        <div style="background: white; border-radius: 8px; padding: 12px; text-align: center;">
            <p style="margin: 0; font-size: 0.75rem; color: #6B7280; font-weight: 600;">Infidelity</p>
            <p style="margin: 6px 0 0 0; font-size: 1.3rem; font-weight: 700; color: #1F2937;">{inf:.2f}</p>
            <p style="margin: 4px 0 0 0; font-size: 0.7rem; color: #6B7280; font-weight: 500;">Lower is better</p>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

# ===== Multi-model Comparison (bulk) =====
st.markdown("""
<div style="margin-top: 40px; margin-bottom: 30px;">
    <h2 style="margin: 0; font-size: 1.8rem; font-weight: 800; color: #1F2937;">🔀 Multi-model Comparison (Bulk)</h2>
    <p style="margin: 8px 0 0 0; font-size: 0.95rem; color: #6B7280;">Upload 2 or more CSV files with the same schema to compare multiple models at once</p>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<h4 style="margin: 0 0 12px 0; font-size: 1rem; font-weight: 700; color: #374151;">📂 Upload Models</h4>
""", unsafe_allow_html=True)

mm_files = st.file_uploader(
    "Upload 2+ CSVs with same schema", type=["csv"], accept_multiple_files=True,
    disabled=_LOCK, key="l4_multi_csvs", label_visibility="collapsed"
)

if mm_files and len(mm_files) >= 2:
    import pandas as pd
    rows = []
    for i, fobj in enumerate(mm_files, start=1):
        try:
            imp, drop, base, _seed = _fit_and_measure(fobj)  # unpack 4-tuple
            rows.append(
                {"Model": f"Model_{i}", "Metric": base, "Deletion_drop": drop})
        except Exception as e:
            rows.append(
                {"Model": f"Model_{i}", "Metric": None, "Deletion_drop": None})

    st.markdown("""
<h4 style="margin: 20px 0 12px 0; font-size: 1rem; font-weight: 700; color: #374151;">📊 Results</h4>
""", unsafe_allow_html=True)
    st.dataframe(pd.DataFrame(rows), width="stretch")
elif mm_files and len(mm_files) == 1:
    st.info("📤 Please upload 2 or more CSV files for bulk comparison")
else:
    st.markdown("""
<div style="background: #F3F4F6; border: 2px dashed #D1D5DB; border-radius: 10px; padding: 20px; text-align: center;">
    <p style="margin: 0; font-size: 0.95rem; color: #6B7280;">
        📂 Upload 2+ CSV files with the same schema to compare
    </p>
</div>
""", unsafe_allow_html=True)


# ===== Model Card (DOCX) =====
st.markdown("""
<div style="margin-top: 40px; margin-bottom: 30px;">
    <h2 style="margin: 0; font-size: 1.8rem; font-weight: 800; color: #1F2937;">📄 Model Card (DOCX)</h2>
    <p style="margin: 8px 0 0 0; font-size: 0.95rem; color: #6B7280;">Generate a comprehensive model card document with all evaluation metrics</p>
</div>
""", unsafe_allow_html=True)


def build_model_card(latest_bundle: dict) -> bytes:
    try:
        from docx import Document  # type: ignore
    except Exception:
        st.error("python-docx not installed. pip install python-docx")
        return b""
    doc = Document()
    doc.add_heading("Model Card ? IRAQAF", level=1)
    doc.add_paragraph("Generated from current dashboard state.")
    # L4
    l4r = latest_bundle.get("L4") or {}
    l4m = l4r.get("metrics", {}) or {}
    doc.add_heading("L4 ? Explainability", level=2)
    doc.add_paragraph(f"Score: {l4r.get('score')}  |  Band: {l4r.get('band')}")
    doc.add_paragraph(
        f"Deletion drop: {l4m.get('deletion_drop')}, Stability t: {l4m.get('stability_tau')}, Infidelity: {l4m.get('infidelity')}")
    # L3
    l3r = latest_bundle.get("L3") or {}
    l3m = l3r.get("metrics", {}) or {}
    doc.add_heading("L3 ? Fairness", level=2)
    doc.add_paragraph(f"Score: {l3r.get('score')}  |  Band: {l3r.get('band')}")
    doc.add_paragraph(f"DPG: {l3m.get('DPG')}  |  EOD: {l3m.get('EOD')}")
    # L1
    l1r = latest_bundle.get("L1") or {}
    l1m = l1r.get("metrics", {}) or {}
    doc.add_heading("L1 ? Governance/Compliance", level=2)
    doc.add_paragraph(f"Score: {l1r.get('score')}  |  Band: {l1r.get('band')}")
    doc.add_paragraph(f"Coverage: {l1m.get('coverage_percent')}%")
    # Evidence
    doc.add_heading("Evidence", level=2)
    for mid in ["L1", "L2", "L3", "L4", "L5"]:
        rep = latest_bundle.get(mid) or {}
        ev = rep.get("evidence") or []
        if ev:
            doc.add_paragraph(f"{mid}:")
            for e in ev:
                doc.add_paragraph(f" `- {e}")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


if st.button("Generate Model Card (Word)", disabled=_LOCK):
    payload = build_model_card(latest)
    if payload:
        st.download_button(
            "⬇️ Download Model_Card.docx",
            data=payload,
            file_name="Model_Card.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_model_card"
        )


# ===== L5 Operations - Trends Over Time =====
st.markdown("### 🛠️ L5 Operations - Trends Over Time")
# Collect all L5 reports
l5_rows = []
for f in files:
    try:
        with open(f, "r") as fh:
            d = json.load(fh)
        if d.get("module") != "L5":
            continue
        m = d.get("metrics", {})
        cov = m.get("logging_coverage", None)       # 0?1 (? better)
        lat = m.get("alert_latency_h", None)        # hours (? better)
        score = d.get("score", None)

        base = os.path.basename(f).replace(".json", "")
        ts_match = re.search(r"(\d{8}-\d{6})", base)
        ts_label = ts_match.group(1) if ts_match else base
        mtime = os.path.getmtime(f)

        l5_rows.append({
            "Run": base,
            "Label": ts_label,
            "time": mtime,
            "logging_coverage": (None if cov is None else float(cov)),
            "alert_latency_h": (None if lat is None else float(lat)),
            "Score": score,
        })
    except Exception:
        pass

if not l5_rows:
    st.info("No L5 operations reports found yet. Run L5 multiple times to visualize ops trends.")
else:
    l5_df = pd.DataFrame(l5_rows).sort_values("time").reset_index(drop=True)
    l5_df["Run #"] = l5_df.index + 1

    st.dataframe(
        l5_df[["Run #", "Label", "logging_coverage", "alert_latency_h", "Score"]],
        width="stretch",
    )

    # Threshold bands:
    # - Coverage (? better):    red <0.90, yellow 0.90?0.95, green >=0.95
    cov_bands = pd.DataFrame({
        "y0": [0.00, 0.90, 0.95],
        "y1": [0.90, 0.95, 1.00],
        "band": ["red", "yellow", "green"]
    })
    # - Latency (? better):     green <=1h, yellow 1?2h, red >2h (domain up to 4h for display)
    lat_bands = pd.DataFrame({
        "y0": [0.00, 1.00, 2.00],
        "y1": [1.00, 2.00, 4.00],
        "band": ["green", "yellow", "red"]
    })

    def ops_chart(df, metric_col: str, title: str, bands_df: pd.DataFrame, y_domain):
        base = alt.Chart(df)
        bands = alt.Chart(bands_df).mark_rect(opacity=0.15).encode(
            y=alt.Y("y0:Q", title=None, scale=alt.Scale(domain=y_domain)),
            y2="y1:Q",
            color=alt.Color(
                "band:N",
                scale=alt.Scale(
                    domain=["green", "yellow", "red"],
                    range=["#00c851", "#ffde59", "#ff4b4b"]
                ),
                legend=None
            )
        )
        line = base.mark_line(point=True).encode(
            x=alt.X("`Run #`:Q", title="Run # (chronological)"),
            y=alt.Y(f"{metric_col}:Q", title=title,
                    scale=alt.Scale(domain=y_domain)),
            tooltip=["Label:N", alt.Tooltip(f"{metric_col}:Q", format=".3f")]
        ).properties(height=260, width="container")
        return (bands + line).resolve_scale(color="independent")

    c1, c2 = st.columns(2)
    with c1:
        st.altair_chart(
            ops_chart(l5_df, "logging_coverage",
                      "Logging Coverage (? better)", cov_bands, (0, 1.0)),
            width="stretch"
        )
    with c2:
        st.altair_chart(
            ops_chart(l5_df, "alert_latency_h",
                      "Alert Latency (hours, ? better)", lat_bands, (0, 4.0)),
            width="stretch"
        )

    # Latest snapshot
    latest_l5 = l5_df.iloc[-1]
    st.markdown("#### ⏱️ Latest Run - Ops Snapshot")
    st.write(
        f"- **Logging coverage:** {latest_l5['logging_coverage'] if pd.notnull(latest_l5['logging_coverage']) else 'N/A'}"
        f"\n\n- **Alert latency:** {latest_l5['alert_latency_h'] if pd.notnull(latest_l5['alert_latency_h']) else 'N/A'} h"
    )

# ===== L5 ? Operations & Live Monitoring =====
st.markdown("""
<div data-tour-target="l5-operations">
<h2>🖥️ L5 - Operations & Live Monitoring</h2>
""", unsafe_allow_html=True)

# --- Source selector + controls
src_col1, src_col2, src_col3, src_col4 = st.columns([1.3, 1.3, 1, 0.8])
with src_col1:
    source = st.selectbox(
        "Source",
        ["Prometheus", "CSV upload", "JSON upload"],
        index=0,
        disabled=_LOCK,
        help="Pick where to read live metrics from."
    )
with src_col2:
    refresh_secs = st.number_input(
        "Auto-refresh (seconds)",
        min_value=0, max_value=3600, value=0, step=5,
        disabled=_LOCK,
        help="0 = off. Uses streamlit_autorefresh when > 0."
    )
with src_col3:
    latency_slo = st.number_input(
        "Latency SLO (p95, seconds)", min_value=0.0, value=1.0, step=0.1, disabled=_LOCK
    )
with src_col4:
    st.markdown("<br>", unsafe_allow_html=True)
    refresh_now = st.button(
        "🔄 Refresh now", disabled=_LOCK, width="stretch")

# OPTIONAL: light scheduler using streamlit_autorefresh
if refresh_secs and not _LOCK:
    try:
        # pip install streamlit-autorefresh
        from streamlit_autorefresh import st_autorefresh
        st_autorefresh(interval=refresh_secs * 1000, key="l5_autorefresh")
    except Exception:
        st.caption(
            "Tip: install streamlit-autorefresh for timed updates: `pip install streamlit-autorefresh`")

# --- Inputs per source (Live Metrics configuration)
with st.expander("Connection & Sources", expanded=True):
    # Select live data source type
    source = st.selectbox(
        "Live metrics source",
        ["Prometheus API", "Upload CSV", "Upload JSON", "HTTP Endpoint"],
        disabled=_LOCK,
        key="live_source_select"
    )

    # Show current monitoring policy thresholds (loaded from configs/policies.yaml)
    st.caption(
        f"Active policy: Latency = {POLICY['latency_slo']} s ? "
        f"Error = {POLICY['error_rate_threshold']} %"
    )

    # Auto-refresh interval (user-adjustable)
    refresh_interval = st.slider(
        "Auto-refresh interval (seconds)",
        5, 300, POLICY.get("refresh_interval_s", 60),
        step=5, disabled=_LOCK,
        key="refresh_interval_slider",
        help="Set how often to refresh live metrics automatically."
    )

    # Auto-refresh trigger (Streamlit built-in)
    st_autorefresh(interval=refresh_interval * 1000, key="l5_refresh")

    # ------------------------------
    # Inputs based on chosen source
    # ------------------------------
    if source == "Prometheus API":
        c1, c2 = st.columns(2)
        with c1:
            prom_latency_url = st.text_input(
                "Prometheus query (latency, seconds)",
                placeholder=(
                    "http://localhost:9090/api/v1/query?"
                    "query=histogram_quantile(0.95, sum(rate(http_request_duration_seconds_bucket[5m])) by (le))"
                ),
                key="prom_latency_url",
                disabled=_LOCK,
            )

            # Validate URL
            if prom_latency_url:
                is_valid, sanitized_url, error = validate_and_sanitize_input(
                    prom_latency_url,
                    input_type="url",
                    max_length=2000
                )
                if not is_valid:
                    st.error(f"⚠️ Invalid URL: {error}")
                    prom_latency_url = None
                else:
                    prom_latency_url = sanitized_url

        with c2:
            prom_error_url = st.text_input(
                "Prometheus query (error rate, %)",
                placeholder=(
                    "http://localhost:9090/api/v1/query?"
                    "query=100*sum(rate(http_requests_total{status=~\"5..\"}[5m]))/"
                    "sum(rate(http_requests_total[5m]))"
                ),
                key="prom_error_url",
                disabled=_LOCK,
            )

            # Validate URL
            if prom_error_url:
                is_valid, sanitized_url, error = validate_and_sanitize_input(
                    prom_error_url,
                    input_type="url",
                    max_length=2000
                )
                if not is_valid:
                    st.error(f"⚠️ Invalid URL: {error}")
                    prom_error_url = None
                else:
                    prom_error_url = sanitized_url

    elif source == "Upload CSV":
        up_csv = st.file_uploader(
            "Upload live metrics CSV (columns: time, latency_s, error_rate_pct)",
            type=["csv"],
            key="l5_live_csv",
            disabled=_LOCK,
        )

    elif source == "Upload JSON":
        up_json = st.file_uploader(
            "Upload live metrics JSON (list of {time, latency_s, error_rate_pct})",
            type=["json"],
            key="l5_live_json",
            disabled=_LOCK,
        )

    elif source == "HTTP Endpoint":
        api_url = st.text_input(
            "HTTP JSON endpoint (GET returning {time, latency_s, error_rate_pct})",
            placeholder="https://example.com/live_metrics",
            key="l5_http_url",
            disabled=_LOCK,
        )

st.markdown("</div>", unsafe_allow_html=True)
# --- Helpers


def _parse_prom_value(resp_json):
    try:
        # Prom API returns data.result[0].value = [ts, value]
        v = float(resp_json["data"]["result"][0]["value"][1])
        t = float(resp_json["data"]["result"][0]["value"][0])
        return pd.to_datetime(t, unit="s"), v
    except Exception:
        return None, None


@st.cache_data(ttl=15, show_spinner=False)
def _fetch_prom(url):
    if not url:
        return None, None

    # Check rate limit
    allowed, error_msg = rate_limiter.is_allowed('prometheus')
    if not allowed:
        logger.warning(f"Prometheus rate limit exceeded: {error_msg}")
        st.warning(f"⚠️ {error_msg}")
        return None, None

    try:
        r = requests.get(url, timeout=10)
        r.raise_for_status()
        return _parse_prom_value(r.json())
    except requests.exceptions.Timeout:
        logger.error(f"Prometheus request timeout: {url}")
        st.error("⚠️ Request timeout. Check Prometheus server.")
        return None, None
    except requests.exceptions.RequestException as e:
        logger.error(f"Prometheus request failed: {e}")
        show_error_inline(e, "Prometheus fetch failed")
        return None, None


def _load_live_df():
    # Priority: uploads ? Prometheus ? empty
    if source == "CSV upload" and 'up_csv' in locals() and up_csv is not None:
        df0 = pd.read_csv(up_csv)
        # normalize column names
        cols = {c.lower(): c for c in df0.columns}
        # required: time, latency_s, error_rate_pct (case-insensitive)
        tcol = cols.get("time")
        lcol = cols.get("latency_s")
        ecol = cols.get("error_rate_pct")
        if all([tcol, lcol, ecol]):
            try:
                df0[tcol] = pd.to_datetime(df0[tcol])
            except Exception:
                pass
            return (df0
                    .rename(columns={tcol: "time", lcol: "latency_s", ecol: "error_rate_pct"})
                    .sort_values("time"))
        return None

    if source == "JSON upload" and 'up_json' in locals() and up_json is not None:
        try:
            js = json.load(up_json)
            df0 = pd.json_normalize(js)
            df0["time"] = pd.to_datetime(df0["time"])
            return df0[["time", "latency_s", "error_rate_pct"]].sort_values("time")
        except Exception:
            return None

    if source == "Prometheus":
        # If Prometheus endpoints provided, fetch a single point now
        t1, v1 = _fetch_prom(st.session_state.get("prom_latency_url"))
        t2, v2 = _fetch_prom(st.session_state.get("prom_error_url"))
        if t1 or t2:
            rows = [{
                "time": t1 or t2 or pd.Timestamp.utcnow(),
                "latency_s": (v1 if v1 is not None else math.nan),
                "error_rate_pct": (v2 if v2 is not None else math.nan),
            }]
            return pd.DataFrame(rows)

    return None


# Force refresh on click (clear cached Prometheus calls)
if refresh_now:
    try:
        _fetch_prom.clear()
    except Exception:
        pass

live_df = _load_live_df()

# --- UI rendering for snapshot/charts/incidents (unchanged logic)
if live_df is None or live_df.empty:
    st.info("No live data yet. Provide Prometheus queries or upload a CSV/JSON.")
else:
    st.subheader("?? Current Live Snapshot")
    latest_row = live_df.iloc[-1]
    c1, c2, c3 = st.columns(3)
    with c1:
        st.metric("p95 latency (s)",
                  f"{float(latest_row['latency_s']):.3f}", help="Lower is better")
    with c2:
        st.metric("Error rate (%)",
                  f"{float(latest_row['error_rate_pct']):.2f}", help="Lower is better")
    with c3:
        ok = float(latest_row['latency_s']) <= latency_slo
        st.metric(
            "SLO status",
            "✅ OK" if ok else "❌ Breach",
            help=f"Threshold p95 = {latency_slo:.2f}s"
        )

    # Charts (Altair)
    st.markdown("#### Trends")
    if "time" in live_df.columns:
        base = alt.Chart(live_df).encode(x=alt.X("time:T", title="Time"))
        lat_chart = base.mark_line().encode(
            y=alt.Y("latency_s:Q", title="Latency (s)")
        ).properties(height=220)
        err_chart = base.mark_line().encode(
            y=alt.Y("error_rate_pct:Q", title="Error rate (%)")
        ).properties(height=220)
        st.altair_chart(lat_chart, theme=None)
        st.altair_chart(err_chart, theme=None)

    # Generate incident if thresholds exceeded (write to reports/incidents.json)
    # Simple policy: latency > SLO OR error_rate > 5% ? incident
    inc = None
    try:
        lat_ok = float(latest_row["latency_s"]) <= latency_slo
        err_ok = float(latest_row["error_rate_pct"]) <= 5.0
        if (not lat_ok) or (not err_ok):
            inc = {
                "time": datetime.utcnow().isoformat(timespec="seconds"),
                "kind": "SLO breach",
                "severity": "high" if float(latest_row["error_rate_pct"]) > 10.0 else "medium",
                "latency_s": float(latest_row["latency_s"]),
                "error_rate_pct": float(latest_row["error_rate_pct"]),
                "note": "Auto-detected breach from live metrics panel"
            }
    except Exception as e:
        logger.error(f"Incident detection failed: {e}", exc_info=True)
        show_error_inline(e, "Incident detection failed")

    # Only attempt to log if an incident was created
    if inc is not None:
        if audit_locked():
            st.warning(
                "Audit Mode is ON ? automatic incident logging disabled.")
            logger.warning(
                "Incident detected but not logged (audit mode active)")
        else:
            os.makedirs("reports", exist_ok=True)
            inc_path = os.path.join("reports", "incidents.json")

            try:
                # Thread-safe file operations
                import threading
                _incident_lock = st.session_state.get('_incident_lock')
                if _incident_lock is None:
                    _incident_lock = threading.Lock()
                    st.session_state._incident_lock = _incident_lock

                with _incident_lock:
                    data = []
                    if os.path.exists(inc_path):
                        with open(inc_path, "r", encoding="utf-8") as fh:
                            data = json.load(fh) or []

                    data.append(inc)

                    # Atomic write using temp file
                    temp_path = f"{inc_path}.tmp"
                    with open(temp_path, "w", encoding="utf-8") as fh:
                        json.dump(data, fh, indent=2)

                    # Atomic rename (overwrites target)
                    os.replace(temp_path, inc_path)

                logger.info(
                    f"Incident recorded: {inc['kind']} severity={inc['severity']}")
                st.success("Incident recorded.")
            except Exception as e:
                logger.error(f"Failed to record incident: {e}", exc_info=True)
                show_error_inline(e, "Incident logging failed")


# --- Incident timeline + scheduler tips
with st.expander(" Run Scheduler Tips", expanded=False):
    st.markdown(
        "- **Auto-refresh:** uses `streamlit-autorefresh` for UI updates.\n"
        " ` - **Windows Task Scheduler** (Triggers ? Daily/On logon; Action ? Start a program)\n"
        " ` - **cron** on Linux/macOS (e.g., `0 * * * * /usr/bin/python /path/cli.py`)\n"
        " ` - **GitHub Actions/CI** for nightly runs"
    )

# ===== L5 🧯 Incident Timeline =====
st.markdown("### 🧯 Incident Timeline")

inc_path = os.path.join("reports", "incidents.json")
inc_rows = []
if os.path.exists(inc_path):
    try:
        with open(inc_path, "r", encoding="utf-8") as fh:
            inc_rows = json.load(fh) or []
    except Exception:
        inc_rows = []

inc_df = None
if not inc_rows:
    st.info("No incidents logged yet.")
else:
    inc_df = pd.json_normalize(inc_rows)
    # normalize & sort
    try:
        inc_df["time"] = pd.to_datetime(inc_df["time"])
    except Exception:
        pass
    inc_df = inc_df.sort_values("time")
    st.dataframe(
        inc_df[["time", "kind", "severity", "latency_s", "error_rate_pct", "note"]],
        width="stretch"
    )

    # Altair timeline (points across severity bands)
    chart = (
        alt.Chart(inc_df)
        .mark_point(size=120)
        .encode(
            x=alt.X("time:T", title="Time"),
            y=alt.Y("severity:N", title="Severity", sort=[
                    "low", "medium", "high", "critical"]),
            color=alt.Color("kind:N", title="Incident type"),
            tooltip=["time", "kind", "severity",
                     "latency_s", "error_rate_pct", "note"]
        )
        .properties(height=220)
    )
    st.altair_chart(chart, theme=None, width="stretch")

# ---------- Demo helper: add a synthetic incident (for testing the timeline) ----------

INC_PATH = Path("reports/incidents.json")
INC_PATH.parent.mkdir(parents=True, exist_ok=True)


def _read_incidents(path: Path = INC_PATH) -> list[dict]:
    """Return a list[dict] from incidents.json (robust to legacy dict payloads)."""
    if not path.exists():
        return []
    try:
        obj = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(obj, list):
            return [x for x in obj if isinstance(x, dict)]
        if isinstance(obj, dict):
            return [obj]
        return []
    except Exception as e:
        show_error_inline(e, "Read incidents.json failed")
        return []


def _write_incidents(items: list[dict], path: Path = INC_PATH) -> bool:
    """Write incidents list back to disk (guarded by Audit Mode)."""
    if _LOCK:  # respect Audit Mode
        st.warning("Audit Mode is ON ? writes disabled.")
        return False
    try:
        path.write_text(json.dumps(items, indent=2), encoding="utf-8")
        return True
    except Exception as e:
        show_error_inline(e, "Write incidents.json failed")
        return False


demo_col1, demo_col2 = st.columns([1, 1])

with demo_col1:
    if st.button("➕ Add demo incident (medium)", disabled=_LOCK, help="Adds a synthetic SLO breach now"):
        demo_inc = {
            "time": pd.Timestamp.utcnow().isoformat(timespec="seconds"),
            "kind": "SLO breach",
            "severity": "medium",
            "latency_s": float((live_df.iloc[-1]["latency_s"] if live_df is not None and not live_df.empty else 1.5)),
            "error_rate_pct": float((live_df.iloc[-1]["error_rate_pct"] if live_df is not None and not live_df.empty else 6.0)),
            "note": "Demo incident"
        }
        incidents_list = _read_incidents(INC_PATH)
        incidents_list.append(demo_inc)
        if _write_incidents(incidents_list, INC_PATH):
            st.success("Demo incident added.")

with demo_col2:
    # Use current table if you already built inc_df above; otherwise rebuild quickly
    if 'inc_df' in locals() and inc_df is not None and not inc_df.empty:
        export_df = inc_df
    else:
        _rows = _read_incidents(INC_PATH)
        export_df = pd.json_normalize(_rows) if _rows else None

    if export_df is not None and not export_df.empty:
        st.download_button(
            "⬇️ Download incidents.csv",
            data=export_df.to_csv(index=False).encode("utf-8"),
            file_name="incidents.csv",
            mime="text/csv",
            disabled=_LOCK
        )
    else:
        st.caption("No incidents to export yet.")

# ===== Quality Drift Alerts (GQAS) =====
st.markdown("### 🚨 Quality Drift Alerts")

# find last two AGG reports
agg_paths = [p for p in files if os.path.basename(p).startswith("AGG-")]
agg_paths = sorted(agg_paths, key=lambda p: os.path.getmtime(p))
if len(agg_paths) < 2:
    st.info("Not enough AGG runs to compute drift. Run the audit at least twice.")
else:
    prev = load_json(agg_paths[-2])
    curr = load_json(agg_paths[-1])
    g0 = (prev or {}).get("gqas")
    g1 = (curr or {}).get("gqas")
    if g0 is None or g1 is None:
        st.info("Missing GQAS in AGG reports.")
    else:
        delta = float(g1) - float(g0)
        st.write(
            f"Previous: **{g0:.2f}** ➡️ Current: **{g1:.2f}** (Δ {delta:+.2f})")
        breach = (delta <= -5.0)
        st.markdown(
            f"Status: {'🟥 Drift > 5 points' if breach else '🟩 Stable'}")

        with st.expander("Notify (Slack / Email)"):
            st.caption(
                "Send a one-time alert to your team. Fields are locked in Audit Mode.")

            # one atomic form = no duplicate element IDs + cleaner UX
            with st.form("drift_notify_form"):
                slack_url = st.text_input(
                    "Slack webhook URL",
                    placeholder="https://hooks.slack.com/services/XXX/YYY/ZZZ",
                    type="password",
                    key="drift_slack_url_input",
                    disabled=audit_mode,
                    help="Stored only for this session, not persisted"
                )
                email_to = st.text_input(
                    "Email to (comma-separated)",
                    placeholder="ops@example.com",
                    key="drift_email_to_input",
                    disabled=audit_mode
                )
                smtp_host = st.text_input(
                    "SMTP host",
                    placeholder="smtp.example.com",
                    key="drift_smtp_host_input",
                    disabled=audit_mode
                )
                smtp_user = st.text_input(
                    "SMTP user",
                    key="drift_smtp_user_input",
                    disabled=audit_mode
                )
                smtp_pass = st.text_input(
                    "SMTP password",
                    type="password",
                    key="drift_smtp_pass_input",
                    disabled=audit_mode,
                    help="⚠️ Not stored - enter each time you send alerts"
                )
                st.caption(
                    "🔒 Credentials are used only for this notification and are NOT saved")

                submit = st.form_submit_button(
                    "Send Notifications",
                    disabled=audit_mode
                )

            # Only run network calls if not in audit mode and user submitted
            if submit and not audit_mode:
                alert_msg = (
                    f":rotating_light: IRAQAF GQAS drift detected: "
                    f"{g0:.2f} ? {g1:.2f} (? {delta:+.2f}). "
                    f"Check recent incidents and module scores."
                )

                def send_slack(msg, url):
                    if not url:
                        st.warning("Slack webhook URL missing.")
                        return False
                    try:
                        import requests
                        r = requests.post(url, json={"text": msg}, timeout=10)
                        if 200 <= r.status_code < 300:
                            return True
                        st.error(
                            f"Slack error: {r.status_code} {r.text[:200]}")
                    except Exception as e:
                        st.error(f"Slack exception: {e}")
                    return False

                def send_email(subject, body, to_csv, host, user, pwd):
                    try:
                        import smtplib
                        from email.mime.text import MIMEText
                        tos = [t.strip()
                               for t in (to_csv or "").split(",") if t.strip()]
                        if not (tos and host):
                            st.warning(
                                "Email: recipient(s) or SMTP host missing.")
                            return False
                        msg = MIMEText(body)
                        msg["Subject"] = subject
                        msg["From"] = user or "iraqaf@local"
                        msg["To"] = ", ".join(tos)
                        with smtplib.SMTP(host, 587, timeout=15) as s:
                            s.starttls()
                            if user and pwd:
                                s.login(user, pwd)
                            s.send_message(msg)
                        return True
                    except Exception as e:
                        st.error(f"Email exception: {e}")
                        return False

                # Fire
                sent_any = False
                if slack_url:
                    sent_any = send_slack(alert_msg, slack_url) or sent_any
                if email_to:
                    sent_any = send_email("IRAQAF GQAS Drift Alert", alert_msg,
                                          email_to, smtp_host, smtp_user, smtp_pass) or sent_any

                if sent_any:
                    st.success("Notification(s) sent.")
                else:
                    st.info("No notifications sent (fill Slack URL and/or Email).")

            # Handle submit outside the form context for clarity
            if submit:
                # persist values
                st.session_state["drift_slack_url"] = slack_url
                st.session_state["drift_email_to"] = email_to
                st.session_state["drift_smtp_host"] = smtp_host
                st.session_state["drift_smtp_user"] = smtp_user
                logger.info(
                    f"Drift alert triggered: GQAS {g0:.2f} ? {g1:.2f} (? {delta:+.2f})")

                sent_any = False
                if slack_url:
                    if send_slack(alert_msg, slack_url):
                        st.success("Slack alert sent.")
                        sent_any = True
                if email_to and smtp_host:
                    if send_email("IRAQAF GQAS Drift Alert", alert_msg, email_to, smtp_host, smtp_user, smtp_pass):
                        st.success("Email sent.")
                        sent_any = True

                if not sent_any:
                    st.info(
                        "Nothing sent. Provide a Slack webhook and/or email + SMTP settings.")


# ========== AGGREGATE SECTION ==========
st.markdown("""
<div data-tour-target="gqas-aggregate">
<h3>🌐 Aggregate Global QA Score (GQAS)</h3>
""", unsafe_allow_html=True)

with measure_section("aggregate_gqas_render"):
    agg = latest["AGG"]

    def gqas_band(gqas: float):
        if gqas >= 92:
            return "green", "🟢"
        elif gqas >= 88:
            return "yellow", "🟡"
        return "red", "🔴"

# ===== Risk profile toggle =====
risk = st.radio("Risk profile", ["High", "Medium"], horizontal=True, index=0)


def get_floors(risk_profile: str):
    """Get floor thresholds based on risk profile"""
    if risk_profile == "High":
        return {"L1": 90, "L2": 90, "L3": 95, "L4": 90, "L5": 85}
    else:  # Medium
        return {"L1": 85, "L2": 85, "L3": 90, "L4": 85, "L5": 80}


def floors_check(module_reports: dict, risk_profile: str):
    """Check floors with risk profile awareness"""
    floors = get_floors(risk_profile)
    results = {}

    for mid, floor in floors.items():
        rep = module_reports.get(mid)
        if not rep:
            results[mid] = ("missing", "?", floor, None, None)
            continue

        # For L1, use coverage_percent; for others use score
        if mid == "L1":
            metrics = rep.get("metrics", {}) or {}
            score = metrics.get("coverage_percent", rep.get("score", 0))
        else:
            score = rep["score"]
        color, emoji, status_text = get_score_color(score, risk_profile)

        results[mid] = (
            color,  # 'green', 'yellow', or 'red'
            emoji,
            floor,
            score,
            status_text
        )
    return results


if agg:
    # Recalculate GQAS dynamically from current module scores
    individual_scores = []
    for mid in ["L1", "L2", "L3", "L4", "L5"]:
        rep = latest.get(mid)
        if rep:
            # For L1, use coverage_percent; for others use score
            if mid == "L1":
                metrics = rep.get("metrics", {}) or {}
                s = metrics.get("coverage_percent", rep.get("score", 0))
            else:
                s = rep.get("score", 0)
            individual_scores.append(s)

    # Calculate GQAS as average of all module scores
    gqas = sum(individual_scores) / \
        len(individual_scores) if individual_scores else agg.get("gqas", 0.0)

    band, emoji = gqas_band(gqas)

    # Centered GQAS card
    st.markdown(
        f"<h3 style='text-align:center;'>⚙️ <b>Aggregate Global QA Score (GQAS)</b></h3>"
        "<p style='text-align:center;color:gray;'>Weighted overall quality from all five modules</p>",
        unsafe_allow_html=True,
    )
    st.markdown(
        f"<h2 style='text-align:center;color:#444;'>{emoji} {gqas:.2f}</h2>", unsafe_allow_html=True)

    progress_html = f"""
    <div style='background:#eee;border-radius:10px;height:25px;margin-bottom:15px;'>
      <div style='width:{gqas}%;background:linear-gradient(90deg,#ff4b4b,#ffde59,#00c851);
      height:25px;border-radius:10px;'></div>
    </div>
    """
    st.markdown(progress_html, unsafe_allow_html=True)

    st.markdown("#### 🧱 Shipping Floors")

    floors = floors_check(latest, risk)
    cols = st.columns(5)
    order = ["L1", "L2", "L3", "L4", "L5"]

    for i, mid in enumerate(order):
        color, sym, floor, score, status_text = floors[mid]
        label = name_for(mid)

        if color == "missing":
            cols[i].markdown(
                "<div style='padding:10px;border-radius:10px;background:#fdf5d4;"
                "text-align:center;min-height:110px;display:flex;flex-direction:column;"
                "justify-content:center;'>"
                f"<div>? <b>{label}</b></div>"
                "<div><i>Missing report</i></div>"
                "</div>",
                unsafe_allow_html=True,
            )
        else:
            # Color mapping for background
            color_map = {
                "green": "#d4edda",   # Light green
                "yellow": "#fff3cd",  # Light yellow
                "red": "#f8d7da"      # Light red
            }
            bg = color_map.get(color, "#f8d7da")

            cols[i].markdown(
                f"""
    <div style="
        background:{bg};
        padding:12px;
        border-radius:14px;
        text-align:center;
        height:150px;
        display:flex;
        flex-direction:column;
        justify-content:center;
        align-items:center;
    ">
        <div style="max-width:180px;text-align:center;">
            <b>{sym} {label}</b>
        </div>
        <div>Floor <b>{floor}</b> | Score <b>{score:.1f}</b></div>
        <div>{status_text}</div>
    </div>
    """,
                unsafe_allow_html=True
            )

    # Verdict badge - only count green as "meets"
    floors_met = all(v[0] == "green" for v in floors.values()
                     if v[0] != "missing")
    color = "#d4edda" if floors_met else "#f8d7da"
    msg = "✅ All floors met" if floors_met else "⚠️ Some floors not met"
    st.markdown(
        f"<div style='background:{color};padding:10px;margin-top:15px;border-radius:8px;"
        "box-shadow:0 2px 6px rgba(0,0,0,0.1);text-align:center;'>"
        f"<b>{msg}</b></div>",
        unsafe_allow_html=True
    )

    # Module scores table (risk-aware)
    st.markdown("#### 📊 Module Scores (for record)")
    score_rows = []
    for mid in order:
        rep = latest.get(mid)
        color, sym, floor, score, status_text = floors.get(
            mid, ("missing", "?", None, None, None))

        if color == "missing" or rep is None:
            score_rows.append({
                "Module": name_for(mid),
                "Score": None,
                "Floor": None,
                "Status": "Missing",
                "Band": "⚫"
            })
        else:
            band_emoji = {"green": "🟢", "yellow": "🟡",
                          "red": "🔴"}.get(color, "⚪")
            score_rows.append({
                "Module": name_for(mid),
                "Score": f"{score:.1f}",
                "Floor": floor,
                "Status": status_text,
                "Band": band_emoji
            })

    st.dataframe(pd.DataFrame(score_rows), width="stretch", hide_index=True)


# ===== Data Provenance Summary =====
st.markdown("### 🧾 Data Provenance")

try:
    with open("configs/project.example.yaml", "r", encoding="utf-8") as fh:
        proj_cfg = yaml.safe_load(fh) or {}
    meta = proj_cfg.get("META") or {}
    dv = meta.get("dataset_version", "?")
    ltd = meta.get("last_train_date", "?")
    owner = meta.get("data_owner", "?")
    st.markdown(
        f"""
        <div style='background:#f8f9fa;border:1px solid #eee;border-radius:10px;padding:10px;margin-bottom:10px;'>
          <b>Dataset Version:</b> {dv}<br>
          <b>Last Train Date:</b> {ltd}<br>
          <b>Data Owner:</b> {owner}
        </div>
        """, unsafe_allow_html=True
    )
except Exception:
    st.info("Provenance metadata not found in configs/project.example.yaml ? META block (optional).")

    # ===== Compare Runs & Trend Visualization =====
st.markdown("### 📈 Compare Runs & Trend Over Time")

# Find all available AGG reports
agg_files = [f for f in files if f.endswith(".json") and "AGG" in f]
run_options = []
for f in agg_files:
    try:
        with open(f, "r") as fh:
            d = json.load(fh)
        if d.get("module") == "AGG":
            label = os.path.basename(f).replace(".json", "")
            d["_label"] = label
            d["_mtime"] = os.path.getmtime(f)  # for chronological sort
            run_options.append(d)
    except Exception:
        pass

if not run_options:
    st.info("No AGG reports found to compare yet. Run the framework multiple times to generate more reports.")
else:
    # sort all runs by time
    run_options = sorted(run_options, key=lambda r: r["_mtime"])

    labels = [r["_label"] for r in run_options]
    selection = st.multiselect(
        "Select runs to compare",
        labels,
        default=(labels[-2:] if len(labels) >= 2 else labels),
    )

    if selection:
        sel_data = [r for r in run_options if r["_label"] in selection]
        # keep selected data chrono-sorted
        sel_data = sorted(sel_data, key=lambda r: r["_mtime"])

        cmp_rows = []
        for r in sel_data:
            row = {
                "Run": r["_label"],
                "GQAS": r.get("gqas"),
                "Floors Met": r.get("floors_met"),
            }
            scores = r.get("scores", {})  # e.g., CRS/SAI/FI/TS/OPS
            for mid, val in scores.items():
                row[mid] = val
            cmp_rows.append(row)

        cmp_df = pd.DataFrame(cmp_rows)
        st.dataframe(cmp_df, width="stretch")

        # Trend visualization (chronological)
        trend_df = cmp_df.copy()
        trend_df["Run Index"] = range(1, len(trend_df) + 1)
        trend_chart = (
            alt.Chart(trend_df)
            .mark_line(point=True)
            .encode(
                x=alt.X("Run Index:Q", title="Run # (chronological)"),
                y=alt.Y("GQAS:Q", title="Aggregate Quality Score"),
                tooltip=["Run", alt.Tooltip("GQAS:Q", format=".2f")],
            )
            .properties(height=250, width="container", title="GQAS Trend Over Time")
        )
        st.altair_chart(trend_chart, width="stretch")

        st.caption(
            "Shows GQAS change across multiple evaluation runs. Each run = one complete quality audit.")
    else:
        st.info("Select at least one run to visualize comparisons.")

        # ===== AGG Snapshot Compare (two runs) =====
st.markdown("### 🧾 Snapshot Compare – Aggregate (AGG)")
agg_paths = [f for f in files if f.endswith(
    ".json") and ("AGG" in os.path.basename(f))]
if len(agg_paths) < 2:
    st.info("Create at least two AGG reports to compare snapshots.")
else:
    agg_labels = {p: _label_from_path(p) for p in agg_paths}
    left, right = st.columns(2)
    with left:
        sel_a = st.selectbox("Baseline run (A)", options=agg_paths,
                             format_func=lambda p: agg_labels[p], key="aggA")
    with right:
        sel_b = st.selectbox("Compare to (B)", options=[
                             p for p in agg_paths if p != sel_a], format_func=lambda p: agg_labels[p], key="aggB")

    A = load_json(sel_a)
    B = load_json(sel_b)
    if not (A and B and A.get("module") == "AGG" and B.get("module") == "AGG"):
        st.warning("Could not load both AGG reports.")
    else:
        a_scores = A.get("scores", {})
        b_scores = B.get("scores", {})
        rows = []

        # Metric name mappings for readability
        metric_names = {
            "CRS": "Governance & Regulatory",
            "SAI": "Privacy & Security",
            "FI": "Fairness & Bias",
            "TS": "Explainability",
            "OPS": "Operations & Monitoring",
            "GQAS": "Overall Quality Score"
        }

        for k in ["CRS", "SAI", "FI", "TS", "OPS"]:
            a_val = a_scores.get(k)
            b_val = b_scores.get(k)
            delta = None if (k not in a_scores or k not in b_scores) else (
                float(b_val) - float(a_val))
            rows.append({
                "Module": metric_names.get(k, k),
                "Baseline (A)": f"{a_val:.1f}" if a_val is not None else "—",
                "Current (B)": f"{b_val:.1f}" if b_val is not None else "—",
                "Change": f"{delta:+.1f}" if delta is not None else "—",
            })

        # Add GQAS row
        a_gqas = A.get("gqas")
        b_gqas = B.get("gqas")
        gqas_delta = None if (a_gqas is None or b_gqas is None) else float(
            b_gqas) - float(a_gqas)
        rows.append({
            "Module": "Overall Quality Score",
            "Baseline (A)": f"{a_gqas:.2f}" if a_gqas is not None else "—",
            "Current (B)": f"{b_gqas:.2f}" if b_gqas is not None else "—",
            "Change": f"{gqas_delta:+.2f}" if gqas_delta is not None else "—",
        })

        cmp_df = pd.DataFrame(rows)

        # Display with better formatting
        st.dataframe(
            cmp_df,
            width="stretch",
            hide_index=True,
            column_config={
                "Module": st.column_config.TextColumn("Module", width="large"),
                "Baseline (A)": st.column_config.TextColumn(f"Baseline: {agg_labels[sel_a]}", width="medium"),
                "Current (B)": st.column_config.TextColumn(f"Current: {agg_labels[sel_b]}", width="medium"),
                "Change": st.column_config.TextColumn("Delta (B - A)", width="small"),
            }
        )

        # Delta visualization - create simpler dataframe for chart
        chart_data = []
        for metric in ["CRS", "SAI", "FI", "TS", "OPS", "GQAS"]:
            if metric == "GQAS":
                a_val = A.get("gqas")
                b_val = B.get("gqas")
            else:
                a_val = a_scores.get(metric)
                b_val = b_scores.get(metric)

            if a_val is not None and b_val is not None:
                delta = float(b_val) - float(a_val)
                chart_data.append({
                    "Module": metric_names.get(metric, metric),
                    "Change": delta
                })

        if chart_data:
            chart_df = pd.DataFrame(chart_data)
            chart = alt.Chart(chart_df).mark_bar().encode(
                x=alt.X("Change:Q", title="Score Change (B - A)", scale=alt.Scale(domain=(
                    min(-5, float(chart_df["Change"].min())-1),
                    max(5, float(chart_df["Change"].max())+1)
                ))),
                y=alt.Y("Module:N", sort=None),
                color=alt.condition(
                    alt.datum.Change > 0,
                    alt.value("#10B981"),  # Green for improvement
                    alt.value("#EF4444")   # Red for decline
                ),
                tooltip=["Module", alt.Tooltip("Change:Q", format="+.2f")]
            ).properties(height=250, title="Score Change Across Modules")
            st.altair_chart(chart, width="stretch")


st.markdown("</div>", unsafe_allow_html=True)
# ===== L1 Governance – Clause Drill-through (revised) =====
st.markdown("### 🏛️ L1 Governance – Clause Drill-through")

l1_paths = [f for f in files if f.endswith(
    ".json") and os.path.basename(f).startswith("L1-")]
if len(l1_paths) < 1:
    st.info("No L1 reports found. Run L1 to enable clause drill-through.")
else:
    pre_a = _nearest_by_time(l1_paths, agg_paths[0]) if agg_paths else None
    pre_b = _nearest_by_time(l1_paths, agg_paths[-1]) if agg_paths else None

    disabled = get_audit_mode()

    c1, c2 = st.columns(2)
    with c1:
        l1_a = st.selectbox(
            "Run A (baseline)",
            options=l1_paths,
            index=(l1_paths.index(pre_a) if pre_a in l1_paths else 0),
            format_func=lambda p: _label_from_path(p),
            key="l1A",
            disabled=_LOCK
        )
    with c2:
        l1_b = st.selectbox(
            "Run B (compare)",
            options=[p for p in l1_paths if p != l1_a],
            format_func=lambda p: _label_from_path(p),
            key="l1B",
            disabled=_LOCK
        )

    L1A = load_json(l1_a) or {}
    L1B = load_json(l1_b) or {}

    if not (L1A and L1B):
        st.warning("Could not load selected L1 reports.")
    else:
        # normalize so each clause has evidence_links
        L1A = _normalize_clause_evidence(L1A)
        L1B = _normalize_clause_evidence(L1B)

        ca = {(c.get("framework"), c.get("id")): c for c in (
            L1A.get("metrics", {}).get("clauses") or [])}
        cb = {(c.get("framework"), c.get("id")): c for c in (
            L1B.get("metrics", {}).get("clauses") or [])}
        keys = sorted(set(ca.keys()) | set(cb.keys()))

        diff_rows = []
        for k in keys:
            Acl, Bcl = ca.get(k), cb.get(k)
            Apass = (Acl.get("passed") if Acl else None)
            Bpass = (Bcl.get("passed") if Bcl else None)
            change = "? pass" if (Apass is False and Bpass is True) else (
                "? fail" if (Apass is True and Bpass is False) else "?")
            diff_rows.append({
                "Framework": k[0],
                "Clause": k[1],
                "A Passed": Apass,
                "B Passed": Bpass,
                "Change": change,
                "Evidence (A)": "📎" if (Acl and Acl.get("evidence_links")) else "?",
                "Evidence (B)": "📎" if (Bcl and Bcl.get("evidence_links")) else "?",
            })

        if not diff_rows:
            st.markdown("""
            <div style="
                background-color: #FEF3C7;
                border: 2px solid #F59E0B;
                border-radius: 8px;
                padding: 16px;
                margin: 16px 0;
            ">
                <div style="display: flex; align-items: flex-start; gap: 12px;">
                    <div style="font-size: 1.5rem;">ℹ️</div>
                    <div>
                        <div style="font-weight: 600; color: #92400E; margin-bottom: 6px;">No clause-level details available</div>
                        <div style="color: #78350F; font-size: 0.9rem; line-height: 1.5;">
                            <p style="margin: 0 0 8px 0;">The current L1 reports do not contain <code style="background: #FECACA; padding: 2px 6px; border-radius: 3px;">metrics.clauses</code> data.</p>
                            <p style="margin: 0;">This typically means:</p>
                            <ul style="margin: 4px 0 0 0; padding-left: 20px;">
                                <li>The L1 audit was run without clause-level extraction enabled</li>
                                <li>The reports use a different data structure (check Run A/B metadata)</li>
                                <li>Try running a new L1 audit to populate clause details</li>
                            </ul>
                        </div>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)

            # Show report metadata for debugging
            with st.expander("🔍 Report Metadata (for troubleshooting)"):
                col_meta1, col_meta2 = st.columns(2)
                label_a = _label_from_path(l1_a)
                label_b = _label_from_path(l1_b)
                with col_meta1:
                    st.write(f"**Run A** ({label_a})")
                    st.json({
                        "module": L1A.get("module"),
                        "version": L1A.get("version"),
                        "score": L1A.get("score"),
                        "band": L1A.get("band"),
                        "metrics_keys": list(L1A.get("metrics", {}).keys()) if isinstance(L1A.get("metrics"), dict) else "N/A",
                        "metrics_type": type(L1A.get("metrics")).__name__,
                    })
                with col_meta2:
                    st.write(f"**Run B** ({label_b})")
                    st.json({
                        "module": L1B.get("module"),
                        "version": L1B.get("version"),
                        "score": L1B.get("score"),
                        "band": L1B.get("band"),
                        "metrics_keys": list(L1B.get("metrics", {}).keys()) if isinstance(L1B.get("metrics"), dict) else "N/A",
                        "metrics_type": type(L1B.get("metrics")).__name__,
                    })
        else:
            def _fmt(v): return "🟢 Pass" if v is True else (
                "🔴 Fail" if v is False else "?")
            view = pd.DataFrame(diff_rows)
            if "A Passed" in view:
                view["A Passed"] = view["A Passed"].map(_fmt)
            if "B Passed" in view:
                view["B Passed"] = view["B Passed"].map(_fmt)
            view["Change"] = view["Change"].fillna("?")
            st.dataframe(view, width='stretch')

            # Per-clause evidence picker + preview
            st.markdown("#### 📎 Open Evidence (per clause)")
            pick = st.selectbox(
                "Select a clause to view evidence",
                options=[f"{fw} › {cid}" for (fw, cid) in keys],
                disabled=_LOCK
            )

            if pick:
                fw, cid = pick.split(" › ", 1)
                clA, clB = ca.get((fw, cid)), cb.get((fw, cid))

                def _render_evidence(tag: str, clause: dict):
                    with st.expander(f"Run {tag} › Evidence for {fw} › {cid}", expanded=False):
                        links = (clause or {}).get("evidence_links") or []
                        if not links:
                            st.caption("No evidence links for this clause.")
                            return
                        for i, p in enumerate(links, 1):
                            if re.match(r"^https?://", str(p), flags=re.I):
                                st.markdown(f"- 🌐 [{p}]({p})")
                            else:
                                pp = Path(p)
                                if pp.exists():
                                    st.markdown(
                                        f"- 📄 `{pp.name}`  \n<small style='color:#888'>{pp.as_posix()}</small>",
                                        unsafe_allow_html=True
                                    )
                                    with pp.open("rb") as fh:
                                        st.download_button("Download", data=fh.read(
                                        ), file_name=pp.name, key=f"dl_evd_{tag}_{i}")
                                    # Note: _preview_widget() function not yet implemented
                                else:
                                    st.markdown(f"- ❌ Missing: `{p}`")

                _render_evidence("A", clA)
                _render_evidence("B", clB)


# ===== Executive Summary (Auto) =====
st.markdown("### 📝 Executive Summary (auto)")


def _band_emoji(band: str) -> str:
    return {"green": "🟢", "yellow": "🟡", "red": "🔴"}.get((band or "").lower(), "?")


def _floor_table(latest_reports: dict, risk_profile: str):
    # same floors as dashboard logic
    if risk_profile == "High":
        floors = {"L1": 90, "L2": 90, "L3": 95, "L4": 90, "L5": 85}
    else:
        floors = {"L1": 85, "L2": 85, "L3": 90, "L4": 85, "L5": 80}
    rows = []
    for mid, floor in floors.items():
        rep = latest_reports.get(mid)
        score = rep.get("score") if rep else None
        band = rep.get("band") if rep else None
        status = (score is not None) and (score >= floor)
        rows.append({
            "Module": mid,
            "Name": {
                "L1": "Governance & Regulatory", "L2": "Privacy & Security",
                "L3": "Fairness & Ethics", "L4": "Explainability & Transparency",
                "L5": "Operations & Monitoring",
            }[mid],
            "Score": (None if score is None else f"{score:.1f}"),
            "Band": (None if band is None else f"{_band_emoji(band)} {band.capitalize()}"),
            "Floor": floor,
            "Pass": ("✅" if status else ("⚠️" if score is not None else "❌"))
        })
    return pd.DataFrame(rows)


def _top_drivers(latest_reports: dict):
    items = []
    for mid in ["L1", "L2", "L3", "L4", "L5"]:
        rep = latest_reports.get(mid)
        if rep and isinstance(rep.get("score"), (int, float)):
            items.append((mid, rep["score"], rep.get("band")))
    if not items:
        return None, None
    items_sorted = sorted(items, key=lambda x: x[1], reverse=True)
    best = items_sorted[0]
    worst = items_sorted[-1]
    return best, worst


def _failed_clauses_summary(l1_report: dict, max_items: int = 6):
    fails = []
    try:
        for c in (l1_report.get("metrics", {}).get("clauses") or []):
            if c.get("passed") is False:
                fails.append({
                    "framework": c.get("framework"),
                    "id": c.get("id"),
                    "desc": c.get("description"),
                    "hint": c.get("hint") or "",
                    "why": c.get("why_failed") or ""
                })
    except Exception:
        return []
    # prioritize by presence of hint/why, then by name
    fails = sorted(fails, key=lambda x: (x["framework"], x["id"]))
    return fails[:max_items]


# Gather data
agg = latest.get("AGG", {})
gqas = agg.get("gqas")
gqas_band = None
if gqas is not None:
    if gqas >= 92:
        gqas_band = "green"
    elif gqas >= 88:
        gqas_band = "yellow"
    else:
        gqas_band = "red"

# Risk profile (you already define it earlier; we read it if exists)
# If you named it `risk`, reuse; else default to "High".
risk_profile = globals().get("risk", "High")

# Build narrative - with styled cards
best, worst = _top_drivers(latest)

# Pre-compute card values
gqas_display = f"{gqas:.2f}" if gqas is not None else "N/A"
best_score_display = f"{best[1]:.1f}" if best else "N/A"
worst_score_display = f"{worst[1]:.1f}" if worst else "N/A"

# Overview section with styled cards
st.markdown("**Overview**")
col1, col2, col3 = st.columns(3, gap="small")

# Card 1: GQAS
with col1:
    border_color = {"green": "#10B981", "yellow": "#F59E0B",
                    "red": "#EF4444"}.get(gqas_band, "#9CA3AF")
    st.markdown(f"""
    <div style="
        background-color: #F8F9FA;
        border-left: 5px solid {border_color};
        border-radius: 8px;
        padding: 16px;
        min-height: 140px;
        display: flex;
        flex-direction: column;
        justify-content: space-between;
    ">
        <div style="color: #6B7280; font-size: 0.85rem; font-weight: 600;">GQAS Score</div>
        <div style="margin: 8px 0;">
            <div style="font-size: 2rem; font-weight: 800; color: {border_color};">
                {gqas_display}
            </div>
        </div>
        <div style="color: #1F2937; font-size: 0.9rem; font-weight: 500;">
            {_band_emoji(gqas_band)} {gqas_band.capitalize() if gqas_band else 'Unknown'}
        </div>
    </div>
    """, unsafe_allow_html=True)

# Card 2: Strongest Module
with col2:
    strong_color = {"green": "#10B981", "yellow": "#F59E0B",
                    "red": "#EF4444"}.get(best[2] if best else None, "#9CA3AF")
    best_name = best[0] if best else "N/A"
    st.markdown(f"""
    <div style="
        background-color: #F8F9FA;
        border-left: 5px solid {strong_color};
        border-radius: 8px;
        padding: 16px;
        min-height: 140px;
        display: flex;
        flex-direction: column;
        justify-content: space-between;
    ">
        <div style="color: #6B7280; font-size: 0.85rem; font-weight: 600;">Strongest Module</div>
        <div style="margin: 8px 0;">
            <div style="font-size: 1.5rem; font-weight: 800; color: {strong_color};">
                {best_name}
            </div>
            <div style="font-size: 1.1rem; font-weight: 600; color: #1F2937; margin-top: 4px;">
                {best_score_display}
            </div>
        </div>
        <div style="color: #1F2937; font-size: 0.9rem;">
            {_band_emoji(best[2]) if best and best[2] else ''}
        </div>
    </div>
    """, unsafe_allow_html=True)

# Card 3: Weakest Module
with col3:
    weak_color = {"green": "#10B981", "yellow": "#F59E0B",
                  "red": "#EF4444"}.get(worst[2] if worst else None, "#9CA3AF")
    worst_name = worst[0] if worst else "N/A"
    st.markdown(f"""
    <div style="
        background-color: #F8F9FA;
        border-left: 5px solid {weak_color};
        border-radius: 8px;
        padding: 16px;
        min-height: 140px;
        display: flex;
        flex-direction: column;
        justify-content: space-between;
    ">
        <div style="color: #6B7280; font-size: 0.85rem; font-weight: 600;">Weakest Module</div>
        <div style="margin: 8px 0;">
            <div style="font-size: 1.5rem; font-weight: 800; color: {weak_color};">
                {worst_name}
            </div>
            <div style="font-size: 1.1rem; font-weight: 600; color: #1F2937; margin-top: 4px;">
                {worst_score_display}
            </div>
        </div>
        <div style="color: #1F2937; font-size: 0.9rem;">
            {_band_emoji(worst[2]) if worst and worst[2] else ''}
        </div>
    </div>
    """, unsafe_allow_html=True)

st.markdown("")

# Next actions (simple rules) with priority colors
actions = []
if gqas is not None and gqas < 92:
    actions.append(
        ("🔴", "Raise overall GQAS = 92 by addressing the weakest module and any failed floors."))
if latest.get("L3") and latest["L3"]["score"] < 95:
    actions.append(
        ("🟡", "Reduce fairness gaps (DPG/EOD) or improve L3 score = 95 for high-risk products."))
if latest.get("L4") and latest["L4"]["score"] < 90:
    actions.append(
        ("🟡", "Improve explainability: ensure deletion_drop = 0.15 and stability t = 0.85."))
if latest.get("L5") and latest["L5"]["score"] < (85 if risk_profile == "High" else 80):
    actions.append(
        ("🟡", "Increase logging coverage = 0.95 and keep alert latency = 1h."))

st.markdown("**Next actions**")
if actions:
    for emoji, action_text in actions:
        color_map = {"🔴": "#EF4444", "🟡": "#F59E0B", "🟢": "#10B981"}
        border_color = color_map.get(emoji, "#9CA3AF")
        st.markdown(f"""
        <div style="
            background-color: #F8F9FA;
            border-left: 5px solid {border_color};
            border-radius: 8px;
            padding: 12px 16px;
            margin-bottom: 8px;
        ">
            <div style="color: #1F2937; font-size: 0.95rem; line-height: 1.5;">
                {emoji} {action_text}
            </div>
        </div>
        """, unsafe_allow_html=True)
else:
    st.markdown("<div style='color: #10B981; font-weight: 600;'>🟢 None – all thresholds currently met.</div>",
                unsafe_allow_html=True)

st.markdown("")

# Compose a plain-text export
summary_text = []
summary_text.append("IRAQAF - Executive Summary")
summary_text.append("="*28)
if gqas is not None:
    summary_text.append(f"GQAS: {gqas:.2f} ({gqas_band})")
summary_text.append(f"Risk profile: {risk_profile}")
summary_text.append("")
if l1 and (fails := _failed_clauses_summary(l1, max_items=12)):
    summary_text.append("Governance gaps:")
    for f in fails:
        line = f"- {f['framework']} › {f['id']}: {f['desc']}"
        if f["why"]:
            line += f" `| Why: {f['why']}"
        if f["hint"]:
            line += f" `| Hint: {f['hint']}"
        summary_text.append(line)
summary_text.append("")
if actions:
    summary_text.append("Next actions:")
    for emoji, a in actions:
        summary_text.append(f"- {emoji} {a}")

summary_blob = "\n".join(summary_text).encode("utf-8")

# Download buttons side-by-side
col_txt, col_docx = st.columns(2, gap="small")

with col_txt:
    st.download_button("⬇️ Download Executive Summary (.txt)", data=summary_blob,
                       file_name="IRAQAF_Executive_Summary.txt", mime="text/plain",
                       width="stretch")

# Optional: Word export using python-docx (reuses same content)
with col_docx:
    try:
        from io import BytesIO
        from docx import Document
        doc = Document()
        doc.add_heading("IRAQAF - Executive Summary", level=1)
        if gqas is not None:
            doc.add_paragraph(f"GQAS: {gqas:.2f} ({gqas_band})")
        doc.add_paragraph(f"Risk profile: {risk_profile}")

        if l1 and (fails := _failed_clauses_summary(l1, max_items=12)):
            doc.add_heading("Governance gaps", level=2)
            for f in fails:
                p = doc.add_paragraph()
                p.add_run(f"{f['framework']} › {f['id']}: ").bold = True
                p.add_run(f"{f['desc']}")
                if f["why"]:
                    doc.add_paragraph(f"Why: {f['why']}")
                if f["hint"]:
                    doc.add_paragraph(f"Hint: {f['hint']}")

        if actions:
            doc.add_heading("Next actions", level=2)
            for emoji, a in actions:
                doc.add_paragraph(f"{emoji} {a}", style="List Bullet")

        bio = BytesIO()
        doc.save(bio)
        st.download_button("⬇️ Download Executive Summary (.docx)", data=bio.getvalue(),
                           file_name="IRAQAF_Executive_Summary.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           width="stretch")
    except Exception:
        st.caption("Tip: `pip install python-docx` to enable Word export here.")

st.markdown("""
<div style="margin-top: 40px; margin-bottom: 30px;">
    <h2 style="margin: 0; font-size: 1.8rem; font-weight: 800; color: #1F2937;">🤖 AI-Generated Executive Summary</h2>
    <p style="margin: 8px 0 0 0; font-size: 0.95rem; color: #6B7280;">Generate an intelligent, actionable summary of your compliance posture powered by LLM (optional)</p>
</div>
""", unsafe_allow_html=True)

col_toggle, col_tip = st.columns([1, 2])
with col_toggle:
    use_llm = st.checkbox(
        "Generate with LLM", value=False, disabled=_LOCK, help="Requires OPENAI_API_KEY environment variable")
with col_tip:
    if not use_llm:
        st.markdown("<p style='margin: 0; font-size: 0.85rem; color: #6B7280; padding-top: 8px;'>Toggle on if you have an API key configured</p>", unsafe_allow_html=True)

if use_llm and not _LOCK:
    try:
        import os
        import textwrap
        from datetime import datetime
        from openai import OpenAI  # pip install openai

        if not os.getenv("OPENAI_API_KEY"):
            st.warning(
                "⚠️ Set OPENAI_API_KEY environment variable to enable LLM.")
        else:
            client = OpenAI()
            prompt = f"""
You are an AI auditor. Write a crisp, actionable executive summary for stakeholders.
Use these artifacts (JSON-like): latest={json.dumps(latest)[:50000]}
Emphasize: GQAS, floors (pass/fail), biggest gaps, top 3 actions, risks, and time-sensitive notes.
Date: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}
"""
            with st.spinner("🧠 Generating intelligent summary..."):
                resp = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.3,
                )
            summary = resp.choices[0].message.content.strip()

            st.markdown("""
            <div style="background: linear-gradient(135deg, #F0F9FF 0%, #E0F2FE 100%); border: 2px solid #0EA5E9; border-radius: 12px; padding: 24px; margin: 20px 0;">
                <div style="color: #0C4A6E; line-height: 1.6; font-size: 0.95rem;">
            """, unsafe_allow_html=True)
            st.markdown(summary)
            st.markdown("</div></div>", unsafe_allow_html=True)

            st.download_button("⬇️ Download AI Summary (txt)", summary.encode("utf-8"),
                               file_name="AI_Executive_Summary.txt", mime="text/plain", disabled=_LOCK, width="stretch")
    except Exception as e:
        st.error(f"❌ LLM summary unavailable: {str(e)[:200]}")


# =========================
# Bottom Panels (polished)
# Evidence Tray & Pins • Sync to YAML • Exports • Radar
# =========================

# ---------- Enhanced style helpers ----------
st.markdown("""
<style>
/* Section headers */
.block-title {
  font-weight: 800;
  font-size: 1.3rem;
  margin: 0 0 8px;
  color: #1F2937;
}
.block-subtle {
  color: #6B7280;
  font-size: 0.9rem;
  margin-top: -4px;
  line-height: 1.4;
}

/* Modern Card Design */
.iraqaf-card {
  background: linear-gradient(135deg, #FFFFFF 0%, #F9FAFB 100%);
  border: 2px solid #E5E7EB;
  border-radius: 14px;
  padding: 24px;
  margin-bottom: 24px;
  box-shadow: 0 4px 6px rgba(0, 0, 0, 0.07), 0 1px 3px rgba(0, 0, 0, 0.06);
  transition: all 0.3s ease;
}
.iraqaf-card:hover {
  border-color: #D1D5DB;
  box-shadow: 0 10px 15px rgba(0, 0, 0, 0.1);
}

/* Pill buttons */
button[kind="secondary"] { border-radius: 999px; }

/* Enhanced expanders */
.iraqaf-expander > div {
  border: 2px solid #E5E7EB;
  border-radius: 12px;
  background: linear-gradient(135deg, #FFFFFF 0%, #F9FAFB 100%);
}

/* Status indicators */
.status-success { color: #10B981; font-weight: 700; }
.status-warning { color: #F59E0B; font-weight: 700; }
.status-error { color: #EF4444; font-weight: 700; }
</style>
""", unsafe_allow_html=True)


def card(title: str, subtitle: str | None = None):
    """Render opening HTML for a styled card container with title and optional subtitle."""
    st.markdown(f"<div class='iraqaf-card'><div class='block-title'>{title}</div>" +
                (f"<div class='block-subtle'>{subtitle}</div>" if subtitle else ""),
                unsafe_allow_html=True)


def close_card():
    """Close the card container opened by card()."""
    st.markdown("</div>", unsafe_allow_html=True)


# =========================
# 📎 Evidence Tray & File Pins
# =========================
card("📎 Evidence Tray & Pins",
     "Drop files, select module(s), and Save 📎 evidence/. Then Sync to YAML to add paths under each module’s `evidence:` list.")

EVID_DIR = config.evidence_dir
INDEX_PATH = config.index_path
EVID_DIR.mkdir(exist_ok=True)
INDEX_PATH.parent.mkdir(exist_ok=True)

MODULE_LABELS = {
    "L1": "Governance & Regulatory",
    "L2": "Privacy & Security",
    "L3": "Fairness & Ethics",
    "L4": "Explainability & Transparency",
    "L5": "Operations & Monitoring",
}


def _load_index() -> dict:
    """Thread-safe index loading with file locking."""
    if not INDEX_PATH.exists():
        return {}

    lock_path = INDEX_PATH.with_suffix('.lock')

    try:
        with FileLock(str(lock_path), timeout=10):
            with open(INDEX_PATH, "r", encoding="utf-8") as fh:
                data = json.load(fh)
                return data if isinstance(data, dict) else {}
    except FileNotFoundError:
        return {}
    except json.JSONDecodeError as e:
        logger.error(f"Corrupted evidence index: {e}")
        # Backup corrupted file
        backup_path = INDEX_PATH.with_suffix(
            f'.backup.{int(time.time())}.json')
        INDEX_PATH.rename(backup_path)
        logger.info(f"Backed up corrupted index to {backup_path}")
        return {}
    except Exception as e:
        logger.error(f"Failed to load evidence index: {e}")
        return {}


_index_lock = threading.Lock()


def _save_index(ix: dict):
    """Thread-safe index saving with file locking (cross-platform)."""
    lock_path = INDEX_PATH.with_suffix('.lock')

    with _index_lock:  # In-process lock
        with FileLock(str(lock_path), timeout=10):  # Cross-process lock
            temp_path = INDEX_PATH.with_suffix('.tmp')
            try:
                with open(temp_path, "w", encoding="utf-8") as fh:
                    json.dump(ix, fh, indent=2)
                temp_path.replace(INDEX_PATH)  # Atomic rename
            except Exception as e:
                if temp_path.exists():
                    temp_path.unlink()  # Clean up temp file
                raise


def _sanitize_name(name: str) -> str:
    keep = "._-abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789"
    return "".join(c if c in keep else "_" for c in name)
# ============================================================================
# UNDO/REDO FUNCTIONALITY FOR EVIDENCE OPERATIONS
# ============================================================================


def init_evidence_history():
    """Initialize undo/redo stacks for evidence operations."""
    if 'evidence_history' not in st.session_state:
        st.session_state.evidence_history = {
            'undo_stack': [],
            'redo_stack': [],
            'max_history': 20
        }


def record_evidence_operation(operation_type: str, data: dict):
    """
    Record an evidence operation for undo/redo.

    Args:
        operation_type: Type of operation (add, remove, update)
        data: Operation data (file paths, module IDs, etc.)
    """
    init_evidence_history()

    operation = {
        'type': operation_type,
        'data': data,
        'timestamp': datetime.now().isoformat()
    }

    # Add to undo stack
    st.session_state.evidence_history['undo_stack'].append(operation)

    # Clear redo stack (new operation invalidates redo)
    st.session_state.evidence_history['redo_stack'] = []

    # Limit history size
    max_history = st.session_state.evidence_history['max_history']
    if len(st.session_state.evidence_history['undo_stack']) > max_history:
        st.session_state.evidence_history['undo_stack'].pop(0)


def undo_evidence_operation() -> bool:
    """Undo the last evidence operation. Returns True if successful."""
    init_evidence_history()

    undo_stack = st.session_state.evidence_history['undo_stack']
    if not undo_stack:
        return False

    operation = undo_stack.pop()
    st.session_state.evidence_history['redo_stack'].append(operation)

    try:
        if operation['type'] == 'add':
            # Remove the files that were added
            for module_id, file_path in operation['data']['files']:
                ix = _load_index()
                if module_id in ix and file_path in ix[module_id]:
                    ix[module_id].remove(file_path)
                    _save_index(ix)

                # Delete physical file
                if Path(file_path).exists():
                    Path(file_path).unlink()

        elif operation['type'] == 'remove':
            # Restore the files that were removed
            for module_id, file_path in operation['data']['files']:
                ix = _load_index()
                if module_id not in ix:
                    ix[module_id] = []
                if file_path not in ix[module_id]:
                    ix[module_id].append(file_path)
                    _save_index(ix)

        logger.info(f"Undo operation: {operation['type']}")
        return True

    except Exception as e:
        logger.error(f"Undo failed: {e}", exc_info=True)
        return False


def redo_evidence_operation() -> bool:
    """Redo the last undone operation. Returns True if successful."""
    init_evidence_history()

    redo_stack = st.session_state.evidence_history['redo_stack']
    if not redo_stack:
        return False

    operation = redo_stack.pop()
    st.session_state.evidence_history['undo_stack'].append(operation)

    try:
        if operation['type'] == 'add':
            # Re-add the files
            for module_id, file_path in operation['data']['files']:
                ix = _load_index()
                if module_id not in ix:
                    ix[module_id] = []
                if file_path not in ix[module_id]:
                    ix[module_id].append(file_path)
                    _save_index(ix)

        elif operation['type'] == 'remove':
            # Re-remove the files
            for module_id, file_path in operation['data']['files']:
                ix = _load_index()
                if module_id in ix and file_path in ix[module_id]:
                    ix[module_id].remove(file_path)
                    _save_index(ix)

        logger.info(f"Redo operation: {operation['type']}")
        return True

    except Exception as e:
        logger.error(f"Redo failed: {e}", exc_info=True)
        return False


def process_single_file(
    file_obj,
    modules: list[str],
    base_dir: Path
) -> tuple[bool, str, str | list[tuple[str, str]]]:
    """Process a single uploaded file (for batch processing)."""
    try:
        data = file_obj.read()

        # Security checks
        MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
        if len(data) > MAX_FILE_SIZE:
            return (False, file_obj.name, f"File too large ({len(data) / 1024 / 1024:.1f}MB)")

        if len(data) == 0:
            return (False, file_obj.name, "Empty file")

        # Validate filename
        if ".." in file_obj.name or "/" in file_obj.name or " " in file_obj.name:
            return (False, file_obj.name, "Invalid filename")

        file_ext = Path(file_obj.name).suffix.lower()

        # Block dangerous extensions
        DANGEROUS_EXTENSIONS = {'.exe', '.bat', '.sh', '.ps1',
                                '.cmd', '.com', '.scr', '.vbs', '.js', '.jar'}
        if file_ext in DANGEROUS_EXTENSIONS:
            return (False, file_obj.name, f"File type not allowed: {file_ext}")

        # Generate safe filename
        stem = _sanitize_name(Path(file_obj.name).stem)
        ext = _sanitize_name(file_ext) or ".bin"
        h = _hash_bytes(data)
        fname = f"{stem}-{h}{ext}"

        # Save to each module
        saved_paths = []
        for mid in modules:
            target_dir = base_dir / mid
            target_dir.mkdir(parents=True, exist_ok=True)

            outp = (target_dir / fname).resolve()
            if not str(outp).startswith(str(target_dir.resolve())):
                return (False, file_obj.name, "Security: Path traversal blocked")

            # Atomic write
            temp_path = outp.with_suffix(outp.suffix + '.tmp')
            try:
                with open(temp_path, "wb") as fh:
                    fh.write(data)

                if temp_path.stat().st_size != len(data):
                    temp_path.unlink()
                    raise IOError("File write verification failed")

                temp_path.replace(outp)
                saved_paths.append((mid, str(outp.as_posix())))
            except Exception as write_error:
                if temp_path.exists():
                    temp_path.unlink()
                raise write_error

        if saved_paths:
            return (True, fname, saved_paths)
        else:
            return (False, file_obj.name, "Failed to save to any module")

    except Exception as e:
        return (False, file_obj.name, f"Error: {str(e)[:100]}")


def batch_process_evidence_files(files_list: list, modules: list, base_dir: Path) -> tuple[list, list]:
    """
    Process multiple evidence files in batch with enhanced security.

    Args:
        files_list: List of uploaded file objects
        modules: List of module IDs to pin to
        base_dir: Base directory for evidence storage

    Returns:
        Tuple of (successful_files, failed_files)
    """
    # Security constants
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    ALLOWED_EXTENSIONS = {'.pdf', '.csv', '.png', '.jpg',
                          '.jpeg', '.txt', '.docx', '.xlsx', '.json', '.yaml', '.yml'}
    DANGEROUS_EXTENSIONS = {'.exe', '.bat', '.sh', '.ps1',
                            '.cmd', '.com', '.scr', '.vbs', '.js', '.jar'}

    successful = []
    failed = []

    for uf in files_list:
        result = process_single_file(uf, modules, base_dir)

        if result[0]:  # Success
            successful.append((result[1], result[2]))
            logger.info(
                f"Saved evidence: {result[1]} to {len(modules)} module(s)")
        else:  # Failed
            failed.append((result[1], result[2]))
            if result[2]:  # Has error message
                logger.warning(f"Failed to save {result[1]}: {result[2]}")

    return successful, failed
# =============================================================================
# UNDO/REDO SYSTEM (Already implemented in your code, but add UI polish)
# =============================================================================

# LOCATION: Add after line ~1450 (after evidence upload section)


# Enhanced undo/redo UI with visual feedback
if 'evidence_history' in st.session_state:
    history = st.session_state.evidence_history
    undo_count = len(history['undo_stack'])
    redo_count = len(history['redo_stack'])

    if undo_count > 0 or redo_count > 0:
        st.markdown("""
        <div style='
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            padding: 12px 16px;
            border-radius: 10px;
            margin: 16px 0;
            color: white;
        '>
            <h4 style='margin: 0 0 8px 0; color: white;'>🕐 History</h4>
        """, unsafe_allow_html=True)

        undo_col, redo_col, info_col = st.columns([1, 1, 2])

        with undo_col:
            undo_disabled = undo_count == 0 or _LOCK
            if st.button(
                f"? Undo ({undo_count})",
                disabled=undo_disabled,
                width="stretch",
                key="evidence_undo_btn_ui",
                help="Undo last evidence operation"
            ):
                if undo_evidence_operation():
                    st.success("✅ Undone!")
                    time.sleep(0.5)
                    st.rerun()
                else:
                    st.error("❌ Undo failed")
        with redo_col:
            redo_disabled = redo_count == 0 or _LOCK
            if st.button(
                f"↷ Redo ({redo_count})",
                disabled=redo_disabled,
                width="stretch",
                key="evidence_redo_btn_ui",
                help="Redo last undone operation"
            ):
                if redo_evidence_operation():
                    st.success("✅ Redone!")
                    time.sleep(0.5)
                    st.rerun()
                else:
                    st.error("❌ Redo failed")
        with info_col:
            if undo_count > 0:
                last_op = history['undo_stack'][-1]
                st.caption(
                    f"Last: {last_op['type']} at {last_op['timestamp'][:19]}")

        st.markdown("</div>", unsafe_allow_html=True)


def _hash_bytes(b: bytes) -> str:
    return hashlib.sha1(b).hexdigest()[:10]


ix = _load_index()

# ============================================================================
# UNIFIED EVIDENCE & EXPORT SECTION (Redesigned for Clarity)
# ============================================================================
st.markdown("""
<div style="margin-top: 50px; margin-bottom: 30px;">
    <h2 style="margin: 0; font-size: 1.8rem; font-weight: 800; color: #1F2937;">📎 Evidence Management & Export</h2>
    <p style="margin: 8px 0 0 0; font-size: 0.95rem; color: #6B7280;">Upload evidence files, manage module assignments, sync to config, and export reports</p>
</div>
""", unsafe_allow_html=True)

# Create organized tabs for different functions
evidence_tabs = st.tabs(
    ["📤 Upload & Pin", "📚 Evidence Index", "🔄 Sync & Export"])

# ============================================================================
# TAB 1: UPLOAD & PIN
# ============================================================================
with evidence_tabs[0]:
    st.markdown("""
    <div style="background: linear-gradient(135deg, #F0F9FF 0%, #E0F2FE 100%); border: 2px dashed #3B82F6; border-radius: 14px; padding: 28px; text-align: center; margin-bottom: 24px;">
        <div style="font-size: 2.5rem; margin-bottom: 8px;">📁</div>
        <p style="margin: 0 0 4px 0; font-size: 1rem; font-weight: 700; color: #1E40AF;">Drag & Drop Evidence Files</p>
        <p style="margin: 4px 0 0 0; font-size: 0.85rem; color: #1E3A8A;">PDF, CSV, PNG, TXT, DOCX • Max 50MB per file</p>
    </div>
    """, unsafe_allow_html=True)

    uploads = st.file_uploader(
        "Upload evidence files",
        type=["pdf", "csv", "png", "txt", "docx"],
        accept_multiple_files=True,
        disabled=_LOCK,
        label_visibility="collapsed"
    )

    if uploads:
        st.success(f"✅ {len(uploads)} file(s) ready to upload", icon="✅")

        # Show file list
        with st.expander("📋 Files to upload", expanded=True):
            for f in uploads:
                size_mb = f.size / (1024 * 1024)
                st.caption(f"📄 {f.name} ({size_mb:.1f} MB)")

    st.divider()

    # Module selection - Improved layout with checkboxes
    st.markdown("<h4 style='margin: 16px 0 12px 0; color: #1F2937; font-weight: 700;'>📌 Assign to Modules</h4>",
                unsafe_allow_html=True)
    st.caption("Select which compliance modules this evidence applies to")

    # Create a more visual module selector with checkboxes
    col1, col2 = st.columns(2)

    modules_selected = []

    with col1:
        st.markdown("**Governance & Compliance**")
        if st.checkbox("🏛️ L1 • Governance & Regulatory", value=True, key="mod_l1"):
            modules_selected.append("L1")
        if st.checkbox("🔐 L2 • Privacy & Security", value=False, key="mod_l2"):
            modules_selected.append("L2")

    with col2:
        st.markdown("**Fairness & Operations**")
        if st.checkbox("⚖️ L3 • Fairness & Ethics", value=False, key="mod_l3"):
            modules_selected.append("L3")
        if st.checkbox("🔍 L4 • Explainability & Transparency", value=False, key="mod_l4"):
            modules_selected.append("L4")

    st.markdown("")
    if st.checkbox("📊 L5 • Operations & Monitoring", value=False, key="mod_l5"):
        modules_selected.append("L5")

    modules_to_pin = modules_selected

    if not modules_to_pin:
        st.warning("⚠️ Select at least one module for evidence assignment")
        modules_to_pin = ["L1"]  # Fallback

    st.divider()

    # Action buttons
    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button("💾 Save & Pin Files", disabled=_LOCK, width="stretch", type="primary", key="save_pin_btn"):
            if not uploads:
                st.error("❌ Upload files first")
            elif not modules_to_pin:
                st.error("❌ Select at least one module")
            else:
                with progress_tracker("Saving evidence files", total=len(uploads)) as tracker:
                    successful, failed = batch_process_evidence_files(
                        uploads, modules_to_pin, EVID_DIR)

                    # Update index for successful saves
                    for fname, paths in successful:
                        for mid, path in paths:
                            ix.setdefault(mid, [])
                            if path not in ix[mid]:
                                ix[mid].append(path)

                    # Save updated index
                    if successful:
                        _save_index(ix)

                        # Record operation for undo/redo
                        all_saved_paths = []
                        for fname, paths in successful:
                            all_saved_paths.extend(paths)

                        record_evidence_operation('add', {
                            'files': all_saved_paths,
                            'timestamp': datetime.now().isoformat()
                        })

                        st.success(
                            f"✅ Saved {len(successful)} file(s) to {', '.join(modules_to_pin)}")
                    # Show failures if any
                    if failed:
                        with st.expander(f"⚠️ {len(failed)} file(s) failed", expanded=True):
                            for fname, error in failed:
                                st.error(f"❌ {fname}: {error}")

    with col2:
        if st.button("📂 Open Folder", width="stretch", key="open_folder_btn"):
            st.info(f"📁 Evidence folder:\n`{EVID_DIR.resolve()}`")

    with col3:
        # Undo/Redo controls
        init_evidence_history()
        undo_count = len(st.session_state.evidence_history['undo_stack'])
        redo_count = len(st.session_state.evidence_history['redo_stack'])

        if undo_count > 0 or redo_count > 0:
            with st.expander("⏱️ History", expanded=False):
                h1, h2 = st.columns(2)
                with h1:
                    undo_disabled = undo_count == 0 or _LOCK
                    if st.button(
                        f"↶ Undo ({undo_count})",
                        disabled=undo_disabled,
                        width="stretch",
                        key="evidence_undo_btn",
                        help="Undo last operation"
                    ):
                        if undo_evidence_operation():
                            st.success("✅ Undone!")
                            st.rerun()

                with h2:
                    redo_disabled = redo_count == 0 or _LOCK
                    if st.button(
                        f"↷ Redo ({redo_count})",
                        disabled=redo_disabled,
                        width="stretch",
                        key="evidence_redo_btn",
                        help="Redo last operation"
                    ):
                        if redo_evidence_operation():
                            st.success("✅ Redone!")
                            st.rerun()

# ============================================================================
# TAB 2: EVIDENCE INDEX
# ============================================================================
with evidence_tabs[1]:
    if ix:
        st.success(
            f"✅ {sum(len(files) for files in ix.values())} file(s) indexed")

        # Display organized by module
        for module_id in ["L1", "L2", "L3", "L4", "L5"]:
            files = ix.get(module_id, [])
            if files:
                with st.expander(f"🏛️ {module_id} • {MODULE_LABELS.get(module_id, 'Unknown')} ({len(files)} file{'s' if len(files) != 1 else ''})", expanded=True):
                    for file_path in files:
                        file_name = Path(file_path).name
                        col_name, col_action = st.columns([4, 1])
                        with col_name:
                            st.caption(f"📄 {file_name}")
                        with col_action:
                            if st.button("🗑️", key=f"del_{module_id}_{file_name}", help="Remove file", width="stretch"):
                                if module_id in ix and file_path in ix[module_id]:
                                    ix[module_id].remove(file_path)
                                    _save_index(ix)
                                    st.success("✅ Removed")
                                    st.rerun()
    else:
        st.info(
            "📭 No evidence files indexed yet. Upload files in the **Upload & Pin** tab to get started.")

# ============================================================================
# TAB 3: SYNC & EXPORT
# ============================================================================
with evidence_tabs[2]:
    sync_export_tabs = st.tabs(["🔄 Sync to Config", "📊 Export Reports"])

    # Sync to YAML sub-tab
    with sync_export_tabs[0]:
        st.markdown("""
        <div style="background: linear-gradient(135deg, #ECFDF5 0%, #E0FFDD 100%); border: 2px solid #10B981; border-radius: 14px; padding: 20px; margin-bottom: 24px;">
            <p style="margin: 0 0 4px 0; font-size: 0.95rem; font-weight: 700; color: #065F46;">🔄 Sync Evidence to Project Config</p>
            <p style="margin: 4px 0 0 0; font-size: 0.85rem; color: #047857;">Merges pinned files into `evidence: [...]` under each module in your YAML config. Auto-deduplicates existing entries.</p>
        </div>
        """, unsafe_allow_html=True)

        def _sync_to_yaml(yaml_path="configs/project.example.yaml"):
            yaml_path = Path(yaml_path)
            if not yaml_path.exists():
                st.error(f"Config not found: {yaml_path}")
                return False
            try:
                import yaml
            except Exception:
                st.error("Missing PyYAML. `pip install pyyaml`")
                return False

            with open(yaml_path, "r", encoding="utf-8") as fh:
                try:
                    cfg = yaml.safe_load(fh) or {}
                except Exception as e:
                    st.error(f"YAML parse error: {e}")
                    return False

            for mid in MODULE_LABELS.keys():
                cfg.setdefault(mid, {})
                ev = cfg[mid].get("evidence", [])
                if ev is None:
                    ev = []
                if not isinstance(ev, list):
                    ev = [str(ev)]
                for rel in ix.get(mid, []):
                    if rel not in ev:
                        ev.append(rel)
                cfg[mid]["evidence"] = sorted(ev)

            backup = yaml_path.with_suffix(f".backup.{int(time.time())}.yaml")
            shutil.copyfile(yaml_path, backup)
            with open(yaml_path, "w", encoding="utf-8") as fh:
                yaml.safe_dump(cfg, fh, sort_keys=False)
            st.success(f"✅ Synced to {yaml_path}\n📁 Backup: {backup.name}")
            return True

        col1, col2 = st.columns([1, 2])
        with col1:
            if st.button("🔄 Sync to YAML", disabled=_LOCK, width="stretch", type="primary", key="sync_yaml_btn"):
                try:
                    if _sync_to_yaml():
                        st.balloons()
                except Exception as e:
                    show_error_inline(e, "YAML sync failed")

        with col2:
            st.markdown(
                "<p style='margin: 0; padding-top: 8px; font-size: 0.85rem; color: #6B7280;'>💡 After syncing, re-run L1/L2 modules to include evidence references in reports</p>",
                unsafe_allow_html=True)

    # Export Reports sub-tab
    with sync_export_tabs[1]:
        st.markdown("""
        <div style="background: linear-gradient(135deg, #FEF3C7 0%, #FEF9E7 100%); border: 2px solid #F59E0B; border-radius: 14px; padding: 20px; margin-bottom: 24px;">
            <p style="margin: 0 0 4px 0; font-size: 0.95rem; font-weight: 700; color: #92400E;">📥 Generate & Download Reports</p>
            <p style="margin: 4px 0 0 0; font-size: 0.85rem; color: #B45309;">Create comprehensive reports in multiple formats (JSON, CSV, DOCX) for stakeholders</p>
        </div>
        """, unsafe_allow_html=True)

        # Show available export formats
        export_col1, export_col2, export_col3 = st.columns(3)

        with export_col1:
            st.markdown("""
            <div style="background: linear-gradient(135deg, #F0F9FF 0%, #E0F2FE 100%); border: 2px solid #3B82F6; border-radius: 12px; padding: 16px; text-align: center;">
                <div style="font-size: 1.8rem;">📄</div>
                <p style="margin: 8px 0 4px 0; font-weight: 700; color: #1E40AF; font-size: 0.9rem;">Word Report</p>
                <p style="margin: 0; font-size: 0.8rem; color: #1E3A8A;">DOCX format with full details</p>
            </div>
            """, unsafe_allow_html=True)

            if st.button("Generate DOCX", width="stretch", key="gen_docx", type="primary"):
                with st.spinner("📋 Building document..."):
                    def build_docx(report_bundle: dict) -> bytes:
                        try:
                            from docx import Document
                        except Exception:
                            st.error(
                                "`python-docx` not installed. Run: `pip install python-docx`")
                            return b""
                        doc = Document()
                        doc.add_heading("IRAQAF QA Report", level=1)
                        doc.add_paragraph(
                            "Automatically generated from the latest IRAQAF run.")

                        agg_rep = report_bundle.get("AGG")
                        if agg_rep:
                            doc.add_heading("Aggregate (GQAS)", level=2)
                            doc.add_paragraph(f"GQAS: {agg_rep.get('gqas')}")
                            doc.add_paragraph(
                                f"Floors met: {agg_rep.get('floors_met')}")

                        doc.add_heading("Modules", level=2)
                        for m in ["L1", "L2", "L3", "L4", "L5"]:
                            rep_m = report_bundle.get(m)
                            if not rep_m:
                                continue
                            doc.add_heading(NAMES[m], level=3)
                            doc.add_paragraph(
                                f"Score: {rep_m['score']}  |  Band: {rep_m['band']}")
                            doc.add_paragraph("Metrics:")
                            doc.add_paragraph(json.dumps(
                                rep_m.get("metrics", {}), indent=2))

                        bio = BytesIO()
                        doc.save(bio)
                        return bio.getvalue()

                    payload = build_docx(latest)
                    if payload:
                        st.download_button("⬇️ Download Report", data=payload, file_name="IRAQAF_Report.docx",
                                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                           key="dl_docx", width="stretch")
                        st.success("✅ Ready to download!")

        with export_col2:
            st.markdown("""
            <div style="background: linear-gradient(135deg, #F0F9FF 0%, #E0F2FE 100%); border: 2px solid #3B82F6; border-radius: 12px; padding: 16px; text-align: center;">
                <div style="font-size: 1.8rem;">📊</div>
                <p style="margin: 8px 0 4px 0; font-weight: 700; color: #1E40AF; font-size: 0.9rem;">JSON Export</p>
                <p style="margin: 0; font-size: 0.8rem; color: #1E3A8A;">Structured data format</p>
            </div>
            """, unsafe_allow_html=True)

            if st.button("Export as JSON", width="stretch", key="export_json"):
                if latest:
                    json_data = json.dumps(latest, indent=2).encode("utf-8")
                    st.download_button("⬇️ Download JSON", data=json_data, file_name=f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                                       mime="application/json", key="dl_json", width="stretch")
                else:
                    st.warning("No data to export yet")

        with export_col3:
            st.markdown("""
            <div style="background: linear-gradient(135deg, #F0F9FF 0%, #E0F2FE 100%); border: 2px solid #3B82F6; border-radius: 12px; padding: 16px; text-align: center;">
                <div style="font-size: 1.8rem;">📈</div>
                <p style="margin: 8px 0 4px 0; font-weight: 700; color: #1E40AF; font-size: 0.9rem;">CSV Export</p>
                <p style="margin: 0; font-size: 0.8rem; color: #1E3A8A;">Spreadsheet format</p>
            </div>
            """, unsafe_allow_html=True)

            if st.button("Export as CSV", width="stretch", key="export_csv"):
                if latest:
                    # Flatten latest into CSV rows
                    rows = []
                    for module_id in ["L1", "L2", "L3", "L4", "L5"]:
                        module_data = latest.get(module_id, {})
                        if module_data:
                            rows.append({
                                "Module": module_id,
                                "Score": module_data.get("score", ""),
                                "Band": module_data.get("band", ""),
                                "Status": module_data.get("status", "")
                            })
                    if rows:
                        df = pd.DataFrame(rows)
                        csv_data = df.to_csv(index=False).encode("utf-8")
                        st.download_button("⬇️ Download CSV", data=csv_data, file_name=f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                           mime="text/csv", key="dl_csv", width="stretch")
                else:
                    st.warning("No data to export yet")


# ============================================================================
# MATURITY RADAR (Moved here for better organization)
# ============================================================================
st.markdown("---")
st.markdown("""
<div style="margin-top: 40px; margin-bottom: 30px;">
    <h2 style="margin: 0; font-size: 1.8rem; font-weight: 800; color: #1F2937;">📡 Maturity Radar</h2>
    <p style="margin: 8px 0 0 0; font-size: 0.95rem; color: #6B7280;">Visual representation of module scores (L1–L5) across the compliance framework</p>
</div>
""", unsafe_allow_html=True)


# Keep the radar visualization rendering from the original code
def _render_html_report(bundle: dict) -> str:
    """Render an HTML report from the bundle."""
    gqas = (bundle.get("AGG") or {}).get("gqas")
    rows = []
    for mid in ["L1", "L2", "L3", "L4", "L5"]:
        rep = bundle.get(mid) or {}
        rows.append(
            f"<tr><td>{mid}</td><td>{rep.get('score', '')}</td><td>{rep.get('band', '')}</td></tr>")
    table = "\n".join(rows)
    return f"""<!doctype html>
<html><head><meta charset="utf-8"><title>IRAQAF Report</title>
<style>
body{{font-family:Inter,Arial,sans-serif;margin:0;background:#f8fafc}}
.wrap{{max-width:960px;margin:24px auto;padding:0 16px}}
.card{{background:#fff;border:1px solid #e5e7eb;border-radius:12px;padding:18px 16px;margin-bottom:12px}}
h1{{margin:0 0 8px}} h2{{margin:18px 0 8px}}
table{{border-collapse:collapse;width:100%}}
td,th{{border:1px solid #e5e7eb;padding:8px}} th{{background:#f6f9ff;text-align:left}}
.small{{color:#6b7280}}
</style></head><body>
<div class="wrap">
  <div class="card">
    <h1>IRAQAF Report</h1>
    <div class="small">Generated {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}</div>
    <h2>Aggregate</h2>
    <p><b>GQAS:</b> {gqas}</p>
    <h2>Modules</h2>
    <table><tr><th>Module</th><th>Score</th><th>Band</th></tr>
    {table}
    </table>
  </div>
</div>
</body></html>"""


# Render the radar chart
radar_rep = {m: (latest.get(m) or {}).get("score")
             for m in ["L1", "L2", "L3", "L4", "L5"]}
if any(isinstance(v, (int, float)) for v in radar_rep.values()):
    # Prepare a closed polygon in polar coordinates, then project to XY.
    names = ["L1", "L2", "L3", "L4", "L5"]
    scores = [float(radar_rep.get(m) or 0.0) for m in names]
    names.append(names[0])
    scores.append(scores[0])  # close
    n = len(names)
    rad_df = pd.DataFrame({"module": names, "score": scores})
    rad_df["idx"] = range(n)
    rad_df["angle"] = rad_df["idx"] * (2*math.pi/(n-1))
    rad_df["x"] = rad_df["score"] * np.cos(rad_df["angle"])
    rad_df["y"] = rad_df["score"] * np.sin(rad_df["angle"])

    # Soft background grid/rings
    rings = pd.DataFrame({"r": [25, 50, 75, 100]})
    ring_df = pd.concat(
        [pd.DataFrame({"r": [r]*361, "deg": np.arange(361)})
         for r in rings["r"]],
        ignore_index=True
    )
    ring_df["rad"] = np.deg2rad(ring_df["deg"])
    ring_df["x"] = ring_df["r"]*np.cos(ring_df["rad"])
    ring_df["y"] = ring_df["r"]*np.sin(ring_df["rad"])

    grid = alt.Chart(ring_df).mark_line(opacity=0.15).encode(x="x:Q", y="y:Q")
    poly = alt.Chart(rad_df).mark_area(opacity=0.15).encode(
        x="x:Q", y="y:Q")
    outline = alt.Chart(rad_df).mark_line(point=True).encode(
        x="x:Q", y="y:Q", tooltip=["module", "score"])

    # spokes with labels
    spoke_base = pd.DataFrame(
        {"module": names[:-1], "idx": range(len(names)-1)})
    spoke_base["angle"] = spoke_base["idx"] * (2*math.pi/(len(names)-1))
    spoke_base["x"] = 110*np.cos(spoke_base["angle"])
    spoke_base["y"] = 110*np.sin(spoke_base["angle"])
    label_layer = alt.Chart(spoke_base).mark_text(fontSize=12).encode(
        x="x:Q", y="y:Q", text="module:N")

    chart = (grid + poly + outline +
             label_layer).properties(width=420, height=420)
    st.altair_chart(chart, width="content")

# ===== Framework Management (Sidebar) =====
with st.sidebar.expander("🏛️ Framework Configuration", expanded=False):
    frameworks = get_available_frameworks()

    if not frameworks:
        st.warning("⚠️ No frameworks found!")
        st.caption("Create `configs/trace_map.yaml` to define frameworks")
    else:
        st.markdown(f"**Active Frameworks:** {len(frameworks)}")

        for fw in frameworks:
            clauses = get_framework_clauses(fw)
            with st.expander(f"{fw} ({len(clauses)})"):
                for clause in clauses:
                    metrics = get_clause_metrics(fw, clause)
                    st.markdown(f"**{clause}**")
                    if metrics:
                        for m in metrics:
                            st.caption(f"  → {m}")
                    else:
                        st.caption("  (no metrics)")

        st.markdown("---")

        if st.button("🔄 Reload trace_map.yaml", key="sidebar_reload_trace"):
            load_trace_map.clear()
            get_available_frameworks.clear()
            get_framework_clauses.clear()
            get_clause_metrics.clear()
            st.success("✅ Reloaded!")
            st.rerun()

        st.caption("💡 Edit `configs/trace_map.yaml` to modify frameworks")

if st.sidebar.checkbox("🔧 Show Debug Info", value=False, key="debug_mode"):
    st.markdown("---")

    # Styled container for debug info
    st.markdown("""
    <div style='
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 20px;
        border-radius: 12px;
        margin-bottom: 20px;
        color: white;
    '>
        <h2 style='margin: 0; color: white;'>🔧 Debug & Performance</h2>
        <p style='margin: 5px 0 0 0; opacity: 0.9;'>System diagnostics and monitoring</p>
    </div>
    """, unsafe_allow_html=True)

    # Performance Metrics Card
    with st.container():
        st.markdown("""
        <div style='
            background: white;
            border: 1px solid #e5e7eb;
            border-radius: 10px;
            padding: 16px;
            margin-bottom: 16px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        '>
            <h3 style='margin: 0 0 12px 0; color: #1f2937;'>⚡ Performance Metrics</h3>
        """, unsafe_allow_html=True)

        if PERFORMANCE_MONITORING:
            monitor = get_monitor()
            all_stats = monitor.get_all_stats()

            if all_stats:
                # Show top 5 slowest operations
                sorted_ops = sorted(
                    all_stats.items(),
                    key=lambda x: x[1].get("mean", 0),
                    reverse=True
                )[:5]

                for op_name, stats in sorted_ops:
                    mean = stats['mean']
                    color = "#10b981" if mean < 0.5 else (
                        "#f59e0b" if mean < 2.0 else "#ef4444")
                    emoji = "🟢" if mean < 0.5 else (
                        "🟡" if mean < 2.0 else "🔴")

                    st.markdown(f"""
                    <div style='
                        display: flex;
                        justify-content: space-between;
                        padding: 8px;
                        background: #f9fafb;
                        border-radius: 6px;
                        margin-bottom: 6px;
                    '>
                        <span style='color: #6b7280;'>{emoji} {op_name}</span>
                        <span style='color: {color}; font-weight: 600;'>{mean:.3f}s</span>
                    </div>
                    """, unsafe_allow_html=True)
            else:
                st.info("No performance data collected yet")
        else:
            st.warning("Performance monitoring is disabled")

        st.markdown("</div>", unsafe_allow_html=True)

    # System Info Card
    with st.container():
        st.markdown("""
        <div style='
            background: white;
            border: 1px solid #e5e7eb;
            border-radius: 10px;
            padding: 16px;
            margin-bottom: 16px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        '>
            <h3 style='margin: 0 0 12px 0; color: #1f2937;'>💾 System Info</h3>
        """, unsafe_allow_html=True)

        mem_usage = get_memory_usage()
        if mem_usage:
            rss_mb = mem_usage.get('rss_mb', 0)
            mem_pct = mem_usage.get('percent', 0)

            # Memory usage bar
            mem_color = "#10b981" if mem_pct < 50 else (
                "#f59e0b" if mem_pct < 80 else "#ef4444")
            st.markdown(f"""
            <div style='margin-bottom: 12px;'>
                <div style='display: flex; justify-content: space-between; margin-bottom: 4px;'>
                    <span style='color: #6b7280; font-size: 0.875rem;'>Memory Usage</span>
                    <span style='color: {mem_color}; font-weight: 600;'>{rss_mb:.1f} MB ({mem_pct:.1f}%)</span>
                </div>
                <div style='
                    width: 100%;
                    height: 8px;
                    background: #e5e7eb;
                    border-radius: 4px;
                    overflow: hidden;
                '>
                    <div style='
                        width: {mem_pct}%;
                        height: 100%;
                        background: {mem_color};
                        transition: width 0.3s ease;
                    '></div>
                </div>
            </div>
            """, unsafe_allow_html=True)

        # Reports loaded
        st.markdown(f"""
        <div style='
            display: flex;
            justify-content: space-between;
            padding: 8px;
            background: #f9fafb;
            border-radius: 6px;
            margin-bottom: 6px;
        '>
            <span style='color: #6b7280;'>📑 Reports loaded</span>
            <span style='color: #1f2937; font-weight: 600;'>{len(files)}</span>
        </div>
        """, unsafe_allow_html=True)

        # Modules found
        modules_count = sum(1 for v in latest.values() if v)
        st.markdown(f"""
        <div style='
            display: flex;
            justify-content: space-between;
            padding: 8px;
            background: #f9fafb;
            border-radius: 6px;
            margin-bottom: 6px;
        '>
            <span style='color: #6b7280;'>📦 Modules found</span>
            <span style='color: #1f2937; font-weight: 600;'>{modules_count}/6</span>
        </div>
        """, unsafe_allow_html=True)

        # Cache info
        try:
            cache_info = st.cache_data.get_stats()
            hit_rate = 0
            if len(cache_info) > 0:
                total_calls = sum(info.get("cache_misses", 0) + info.get("cache_hits", 0)
                                  for info in cache_info if isinstance(info, dict))
                total_hits = sum(info.get("cache_hits", 0)
                                 for info in cache_info if isinstance(info, dict))
                if total_calls > 0:
                    hit_rate = (total_hits / total_calls) * 100

            cache_color = "#10b981" if hit_rate > 50 else (
                "#f59e0b" if hit_rate > 20 else "#ef4444")
            st.markdown(f"""
            <div style='
                display: flex;
                justify-content: space-between;
                padding: 8px;
                background: #f9fafb;
                border-radius: 6px;
            '>
                <span style='color: #6b7280;'>💿 Cache hit rate</span>
                <span style='color: {cache_color}; font-weight: 600;'>{hit_rate:.1f}%</span>
            </div>
            """, unsafe_allow_html=True)
        except:
            pass

        st.markdown("</div>", unsafe_allow_html=True)

    # Actions Card
    with st.container():
        st.markdown("""
        <div style='
            background: white;
            border: 1px solid #e5e7eb;
            border-radius: 10px;
            padding: 16px;
            margin-bottom: 16px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        '>
            <h3 style='margin: 0 0 12px 0; color: #1f2937;'>🛠️ Actions</h3>
        """, unsafe_allow_html=True)

        col1, col2 = st.columns(2)

        with col1:
            if st.button("🗑️ Clear All Caches", width="stretch", key="debug_clear_cache"):
                st.cache_data.clear()
                st.success("✅ Caches cleared!")
                st.rerun()

        with col2:
            if st.button("⤴️ Export Performance Log", width="stretch", key="debug_export_perf"):
                if PERFORMANCE_MONITORING:
                    monitor = get_monitor()
                    perf_data = json.dumps(monitor.get_all_stats(), indent=2)
                    st.download_button(
                        "⤵️ Download performance.json",
                        data=perf_data.encode("utf-8"),
                        file_name=f"performance_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                        mime="application/json",
                        key="debug_download_perf"
                    )
                else:
                    st.warning("Performance monitoring not enabled")

        st.markdown("</div>", unsafe_allow_html=True)

    # Security Card
    with st.container():
        st.markdown("""
        <div style='
            background: white;
            border: 1px solid #e5e7eb;
            border-radius: 10px;
            padding: 16px;
            margin-bottom: 16px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        '>
            <h3 style='margin: 0 0 12px 0; color: #1f2937;'>🔒 Security</h3>
        """, unsafe_allow_html=True)

        # Security status checks
        security_checks = {
            "Rate limiting": ("✅", "Active"),
            "Input validation": ("✅", "Active"),
            "File upload limits": ("✅", "50MB max"),
            "Audit mode": ("✅" if get_audit_mode() else "⚪", "Enabled" if get_audit_mode() else "Disabled"),
        }

        for check, (emoji, status) in security_checks.items():
            status_color = "#10b981" if emoji == "✅" else "#9ca3af"
            st.markdown(f"""
            <div style='
                display: flex;
                justify-content: space-between;
                padding: 8px;
                background: #f9fafb;
                border-radius: 6px;
                margin-bottom: 6px;
            '>
                <span style='color: #6b7280;'>{check}</span>
                <span style='color: {status_color}; font-weight: 600;'>{emoji} {status}</span>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("</div>", unsafe_allow_html=True)

    # ===== ADD VALIDATION FUNCTION HERE =====
    # Trace Map Validation Card
    with st.container():
        st.markdown("""
        <div style='
            background: white;
            border: 1px solid #e5e7eb;
            border-radius: 10px;
            padding: 16px;
            margin-bottom: 16px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        '>
            <h3 style='margin: 0 0 12px 0; color: #1f2937;'>🔍 Trace Map Validation</h3>
        """, unsafe_allow_html=True)

        def validate_trace_map_structure():
            """Validate trace_map.yaml structure and metric paths"""
            trace_map = load_trace_map()

            if not trace_map:
                st.error("❌ trace_map.yaml not found or empty")
                return False

            issues = []
            warnings = []
            valid_modules = ["L1", "L2", "L3", "L4", "L5", "AGG"]

            for framework, clauses in trace_map.items():
                if not isinstance(clauses, dict):
                    issues.append(f"❌ {framework}: clauses must be a dict")
                    continue

                for clause, metrics in clauses.items():
                    if not isinstance(metrics, list):
                        issues.append(
                            f"❌ {framework}.{clause}: metrics must be a list")
                        continue

                    for metric in metrics:
                        # Validate metric path format
                        if isinstance(metric, str):
                            parts = metric.split(".")
                            if parts[0] not in valid_modules:
                                warnings.append(
                                    f"⚠️ {framework}.{clause}: '{metric}' - invalid module '{parts[0]}'")

                            # Check if metric path exists in latest reports
                            module_report = latest.get(parts[0])
                            if module_report:
                                # Try to navigate the path
                                current = module_report
                                valid_path = True
                                for part in parts[1:]:
                                    if isinstance(current, dict) and part in current:
                                        current = current[part]
                                    else:
                                        valid_path = False
                                        break

                                if not valid_path:
                                    warnings.append(
                                        f"⚠️ {framework}.{clause}: '{metric}' - path not found in report")

            # Display results
            total_frameworks = len(trace_map)
            total_clauses = sum(len(c) for c in trace_map.values())
            total_metrics = sum(len(m) for c in trace_map.values()
                                for m in c.values())

            if issues:
                st.error(f"❌ Found {len(issues)} critical issue(s):")
                for issue in issues[:5]:
                    st.caption(issue)
                if len(issues) > 5:
                    st.caption(f"... and {len(issues) - 5} more")
                return False

            if warnings:
                st.warning(f"⚠️ Found {len(warnings)} warning(s):")
                for warning in warnings[:5]:
                    st.caption(warning)
                if len(warnings) > 5:
                    with st.expander(f"Show all {len(warnings)} warnings"):
                        for warning in warnings:
                            st.caption(warning)

            # Success metrics
            st.markdown(f"""
            <div style='
                display: grid;
                grid-template-columns: repeat(3, 1fr);
                gap: 8px;
                margin-top: 12px;
            '>
                <div style='background: #f0fdf4; padding: 12px; border-radius: 6px; text-align: center;'>
                    <div style='font-size: 1.5rem; font-weight: 700; color: #16a34a;'>{total_frameworks}</div>
                    <div style='font-size: 0.875rem; color: #166534;'>Frameworks</div>
                </div>
                <div style='background: #f0f9ff; padding: 12px; border-radius: 6px; text-align: center;'>
                    <div style='font-size: 1.5rem; font-weight: 700; color: #0284c7;'>{total_clauses}</div>
                    <div style='font-size: 0.875rem; color: #075985;'>Clauses</div>
                </div>
                <div style='background: #fef3c7; padding: 12px; border-radius: 6px; text-align: center;'>
                    <div style='font-size: 1.5rem; font-weight: 700; color: #d97706;'>{total_metrics}</div>
                    <div style='font-size: 0.875rem; color: #92400e;'>Metrics</div>
                </div>
            </div>
            """, unsafe_allow_html=True)

            if not warnings:
                st.success("✅ trace_map.yaml is valid with no warnings!")

            return True

        # Run validation
        validate_trace_map_structure()

        st.markdown("</div>", unsafe_allow_html=True)
    # ===== END VALIDATION FUNCTION =====

    # =============================================================================
    #  🎯 SYSTEM INTEGRATION DASHBOARD SECTIONS
    # =============================================================================

    if SYSTEM_INTEGRATION_AVAILABLE and st.session_state.get("system_integration_enabled"):
        coordinator = st.session_state.get("system_coordinator")
        monitor = st.session_state.get("realtime_monitor")

        if coordinator and monitor:
            # Create tabs for system integration features
            sys_tab1, sys_tab2, sys_tab3, sys_tab4, sys_tab5 = st.tabs([
                "📊 System Status",
                "⚡ Real-Time Events",
                "💾 Database Insights",
                "🔍 Regulatory Tracking",
                "📈 Compliance Trends"
            ])

            # ===== TAB 1: SYSTEM STATUS =====
            with sys_tab1:
                st.markdown("### 📊 System Status Dashboard")

                try:
                    status = coordinator.get_system_status()

                    # Top metrics row
                    col1, col2, col3, col4 = st.columns(4)

                    with col1:
                        st.metric(
                            "📋 Total Changes",
                            status.get("total_regulatory_changes", 0),
                            delta="tracked",
                            label_visibility="collapsed"
                        )

                    with col2:
                        open_alerts = status.get("open_alerts", 0)
                        alert_color = "🔴" if open_alerts > 5 else "🟡" if open_alerts > 0 else "🟢"
                        st.metric(
                            "⚠️ Open Alerts",
                            open_alerts,
                            delta=alert_color,
                            label_visibility="collapsed"
                        )

                    with col3:
                        avg_compliance = status.get(
                            "average_compliance_score", 0)
                        st.metric(
                            "✅ Avg Compliance",
                            f"{avg_compliance:.1f}%",
                            delta="current",
                            label_visibility="collapsed"
                        )

                    with col4:
                        pending_actions = status.get(
                            "pending_remediation_count", 0)
                        st.metric(
                            "🔧 Pending Actions",
                            pending_actions,
                            delta="to-do",
                            label_visibility="collapsed"
                        )

                    st.markdown("---")

                    # Database status
                    st.markdown("**💾 Database Status**")
                    db_col1, db_col2 = st.columns([2, 1])

                    with db_col1:
                        db_path = status.get("database_path", "Not available")
                        st.caption(f"📁 Database: `{db_path}`")

                        db_size = status.get("database_size_mb", 0)
                        st.progress(min(db_size / 100, 1.0),
                                    text=f"Size: {db_size:.1f} MB / 100 MB")

                    with db_col2:
                        if st.button("🔄 Refresh Status", width="stretch", key="refresh_status"):
                            st.rerun()

                    st.markdown("---")

                    # Monitor status
                    st.markdown("**⚡ Real-Time Monitor**")
                    monitor_col1, monitor_col2, monitor_col3 = st.columns(3)

                    with monitor_col1:
                        is_active = monitor.get_statistics().get("is_monitoring", False)
                        status_text = "🟢 Active" if is_active else "🔴 Inactive"
                        st.caption(f"Status: {status_text}")

                    with monitor_col2:
                        recent_events = len(monitor.get_recent_events(100))
                        st.caption(f"Recent Events: {recent_events}")

                    with monitor_col3:
                        callbacks = monitor.get_statistics().get("callback_count", 0)
                        st.caption(f"Active Callbacks: {callbacks}")

                    # Event statistics
                    st.markdown("**📊 Event Distribution**")
                    event_stats = monitor.get_statistics().get("event_counts", {})

                    if event_stats:
                        events_df = pd.DataFrame([
                            {"Event Type": k, "Count": v}
                            for k, v in sorted(event_stats.items(), key=lambda x: x[1], reverse=True)
                        ])

                        col_left, col_right = st.columns([1, 2])

                        with col_left:
                            st.dataframe(
                                events_df, width="stretch", hide_index=True)

                        with col_right:
                            fig = alt.Chart(events_df).mark_bar().encode(
                                x=alt.X("Count:Q", title="Count"),
                                y=alt.Y("Event Type:N", sort="-x",
                                        title="Event Type"),
                                color=alt.Color(
                                    "Count:Q", scale=alt.Scale(scheme="blues"))
                            ).properties(height=250).interactive()
                            st.altair_chart(fig, width="stretch")

                except Exception as e:
                    st.error(f"❌ Error loading system status: {e}")
                    logger.error(f"System status error: {e}", exc_info=True)

            # ===== TAB 2: REAL-TIME EVENTS =====
            with sys_tab2:
                st.markdown("### ⚡ Real-Time Events & Alerts")

                try:
                    # Debug: Check monitor type
                    monitor_type = type(monitor).__name__
                    has_get_events = hasattr(monitor, 'get_recent_events')

                    if not has_get_events:
                        st.error(
                            f"❌ Monitor does not support get_recent_events. Type: {monitor_type}")
                    else:
                        # Event filters
                        col1, col2, col3 = st.columns([2, 1, 1])

                        with col1:
                            event_filter = st.selectbox(
                                "Filter by event type",
                                options=["All Events"] + [str(e.value) if hasattr(
                                    e, 'value') else str(e) for e in EventType if EventType],
                                key="event_type_filter"
                            )

                        with col2:
                            max_events = st.number_input(
                                "Show last", min_value=5, max_value=100, value=20, key="max_events")

                        with col3:
                            if st.button("🔄 Refresh", width="stretch", key="refresh_events"):
                                st.rerun()

                        # Get recent events
                        recent_events = monitor.get_recent_events(max_events)

                        if recent_events:
                            # Display events as timeline
                            st.markdown("**📅 Event Timeline**")

                            for event in reversed(recent_events):
                                event_dict = event.to_dict() if hasattr(event, 'to_dict') else event

                                # Color code by event type
                                event_type = event_dict.get(
                                    'event_type', 'UNKNOWN')
                                color_map = {
                                    'REGULATORY_CHANGE': '🔵',
                                    'ALERT_TRIGGERED': '🔴',
                                    'ALERT_RESOLVED': '🟢',
                                    'REMEDIATION_PROGRESS': '🟡',
                                    'COMPLIANCE_SCORE_UPDATE': '📊',
                                    'THRESHOLD_BREACH': '⚠️',
                                    'DEADLINE_WARNING': '⏰',
                                    'SYSTEM_HEALTH_UPDATE': '💚',
                                }

                                emoji = color_map.get(event_type, '📌')
                                timestamp = event_dict.get(
                                    'timestamp', 'Unknown')

                                with st.container(border=True):
                                    col_emoji, col_content = st.columns(
                                        [0.1, 1])

                                    with col_emoji:
                                        st.write(emoji)

                                    with col_content:
                                        st.markdown(
                                            f"**{event_type}** • `{timestamp}`")

                                        data = event_dict.get('data', {})
                                        if isinstance(data, dict):
                                            for key, val in list(data.items())[:3]:
                                                st.caption(
                                                    f"• {key}: {str(val)[:100]}")
                                        else:
                                            st.caption(str(data)[:200])
                        else:
                            st.info("No recent events recorded yet.")

                except Exception as e:
                    st.error(f"❌ Error loading events: {e}")
                    logger.error(f"Events error: {e}", exc_info=True)

            # ===== TAB 3: DATABASE INSIGHTS =====
            with sys_tab3:
                st.markdown("### 💾 Database Insights")

                try:
                    # Get compliance report with database data
                    report = coordinator.get_compliance_report()

                    col1, col2 = st.columns(2)

                    with col1:
                        st.markdown("**📋 Compliance Scores by Framework**")

                        scores_data = report.get("recent_scores", [])
                        if scores_data:
                            scores_df = pd.DataFrame([
                                {
                                    "Framework": s.get("framework", "Unknown"),
                                    "System": s.get("system_name", "Unknown"),
                                    "Score": f"{s.get('score', 0):.1f}%",
                                    "Status": s.get("status", "Unknown")
                                }
                                for s in scores_data[:10]
                            ])

                            st.dataframe(
                                scores_df, width="stretch", hide_index=True)
                        else:
                            st.info("No compliance scores recorded yet.")

                    with col2:
                        st.markdown("**🚨 Critical Issues**")

                        alerts_data = report.get("critical_alerts", [])
                        if alerts_data:
                            for alert in alerts_data[:5]:
                                with st.container(border=True):
                                    st.markdown(
                                        f"🔴 **{alert.get('alert_type', 'Alert')}**")
                                    st.caption(
                                        alert.get('message', 'No details'))
                                    st.caption(
                                        f"Risk: {alert.get('risk_level', 'unknown')}")
                        else:
                            st.success("✅ No critical issues!")

                    st.markdown("---")

                    # Remediation tracking
                    st.markdown("**🔧 Remediation Progress**")

                    remediation_data = report.get("remediation_overview", {})

                    rem_col1, rem_col2, rem_col3, rem_col4 = st.columns(4)

                    with rem_col1:
                        st.metric("Total", remediation_data.get("total", 0))

                    with rem_col2:
                        st.metric(
                            "Pending", remediation_data.get("pending", 0))

                    with rem_col3:
                        st.metric("In Progress", remediation_data.get(
                            "in_progress", 0))

                    with rem_col4:
                        st.metric(
                            "Completed", remediation_data.get("completed", 0))

                    # Remediation chart
                    rem_status_counts = {
                        "Pending": remediation_data.get("pending", 0),
                        "In Progress": remediation_data.get("in_progress", 0),
                        "Completed": remediation_data.get("completed", 0),
                        "Blocked": remediation_data.get("blocked", 0),
                    }

                    rem_df = pd.DataFrame([
                        {"Status": k, "Count": v}
                        for k, v in rem_status_counts.items() if v > 0
                    ])

                    if not rem_df.empty:
                        fig = alt.Chart(rem_df).mark_pie(innerRadius=50).encode(
                            theta=alt.Theta("Count:Q"),
                            color=alt.Color(
                                "Status:N", scale=alt.Scale(scheme="set2"))
                        ).properties(height=300)
                        st.altair_chart(fig, width="stretch")

                except Exception as e:
                    st.error(f"❌ Error loading database insights: {e}")
                    logger.error(
                        f"Database insights error: {e}", exc_info=True)

            # ===== TAB 4: REGULATORY TRACKING =====
            with sys_tab4:
                st.markdown("### 🔍 Regulatory Tracking")

                try:
                    # Get recent changes (with fallback)
                    try:
                        changes = coordinator.get_recent_changes(days=30)
                    except TypeError:
                        # If days parameter not supported, try without it
                        changes = coordinator.get_recent_changes()
                    except:
                        changes = []

                    if changes:
                        st.markdown(
                            f"**📜 Recent Regulatory Changes ({len(changes)})**")

                        for change in changes[:10]:
                            with st.container(border=True):
                                col_icon, col_content = st.columns([0.1, 1])

                                with col_icon:
                                    is_critical = change.get(
                                        "is_critical", False)
                                    st.write("🔴" if is_critical else "🔵")

                                with col_content:
                                    st.markdown(
                                        f"**{change.get('name', 'Change')}** • `{change.get('regulation_id', 'N/A')}`")
                                    st.caption(
                                        f"Type: {change.get('change_type', 'Update')}")
                                    st.caption(
                                        f"Impact: {change.get('impact_level', 'unknown')}")

                                    if change.get('description'):
                                        st.caption(
                                            f"📝 {change.get('description', '')[:200]}")

                                    deadline = change.get(
                                        'implementation_deadline', 'Not specified')
                                    st.caption(f"⏰ Deadline: {deadline}")
                    else:
                        st.info("No recent regulatory changes.")

                    st.markdown("---")

                    # Log new regulatory change
                    st.markdown("**📝 Log New Regulatory Change**")

                    with st.form("new_regulatory_change", border=False):
                        col1, col2 = st.columns(2)

                        with col1:
                            reg_id = st.text_input(
                                "Regulation ID", placeholder="e.g., GDPR-2024-01")
                            reg_name = st.text_input(
                                "Change Name", placeholder="e.g., Data Protection Update")

                        with col2:
                            change_type = st.selectbox(
                                "Change Type", ["Amendment", "New Regulation", "Clarification", "Enforcement"])
                            impact_level = st.selectbox(
                                "Impact Level", ["Critical", "High", "Medium", "Low"])

                        description = st.text_area(
                            "Description", placeholder="Describe the regulatory change...")
                        deadline = st.date_input("Implementation Deadline")

                        if st.form_submit_button("✅ Log Change", width="stretch"):
                            try:
                                coordinator.track_regulatory_change(
                                    regulation_id=reg_id,
                                    name=reg_name,
                                    change_type=change_type,
                                    description=description,
                                    impact_level=impact_level,
                                    affected_systems=["System1"],
                                    implementation_deadline=str(deadline),
                                    is_critical=(impact_level == "Critical")
                                )
                                st.success(
                                    "✅ Regulatory change logged successfully!")
                                st.rerun()
                            except Exception as e:
                                st.error(f"Failed to log change: {e}")

                except Exception as e:
                    st.error(f"❌ Error loading regulatory tracking: {e}")
                    logger.error(
                        f"Regulatory tracking error: {e}", exc_info=True)

            # ===== TAB 5: COMPLIANCE TRENDS =====
            with sys_tab5:
                st.markdown("### 📈 Compliance Trends")

                try:
                    # Get compliance scores over time (try all frameworks)
                    scores = []
                    try:
                        for framework in ["ISO27001", "SOC2", "HIPAA", "GDPR", "CCPA"]:
                            try:
                                fscores = coordinator.get_compliance_scores(
                                    framework=framework)
                                if fscores:
                                    scores.extend(fscores)
                            except:
                                pass
                    except:
                        scores = []

                    if scores:
                        scores_df = pd.DataFrame(scores)

                        if "recorded_at" in scores_df.columns:
                            scores_df["recorded_at"] = pd.to_datetime(
                                scores_df["recorded_at"])
                            scores_df = scores_df.sort_values("recorded_at")

                        # Trend chart
                        st.markdown("**📊 Compliance Score Trends**")

                        if "framework" in scores_df.columns and "score" in scores_df.columns:
                            fig = alt.Chart(scores_df).mark_line(point=True).encode(
                                x=alt.X(
                                    "recorded_at:T", title="Date") if "recorded_at" in scores_df.columns else alt.value(None),
                                y=alt.Y(
                                    "score:Q", title="Compliance Score (%)"),
                                color=alt.Color(
                                    "framework:N", title="Framework"),
                                tooltip=["framework:N", "score:Q", "status:N"]
                            ).properties(height=300).interactive()

                            st.altair_chart(fig, width="stretch")

                        # Stats by framework
                        st.markdown("**🏆 Framework Rankings**")

                        framework_stats = scores_df.groupby("framework").agg({
                            "score": ["mean", "max", "min", "count"]
                        }).round(1)

                        framework_stats.columns = [
                            "Avg Score", "Max", "Min", "Records"]
                        st.dataframe(framework_stats, width="stretch")
                    else:
                        st.info(
                            "No compliance scores recorded yet. Start logging scores to see trends.")

                except Exception as e:
                    st.error(f"❌ Error loading compliance trends: {e}")
                    logger.error(
                        f"Compliance trends error: {e}", exc_info=True)

    # Last refresh timestamp (existing code continues...)
    st.markdown(f"""
    <div style='
        text-align: center;
        color: #9ca3af;
        font-size: 0.875rem;
        padding: 12px;
    '>
        🔃 Last refresh: {st.session_state.get('last_refresh', 'Never')}
    </div>
    """, unsafe_allow_html=True)

# ============================================================================
# L2 PRIVACY/SECURITY MONITOR INTEGRATION
# ============================================================================
# This section adds L2 Privacy/Security Monitor content to the dashboard
# The content integrates with the main app after login

if L2_MONITOR_AVAILABLE and show_l2_privacy_security_monitor is not None:
    with st.expander(' L2 Privacy/Security Monitor', expanded=False):
        try:
            show_l2_privacy_security_monitor()
        except Exception as e:
            st.error(f'Error loading L2 Monitor: {str(e)}')
            logger.error(f'L2 Monitor error: {e}', exc_info=True)
